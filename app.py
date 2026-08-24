import html
import streamlit as st
import streamlit.components.v1 as components
from datetime import datetime, date, timezone, timedelta
from decimal import Decimal, ROUND_HALF_UP, InvalidOperation
from io import BytesIO
from pathlib import Path
import os
import base64
import io, csv, zipfile
try:
    from pypdf import PdfReader
except Exception:
    PdfReader=None
try:
    from docx import Document as DocxDocument
except Exception:
    DocxDocument=None
import json, re, textwrap, urllib.parse, hashlib, hmac
import pandas as pd
import requests
from supabase import create_client
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import Flow
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseUpload
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.utils import ImageReader
from icalendar import Calendar
import recurring_ical_events

APP_DIR=Path(__file__).parent
ASSET_DIR=APP_DIR/"assets"
st.set_page_config(page_title='NBS Philo Hub', page_icon=str(ASSET_DIR/'philo_affiliates_logo.png'), layout='wide', initial_sidebar_state='collapsed')
def image_data_uri(name):
    p=ASSET_DIR/name
    return "data:image/png;base64,"+base64.b64encode(p.read_bytes()).decode("ascii") if p.exists() else ""
PHILO_LOGO_URI=image_data_uri("philo_affiliates_logo.png")
PHILO_CREST_URI=image_data_uri("philo_crest.png")
CARNATION_LEFT_URI=image_data_uri("carnation_left.png")
CARNATION_RIGHT_URI=image_data_uri("carnation_right.png")
st.markdown(f"""
<style>
:root{{
  --gold:#c99a20;
  --gold2:#e5c35c;
  --gold3:#8f6a08;
  --ivory:#fffdf8;
  --paper:#fffefa;
  --black:#121212;
  --soft:#746a57;
  --line:#d7af43;
  --line-soft:#ead9a9;
}}
.stApp{{
  background:
    linear-gradient(rgba(255,255,255,.94),rgba(255,255,255,.94)),
    url("{CARNATION_LEFT_URI}") left 340px / 260px auto no-repeat fixed,
    url("{CARNATION_RIGHT_URI}") right 340px / 260px auto no-repeat fixed,
    linear-gradient(180deg,#fffefa,#fff 52%,#fdf8ec);
  color:var(--black);
}}
.block-container{{
  max-width:1260px!important;
  padding-top:1.15rem!important;
  padding-bottom:3rem!important;
}}
h1,h2,h3,h4{{
  font-family:Georgia,"Times New Roman",serif!important;
  color:var(--black)!important;
  font-weight:500!important;
}}
p,label,span,div,input,textarea,button{{
  font-family:Georgia,"Times New Roman",serif;
}}
[data-testid="stCaptionContainer"],small{{
  font-family:"Segoe UI",Arial,sans-serif!important;
}}

.philo-header{{
  background:rgba(255,255,255,.91);
  color:var(--black);
  border-bottom:2px solid var(--line);
  border-radius:0;
  padding:18px 18px 22px;
  margin:-18px -18px 26px;
  box-shadow:none;
  display:grid;
  grid-template-columns:220px 1fr 190px;
  align-items:center;
  gap:24px;
  text-align:center;
}}
.philo-header .brand-logo{{
  width:205px;height:205px;object-fit:contain;background:transparent;border-radius:0;padding:0;box-shadow:none;
}}
.philo-header .brand-crest{{
  width:165px;height:165px;object-fit:contain;filter:none;justify-self:center;
}}
.philo-header h1{{
  margin:0;font-size:50px;line-height:1.05;letter-spacing:1.4px;text-transform:uppercase;color:var(--gold)!important;font-weight:500!important;
}}
.philo-header p{{
  margin:18px auto 0;max-width:620px;color:#141414;font-size:20px;line-height:1.45;font-style:italic;
}}
.gold-rule{{
  width:74%;height:1px;background:var(--line);margin:16px auto 0;position:relative;
}}
.gold-rule:after{{
  content:"◇";position:absolute;left:50%;top:50%;transform:translate(-50%,-53%);color:var(--gold);background:#fff;padding:0 8px;font-size:15px;
}}

div[data-testid="stForm"]{{
  background:rgba(255,255,255,.96)!important;
  border:1.5px solid var(--line)!important;
  border-radius:24px!important;
  padding:26px 30px 28px!important;
  box-shadow:0 8px 24px rgba(111,84,15,.07)!important;
  max-width:760px;
  margin-left:auto;
  margin-right:auto;
}}
div[data-testid="stForm"]:before{{
  content:"Welcome to";
  display:block;
  text-align:center;
  color:var(--gold);
  font-family:"Brush Script MT","Segoe Script",cursive;
  font-size:32px;
  margin:-4px 0 4px;
}}
div[data-testid="stForm"] label{{
  color:#171717!important;
  font-size:17px!important;
}}
input,textarea,[data-baseweb="select"]>div{{
  background:#fff!important;
  border-color:#d9d9d9!important;
  border-radius:10px!important;
}}
div.stButton>button,div[data-testid="stFormSubmitButton"] button{{
  border:1px solid var(--line)!important;
  border-radius:10px!important;
  background:linear-gradient(90deg,#d7a92d,#f0d17c,#d5a42a)!important;
  color:#111!important;
  font-family:Georgia,"Times New Roman",serif!important;
  font-size:17px!important;
  min-height:48px;
}}
div.stButton>button:hover,div[data-testid="stFormSubmitButton"] button:hover{{
  border-color:var(--gold3)!important;color:#111!important;filter:brightness(.98);
}}

section[data-testid="stSidebar"]{{
  background:rgba(255,255,255,.97)!important;
  border-right:1.5px solid var(--line)!important;
}}
section[data-testid="stSidebar"] *{{color:#171717!important;}}
section[data-testid="stSidebar"] [data-baseweb="radio"] label{{
  padding:8px 9px!important;border-radius:8px!important;
}}
section[data-testid="stSidebar"] [data-baseweb="radio"] label:hover{{background:#fff8df!important;}}
section[data-testid="stSidebar"] [aria-checked="true"]{{
  background:#fff4cd!important;border-left:3px solid var(--gold)!important;
}}

.card-grid{{display:grid;grid-template-columns:repeat(3,minmax(0,1fr));gap:18px;margin:18px 0 26px;}}
.card{{
  background:rgba(255,255,255,.97);border:1px solid var(--line-soft);border-top:2px solid var(--gold);
  border-radius:18px;padding:20px;box-shadow:0 4px 14px rgba(100,75,10,.05);
}}
.card h3{{margin:0 0 9px;font-size:21px;}}
.gold-chip{{
  display:inline-block;background:#fff8dc;color:#5b4300;padding:5px 10px;border:1px solid #ead37e;border-radius:999px;
  font-family:"Segoe UI",Arial,sans-serif;font-weight:700;font-size:12px;
}}
div[data-testid="stMetric"],div[data-testid="stExpander"],div[data-testid="stDataFrame"]{{
  background:rgba(255,255,255,.96)!important;border:1px solid var(--line-soft)!important;border-radius:15px!important;
}}
.carnation-note{{
  background:rgba(255,255,255,.92);border-top:1px solid var(--line);border-bottom:1px solid var(--line);
  padding:12px 16px;text-align:center;color:#7f641c;font-style:italic;margin:0 0 18px;
}}
h1:after,h2:after{{
  content:"";display:block;width:95px;height:1px;background:var(--gold);margin:9px 0 12px;
}}
.philo-header h1:after{{display:none}}
.block-container:after{{
  content:"Nu Beta Sigma Philo Affiliate";
  display:block;text-align:center;color:var(--gold);
  font-family:"Brush Script MT","Segoe Script",cursive;font-size:25px;margin:36px 0 4px;
}}

@media(max-width:950px){{
  .philo-header{{grid-template-columns:120px 1fr 100px;gap:10px}}
  .philo-header .brand-logo{{width:115px;height:115px}}
  .philo-header .brand-crest{{width:95px;height:95px}}
  .philo-header h1{{font-size:36px}}
  .philo-header p{{font-size:17px;margin-top:10px}}
  .card-grid{{grid-template-columns:1fr}}
}}
@media(max-width:680px){{
  .stApp{{background:linear-gradient(180deg,#fffefa,#fff)}}
  .philo-header{{grid-template-columns:80px 1fr;padding:12px 4px 18px}}
  .philo-header .brand-logo{{width:76px;height:76px}}
  .philo-header .brand-crest{{display:none}}
  .philo-header h1{{font-size:29px}}
  .philo-header p{{font-size:14px}}
  div[data-testid="stForm"]{{padding:20px 18px 22px!important}}
}}

/* v2 app-style navigation */
[data-testid="stSidebar"], [data-testid="collapsedControl"]{{display:none!important;}}
.nav-note{{font-family:"Segoe UI",Arial,sans-serif;color:#6e624d;font-size:13px;margin:-8px 0 10px;}}
.officer-banner{{background:#fff9e8;border:1px solid var(--line);border-radius:16px;padding:14px 18px;margin:8px 0 18px;}}
.role-pill{{display:inline-block;background:#111;color:#fff;border-radius:999px;padding:5px 11px;font-family:"Segoe UI",Arial,sans-serif;font-size:12px;margin-right:6px;}}
.private-note{{background:#fff5f5;border-left:4px solid #a51d1d;padding:10px 13px;border-radius:8px;}}
.status-ok{{background:#f3fbf4;border-left:4px solid #3b7f47;padding:10px 13px;border-radius:8px;}}
[data-testid="stFileUploader"]{{background:rgba(255,255,255,.94);border:1px solid var(--line-soft);border-radius:14px;padding:8px;}}


/* v2.8.2 — Mobile login/input readability */
@media (max-width: 768px) {{
    .stTextInput input,
    .stTextArea textarea,
    .stNumberInput input,
    div[data-baseweb="input"] input,
    div[data-baseweb="textarea"] textarea {{
        color: #000000 !important;
        -webkit-text-fill-color: #000000 !important;
        caret-color: #000000 !important;
        opacity: 1 !important;
    }}

    .stTextInput label,
    .stTextArea label,
    .stNumberInput label,
    .stSelectbox label,
    .stDateInput label,
    .stTimeInput label,
    div[data-testid="stForm"] label {{
        color: #000000 !important;
    }}

    .stTextInput input::placeholder,
    .stTextArea textarea::placeholder,
    div[data-baseweb="input"] input::placeholder,
    div[data-baseweb="textarea"] textarea::placeholder {{
        color: #555555 !important;
        -webkit-text-fill-color: #555555 !important;
        opacity: 1 !important;
    }}

    div[data-baseweb="select"] *,
    div[data-baseweb="select"] input {{
        color: #000000 !important;
        -webkit-text-fill-color: #000000 !important;
    }}
}}


/* v3.1 pearls + carnations */
.pearl-divider{{height:15px;width:min(340px,78%);margin:10px auto 18px;background:radial-gradient(circle at 7px 7px,#fff 0 4.4px,#d9d3c8 4.8px 5.5px,rgba(255,255,255,0) 5.9px) 0 0/17px 14px repeat-x;filter:drop-shadow(0 1px 1px rgba(105,88,54,.18));}}
.officer-banner{{position:relative}}.officer-banner:after{{content:"○ ○ ○";position:absolute;right:14px;top:10px;color:#d8d1c3;letter-spacing:2px}}
.phone-install-card{{background:rgba(255,255,255,.97);border:1px solid var(--line-soft);border-radius:18px;padding:16px 18px;margin:8px 0 18px;box-shadow:0 4px 14px rgba(100,75,10,.05)}}
.phone-install-card img{{width:68px;height:68px;object-fit:contain;border-radius:16px;float:left;margin-right:14px}}.phone-install-card h4{{margin:0 0 4px}}.phone-install-card p{{margin:0;color:#4c4538}}
h1:after,h2:after{{content:""!important;display:block!important;width:115px!important;height:14px!important;background:radial-gradient(circle at 7px 7px,#fff 0 4.3px,#d5cec1 4.8px 5.5px,rgba(255,255,255,0) 5.9px) 0 0/17px 14px repeat-x!important;margin:9px 0 12px!important}}.philo-header h1:after{{display:none!important}}


/* v3.6.23 — large, obvious dashboard tabs */
div[data-baseweb="tab-list"]{{
  gap:10px!important;
  flex-wrap:wrap!important;
  border-bottom:none!important;
  margin:8px 0 18px!important;
}}
button[data-baseweb="tab"]{{
  background:#fffdf8!important;
  border:2px solid #d7af43!important;
  border-radius:14px!important;
  padding:11px 17px!important;
  min-height:46px!important;
  color:#17120a!important;
  box-shadow:0 3px 10px rgba(111,84,15,.08)!important;
  cursor:pointer!important;
}}
button[data-baseweb="tab"] p,
button[data-baseweb="tab"] span,
button[data-baseweb="tab"] div{{
  color:#17120a!important;
  font-weight:700!important;
  font-size:16px!important;
}}
button[data-baseweb="tab"]:hover{{
  background:#fff4ce!important;
  border-color:#b88508!important;
  transform:translateY(-1px)!important;
}}
button[data-baseweb="tab"][aria-selected="true"]{{
  background:linear-gradient(90deg,#d7a92d,#f2d57b,#d7a92d)!important;
  border-color:#8f6a08!important;
  box-shadow:0 5px 13px rgba(111,84,15,.16)!important;
}}
div[data-baseweb="tab-highlight"]{{display:none!important;}}
@media(max-width:680px){{
  button[data-baseweb="tab"]{{padding:9px 11px!important;min-height:42px!important;}}
  button[data-baseweb="tab"] p,
  button[data-baseweb="tab"] span,
  button[data-baseweb="tab"] div{{font-size:14px!important;}}
}}

</style>
""", unsafe_allow_html=True)

def sec(name):
    try:return st.secrets[name]
    except:return None

def cfg():
    s=sec('supabase')
    if not s:return None
    u=str(s.get('url','')).strip(); k=str(s.get('service_role_key','')).strip(); b=str(s.get('bucket','philo-private')).strip() or 'philo-private'
    return {'url':u,'key':k,'bucket':b} if u and k else None

@st.cache_resource
def sb():
    c=cfg(); return create_client(c['url'],c['key']) if c else None

@st.cache_resource
def ensure_bucket():
    if not cfg():return False
    try:sb().storage.create_bucket(cfg()['bucket'],options={'public':False})
    except:pass
    return True

def test_preview_active():
    return bool(st.session_state.get('test_preview_kind')) or bool(st.session_state.get('demo_admin_mode'))

class _DryRunResult:
    def __init__(self,data=None):
        self.data=data if data is not None else []

class _DryRunMutation:
    def __init__(self,action,payload=None):
        self.action=action
        self.payload=payload
    def __getattr__(self,name):
        def method(*args,**kwargs):
            return self
        return method
    def execute(self):
        st.session_state['test_preview_last_action']=f"{self.action.title()} simulated — nothing was saved."
        if self.action in {'insert','upsert'}:
            if isinstance(self.payload,list):
                data=[]
                for i,p in enumerate(self.payload):
                    row=dict(p or {});row.setdefault('id',-9000-i);data.append(row)
                return _DryRunResult(data)
            row=dict(self.payload or {});row.setdefault('id',-9999)
            return _DryRunResult([row])
        return _DryRunResult([])

class _DryRunTable:
    def __init__(self,name):
        self.name=name
        self.real=sb().table(name)
    def select(self,*args,**kwargs):
        return self.real.select(*args,**kwargs)
    def insert(self,payload,*args,**kwargs):
        return _DryRunMutation('insert',payload)
    def update(self,payload,*args,**kwargs):
        return _DryRunMutation('update',payload)
    def delete(self,*args,**kwargs):
        return _DryRunMutation('delete')
    def upsert(self,payload,*args,**kwargs):
        return _DryRunMutation('upsert',payload)

def table(name):
    return _DryRunTable(name) if test_preview_active() else sb().table(name)
def rows(name,**eq):
    q=table(name).select('*')
    for k,v in eq.items():q=q.eq(k,v)
    return q.execute().data or []
def setting(key,default=''):
    r=table('settings').select('value').eq('key',key).execute().data or []
    return r[0]['value'] if r else default
def save_setting(k,v):table('settings').upsert({'key':k,'value':str(v)}).execute()
def validate_image_upload(uploaded,max_mb=8):
    if uploaded is None:
        return None
    data=bytes(uploaded.getbuffer())
    if not data:
        raise ValueError('The selected image file is empty.')
    if len(data)>max_mb*1024*1024:
        raise ValueError(f'Image is too large. Maximum size is {max_mb} MB.')
    ctype=str(getattr(uploaded,'type','') or '').lower()
    allowed={'image/png','image/jpeg'}
    if ctype and ctype not in allowed:
        raise ValueError('Use a PNG or JPEG image.')
    return data

def upload_private(data,path,content_type='application/octet-stream'):
    if test_preview_active():
        st.session_state['test_preview_last_action']='Upload simulated — no file was stored.'
        return f"preview://{path}"
    if not cfg():return ''
    ensure_bucket()
    try:
        sb().storage.from_(cfg()['bucket']).upload(path=path,file=data,file_options={'upsert':'true','content-type':content_type,'cache-control':'0'})
        return path
    except Exception as e:
        st.error(f'Could not upload file: {e}'); return ''
def signed_url(path,seconds=3600):
    if not path:return ''
    if str(path).startswith('preview://'):return ''
    try:
        r=sb().storage.from_(cfg()['bucket']).create_signed_url(path,seconds)
        return r.get('signedURL') or r.get('signedUrl') or ''
    except:return ''


# =========================================================
# v3.6.7 — Previous Sorority Year Records
# Affiliate established March 19, 2022.
# This is an officer records archive, not the Historian photo archive.
# =========================================================
AFFILIATE_ESTABLISHED_DATE=date(2022,3,19)

PREVIOUS_YEAR_DOCUMENT_TYPES=[
    "Officer Report",
    "Monthly Report",
    "Annual / End-of-Year Report",
    "Meeting Minutes",
    "Agenda",
    "Budget",
    "Treasurer Report",
    "Financial Secretary Report",
    "Dues / Assessment Record",
    "Reimbursement / Voucher Record",
    "Bank / Reconciliation Record",
    "Receipt / Expense Support",
    "Membership Record",
    "Correspondence",
    "Program / Event Report",
    "Committee Report",
    "Governance / Procedure Record",
    "Planning Document",
    "Other"
]

PREVIOUS_YEAR_ARCHIVE_CATEGORIES=[
    "Officer Administration",
    "Reports",
    "Meeting Records",
    "Financial Records",
    "Membership Records",
    "Programs & Events",
    "Governance & Procedures",
    "Correspondence",
    "Committee Records",
    "Other"
]

FINANCIAL_DOCUMENT_TYPES={
    "Budget",
    "Treasurer Report",
    "Financial Secretary Report",
    "Dues / Assessment Record",
    "Reimbursement / Voucher Record",
    "Bank / Reconciliation Record",
    "Receipt / Expense Support",
}

PREVIOUS_YEAR_ALLOWED_EXTENSIONS={
    ".pdf",".doc",".docx",".xls",".xlsx",".csv",".txt",".rtf",".png",".jpg",".jpeg"
}

PREVIOUS_YEAR_ALLOWED_MIME_TYPES={
    "application/pdf",
    "application/msword",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/vnd.ms-excel",
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "text/csv",
    "text/plain",
    "application/rtf",
    "text/rtf",
    "image/png",
    "image/jpeg",
    "application/octet-stream",  # browser fallback; extension is still checked
}

def sorority_year_options():
    """Return Sorority Year labels beginning with the affiliate's 2022 establishment year."""
    current=date.today()
    # Use July-June labeling already consistent with the app's fiscal/sorority-year patterns.
    end_start=current.year if current.month>=7 else current.year-1
    years=[]
    for start in range(2022,end_start+1):
        years.append(f"{start}-{start+1}")
    return list(reversed(years))

def previous_year_is_financial(document_type,archive_category):
    return document_type in FINANCIAL_DOCUMENT_TYPES or archive_category=="Financial Records"

def validate_previous_year_file(uploaded,max_mb=20):
    if uploaded is None:
        raise ValueError("Choose a file to upload.")
    filename=Path(str(uploaded.name or "")).name
    ext=Path(filename).suffix.lower()
    if ext not in PREVIOUS_YEAR_ALLOWED_EXTENSIONS:
        raise ValueError(
            "Unsupported file type. Use PDF, Word, Excel, CSV, TXT/RTF, PNG, or JPEG."
        )
    raw=bytes(uploaded.getbuffer())
    if not raw:
        raise ValueError("The selected file is empty.")
    if len(raw)>max_mb*1024*1024:
        raise ValueError(f"File is too large. Maximum size is {max_mb} MB.")
    mime=str(getattr(uploaded,"type","") or "application/octet-stream").lower()
    if mime not in PREVIOUS_YEAR_ALLOWED_MIME_TYPES:
        raise ValueError(f"Unsupported file content type: {mime}")
    # Lightweight signature checks for the most common formats.
    if ext==".pdf" and not raw.startswith(b"%PDF"):
        raise ValueError("This file has a .pdf extension but does not appear to be a valid PDF.")
    if ext in {".png"} and not raw.startswith(b"\x89PNG\r\n\x1a\n"):
        raise ValueError("This file has a .png extension but does not appear to be a valid PNG.")
    if ext in {".jpg",".jpeg"} and not raw.startswith(b"\xff\xd8\xff"):
        raise ValueError("This file has a JPEG extension but does not appear to be a valid JPEG.")
    return {
        "bytes":raw,
        "filename":filename,
        "extension":ext,
        "mime_type":mime,
        "size":len(raw),
        "sha256":hashlib.sha256(raw).hexdigest(),
    }

def previous_year_record_rows():
    try:
        return table("previous_sorority_year_documents").select("*").order(
            "document_date",desc=True
        ).order("created_at",desc=True).execute().data or []
    except Exception as ex:
        raise RuntimeError(
            "Previous Sorority Year Records is not ready in Supabase. "
            "Apply migration_v3_6_7_previous_sorority_year_records.sql first. "
            f"Database message: {ex}"
        )

def current_position_names(mid):
    return {x.get("position") for x in member_offices(mid) if x.get("position")}

def can_view_previous_year_record(record,mid,is_admin_flag=False):
    if is_admin_flag:
        return True
    positions=current_position_names(mid)
    office=str(record.get("office") or "")
    if int(record.get("uploaded_by_member_id") or 0)==int(mid or 0):
        return True
    if record.get("is_financial"):
        # Financial archives are never general-member records.
        return bool(positions & {"President","Financial Secretary","Treasurer"})
    # Non-financial officer records: current holder of that office or President.
    return office in positions or "President" in positions

def can_manage_previous_year_record(record,mid,is_admin_flag=False):
    if is_admin_flag:
        return True
    positions=current_position_names(mid)
    if int(record.get("uploaded_by_member_id") or 0)==int(mid or 0):
        return True
    return str(record.get("office") or "") in positions or "President" in positions

def download_private_bytes(path):
    if not path:
        return b""
    try:
        return sb().storage.from_(cfg()["bucket"]).download(path)
    except Exception as ex:
        raise RuntimeError(f"Could not download archived file: {ex}")

def previous_year_duplicate(sha256_value,office,sorority_year):
    rows_=table("previous_sorority_year_documents").select(
        "id,title,file_name,created_at"
    ).eq("file_sha256",sha256_value).eq(
        "office",office
    ).eq("sorority_year",sorority_year).execute().data or []
    return rows_[0] if rows_ else None


def _clean_finance_number(value):
    if value is None:
        return None
    if isinstance(value,float) and pd.isna(value):
        return None
    s=str(value).strip().replace("$","").replace(",","")
    if not s or s.lower() in {"nan","none","n/a","na","source needed"}:
        return None
    try:
        return float(Decimal(s).quantize(Decimal("0.01"),rounding=ROUND_HALF_UP))
    except Exception:
        return None

def _normalize_year_label(value):
    s=str(value or "").strip().replace("–","-").replace("—","-")
    m=re.search(r"(20\d{2})\s*-\s*(20\d{2})",s)
    return f"{m.group(1)}-{m.group(2)}" if m else ""

def _finance_rows_from_dataframe(df):
    """Find a usable header row and return normalized prior-year finance summaries."""
    if df is None or df.empty:
        return []

    # locate a row containing a Sorority Year heading
    header_idx=None
    for i in range(min(len(df),20)):
        vals=[str(x).strip() for x in df.iloc[i].tolist()]
        if any(v.lower()=="sorority year" for v in vals):
            header_idx=i
            break
    if header_idx is None:
        return []

    headers=[str(x).strip() for x in df.iloc[header_idx].tolist()]
    body=df.iloc[header_idx+1:].copy()
    body.columns=headers
    records=[]

    aliases={
        "opening_balance":["Opening Balance","Opening"],
        "deposits":["Deposits","Total Deposits","FY Deposits"],
        "withdrawals":["Withdrawals","Total Withdrawals","FY Withdrawals"],
        "net_activity":["Net Activity","Net"],
        "closing_balance":["Closing Balance","Calculated Closing","Calculated Closing Balance","Ending Balance"],
        "bank_statement_date":["Bank Statement Date","Bank Statement As Of"],
        "bank_statement_balance":["Bank Statement Balance","Bank Ending Balance"],
        "status":["Verification Status","Status","Archive Status"],
        "source":["Source File / Report","Source / Evidence","Source File"],
        "notes":["Notes","Source Note"],
    }

    def pick(row,names):
        for name in names:
            if name in row.index:
                v=row.get(name)
                if not (isinstance(v,float) and pd.isna(v)):
                    return v
        return None

    for _,row in body.iterrows():
        year=_normalize_year_label(row.get("Sorority Year"))
        if not year:
            continue
        status=str(pick(row,aliases["status"]) or "").strip()
        # Never import unverified/source-needed years.
        if status and any(x in status.casefold() for x in ["source needed","unverified","needs source"]):
            continue

        opening=_clean_finance_number(pick(row,aliases["opening_balance"]))
        deposits=_clean_finance_number(pick(row,aliases["deposits"]))
        withdrawals=_clean_finance_number(pick(row,aliases["withdrawals"]))
        net=_clean_finance_number(pick(row,aliases["net_activity"]))
        closing=_clean_finance_number(pick(row,aliases["closing_balance"]))
        bank_bal=_clean_finance_number(pick(row,aliases["bank_statement_balance"]))

        # Require actual supported finance values; don't create empty year shells.
        if all(v is None for v in [opening,deposits,withdrawals,net,closing,bank_bal]):
            continue

        if net is None and deposits is not None and withdrawals is not None:
            net=float(Decimal(str(deposits))-Decimal(str(withdrawals)))
        if closing is None and opening is not None and net is not None:
            closing=float(Decimal(str(opening))+Decimal(str(net)))

        bank_date=pick(row,aliases["bank_statement_date"])
        if bank_date is not None and not (isinstance(bank_date,float) and pd.isna(bank_date)):
            try:
                bank_date=pd.to_datetime(bank_date).date().isoformat()
            except Exception:
                bank_date=str(bank_date).strip()
        else:
            bank_date=None

        records.append({
            "sorority_year":year,
            "opening_balance":opening,
            "deposits":deposits,
            "withdrawals":withdrawals,
            "net_activity":net,
            "closing_balance":closing,
            "bank_statement_date":bank_date,
            "bank_statement_balance":bank_bal,
            "verification_status":status or "Verified from uploaded source",
            "source_note":str(pick(row,aliases["source"]) or "").strip(),
            "notes":str(pick(row,aliases["notes"]) or "").strip(),
        })
    return records

def parse_prior_year_finance_upload(uploaded):
    if uploaded is None:
        raise ValueError("Choose a finance workbook or CSV to import.")
    raw=bytes(uploaded.getbuffer())
    if not raw:
        raise ValueError("The selected finance file is empty.")
    if len(raw)>20*1024*1024:
        raise ValueError("Finance import file is too large. Maximum size is 20 MB.")

    name=Path(str(uploaded.name or "")).name
    ext=Path(name).suffix.lower()
    if ext not in {".xlsx",".xls",".csv"}:
        raise ValueError("Finance import supports Excel (.xlsx/.xls) or CSV files.")

    sha=hashlib.sha256(raw).hexdigest()
    summaries=[]

    if ext==".csv":
        df=pd.read_csv(BytesIO(raw),header=None)
        summaries.extend(_finance_rows_from_dataframe(df))
    else:
        excel=pd.ExcelFile(BytesIO(raw))
        preferred=[
            "All Years Summary",
            "Prior Year Intake",
            "2025-2026 Snapshot",
            "2025–2026 Snapshot",
        ]
        ordered=[s for s in preferred if s in excel.sheet_names]
        ordered += [s for s in excel.sheet_names if s not in ordered]
        for sheet in ordered:
            try:
                df=pd.read_excel(BytesIO(raw),sheet_name=sheet,header=None)
                found=_finance_rows_from_dataframe(df)
                if found:
                    summaries.extend(found)
            except Exception:
                continue

    # De-dupe by Sorority Year, preferring the row with the most populated financial fields.
    by_year={}
    def completeness(r):
        return sum(r.get(k) is not None for k in [
            "opening_balance","deposits","withdrawals","net_activity",
            "closing_balance","bank_statement_balance"
        ])
    for rec in summaries:
        yr=rec["sorority_year"]
        if yr not in by_year or completeness(rec)>completeness(by_year[yr]):
            by_year[yr]=rec

    if not by_year:
        raise ValueError(
            "No verified Sorority Year finance summary was found. "
            "Use the archive workbook format with a Sorority Year row and verified amounts."
        )
    return {
        "bytes":raw,
        "filename":name,
        "sha256":sha,
        "mime_type":str(getattr(uploaded,"type","") or "application/octet-stream"),
        "summaries":list(by_year.values()),
    }

def prior_year_finance_summary_rows():
    try:
        return table("previous_sorority_year_finance_summaries").select("*").order(
            "sorority_year",desc=True
        ).execute().data or []
    except Exception as ex:
        raise RuntimeError(
            "Prior-Year Finance Summary table is not ready. "
            "Apply migration_v3_6_11_prior_year_finance_import.sql first. "
            f"Database message: {ex}"
        )

def import_prior_year_finance(uploaded,mid,uploader_name):
    parsed=parse_prior_year_finance_upload(uploaded)

    # Archive source file once under Financial Records.
    existing_doc=table("previous_sorority_year_documents").select(
        "id,file_path"
    ).eq("file_sha256",parsed["sha256"]).execute().data or []

    if existing_doc:
        source_document_id=existing_doc[0]["id"]
        source_path=existing_doc[0].get("file_path")
    else:
        safe_name=re.sub(r"[^A-Za-z0-9._-]+","_",parsed["filename"])
        source_path=f"previous-sorority-year/finance-imports/{parsed['sha256'][:12]}_{safe_name}"
        uploaded_path=upload_private(
            parsed["bytes"],source_path,parsed["mime_type"]
        )
        if not uploaded_path:
            raise RuntimeError("Source finance workbook could not be stored.")
        source_path=uploaded_path

        newest_year=sorted(
            [x["sorority_year"] for x in parsed["summaries"]]
        )[-1]
        sy_start=int(newest_year.split("-")[0])
        doc_date=date(sy_start+1,6,30).isoformat()
        inserted=table("previous_sorority_year_documents").insert({
            "sorority_year":newest_year,
            "document_date":doc_date,
            "document_type":"Annual / End-of-Year Report",
            "office":"Treasurer",
            "title":f"Imported Prior-Year Finance Source — {parsed['filename']}",
            "description":"Source workbook used by Import Prior-Year Finance Data.",
            "archive_category":"Financial Records",
            "file_path":source_path,
            "file_name":parsed["filename"],
            "mime_type":parsed["mime_type"],
            "file_size":len(parsed["bytes"]),
            "file_sha256":parsed["sha256"],
            "is_financial":True,
            "visibility":"Restricted - Financial Leadership & Advisor",
            "uploaded_by_member_id":mid,
            "uploaded_by_name":uploader_name,
            "created_at":datetime.now(timezone.utc).isoformat(),
        }).execute().data or []
        if not inserted:
            raise RuntimeError("Source document metadata was not created.")
        source_document_id=inserted[0]["id"]

    imported=[]
    for rec in parsed["summaries"]:
        payload=dict(rec)
        payload.update({
            "source_document_id":source_document_id,
            "source_file_sha256":parsed["sha256"],
            "imported_by_member_id":mid,
            "imported_by_name":uploader_name,
            "updated_at":datetime.now(timezone.utc).isoformat(),
        })
        table("previous_sorority_year_finance_summaries").upsert(
            payload,on_conflict="sorority_year"
        ).execute()
        imported.append(rec["sorority_year"])

    return imported


FINANCIAL_HISTORY_VIEW_ROLES={"President","Financial Secretary","Treasurer"}

def can_view_financial_history(mid,is_admin_flag=False):
    return bool(is_admin_flag or (current_position_names(mid) & FINANCIAL_HISTORY_VIEW_ROLES))

def financial_history_month_rows():
    try:
        return table("previous_sorority_year_finance_months").select("*").order(
            "month_start"
        ).execute().data or []
    except Exception as ex:
        raise RuntimeError(
            "Monthly financial history is not ready. Apply "
            "migration_v3_6_12_read_only_financial_history.sql first. "
            f"Database message: {ex}"
        )



TREASURER_CONTINUITY_ROLES={"President","Financial Secretary","Treasurer"}

def continuity_preference(member_id_,section,key,default=False):
    try:
        rows_=table("continuity_preferences").select("preference_value").eq(
            "member_id",member_id_
        ).eq("section",section).eq("preference_key",key).execute().data or []
        if rows_:
            return str(rows_[0].get("preference_value") or "").lower()=="true"
    except Exception:
        pass
    return bool(default)

def save_continuity_preference(member_id_,section,key,value):
    try:
        table("continuity_preferences").upsert({
            "member_id":member_id_,
            "section":section,
            "preference_key":key,
            "preference_value":"true" if bool(value) else "false",
            "updated_at":datetime.now(timezone.utc).isoformat(),
        },on_conflict="member_id,section,preference_key").execute()
    except Exception:
        # Preference persistence should never block the actual workflow.
        pass

CURRENT_SORORITY_YEAR="2026-2027"

def can_use_treasurer_continuity(mid,is_admin_flag=False):
    return bool(is_admin_flag or (current_position_names(mid) & TREASURER_CONTINUITY_ROLES))

def current_help_enabled():
    if "treasurer_help_guide" not in st.session_state:
        st.session_state["treasurer_help_guide"]=continuity_preference(member_id,"treasurer","help_guide",True)
    return bool(st.session_state["treasurer_help_guide"])

def current_easy_view():
    if "treasurer_easy_view" not in st.session_state:
        st.session_state["treasurer_easy_view"]=continuity_preference(member_id,"treasurer","easy_view",False)
    return bool(st.session_state["treasurer_easy_view"])

def continuity_rows(table_name,**filters):
    q=table(table_name).select("*")
    for k,v in filters.items():
        q=q.eq(k,v)
    return q.order("created_at",desc=True).execute().data or []

def finance_duplicate_fingerprint(transaction_date,amount,direction,description,reference=""):
    desc=re.sub(r"\s+"," ",str(description or "").strip()).casefold()
    ref=re.sub(r"\s+"," ",str(reference or "").strip()).casefold()
    raw=f"{transaction_date}|{Decimal(str(amount or 0)).quantize(Decimal('0.01'))}|{direction}|{desc}|{ref}"
    return hashlib.sha256(raw.encode("utf-8")).hexdigest()


def current_sorority_month_labels():
    start=date(2026,7,1)
    end=date(2027,6,1)
    today=date.today()
    labels=[]
    cur=start
    while cur<=end and cur<=date(today.year,today.month,1):
        labels.append(cur.strftime("%Y-%m"))
        cur=(cur.replace(day=28)+timedelta(days=4)).replace(day=1)
    return labels

def finance_possible_duplicate(transaction_date,amount,direction,description="",reference=""):
    try:
        candidates=table("finance_transactions").select(
            "id,transaction_date,amount,direction,payer_payee,notes,reference_number"
        ).eq("transaction_date",str(transaction_date)).eq(
            "direction",direction
        ).eq("amount",float(amount)).execute().data or []
    except Exception:
        return None
    target_desc=re.sub(r"\s+"," ",str(description or "").strip()).casefold()
    target_ref=re.sub(r"\s+"," ",str(reference or "").strip()).casefold()
    for r in candidates:
        desc=re.sub(r"\s+"," ",str(r.get("payer_payee") or r.get("notes") or "").strip()).casefold()
        ref=re.sub(r"\s+"," ",str(r.get("reference_number") or "").strip()).casefold()
        if (target_ref and ref==target_ref) or (target_desc and desc==target_desc) or (not target_desc and not target_ref):
            return r
    return None

def add_duplicate_review(transaction_date,amount,direction,description="",reference="",note=""):
    fp=finance_duplicate_fingerprint(transaction_date,amount,direction,description,reference)
    table("finance_duplicate_review").upsert({
        "transaction_fingerprint":fp,
        "transaction_date":str(transaction_date),
        "amount":float(amount),
        "direction":direction,
        "description":str(description or ""),
        "reference_number":str(reference or ""),
        "review_note":str(note or "Possible duplicate detected before posting."),
        "resolved":False,
        "created_at":datetime.now(timezone.utc).isoformat(),
    },on_conflict="transaction_fingerprint").execute()

def annual_finance_archive_bytes(year):
    out=BytesIO()
    with zipfile.ZipFile(out,"w",zipfile.ZIP_DEFLATED) as z:
        summary=[r for r in prior_year_finance_summary_rows() if str(r.get("sorority_year"))==year]
        monthly=[r for r in financial_history_month_rows() if str(r.get("sorority_year"))==year]
        rec=continuity_rows("finance_reconciliation_log",sorority_year=year)
        hand=continuity_rows("finance_handoff_notes",sorority_year=year)
        close=continuity_rows("finance_month_close",sorority_year=year)
        def csv_bytes(rows_):
            if not rows_: return b""
            return pd.DataFrame(rows_).to_csv(index=False).encode("utf-8")
        z.writestr(f"{year}_annual_summary.csv",csv_bytes(summary))
        z.writestr(f"{year}_monthly_history.csv",csv_bytes(monthly))
        z.writestr(f"{year}_reconciliation_log.csv",csv_bytes(rec))
        z.writestr(f"{year}_handoff_notes.csv",csv_bytes(hand))
        z.writestr(f"{year}_monthly_close.csv",csv_bytes(close))
        index={
            "sorority_year":year,
            "generated_at":datetime.now(timezone.utc).isoformat(),
            "generated_by":member_name,
            "warning":"This package exports app-held records. Supporting bank statements/receipts remain separate source documents unless individually archived.",
        }
        z.writestr("README.json",json.dumps(index,indent=2))
    return out.getvalue()

def continuity_audit_status():
    issues=[]
    rec=continuity_rows("finance_reconciliation_log",resolved=False)
    if rec:
        issues.append(("Reconciliation Required",f"{len(rec)} unresolved reconciliation item(s)"))

    close_rows=continuity_rows("finance_month_close",sorority_year=CURRENT_SORORITY_YEAR)
    expected=current_sorority_month_labels()
    represented={str(r.get("month_label") or "") for r in close_rows}
    missing=[m for m in expected if m not in represented]
    if missing:
        issues.append(("Needs Attention",f"{len(missing)} elapsed month(s) have no Monthly Close record"))
    incomplete=[r for r in close_rows if str(r.get("status") or "").lower()!="complete"]
    if incomplete:
        issues.append(("Needs Attention",f"{len(incomplete)} current-year month(s) are not complete"))

    try:
        prior=prior_year_finance_summary_rows()
        flagged=[r for r in prior if any(x in str(r.get("verification_status") or "").casefold() for x in ["review required","reconciliation"])]
        if flagged:
            issues.append(("Reconciliation Required",f"{len(flagged)} past Sorority Year summary record(s) are flagged for reconciliation"))
    except Exception:
        pass

    dup=continuity_rows("finance_duplicate_review",resolved=False)
    if dup:
        issues.append(("Needs Attention",f"{len(dup)} possible duplicate transaction(s) need review"))
    hand=continuity_rows("finance_handoff_notes",resolved=False)
    if hand:
        issues.append(("Needs Attention",f"{len(hand)} unresolved handoff note(s)"))
    return issues

def render_help_callout(title,body):
    if current_help_enabled():
        st.info(f"**{title}**\n\n{body}")

def render_treasurer_continuity_center(mid,is_admin_flag=False):
    if not can_use_treasurer_continuity(mid,is_admin_flag):
        return

    st.markdown("## Treasurer Continuity Center")
    st.caption("Designed so future financial officers can understand what to do, preserve an audit trail, and continue the work without guessing.")

    c1,c2=st.columns(2)
    with c1:
        st.toggle(
            "Show Help Guide",
            key="treasurer_help_guide",
            help="Turn on step-by-step instructions whenever you need help.",
            on_change=lambda: save_continuity_preference(mid,"treasurer","help_guide",st.session_state.get("treasurer_help_guide"))
        )
    with c2:
        st.toggle(
            "Easy View",
            key="treasurer_easy_view",
            help="Uses simpler labels and fewer side-by-side controls.",
            on_change=lambda: save_continuity_preference(mid,"treasurer","easy_view",st.session_state.get("treasurer_easy_view"))
        )

    if current_easy_view():
        st.info("Easy View is ON — screens use simpler wording and one main action at a time.")

    issues=continuity_audit_status()
    if not issues:
        st.success("Audit Readiness: READY — no unresolved continuity items were found.")
    else:
        labels={x[0] for x in issues}
        if "Reconciliation Required" in labels:
            st.error("Audit Readiness: RECONCILIATION REQUIRED")
        else:
            st.warning("Audit Readiness: NEEDS ATTENTION")
        for status,msg in issues:
            st.write(f"- **{status}:** {msg}")

    tabs=st.tabs([
        "Start Here",
        "Monthly Close",
        "Reconciliation Log",
        "Duplicate Review",
        "Audit Readiness",
        "Year-End Handoff",
        "Annual Archive"
    ])

    with tabs[0]:
        render_help_callout(
            "What this page does",
            "Use this page when you are not sure where to begin. Choose the task that matches what you need to do today."
        )
        st.markdown("### What do you need to do today?")
        actions=[
            ("Record Money In","Use the current 2026–2027 finance workflow to enter deposits, dues, donations, fundraising, and other income."),
            ("Record Money Out","Use the current 2026–2027 finance workflow to enter expenses, reimbursements, assessments, and other withdrawals."),
            ("Match Bank Statement","Use Monthly Close to compare the books to the bank statement."),
            ("Review This Month","Use Monthly Close to see what is missing before the month is complete."),
            ("View Past Years","Use Financial History Over Time to compare prior Sorority Years."),
            ("Prepare for Audit","Use Audit Readiness to see what still needs attention."),
            ("Prepare Handoff","Use Year-End Handoff to leave clear notes for the next officer."),
        ]
        for title,desc in actions:
            with st.container(border=True):
                st.markdown(f"**{title}**")
                st.write(desc)

    with tabs[1]:
        render_help_callout(
            "Monthly Close",
            "At the end of each month, confirm the bank statement, deposits, withdrawals, receipts, classifications, and ending balance. You may save an incomplete month, but it will remain Needs Attention."
        )
        month_options=current_sorority_month_labels()
        if not month_options:
            st.info("No 2026–2027 month is eligible to close yet.")
        else:
            month=st.selectbox("Month",month_options,key="continuity_close_month")
            existing=continuity_rows("finance_month_close",sorority_year=CURRENT_SORORITY_YEAR)
            cur=next((r for r in existing if r.get("month_label")==month),{})
            checks=[
                ("bank_statement_obtained","Bank statement obtained"),
                ("all_deposits_entered","All deposits entered"),
                ("all_withdrawals_entered","All withdrawals entered"),
                ("receipts_complete","Receipts/source documents complete"),
                ("categories_reviewed","Donations/fundraising/categories reviewed"),
            ]
            values={}
            for key,label in checks:
                values[key]=st.checkbox(label,value=bool(cur.get(key)),key=f"close_{month}_{key}")

            statement_balance=st.number_input(
                "Bank statement ending balance",
                value=float(cur.get("statement_ending_balance") or 0),
                step=0.01,format="%.2f",key=f"statement_bal_{month}"
            )
            ledger_balance=st.number_input(
                "Ledger ending balance",
                value=float(cur.get("ledger_ending_balance") or 0),
                step=0.01,format="%.2f",key=f"ledger_bal_{month}"
            )
            variance=round(float(statement_balance)-float(ledger_balance),2)
            matched=abs(variance)<=0.01
            st.metric("Statement vs. Ledger Difference",f"${variance:,.2f}")
            if matched:
                st.success("Balances match within one cent.")
            else:
                st.warning("Balances do not match. Explain the difference before saving.")

            source_reference=st.text_input(
                "Bank statement / source reference",
                value=str(cur.get("source_reference") or ""),
                key=f"close_source_{month}"
            )
            notes=st.text_area("Monthly close notes / difference explanation",value=str(cur.get("notes") or ""),key=f"close_notes_{month}")

            if st.button("Save Monthly Close",use_container_width=True,key=f"save_close_{month}"):
                if not source_reference.strip():
                    st.error("Enter a bank statement/source reference.")
                elif not matched and not notes.strip():
                    st.error("Explain the balance difference before saving.")
                else:
                    complete=all(values.values()) and matched
                    payload={
                        "sorority_year":CURRENT_SORORITY_YEAR,
                        "month_label":month,
                        **values,
                        "ending_balance_matched":matched,
                        "difference_explained":bool(notes.strip()) if not matched else True,
                        "statement_ending_balance":float(statement_balance),
                        "ledger_ending_balance":float(ledger_balance),
                        "variance":variance,
                        "source_reference":source_reference.strip(),
                        "status":"Complete" if complete else "Needs Attention",
                        "notes":notes.strip(),
                        "updated_by_member_id":mid,
                        "updated_by_name":member_name,
                        "updated_at":datetime.now(timezone.utc).isoformat(),
                    }
                    table("finance_month_close").upsert(payload,on_conflict="sorority_year,month_label").execute()
                    st.success("Monthly close saved.")
                    st.rerun()

    with tabs[2]:
        render_help_callout(
            "Reconciliation Adjustment Log",
            "Never silently overwrite an old financial figure. Record the original amount, corrected amount, reason, and source document so an auditor can see exactly what changed and why."
        )
        with st.form("reconciliation_log_form",clear_on_submit=True):
            year=st.selectbox("Sorority Year",["2023-2024","2024-2025","2025-2026",CURRENT_SORORITY_YEAR])
            field=st.text_input("Field / Item being reconciled",placeholder="Example: June 30 ending balance")
            original=st.number_input("Original amount",step=0.01,format="%.2f")
            corrected=st.number_input("Corrected amount",step=0.01,format="%.2f")
            reason=st.text_area("Reason for correction")
            source=st.text_input("Supporting source / document")
            save=st.form_submit_button("Add Reconciliation Record",use_container_width=True)
        if save:
            if not field.strip() or not reason.strip() or not source.strip():
                st.error("Field, reason, and supporting source are required.")
            else:
                table("finance_reconciliation_log").insert({
                    "sorority_year":year,
                    "field_name":field.strip(),
                    "original_amount":float(original),
                    "corrected_amount":float(corrected),
                    "difference":float(corrected-original),
                    "reason":reason.strip(),
                    "source_reference":source.strip(),
                    "resolved":False,
                    "created_by_member_id":mid,
                    "created_by_name":member_name,
                    "created_at":datetime.now(timezone.utc).isoformat(),
                }).execute()
                st.success("Reconciliation record added.")
                st.rerun()
        rows_=continuity_rows("finance_reconciliation_log")
        if rows_:
            st.dataframe(pd.DataFrame(rows_),hide_index=True,use_container_width=True)
            unresolved=[r for r in rows_ if not r.get("resolved")]
            if unresolved:
                st.markdown("##### Resolve a Reconciliation Item")
                rmap={r["id"]:f"{r.get('sorority_year')} • {r.get('field_name')}" for r in unresolved}
                rid=st.selectbox("Item",list(rmap.keys()),format_func=lambda x:rmap[x],key="resolve_reconciliation_id")
                resolution=st.text_area("Resolution / what source proved the answer",key="resolve_reconciliation_note")
                if st.button("Mark Reconciliation Resolved",use_container_width=True):
                    if not resolution.strip():
                        st.error("Enter the resolution and supporting source.")
                    else:
                        table("finance_reconciliation_log").update({
                            "resolved":True,
                            "resolution_note":resolution.strip(),
                            "resolved_by_member_id":mid,
                            "resolved_by_name":member_name,
                            "resolved_at":datetime.now(timezone.utc).isoformat(),
                        }).eq("id",rid).execute()
                        st.success("Reconciliation item marked resolved.")
                        st.rerun()

    with tabs[3]:
        render_help_callout(
            "Duplicate Review",
            "Financial duplicates can double-count money. Suspected duplicates should be reviewed, not automatically deleted."
        )
        st.caption("The app can store possible duplicate fingerprints here for review before any cleanup.")
        rows_=continuity_rows("finance_duplicate_review")
        if rows_:
            st.dataframe(pd.DataFrame(rows_),hide_index=True,use_container_width=True)
            unresolved=[r for r in rows_ if not r.get("resolved")]
            if unresolved:
                dmap={r["id"]:f"{r.get('transaction_date')} • ${float(r.get('amount') or 0):,.2f} • {r.get('description') or ''}" for r in unresolved}
                did=st.selectbox("Possible duplicate",list(dmap.keys()),format_func=lambda x:dmap[x],key="duplicate_review_id")
                decision=st.selectbox("Decision",["Keep Both — not a duplicate","Duplicate — do not post another copy","Needs More Research"],key="duplicate_decision")
                note=st.text_area("Review note",key="duplicate_resolution_note")
                if st.button("Save Duplicate Review Decision",use_container_width=True):
                    table("finance_duplicate_review").update({
                        "resolved":decision!="Needs More Research",
                        "review_decision":decision,
                        "resolution_note":note.strip(),
                        "resolved_by_member_id":mid,
                        "resolved_by_name":member_name,
                        "resolved_at":datetime.now(timezone.utc).isoformat() if decision!="Needs More Research" else None,
                    }).eq("id",did).execute()
                    st.success("Duplicate review decision saved.")
                    st.rerun()
        else:
            st.success("No duplicate transactions are waiting for review.")

    with tabs[4]:
        render_help_callout(
            "Audit Readiness",
            "This section answers: if someone asked for the books today, what would still need attention?"
        )
        issues=continuity_audit_status()
        if not issues:
            st.success("READY — no unresolved continuity items were found.")
        else:
            for status,msg in issues:
                if status=="Reconciliation Required":
                    st.error(f"{status}: {msg}")
                else:
                    st.warning(f"{status}: {msg}")
        st.markdown("### Source continuity checks")
        st.write("- 2023–2024 ending vs. 2024–2025 opening: review documented $400 difference.")
        st.write("- 2024–2025 ending vs. 2025–2026 opening: review documented $1,664.50 difference.")
        st.write("- 2024–2025 line items vs. reported balances: reconciliation required.")

    with tabs[5]:
        render_help_callout(
            "Year-End Handoff",
            "Use this before leaving office. Record anything the next financial officer needs to know so she is not forced to reconstruct your work."
        )
        with st.form("handoff_note_form",clear_on_submit=True):
            category=st.selectbox("Handoff Category",[
                "Outstanding Reconciliation","Recurring Expense","Expected Deposit",
                "Access / Account","Source Documents","Deadline","Other"
            ])
            title=st.text_input("Handoff item")
            note=st.text_area("What the next officer needs to know")
            due=st.date_input("Due / follow-up date",value=date.today())
            save=st.form_submit_button("Add Handoff Item",use_container_width=True)
        if save:
            if not title.strip() or not note.strip():
                st.error("Handoff item and explanation are required.")
            else:
                table("finance_handoff_notes").insert({
                    "sorority_year":CURRENT_SORORITY_YEAR,
                    "category":category,
                    "title":title.strip(),
                    "note":note.strip(),
                    "follow_up_date":due.isoformat(),
                    "resolved":False,
                    "created_by_member_id":mid,
                    "created_by_name":member_name,
                    "created_at":datetime.now(timezone.utc).isoformat(),
                }).execute()
                st.success("Handoff item saved.")
                st.rerun()
        rows_=continuity_rows("finance_handoff_notes")
        if rows_:
            st.dataframe(pd.DataFrame(rows_),hide_index=True,use_container_width=True)
            unresolved=[r for r in rows_ if not r.get("resolved")]
            if unresolved:
                hmap={r["id"]:f"{r.get('category')} • {r.get('title')}" for r in unresolved}
                hid=st.selectbox("Resolve handoff item",list(hmap.keys()),format_func=lambda x:hmap[x],key="finance_handoff_resolve_id")
                hnote=st.text_input("Resolution note",key="finance_handoff_resolve_note")
                if st.button("Mark Handoff Item Resolved",use_container_width=True):
                    if not hnote.strip():
                        st.error("Enter a short resolution note.")
                    else:
                        table("finance_handoff_notes").update({
                            "resolved":True,
                            "resolution_note":hnote.strip(),
                            "resolved_at":datetime.now(timezone.utc).isoformat(),
                        }).eq("id",hid).execute()
                        st.success("Handoff item resolved.")
                        st.rerun()

        st.markdown("### Year-End Handoff Checklist")
        checklist=[
            "Final June reconciliation complete",
            "Annual financial report generated",
            "Final dashboard exported",
            "Source workbook archived",
            "Previous Sorority Year Word record uploaded",
            "Ending balance documented",
            "Next-year opening balance prepared",
            "Access/password review completed",
            "Successor walkthrough completed",
            "Outstanding reconciliation items documented",
        ]
        for i,item in enumerate(checklist):
            st.checkbox(item,key=f"handoff_check_{i}")

    with tabs[6]:
        render_help_callout(
            "Annual Archive Package",
            "Use this at year-end to make an independent copy of the records so the affiliate still has a usable archive even if the app changes later."
        )
        archive_year=st.selectbox(
            "Archive Sorority Year",
            ["2023-2024","2024-2025","2025-2026",CURRENT_SORORITY_YEAR],
            key="archive_year_select"
        )
        st.write("The package should include:")
        st.write("- annual financial summary")
        st.write("- monthly ledger/history")
        st.write("- reconciliation log")
        st.write("- audit-readiness status")
        st.write("- handoff notes")
        st.write("- source-document index")
        archive_bytes=annual_finance_archive_bytes(archive_year)
        st.download_button(
            "⬇️ Download Annual Finance Archive ZIP",
            archive_bytes,
            file_name=f"NBS_Finance_Archive_{archive_year}.zip",
            mime="application/zip",
            use_container_width=True
        )
        st.caption(
            "This ZIP contains app-held summary/history/reconciliation/handoff/month-close records. "
            "Bank statements, receipts, and other source files remain separate unless they were archived elsewhere in the Hub."
        )

# =========================================================
# v3.6.15 — Historian Continuity Center
# =========================================================
HISTORIAN_CONTINUITY_ROLES={"Historian"}
HISTORIAN_CURRENT_SORORITY_YEAR="2026-2027"

HISTORIAN_PUBLICATION_STATUS_OPTIONS=[
    "Archive Only",
    "Approved for Public Use",
    "Internal Use",
    "Needs Review",
]

HISTORIAN_RECORD_STATUS_OPTIONS=[
    "Archive Ready",
    "Needs Attention",
    "Needs Identification",
    "Date Unknown / Needs Review",
    "Source Conflict / Needs Review",
]

HISTORIAN_MILESTONE_TYPES=[
    "Affiliate Milestone",
    "Anniversary",
    "Award",
    "Recognition",
    "First / Historic First",
    "Community Honor",
    "Regional Recognition",
    "National Recognition",
    "Leadership Milestone",
    "Program Milestone",
    "Other",
]

def can_use_historian_continuity(mid,is_admin_flag=False):
    return bool(
        is_admin_flag
        or ("Historian" in current_position_names(mid))
        or historian_delegate_rows(mid)
    )


def historian_help_enabled():
    if "historian_help_guide" not in st.session_state:
        st.session_state["historian_help_guide"]=historian_preference(member_id,"help_guide",True)
    return bool(st.session_state["historian_help_guide"])


def historian_easy_view():
    if "historian_easy_view" not in st.session_state:
        st.session_state["historian_easy_view"]=historian_preference(member_id,"easy_view",False)
    return bool(st.session_state["historian_easy_view"])


def historian_rows(table_name,**filters):
    q=table(table_name).select("*")
    for k,v in filters.items():
        q=q.eq(k,v)
    try:
        return q.order("created_at",desc=True).execute().data or []
    except Exception as ex:
        raise RuntimeError(
            f"Historian table '{table_name}' is not ready. "
            "Apply migration_v3_6_19_historian_officer_hardening.sql after the earlier Historian migration. "
            f"Database message: {ex}"
        )


def historian_event_fingerprint(event_date,title,location=""):
    title_norm=re.sub(r"[^a-z0-9]+"," ",str(title or "").casefold()).strip()
    loc_norm=re.sub(r"[^a-z0-9]+"," ",str(location or "").casefold()).strip()
    raw=f"{event_date or 'UNKNOWN'}|{title_norm}|{loc_norm}"
    return hashlib.sha256(raw.encode("utf-8")).hexdigest()


def historian_photo_fingerprint(file_bytes):
    return hashlib.sha256(file_bytes).hexdigest()

def historian_completeness_issues(record):
    issues=[]
    if not record.get("event_date"):
        issues.append("Missing date")
    if not str(record.get("title") or "").strip():
        issues.append("Missing title")
    if not str(record.get("description") or "").strip():
        issues.append("Missing description")
    if not str(record.get("source_note") or "").strip() and not record.get("source_document_id"):
        issues.append("Missing source/evidence")
    if not record.get("has_photos"):
        issues.append("No photos")
    if not record.get("has_flyer_program"):
        issues.append("No flyer/program")
    if record.get("needs_identification"):
        issues.append("People need identification")
    if record.get("source_conflict"):
        issues.append("Source conflict")
    return issues



def historian_events_equivalent(a,b):
    ta=re.sub(r"[^a-z0-9]+"," ",str(a.get("title") or "").casefold()).strip()
    tb=re.sub(r"[^a-z0-9]+"," ",str(b.get("title") or "").casefold()).strip()
    if not ta or ta!=tb:
        return False
    da=str(a.get("event_date") or "")
    db=str(b.get("event_date") or "")
    if da and db and da!=db:
        return False
    la=re.sub(r"[^a-z0-9]+"," ",str(a.get("location") or "").casefold()).strip()
    lb=re.sub(r"[^a-z0-9]+"," ",str(b.get("location") or "").casefold()).strip()
    return (not la or not lb or la==lb)


def historian_year_score(year):
    rows_=[r for r in historian_event_rows() if r.get("sorority_year")==year]
    if not rows_:
        return 0,0,0
    complete=sum(1 for r in rows_ if not historian_completeness_issues(r))
    return complete,len(rows_),round(100*complete/len(rows_))


def historian_archive_bytes(year):
    out=BytesIO()
    with zipfile.ZipFile(out,"w",zipfile.ZIP_DEFLATED) as z:
        events=[r for r in historian_event_rows(include_archived=True) if r.get("sorority_year")==year]
        event_ids={r.get("id") for r in events}
        payloads={
            "events":events,
            "photos":[r for r in historian_rows("historian_photo_records") if r.get("event_memory_id") in event_ids],
            "milestones":[r for r in historian_rows("historian_milestones") if r.get("sorority_year")==year],
            "leadership":[r for r in historian_rows("historian_leadership_history") if r.get("sorority_year")==year],
            "sources":[r for r in historian_rows("historian_source_documents") if r.get("sorority_year")==year],
            "handoff":[r for r in historian_rows("historian_handoff_notes") if r.get("sorority_year")==year],
            "assets":[r for r in historian_rows("historian_event_assets") if r.get("event_memory_id") in event_ids],
        }
        for name,rows_ in payloads.items():
            data=pd.DataFrame(rows_).to_csv(index=False).encode("utf-8") if rows_ else b""
            z.writestr(f"{year}_{name}.csv",data)
        z.writestr("README.json",json.dumps({
            "sorority_year":year,
            "generated_at":datetime.now(timezone.utc).isoformat(),
            "generated_by":member_name,
            "note":"Archive index/metadata export. Protected source files and photos remain in private storage unless separately downloaded.",
        },indent=2))
    return out.getvalue()


def historian_archive_status():
    events=historian_event_rows()
    photos=historian_rows("historian_photo_records")
    issues=[]
    if not events and not photos:
        return "Needs Attention",["No event-memory or photo records have been added yet."]

    for r in events:
        row_issues=historian_completeness_issues(r)
        if row_issues:
            issues.append(f"{r.get('title') or 'Untitled record'}: "+", ".join(row_issues))

    orphan=[p for p in photos if not p.get("event_memory_id")]
    if orphan:
        issues.append(f"{len(orphan)} photo(s) are not linked to an event")
    unidentified=[p for p in photos if p.get("needs_identification")]
    if unidentified:
        issues.append(f"{len(unidentified)} photo(s) still need identification")

    try:
        unresolved=[r for r in historian_rows("historian_handoff_notes") if not r.get("resolved")]
        if unresolved:
            issues.append(f"{len(unresolved)} Historian handoff item(s) are still open")
    except Exception:
        pass

    if not issues:
        return "Archive Complete",[]
    if len(issues)>=5:
        return "Major Records Missing",issues
    return "Needs Attention",issues


def historian_help_callout(title,body):
    if historian_help_enabled():
        st.info(f"**{title}**\n\n{body}")

def render_historian_continuity_center(mid,is_admin_flag=False):
    if not can_use_historian_continuity(mid,is_admin_flag):
        return

    st.markdown("## Historian Continuity Center")
    st.caption(
        "Preserve the affiliate's HERstory with source evidence, clear verification status, "
        "easy successor handoff, and no guessing when dates or details are unknown."
    )

    c1,c2=st.columns(2)
    with c1:
        st.toggle(
            "Show Historian Help Guide",
            key="historian_help_guide",
            help="Turn step-by-step Historian guidance on or off.",
            on_change=lambda: save_historian_preference(mid,"help_guide",st.session_state.get("historian_help_guide"))
        )
    with c2:
        st.toggle(
            "Historian Easy / Large View",
            key="historian_easy_view",
            help="Uses larger controls and simpler layouts.",
            on_change=lambda: save_historian_preference(mid,"easy_view",st.session_state.get("historian_easy_view"))
        )

    if historian_easy_view():
        st.markdown(
            "<style>"
            "div.stButton > button {min-height:54px;font-size:18px;font-weight:700;}"
            "div[data-baseweb='input'] input, textarea {font-size:18px !important;}"
            "</style>",
            unsafe_allow_html=True
        )
        st.info("Easy / Large View is ON.")

    status,issues=historian_archive_status()
    if status=="Archive Complete":
        st.success("Historian Completeness: ARCHIVE COMPLETE")
    elif status=="Major Records Missing":
        st.error("Historian Completeness: MAJOR RECORDS MISSING")
    else:
        st.warning("Historian Completeness: NEEDS ATTENTION")
    for item in issues[:10]:
        st.write(f"- {item}")

    tabs=st.tabs([
        "Start Here",
        "Source Documents",
        "Event Memory",
        "Photos",
        "Approved Flyers",
        "Missing / Review Queue",
        "Timeline & Search",
        "Leadership & Milestones",
        "Year in Review",
        "Handoff / Archive",
        "Access",
    ])

    # 1. START HERE
    with tabs[0]:
        historian_help_callout(
            "Start Here",
            "Choose the task that matches what you have in front of you. Original records should be uploaded before entering facts from them whenever practical."
        )
        choices=[
            ("I found an old Historian report","Upload it under Source Documents, then connect events/milestones to that source."),
            ("I need to record an event","Use Event Memory and select the source that proves it happened."),
            ("I have photos","Upload them under Photos and connect them to the event."),
            ("A flyer was approved","Use Approved Flyers to connect the final approved flyer to the event archive."),
            ("Something is missing or uncertain","Leave the unknown field blank and use Needs Review instead of guessing."),
            ("I am preparing the next Historian","Use Handoff / Archive to document unresolved items and export the year index."),
        ]
        for title,desc in choices:
            with st.container(border=True):
                st.markdown(f"**{title}**")
                st.write(desc)

        st.markdown("### Sorority-Year Completeness")
        score=[]
        for sy in sorority_year_options():
            complete,total,pct=historian_year_score(sy)
            score.append({"Sorority Year":sy,"Complete Event Records":complete,"Total Event Records":total,"Completeness %":pct})
        st.dataframe(pd.DataFrame(score),hide_index=True,use_container_width=True)

    # 2. SOURCE DOCUMENTS
    with tabs[1]:
        historian_help_callout(
            "Source Documents",
            "This is where original Historian reports, newsletters, flyers, minutes, certificates, and other evidence are preserved. Facts entered later can point back to these files."
        )
        with st.form("historian_source_upload",clear_on_submit=True):
            sy=st.selectbox("Sorority Year",sorority_year_options(),key="hist_src_sy")
            source_type=st.selectbox("Document Type",HISTORIAN_SOURCE_TYPES)
            date_known=st.checkbox("Document date is known",value=True,key="hist_src_date_known")
            source_date=st.date_input(
                "Document Date",
                value=date.today(),
                min_value=AFFILIATE_ESTABLISHED_DATE,
                max_value=date.today(),
                disabled=not date_known
            )
            title=st.text_input("Source Document Title")
            description=st.text_area("What this source supports")
            visibility=st.selectbox("Visibility / Privacy",HISTORIAN_VISIBILITY_OPTIONS)
            upload=st.file_uploader(
                "Source File",
                type=["pdf","docx","xlsx","xls","csv","png","jpg","jpeg"],
                key="hist_source_file"
            )
            submit=st.form_submit_button("Save Source Document",use_container_width=True)

        if submit:
            try:
                if not title.strip():
                    raise ValueError("Enter a source document title.")
                raw,ext=historian_file_validate(upload)
                digest=hashlib.sha256(raw).hexdigest()
                dup=table("historian_source_documents").select("id,title").eq("file_sha256",digest).execute().data or []
                if dup:
                    raise ValueError(f"This exact source file is already stored as '{dup[0].get('title')}'.")
                safe_name=re.sub(r"[^A-Za-z0-9._-]+","_",Path(upload.name or f"source{ext}").name)
                folder=str(source_date.year) if date_known else "unknown-date"
                path=f"historian-sources/{folder}/{digest[:12]}_{safe_name}"
                stored=upload_private(raw,path,upload.type or "application/octet-stream")
                if not stored:
                    raise RuntimeError("Source file upload failed.")
                row=table("historian_source_documents").insert({
                    "sorority_year":sy,
                    "source_type":source_type,
                    "source_date":source_date.isoformat() if date_known else None,
                    "title":title.strip(),
                    "description":description.strip(),
                    "visibility":visibility,
                    "file_path":stored,
                    "file_name":upload.name or safe_name,
                    "file_sha256":digest,
                    "uploaded_by_member_id":mid,
                    "uploaded_by_name":member_name,
                    "created_at":datetime.now(timezone.utc).isoformat(),
                }).execute().data or []
                historian_change_log("CREATE","source_document",row[0]["id"] if row else "",title,mid)
                st.success("Source document saved.")
                st.rerun()
            except Exception as ex:
                st.error(f"Source document was not saved: {ex}")

        sources=historian_rows("historian_source_documents")
        if sources:
            st.dataframe(pd.DataFrame([{
                "Sorority Year":r.get("sorority_year"),
                "Date":r.get("source_date") or "Unknown",
                "Type":r.get("source_type"),
                "Title":r.get("title"),
                "Visibility":r.get("visibility"),
                "File":r.get("file_name"),
            } for r in sources]),hide_index=True,use_container_width=True)

        st.markdown("### Legacy History CSV Import")
        st.download_button(
            "Download Historian Import Template",
            historian_bulk_template_bytes(),
            file_name="NBS_Historian_Legacy_Event_Import_Template.csv",
            mime="text/csv",
            use_container_width=True
        )
        bulk=st.file_uploader("Upload Completed Historian Event CSV",type=["csv"],key="historian_bulk_csv")
        if st.button("Import Legacy Event Records",use_container_width=True):
            try:
                if bulk is None:
                    raise ValueError("Choose the completed Historian CSV.")
                df=pd.read_csv(bulk)
                required={"sorority_year","date_known","event_date","record_type","title","location","description","people_involved","verification_status","source_note"}
                missing_cols=required-set(df.columns)
                if missing_cols:
                    raise ValueError("Missing columns: "+", ".join(sorted(missing_cols)))
                if len(df)>500:
                    raise ValueError("Import limit is 500 rows at a time.")
                created=0; skipped=0; errors=[]
                existing=historian_event_rows(include_archived=True)
                for idx,row in df.iterrows():
                    try:
                        title=str(row.get("title") or "").strip()
                        if not title or title.casefold()=="nan":
                            raise ValueError("title is required")
                        known=str(row.get("date_known") or "").strip().casefold() in {"true","1","yes","y"}
                        ev_date=None
                        if known:
                            ev_date=pd.to_datetime(row.get("event_date"),errors="raise").date().isoformat()
                        proposed={"title":title,"event_date":ev_date,"location":str(row.get("location") or "")}
                        if any(historian_events_equivalent(x,proposed) for x in existing):
                            skipped+=1
                            continue
                        fp=historian_event_fingerprint(ev_date,title,proposed["location"])
                        payload={
                            "sorority_year":str(row.get("sorority_year") or "").strip(),
                            "event_date":ev_date,
                            "record_type":str(row.get("record_type") or "Other").strip(),
                            "title":title,
                            "location":str(row.get("location") or "").strip(),
                            "description":str(row.get("description") or "").strip(),
                            "people_involved":str(row.get("people_involved") or "").strip(),
                            "verification_status":str(row.get("verification_status") or "Previously Documented — Needs Source").strip(),
                            "source_note":str(row.get("source_note") or "").strip(),
                            "has_photos":str(row.get("has_photos") or "").strip().casefold() in {"true","1","yes","y"},
                            "has_flyer_program":str(row.get("has_flyer_program") or "").strip().casefold() in {"true","1","yes","y"},
                            "needs_identification":str(row.get("needs_identification") or "").strip().casefold() in {"true","1","yes","y"},
                            "source_conflict":str(row.get("source_conflict") or "").strip().casefold() in {"true","1","yes","y"},
                            "record_status":"Needs Attention",
                            "event_fingerprint":fp,
                            "is_archived":False,
                            "created_by_member_id":mid,
                            "created_by_name":member_name,
                            "created_at":datetime.now(timezone.utc).isoformat(),
                        }
                        table("historian_event_memory").insert(payload).execute()
                        existing.append(payload); created+=1
                    except Exception as ex:
                        errors.append(f"Row {idx+2}: {ex}")
                st.success(f"Import complete: {created} created, {skipped} duplicate(s) skipped.")
                if errors:
                    st.warning("\n".join(errors[:20]))
                st.rerun()
            except Exception as ex:
                st.error(f"Import was not completed: {ex}")

    # 3. EVENT MEMORY
    with tabs[2]:
        historian_help_callout(
            "Event Memory",
            "Record what happened and connect it to the source that supports it. Unknown old dates should stay unknown rather than defaulting to today."
        )
        sources=historian_rows("historian_source_documents")
        source_map={r["id"]:f"{r.get('source_date') or 'Unknown date'} • {r.get('title')}" for r in sources}

        with st.form("historian_event_memory_form",clear_on_submit=True):
            sy=st.selectbox("Sorority Year",sorority_year_options(),key="hist_event_sy")
            date_known=st.checkbox("Event date is known",value=True,key="hist_event_date_known")
            event_date=st.date_input(
                "Event Date",
                value=date.today(),
                min_value=AFFILIATE_ESTABLISHED_DATE,
                max_value=date.today(),
                disabled=not date_known
            )
            record_type=st.selectbox("Record Type",HISTORIAN_RECORD_TYPES)
            title=st.text_input("Event / Record Title",max_chars=180)
            location=st.text_input("Location")
            description=st.text_area("Description / What Happened",max_chars=2500)
            people=st.text_area("People / Groups Involved")
            verification=st.selectbox("Verification Status",HISTORIAN_VERIFICATION_OPTIONS)
            source_document_id=st.selectbox(
                "Linked Source Document",
                [None]+list(source_map.keys()),
                format_func=lambda x:"No linked file" if x is None else source_map[x]
            )
            source_note=st.text_input("Additional Source / Evidence Note")
            c1,c2=st.columns(2)
            with c1:
                has_photos=st.checkbox("Photos are available")
                has_flyer_program=st.checkbox("Flyer / program is available")
            with c2:
                needs_identification=st.checkbox("People still need identification")
                source_conflict=st.checkbox("Sources disagree / needs review")
            status=st.selectbox("Record Status",HISTORIAN_RECORD_STATUS_OPTIONS)
            b1,b2=st.columns(2)
            save_draft=b1.form_submit_button("Save Draft",use_container_width=True)
            save_record=b2.form_submit_button("Save Event Record",use_container_width=True)

        if save_draft or save_record:
            try:
                if not title.strip():
                    raise ValueError("Enter a title.")
                ev_date=event_date.isoformat() if date_known else None
                if status=="Archive Ready" and (
                    not date_known or not description.strip() or (not source_note.strip() and not source_document_id)
                ):
                    raise ValueError("Archive Ready requires a known date, description, and source/evidence.")
                proposed={"title":title,"event_date":ev_date,"location":location}
                candidates=historian_event_rows(include_archived=True)
                dup=next((r for r in candidates if historian_events_equivalent(r,proposed)),None)
                if dup:
                    raise ValueError(f"Possible duplicate already exists: {dup.get('title')}.")
                final_status="Draft" if save_draft else status
                fp=historian_event_fingerprint(ev_date,title,location)
                row=table("historian_event_memory").insert({
                    "sorority_year":sy,
                    "event_date":ev_date,
                    "record_type":record_type,
                    "title":title.strip(),
                    "location":location.strip(),
                    "description":description.strip(),
                    "people_involved":people.strip(),
                    "verification_status":verification,
                    "source_document_id":source_document_id,
                    "source_note":source_note.strip(),
                    "has_photos":has_photos,
                    "has_flyer_program":has_flyer_program,
                    "needs_identification":needs_identification,
                    "source_conflict":source_conflict,
                    "record_status":final_status,
                    "event_fingerprint":fp,
                    "is_archived":False,
                    "created_by_member_id":mid,
                    "created_by_name":member_name,
                    "created_at":datetime.now(timezone.utc).isoformat(),
                }).execute().data or []
                historian_change_log("CREATE","event_memory",row[0]["id"] if row else "",title,mid)
                st.success(f"Event record saved as {final_status}.")
                st.rerun()
            except Exception as ex:
                st.error(f"Event was not saved: {ex}")

        events=historian_event_rows()
        if events:
            st.markdown("### Edit / Archive Existing Event")
            emap={r["id"]:f"{r.get('event_date') or 'Date Unknown'} • {r.get('title')} • {r.get('record_status')}" for r in events}
            eid=st.selectbox("Event",list(emap.keys()),format_func=lambda x:emap[x],key="hist_edit_event_id")
            erow=next(r for r in events if r["id"]==eid)
            new_desc=st.text_area("Description",value=str(erow.get("description") or ""),key="hist_edit_desc")
            new_people=st.text_area("People / Groups",value=str(erow.get("people_involved") or ""),key="hist_edit_people")
            new_status=st.selectbox(
                "Status",
                HISTORIAN_RECORD_STATUS_OPTIONS,
                index=HISTORIAN_RECORD_STATUS_OPTIONS.index(erow.get("record_status")) if erow.get("record_status") in HISTORIAN_RECORD_STATUS_OPTIONS else 1,
                key="hist_edit_status"
            )
            edit_note=st.text_input("Why are you changing this record?",key="hist_edit_reason")
            c1,c2=st.columns(2)
            if c1.button("Save Event Changes",use_container_width=True):
                if not edit_note.strip():
                    st.error("Enter a reason for the change so the audit trail remains clear.")
                else:
                    table("historian_event_memory").update({
                        "description":new_desc.strip(),
                        "people_involved":new_people.strip(),
                        "record_status":new_status,
                    }).eq("id",eid).execute()
                    historian_change_log("UPDATE","event_memory",eid,edit_note.strip(),mid)
                    st.success("Event changes saved.")
                    st.rerun()

            if c2.button("Archive / Remove from Active View",use_container_width=True):
                st.session_state[f"archive_event_confirm_{eid}"]=True
            if st.session_state.get(f"archive_event_confirm_{eid}"):
                reason=st.text_input("Archive reason",key=f"archive_reason_{eid}")
                a1,a2=st.columns(2)
                if a1.button("Confirm Archive",use_container_width=True,key=f"archive_yes_{eid}"):
                    if not reason.strip():
                        st.error("Enter the archive reason.")
                    else:
                        table("historian_event_memory").update({
                            "is_archived":True,
                            "archived_at":datetime.now(timezone.utc).isoformat(),
                            "archived_by_name":member_name,
                            "archive_reason":reason.strip(),
                        }).eq("id",eid).execute()
                        historian_change_log("ARCHIVE","event_memory",eid,reason.strip(),mid)
                        st.session_state.pop(f"archive_event_confirm_{eid}",None)
                        st.success("Event archived from active view; history remains preserved.")
                        st.rerun()
                if a2.button("Cancel",use_container_width=True,key=f"archive_no_{eid}"):
                    st.session_state.pop(f"archive_event_confirm_{eid}",None)
                    st.rerun()

    # 4. PHOTOS
    with tabs[3]:
        historian_help_callout(
            "Photos",
            "Keep photos private by default unless public use is actually approved. Unknown dates and unknown people should stay marked for review."
        )
        events=historian_event_rows()
        event_map={r["id"]:f"{r.get('event_date') or 'Date Unknown'} • {r.get('title')}" for r in events}
        with st.form("historian_photo_record_form",clear_on_submit=True):
            event_id=st.selectbox("Connect to Event",[None]+list(event_map.keys()),format_func=lambda x:"Not connected yet" if x is None else event_map[x])
            date_known=st.checkbox("Photo date / approximate date is known",value=True,key="hist_photo_date_known")
            photo_date=st.date_input(
                "Photo Date / Approximate Date",
                value=date.today(),
                min_value=AFFILIATE_ESTABLISHED_DATE,
                max_value=date.today(),
                disabled=not date_known
            )
            caption=st.text_area("Caption")
            people=st.text_area("People Identified")
            publication_status=st.selectbox("Photo Use / Publication Status",HISTORIAN_PUBLICATION_STATUS_OPTIONS)
            publication_basis=st.text_input("Approval / Source Note")
            photo=st.file_uploader("Photo",type=["jpg","jpeg","png"],key="hist_photo_upload")
            save=st.form_submit_button("Save Photo Record",use_container_width=True)

        if save:
            stored=None
            try:
                if publication_status=="Approved for Public Use" and not publication_basis.strip():
                    raise ValueError("Approved for Public Use requires an approval/source note.")
                raw=validate_image_upload(photo,max_mb=10)
                digest=historian_photo_fingerprint(raw)
                dup=table("historian_photo_records").select("id,file_name").eq("file_sha256",digest).execute().data or []
                if dup:
                    raise ValueError("This exact photo is already in the Historian archive.")
                safe_name=re.sub(r"[^A-Za-z0-9._-]+","_",Path(photo.name or "photo.jpg").name)
                folder=str(photo_date.year) if date_known else "unknown-date"
                path=f"historian-continuity/{folder}/{digest[:12]}_{safe_name}"
                stored=upload_private(raw,path,photo.type or "image/jpeg")
                table("historian_photo_records").insert({
                    "event_memory_id":event_id,
                    "photo_date":photo_date.isoformat() if date_known else None,
                    "caption":caption.strip(),
                    "people_identified":people.strip(),
                    "publication_status":publication_status,
                    "publication_basis":publication_basis.strip(),
                    "needs_identification":not bool(people.strip()),
                    "file_path":stored,
                    "file_name":photo.name or safe_name,
                    "file_sha256":digest,
                    "uploaded_by_member_id":mid,
                    "uploaded_by_name":member_name,
                    "created_at":datetime.now(timezone.utc).isoformat(),
                }).execute()
                if event_id:
                    table("historian_event_memory").update({"has_photos":True}).eq("id",event_id).execute()
                st.success("Photo saved.")
                st.rerun()
            except Exception as ex:
                st.error(f"Photo was not saved: {ex}")

        photos=historian_rows("historian_photo_records")
        if photos:
            st.dataframe(pd.DataFrame([{
                "Date":r.get("photo_date") or "Unknown",
                "Caption":r.get("caption"),
                "People":r.get("people_identified"),
                "Publication":r.get("publication_status"),
                "Needs Identification":r.get("needs_identification"),
                "Linked Event":event_map.get(r.get("event_memory_id"),"Not linked"),
                "File":r.get("file_name"),
            } for r in photos]),hide_index=True,use_container_width=True)

            pmap={r["id"]:f"{r.get('photo_date') or 'Unknown'} • {r.get('file_name')}" for r in photos}
            pid=st.selectbox("Update Photo Metadata",list(pmap.keys()),format_func=lambda x:pmap[x],key="hist_update_photo_id")
            prow=next(r for r in photos if r["id"]==pid)
            new_event=st.selectbox(
                "Link / Relink Event",
                [None]+list(event_map.keys()),
                index=([None]+list(event_map.keys())).index(prow.get("event_memory_id")) if prow.get("event_memory_id") in event_map else 0,
                format_func=lambda x:"Not connected yet" if x is None else event_map[x],
                key="hist_photo_event_update"
            )
            new_caption=st.text_area("Caption",value=str(prow.get("caption") or ""),key="hist_photo_caption_update")
            new_people=st.text_area("People Identified",value=str(prow.get("people_identified") or ""),key="hist_photo_people_update")
            new_pub=st.selectbox(
                "Publication Status",
                HISTORIAN_PUBLICATION_STATUS_OPTIONS,
                index=HISTORIAN_PUBLICATION_STATUS_OPTIONS.index(prow.get("publication_status")) if prow.get("publication_status") in HISTORIAN_PUBLICATION_STATUS_OPTIONS else 0,
                key="hist_photo_pub_update"
            )
            new_basis=st.text_input("Approval / Source Note",value=str(prow.get("publication_basis") or ""),key="hist_photo_basis_update")
            if st.button("Save Photo Metadata Update",use_container_width=True):
                if new_pub=="Approved for Public Use" and not new_basis.strip():
                    st.error("Approved for Public Use requires an approval/source note.")
                else:
                    table("historian_photo_records").update({
                        "event_memory_id":new_event,
                        "caption":new_caption.strip(),
                        "people_identified":new_people.strip(),
                        "publication_status":new_pub,
                        "publication_basis":new_basis.strip(),
                        "needs_identification":not bool(new_people.strip()),
                    }).eq("id",pid).execute()
                    historian_change_log("UPDATE","photo",pid,"Photo metadata updated",mid)
                    if new_event:
                        table("historian_event_memory").update({"has_photos":True}).eq("id",new_event).execute()
                    st.success("Photo metadata updated.")
                    st.rerun()

    # 5. APPROVED FLYERS -> HISTORIAN
    with tabs[4]:
        historian_help_callout(
            "Approved Flyers",
            "After Chapter approval, connect the final flyer to the permanent event archive instead of leaving it only in the communications workflow."
        )
        try:
            approved=table("philo_flyer_submissions").select("*").in_("status",["Approved by Chapter","Published / Final"]).order("updated_at",desc=True).execute().data or []
        except Exception:
            approved=[]
        assets=historian_rows("historian_event_assets")
        archived_flyer_ids={r.get("flyer_submission_id") for r in assets if r.get("flyer_submission_id")}
        pending=[r for r in approved if r.get("id") not in archived_flyer_ids]
        events=historian_event_rows()
        event_map={r["id"]:f"{r.get('event_date') or 'Date Unknown'} • {r.get('title')}" for r in events}

        if not pending:
            st.success("No approved flyers are waiting to be connected to the Historian archive.")
        elif not events:
            st.warning("Create the Event Memory record first, then connect the approved flyer.")
        else:
            fmap={r["id"]:f"{r.get('title')} • {r.get('chapter_decision_date') or 'Approved'}" for r in pending}
            flyer_id=st.selectbox("Approved Flyer",list(fmap.keys()),format_func=lambda x:fmap[x],key="hist_approved_flyer")
            event_id=st.selectbox("Connect to Event",list(event_map.keys()),format_func=lambda x:event_map[x],key="hist_flyer_event")
            if st.button("Archive Approved Flyer with Event",use_container_width=True):
                row=table("historian_event_assets").insert({
                    "event_memory_id":event_id,
                    "asset_type":"Approved Flyer",
                    "flyer_submission_id":flyer_id,
                    "created_by_member_id":mid,
                    "created_by_name":member_name,
                    "created_at":datetime.now(timezone.utc).isoformat(),
                }).execute().data or []
                table("historian_event_memory").update({"has_flyer_program":True}).eq("id",event_id).execute()
                historian_change_log("LINK","approved_flyer",row[0]["id"] if row else "",f"Flyer {flyer_id} linked to event {event_id}",mid)
                st.success("Approved flyer added to the Historian event archive.")
                st.rerun()

    # 6. MISSING / REVIEW QUEUE
    with tabs[5]:
        historian_help_callout(
            "Missing / Review Queue",
            "This queue tells the Historian what still needs attention. It is better to show Unknown or Needs Review than to invent historical facts."
        )
        events=historian_event_rows()
        missing=[]
        for r in events:
            issues_=historian_completeness_issues(r)
            if issues_:
                missing.append({
                    "Sorority Year":r.get("sorority_year"),
                    "Event":r.get("title"),
                    "Date":r.get("event_date") or "Unknown",
                    "Needs Attention":", ".join(issues_),
                    "Verification":r.get("verification_status"),
                    "Status":r.get("record_status"),
                })
        if missing:
            st.dataframe(pd.DataFrame(missing),hide_index=True,use_container_width=True)
        else:
            st.success("No incomplete active event-memory records were found.")

        photos=historian_rows("historian_photo_records")
        orphan=[r for r in photos if not r.get("event_memory_id")]
        if orphan:
            st.warning(f"{len(orphan)} photo(s) are not connected to an event.")
        unidentified=[r for r in photos if r.get("needs_identification")]
        if unidentified:
            st.warning(f"{len(unidentified)} photo(s) still need identification.")

    # 7. TIMELINE & SEARCH
    with tabs[6]:
        historian_help_callout(
            "Timeline & Search",
            "Search across events, milestones, leadership, source documents, photos, and Historian handoff items."
        )
        st.markdown("### Sorority-Year Timeline")
        events=historian_event_rows()
        for sy in reversed(sorority_year_options()):
            current=(sy==HISTORIAN_CURRENT_SORORITY_YEAR)
            with st.expander(f"{sy} — {'CURRENT YEAR' if current else 'Previous Sorority Year'}",expanded=current):
                rows_=[r for r in events if r.get("sorority_year")==sy]
                if not rows_:
                    st.caption("No event-memory records yet.")
                for r in sorted(rows_,key=lambda x:str(x.get("event_date") or "9999")):
                    st.write(f"**{r.get('event_date') or 'Date Unknown'} — {r.get('title')}**")
                    st.caption(
                        f"{r.get('record_type') or 'Event'} • {r.get('verification_status') or 'Needs source'} • "
                        f"Source: {r.get('source_note') or ('Linked source file' if r.get('source_document_id') else 'Not listed')}"
                    )

        st.markdown("### Search Archive")
        year_filter=st.selectbox("Sorority Year",["All"]+sorority_year_options(),key="hist_search_year")
        keyword=st.text_input("Keyword",key="hist_search_keyword").strip().casefold()
        results=[]

        for r in historian_event_rows(include_archived=True):
            if year_filter!="All" and r.get("sorority_year")!=year_filter: continue
            hay=" ".join(str(r.get(k) or "") for k in ["title","description","people_involved","location","source_note","record_type","verification_status"]).casefold()
            if keyword and keyword not in hay: continue
            results.append({"Type":"Event","Sorority Year":r.get("sorority_year"),"Date":r.get("event_date"),"Title":r.get("title"),"Details":r.get("description"),"Status":r.get("record_status")})

        for r in historian_rows("historian_milestones"):
            if year_filter!="All" and r.get("sorority_year")!=year_filter: continue
            hay=" ".join(str(r.get(k) or "") for k in ["title","detail","milestone_type","source_note"]).casefold()
            if keyword and keyword not in hay: continue
            results.append({"Type":"Milestone","Sorority Year":r.get("sorority_year"),"Date":r.get("milestone_date"),"Title":r.get("title"),"Details":r.get("detail"),"Status":r.get("milestone_type")})

        for r in historian_rows("historian_leadership_history"):
            if year_filter!="All" and r.get("sorority_year")!=year_filter: continue
            hay=" ".join(str(r.get(k) or "") for k in ["office","person_name","notes","source_note"]).casefold()
            if keyword and keyword not in hay: continue
            results.append({"Type":"Leadership","Sorority Year":r.get("sorority_year"),"Date":"","Title":f"{r.get('office')} — {r.get('person_name')}","Details":r.get("notes"),"Status":"Source linked" if r.get("source_note") else "Needs source"})

        for r in historian_rows("historian_source_documents"):
            if year_filter!="All" and r.get("sorority_year")!=year_filter: continue
            hay=" ".join(str(r.get(k) or "") for k in ["title","description","source_type","file_name"]).casefold()
            if keyword and keyword not in hay: continue
            results.append({"Type":"Source Document","Sorority Year":r.get("sorority_year"),"Date":r.get("source_date"),"Title":r.get("title"),"Details":r.get("description"),"Status":r.get("visibility")})

        photos=historian_rows("historian_photo_records")
        event_by_id={r.get("id"):r for r in historian_event_rows(include_archived=True)}
        for r in photos:
            er=event_by_id.get(r.get("event_memory_id"),{})
            sy=er.get("sorority_year")
            if year_filter!="All" and sy!=year_filter: continue
            hay=" ".join(str(r.get(k) or "") for k in ["caption","people_identified","publication_status","publication_basis","file_name"]).casefold()
            if keyword and keyword not in hay: continue
            results.append({"Type":"Photo","Sorority Year":sy,"Date":r.get("photo_date"),"Title":r.get("file_name"),"Details":r.get("caption"),"Status":r.get("publication_status")})

        if results:
            st.dataframe(pd.DataFrame(results),hide_index=True,use_container_width=True)
        else:
            st.info("No matching archive records were found.")

    # 8. LEADERSHIP & MILESTONES
    with tabs[7]:
        historian_help_callout(
            "Leadership & Milestones",
            "Preserve Philo officer service, awards, recognition, anniversaries, and firsts with a supporting source."
        )
        with st.form("historian_leadership_form",clear_on_submit=True):
            sy=st.selectbox("Sorority Year",sorority_year_options(),key="hist_leader_sy")
            office=st.selectbox("Philo Office",OFFICER_POSITIONS)
            person=st.text_input("Philo Name")
            notes=st.text_area("Leadership Notes / Accomplishments")
            source=st.text_input("Source / Evidence")
            save=st.form_submit_button("Save Leadership History",use_container_width=True)
        if save:
            if not person.strip() or not source.strip():
                st.error("Philo name and source/evidence are required.")
            else:
                table("historian_leadership_history").upsert({
                    "sorority_year":sy,
                    "office":office,
                    "person_name":person.strip(),
                    "notes":notes.strip(),
                    "source_note":source.strip(),
                    "created_by_member_id":mid,
                    "created_by_name":member_name,
                    "created_at":datetime.now(timezone.utc).isoformat(),
                },on_conflict="sorority_year,office,person_name").execute()
                st.success("Leadership history saved.")
                st.rerun()

        with st.form("historian_milestone_form",clear_on_submit=True):
            sy=st.selectbox("Sorority Year",sorority_year_options(),key="hist_milestone_sy")
            date_known=st.checkbox("Milestone date is known",value=True,key="hist_milestone_date_known")
            mdate=st.date_input("Milestone Date",value=date.today(),min_value=AFFILIATE_ESTABLISHED_DATE,max_value=date.today(),disabled=not date_known)
            mtype=st.selectbox("Milestone Type",[
                "Affiliate Milestone","Anniversary","Award","Recognition","First / Historic First",
                "Community Honor","Regional Recognition","National Recognition","Leadership Milestone","Program Milestone","Other"
            ])
            title=st.text_input("Milestone / Recognition")
            detail=st.text_area("What Happened / Why It Matters")
            source=st.text_input("Supporting Source")
            save_m=st.form_submit_button("Save Milestone",use_container_width=True)
        if save_m:
            if not title.strip() or not source.strip():
                st.error("Milestone title and supporting source are required.")
            else:
                d=mdate.isoformat() if date_known else None
                fp=hashlib.sha256(f"{sy}|{d or 'UNKNOWN'}|{mtype}|{title.strip().casefold()}".encode("utf-8")).hexdigest()
                dup=table("historian_milestones").select("id").eq("milestone_fingerprint",fp).execute().data or []
                if dup:
                    st.error("A matching milestone record already exists.")
                else:
                    table("historian_milestones").insert({
                        "sorority_year":sy,
                        "milestone_date":d,
                        "milestone_type":mtype,
                        "title":title.strip(),
                        "detail":detail.strip(),
                        "source_note":source.strip(),
                        "milestone_fingerprint":fp,
                        "created_by_member_id":mid,
                        "created_by_name":member_name,
                        "created_at":datetime.now(timezone.utc).isoformat(),
                    }).execute()
                    st.success("Milestone saved.")
                    st.rerun()

    # 9. YEAR IN REVIEW
    with tabs[8]:
        historian_help_callout(
            "Year in Review",
            "Generate a draft only from records that are already in the archive. Missing information remains missing; the draft does not invent history."
        )
        sy=st.selectbox("Sorority Year",sorority_year_options(),key="hist_year_review")
        events=[r for r in historian_event_rows() if r.get("sorority_year")==sy]
        milestones=[r for r in historian_rows("historian_milestones") if r.get("sorority_year")==sy]
        leaders=[r for r in historian_rows("historian_leadership_history") if r.get("sorority_year")==sy]
        photos=historian_rows("historian_photo_records")
        event_ids={r.get("id") for r in events}
        photos=[p for p in photos if p.get("event_memory_id") in event_ids]

        if st.button("Create Year-in-Review Draft",use_container_width=True):
            complete,total,pct=historian_year_score(sy)
            parts=[f"Nu Beta Sigma Philo Affiliate — {sy} Year in Review"]
            parts.append(f"\nArchive completeness: {complete} of {total} event records complete ({pct}%).")
            if any(historian_completeness_issues(r) for r in events):
                parts.append("Some records still need review. This draft does not fill missing facts.")
            if events:
                parts.append("\nEvents and Activities:")
                for r in sorted(events,key=lambda x:str(x.get("event_date") or "9999")):
                    parts.append(f"- {r.get('event_date') or 'Date Unknown'}: {r.get('title')} — {r.get('description') or 'Description needs completion.'}")
            if milestones:
                parts.append("\nMilestones and Recognition:")
                for r in milestones:
                    parts.append(f"- {r.get('milestone_date') or 'Date Unknown'}: {r.get('title')} ({r.get('milestone_type')})")
            if leaders:
                parts.append("\nPhilo Leadership:")
                for r in leaders:
                    parts.append(f"- {r.get('office')}: {r.get('person_name')}")
            parts.append(f"\nPhotos indexed: {len(photos)}.")
            unidentified=sum(1 for p in photos if p.get("needs_identification"))
            if unidentified:
                parts.append(f"Photos still needing identification: {unidentified}.")
            st.session_state["historian_year_review_draft"]="\n".join(parts)

        st.text_area("Editable Year-in-Review Draft",key="historian_year_review_draft",height=420)

    # 10. HANDOFF / ARCHIVE
    with tabs[9]:
        historian_help_callout(
            "Handoff / Archive",
            "Document unresolved history before leaving office and export a year index so the next Historian has an independent reference."
        )
        with st.form("historian_handoff_form",clear_on_submit=True):
            category=st.selectbox("Handoff Category",[
                "Missing Photo Identification","Missing Event Record","Source Conflict",
                "Outstanding Upload","Milestone / Anniversary Follow-up","Storage / Access","Other"
            ])
            title=st.text_input("Handoff Item")
            note=st.text_area("What the next Historian needs to know")
            save_h=st.form_submit_button("Save Handoff Item",use_container_width=True)
        if save_h:
            if not title.strip() or not note.strip():
                st.error("Handoff item and explanation are required.")
            else:
                table("historian_handoff_notes").insert({
                    "sorority_year":HISTORIAN_CURRENT_SORORITY_YEAR,
                    "category":category,
                    "title":title.strip(),
                    "note":note.strip(),
                    "resolved":False,
                    "created_by_member_id":mid,
                    "created_by_name":member_name,
                    "created_at":datetime.now(timezone.utc).isoformat(),
                }).execute()
                st.success("Handoff item saved.")
                st.rerun()

        hrows=historian_rows("historian_handoff_notes")
        if hrows:
            st.dataframe(pd.DataFrame(hrows),hide_index=True,use_container_width=True)
            open_rows=[r for r in hrows if not r.get("resolved")]
            if open_rows:
                hmap={r["id"]:f"{r.get('category')} • {r.get('title')}" for r in open_rows}
                hid=st.selectbox("Resolve Handoff Item",list(hmap.keys()),format_func=lambda x:hmap[x],key="hist_handoff_resolve")
                resolution=st.text_input("Resolution / source note",key="hist_handoff_resolution")
                if st.button("Mark Handoff Item Resolved",use_container_width=True):
                    if not resolution.strip():
                        st.error("Enter the resolution/source note.")
                    else:
                        table("historian_handoff_notes").update({
                            "resolved":True,
                            "resolution_note":resolution.strip(),
                            "resolved_at":datetime.now(timezone.utc).isoformat(),
                        }).eq("id",hid).execute()
                        st.success("Handoff item resolved.")
                        st.rerun()

        sy=st.selectbox("Archive Sorority Year",sorority_year_options(),key="hist_archive_sy")
        st.download_button(
            "Download Historian Archive Index ZIP",
            historian_archive_bytes(sy),
            file_name=f"NBS_Historian_Archive_{sy}.zip",
            mime="application/zip",
            use_container_width=True
        )
        st.caption("The ZIP contains archive indexes/metadata. Protected files remain in private storage unless separately downloaded.")

    # 11. ACCESS / DELEGATE
    with tabs[10]:
        historian_help_callout(
            "Historian Access",
            "A joint Historian or designated helper may need access even if her primary officer title is not Historian. Delegate access is specific and revocable; it does not expose Sigma/Soror-only Advisor data."
        )
        if is_admin_flag or ("Historian" in current_position_names(mid)):
            try:
                people=philo_dropdown_people()
            except Exception:
                people=[]
            pmap={r["id"]:r["full_name"] for r in people if r.get("id") and r.get("full_name")}
            if pmap:
                delegate=st.selectbox("Grant Historian Delegate Access",list(pmap.keys()),format_func=lambda x:pmap[x],key="hist_delegate_member")
                note=st.text_input("Reason / role",placeholder="Example: Joint Historian 2024–2026")
                if st.button("Grant / Renew Delegate Access",use_container_width=True):
                    table("historian_delegate_access").upsert({
                        "member_id":delegate,
                        "member_name":pmap[delegate],
                        "access_note":note.strip(),
                        "active":True,
                        "granted_by_member_id":mid,
                        "granted_by_name":member_name,
                        "updated_at":datetime.now(timezone.utc).isoformat(),
                    },on_conflict="member_id").execute()
                    st.success("Historian delegate access granted.")
                    st.rerun()

            delegates=historian_delegate_rows()
            if delegates:
                st.dataframe(pd.DataFrame(delegates),hide_index=True,use_container_width=True)
                dmap={r["id"]:f"{r.get('member_name')} • {r.get('access_note')}" for r in delegates}
                did=st.selectbox("Revoke Delegate",list(dmap.keys()),format_func=lambda x:dmap[x],key="hist_revoke_delegate")
                if st.button("Revoke Historian Delegate Access",use_container_width=True):
                    table("historian_delegate_access").update({"active":False,"updated_at":datetime.now(timezone.utc).isoformat()}).eq("id",did).execute()
                    st.success("Historian delegate access revoked.")
                    st.rerun()
        else:
            st.info("Only the Historian or Advisor/Admin can manage Historian delegate access.")

def render_read_only_financial_history(mid,is_admin_flag=False):
    """Protected past-year financial history for authorized financial leadership."""
    if not can_view_financial_history(mid,is_admin_flag):
        return
    st.markdown("#### Financial History Over Time")
    st.caption(
        "Past Sorority Years are protected from editing in the app. The current 2026-2027 "
        "financial workflows remain interactive elsewhere in the Hub."
    )
    try:
        yearly=prior_year_finance_summary_rows()
        monthly=financial_history_month_rows()
    except Exception as ex:
        st.error(str(ex)); return
    if not yearly and not monthly:
        st.info("No verified past financial history has been imported yet.")
        return

    if yearly:
        sorted_years=sorted(yearly,key=lambda r:str(r.get("sorority_year") or ""))
        comparison=[]
        prior_close=None
        for r in sorted_years:
            close=r.get("closing_balance")
            yoy=None
            yoy_pct=None
            direction="BASE YEAR" if prior_close is None else ""
            if close is not None and prior_close not in (None,0):
                yoy=float(close)-float(prior_close)
                yoy_pct=yoy/float(prior_close)
                direction="UP" if yoy>0 else ("DOWN" if yoy<0 else "FLAT")
            comparison.append({
                "Sorority Year":r.get("sorority_year"),
                "Opening Balance":r.get("opening_balance"),
                "Deposits":r.get("deposits"),
                "Withdrawals":r.get("withdrawals"),
                "Recorded Net":r.get("net_activity"),
                "Ending Balance":close,
                "YOY $ Change":yoy,
                "YOY %":yoy_pct,
                "Fund Direction":direction,
                "Verification / Audit Status":r.get("verification_status"),
                "Source":r.get("source_note") or "",
            })
            if close is not None:
                prior_close=float(close)

        st.markdown("##### Year-over-Year Fund Position")
        st.dataframe(pd.DataFrame(comparison),hide_index=True,use_container_width=True)
        latest=[r for r in comparison if r.get("Ending Balance") is not None]
        if latest:
            last=latest[-1]
            if last.get("YOY $ Change") is not None:
                if last["YOY $ Change"]<0:
                    st.error(f"Fund position ended ${abs(last['YOY $ Change']):,.2f} below the previous Sorority Year.")
                elif last["YOY $ Change"]>0:
                    st.success(f"Fund position ended ${last['YOY $ Change']:,.2f} above the previous Sorority Year.")

    if monthly:
        years=sorted({str(r.get("sorority_year")) for r in monthly},reverse=True)
        selected=st.selectbox("View Past Sorority Year",years,key="financial_history_year")
        data=[r for r in monthly if str(r.get("sorority_year"))==selected]
        st.markdown("##### Monthly Deposits, Withdrawals & Source Categories")
        display=[]
        for r in data:
            display.append({
                "Month":r.get("month_start"),
                "Beginning Balance":r.get("beginning_balance"),
                "Dues":r.get("dues"),
                "Donations":r.get("donations"),
                "Zeffy / Online Giving":r.get("zeffy"),
                "Fundraising":r.get("fundraising"),
                "Total Deposits":r.get("total_deposits"),
                "Assessments":r.get("assessments"),
                "Reimbursements / Club Expenses":r.get("reimbursements"),
                "Gifts":r.get("gifts"),
                "Other Withdrawals":r.get("other_withdrawals"),
                "Total Withdrawals":r.get("total_withdrawals"),
                "Recorded Net":r.get("net_change"),
                "Reported Ending Balance":r.get("ending_balance"),
                "Reconciliation Variance":r.get("reconciliation_variance"),
                "Reconciliation Status":r.get("reconciliation_status"),
            })
        st.dataframe(pd.DataFrame(display),hide_index=True,use_container_width=True)
        if selected=="2024-2025":
            st.warning(
                "The 2024-2025 source workbook contains balance/line-item differences. "
                "The Hub preserves both the reported balances and recorded transaction categories "
                "and flags the variance for reconciliation; it does not rewrite the source history."
            )
        st.info("🔒 Past-year financial history is view-only in the Hub for audit protection.")

def render_prior_year_finance_import(mid,is_admin_flag=False):
    positions=current_position_names(mid)
    allowed=is_admin_flag or bool(
        positions & {"President","Financial Secretary","Treasurer"}
    )
    if not allowed:
        return

    st.markdown("#### Prior-Year Finance Summary")
    st.caption(
        "Import verified finance totals from the archive workbook or another structured Excel/CSV source. "
        "Rows marked SOURCE NEEDED are skipped automatically."
    )

    with st.expander("📥 Import Prior-Year Finance Data",expanded=False):
        finance_file=st.file_uploader(
            "Finance Workbook / CSV",
            type=["xlsx","xls","csv"],
            key="prior_year_finance_import_file"
        )
        if st.button(
            "Import Verified Finance Data",
            key="prior_year_finance_import_button",
            use_container_width=True
        ):
            try:
                imported=import_prior_year_finance(finance_file,mid,member_name)
                st.success(
                    "Imported verified finance data for: "+", ".join(imported)
                )
                st.rerun()
            except Exception as ex:
                st.error(f"Finance import was not completed: {ex}")

    try:
        summaries=prior_year_finance_summary_rows()
    except Exception as ex:
        st.error(str(ex))
        return

    if not summaries:
        st.info("No prior-year finance summaries have been imported yet.")
        return

    display=[]
    for r in summaries:
        display.append({
            "Sorority Year":r.get("sorority_year"),
            "Opening Balance":r.get("opening_balance"),
            "Deposits":r.get("deposits"),
            "Withdrawals":r.get("withdrawals"),
            "Net Activity":r.get("net_activity"),
            "Closing Balance":r.get("closing_balance"),
            "Bank Statement Date":r.get("bank_statement_date"),
            "Bank Statement Balance":r.get("bank_statement_balance"),
            "Verification Status":r.get("verification_status"),
            "Source":r.get("source_note") or "",
        })
    st.dataframe(pd.DataFrame(display),hide_index=True,use_container_width=True)

def render_previous_sorority_year_records(position,mid,is_admin_flag=False):
    st.markdown("### Previous Sorority Year Records")
    st.caption(
        "Upload reports and records from prior Sorority Years. "
        "The Nu Beta Sigma Philo Affiliate was established March 19, 2022."
    )

    years=sorority_year_options()
    if not years:
        st.info("No Sorority Year options are available yet.")
        return

    # Officers may file records only for an office they currently hold.
    # President may cover a vacant office; Advisor/Admin may manage every Philo office.
    authorized=[x.get("position") for x in member_offices(mid) if x.get("position")]
    if "President" in authorized:
        authorized += [x for x in vacant_offices() if x not in authorized]
    if is_admin_flag:
        authorized=list(OFFICER_POSITIONS)
    if position and position not in authorized and not is_admin_flag:
        authorized.append(position)
    authorized=list(dict.fromkeys([x for x in authorized if x]))

    with st.expander("⬆️ Upload a Previous Sorority Year Record",expanded=False):
        with st.form(f"previous_year_upload_{position}",clear_on_submit=True):
            c1,c2=st.columns(2)
            with c1:
                sy=st.selectbox("Sorority Year",years)
                doc_date=st.date_input(
                    "Document Date",
                    value=AFFILIATE_ESTABLISHED_DATE,
                    min_value=AFFILIATE_ESTABLISHED_DATE,
                    max_value=date.today()
                )
                doc_type=st.selectbox("Document Type",PREVIOUS_YEAR_DOCUMENT_TYPES)
                office=st.selectbox(
                    "Officer / Office",
                    authorized if authorized else [position],
                    index=(authorized.index(position) if position in authorized else 0)
                )
            with c2:
                archive_category=st.selectbox(
                    "Archive Category",PREVIOUS_YEAR_ARCHIVE_CATEGORIES
                )
                title=st.text_input("Title",max_chars=180)
                description=st.text_area("Description",max_chars=1200,height=110)
                uploaded=st.file_uploader(
                    "File Upload",
                    type=["pdf","doc","docx","xls","xlsx","csv","txt","rtf","png","jpg","jpeg"],
                    key=f"previous_year_file_{position}"
                )

            financial=previous_year_is_financial(doc_type,archive_category)
            if financial:
                st.info(
                    "🔒 Financial record: visibility will be restricted automatically "
                    "to Advisor/Admin and authorized Philo financial leadership "
                    "(President, Financial Secretary, and Treasurer)."
                )
            else:
                st.caption(
                    "Non-financial officer records are limited to Advisor/Admin, "
                    "the current holder of the selected office, the President, and the uploader."
                )
            submit=st.form_submit_button(
                "Upload to Previous Sorority Year Records",use_container_width=True
            )

        if submit:
            try:
                if not str(title or "").strip():
                    raise ValueError("Enter a document title.")
                if not office:
                    raise ValueError("Select an officer/office.")
                # A date must belong to the selected Sorority Year window.
                sy_start=int(str(sy).split("-")[0])
                sy_start_date=date(sy_start,7,1)
                sy_end_date=date(sy_start+1,6,30)
                # Special first year accepts establishment date through June 30, 2023.
                if sy_start==2022:
                    sy_start_date=AFFILIATE_ESTABLISHED_DATE
                if not (sy_start_date<=doc_date<=sy_end_date):
                    raise ValueError(
                        f"Document Date must fall within the selected {sy} Sorority Year."
                    )
                info=validate_previous_year_file(uploaded,max_mb=20)
                duplicate=previous_year_duplicate(info["sha256"],office,sy)
                if duplicate:
                    raise ValueError(
                        "Duplicate detected. This exact file is already archived for "
                        f"{office}, {sy} as “{duplicate.get('title') or duplicate.get('file_name')}”."
                    )

                financial=previous_year_is_financial(doc_type,archive_category)
                visibility=(
                    "Restricted - Financial Leadership & Advisor"
                    if financial else
                    "Restricted - Office Leadership & Advisor"
                )
                safe_name=re.sub(
                    r"[^A-Za-z0-9._-]+","_",Path(info["filename"]).name
                )
                storage_path=(
                    f"previous-sorority-year/{sy}/{office.replace(' ','_')}/"
                    f"{info['sha256'][:12]}_{safe_name}"
                )
                uploaded_path=upload_private(
                    info["bytes"],storage_path,info["mime_type"]
                )
                if not uploaded_path:
                    raise RuntimeError("File storage upload failed; no archive record was created.")

                payload={
                    "sorority_year":sy,
                    "document_date":doc_date.isoformat(),
                    "document_type":doc_type,
                    "office":office,
                    "title":str(title).strip(),
                    "description":str(description or "").strip(),
                    "archive_category":archive_category,
                    "file_path":uploaded_path,
                    "file_name":info["filename"],
                    "mime_type":info["mime_type"],
                    "file_size":info["size"],
                    "file_sha256":info["sha256"],
                    "is_financial":financial,
                    "visibility":visibility,
                    "uploaded_by_member_id":mid,
                    "uploaded_by_name":member_name,
                    "created_at":datetime.now(timezone.utc).isoformat(),
                }
                try:
                    table("previous_sorority_year_documents").insert(payload).execute()
                except Exception:
                    # Compensate so failed metadata writes do not leave an untracked file.
                    try:
                        (None if test_preview_active() else sb().storage.from_(cfg()["bucket"]).remove([uploaded_path]))
                    finally:
                        raise

                st.success("Previous Sorority Year record uploaded.")
                st.rerun()
            except Exception as ex:
                st.error(f"Upload was not completed: {ex}")

    render_prior_year_finance_import(mid,is_admin_flag)

    render_read_only_financial_history(mid,is_admin_flag)

    render_treasurer_continuity_center(mid,is_admin_flag)

    st.markdown("#### Records I Can Access")
    try:
        all_records=previous_year_record_rows()
    except Exception as ex:
        st.error(str(ex))
        return

    visible=[
        r for r in all_records
        if can_view_previous_year_record(r,mid,is_admin_flag)
    ]
    if not visible:
        st.info("No Previous Sorority Year records are available for your current access.")
        return

    filter_year=st.selectbox(
        "Filter by Sorority Year",
        ["All"]+years,
        key=f"previous_year_filter_{position}"
    )
    if filter_year!="All":
        visible=[r for r in visible if r.get("sorority_year")==filter_year]

    for r in visible:
        lock="🔒 " if r.get("is_financial") else ""
        label=(
            f"{lock}{r.get('sorority_year')} • {r.get('office')} • "
            f"{r.get('document_type')} • {r.get('title')}"
        )
        with st.expander(label):
            st.write(f"**Document Date:** {r.get('document_date') or 'Not listed'}")
            st.write(f"**Archive Category:** {r.get('archive_category') or ''}")
            st.write(f"**Visibility:** {r.get('visibility') or ''}")
            if r.get("description"):
                st.write(r.get("description"))
            st.caption(
                f"Uploaded by {r.get('uploaded_by_name') or 'Unknown'} • "
                f"File: {r.get('file_name') or ''}"
            )
            c1,c2=st.columns(2)
            try:
                file_bytes=download_private_bytes(r.get("file_path"))
                c1.download_button(
                    "Download File",
                    file_bytes,
                    file_name=r.get("file_name") or "archived_document",
                    mime=r.get("mime_type") or "application/octet-stream",
                    key=f"previous_year_download_{r.get('id')}",
                    use_container_width=True
                )
            except Exception as ex:
                c1.error(str(ex))

            if can_manage_previous_year_record(r,mid,is_admin_flag):
                confirm=c2.checkbox(
                    "Confirm delete",
                    key=f"previous_year_delete_confirm_{r.get('id')}"
                )
                if c2.button(
                    "Delete Record",
                    key=f"previous_year_delete_{r.get('id')}",
                    disabled=not confirm,
                    use_container_width=True
                ):
                    try:
                        path=r.get("file_path")
                        table("previous_sorority_year_documents").delete().eq(
                            "id",r.get("id")
                        ).execute()
                        if path:
                            (None if test_preview_active() else sb().storage.from_(cfg()["bucket"]).remove([path]))
                        st.success("Record deleted.")
                        st.rerun()
                    except Exception as ex:
                        st.error(f"Record was not deleted: {ex}")



OFFICER_POSITIONS=[
    "President","Vice President","Recording Secretary","Financial Secretary",
    "Treasurer","Historian","Chaplain","Parliamentarian","Sergeant-at-Arms"
]

# v3.6.4 — terminology and visibility rules.
# Philo officer titles stay the same at local, regional, and national Philo levels.
# Sigma Gamma Rho officer titles/names are NOT member-facing data. Only Advisor/Admin
# code paths may request them through the protected accessors below.
PHILO_OFFICER_POSITIONS=list(OFFICER_POSITIONS)

_SIGMA_NER_OFFICER_TITLES=(
    "Regional Syntaktes",
    "Regional Undergraduate Chapter Coordinator",
    "Regional Undergraduate Student Coordinator",
    "Regional Grammateus",
    "Regional Anti-Grammateus",
    "Regional Tamiochus",
    "Regional Epistoleus",
    "Regional Nominations Committee Chair",
    "Regional Elections Committee Chair",
    "Regional Legal Advisor",
    "Regional Parliamentarian",
    "Regional Philo Coordinator",
    "Regional Rhoer Coordinator",
    "Regional Programs Coordinator",
    "Area Coordinator",
    "Regional Chaplain",
    "Regional Sergeant-at-Arms",
)

_SIGMA_INTERNATIONAL_OFFICER_TITLES=(
    "International Grand Basileus",
    "International First Grand Anti-Basileus",
    "International Second Grand Anti-Basileus",
    "International Grand Grammateus",
    "International Grand Anti-Grammateus",
    "International Editor-in-Chief of The AURORA",
    "International Grand Tamiochus",
    "International Grand Epistoleus",
    "International Legal Advisor",
    "International Parliamentarian",
    "International Philo Coordinator",
    "International Rhoer Coordinator",
    "International Programs Coordinator",
)

NER_STANDING_COMMITTEES=[
    "Achievements and Awards Committee",
    "Budget and Finance Committee",
    "Bylaws, Policy and Procedures Committee",
    "Elections Committee",
    "History and Archival Committee",
    "Legislative Action Committee",
    "Membership Committee",
    "Memorial Committee",
    "Music Committee",
    "Nominations Committee",
    "Programs Committee",
    "Protocol Committee",
    "Public Relations Committee",
    "Recommendations Committee",
    "Scholarship Committee",
    "Social Action Committee",
    "Time and Place Committee",
    "Undergraduate Committee",
    "Youth Projects Committee",
]

INTERNATIONAL_STANDING_COMMITTEES=[
    "Achievement and Awards",
    "Board",
    "Boule",
    "Constitution and Bylaws",
    "Corporate Initiatives",
    "Cultured Pearls",
    "Education",
    "Election",
    "History and Archive Committee",
    "Honorary Member",
    "Legislative Action",
    "Membership",
    "Music and Arts",
    "International Programs",
    "Nominations",
    "Planning and Budget",
    "Protocol",
    "Public Relations",
    "Recommendations",
    "Scholarship",
    "Sigma Activities",
    "Social Action",
    "Time and Place",
]

GOVERNANCE_SOURCE_OPTIONS=[
    "Local Philo Bylaws",
    "Local Philo SOP / Procedures",
    "Northeastern Region Bylaws",
    "Northeastern Region Standard Operating Procedures",
    "Sigma Gamma Rho National Constitution and Bylaws",
    "Sigma Gamma Rho National Standard Operating Procedures",
    "National Philo Handbook",
    "Robert's Rules of Order Newly Revised",
]

ALIGNMENT_REVIEW_OPTIONS=[
    "Northeastern Region Bylaws",
    "Northeastern Region Standard Operating Procedures",
    "Sigma Gamma Rho National Constitution and Bylaws",
    "Sigma Gamma Rho National Standard Operating Procedures",
    "National Philo Handbook",
]
POSITION_EMAILS={
    "President":"philopresident.nbs@gmail.com",
    "Vice President":"philovicepresident.nbs@gmail.com",
    "Recording Secretary":"philorecordingsecretary.nbs@gmail.com",
    "Financial Secretary":"philofinancialsecretary.nbs@gmail.com",
    "Treasurer":"philotreasurer.nbs@gmail.com",
    "Historian":"philohistorian.nbs@gmail.com",
    "Chaplain":"philochaplain.nbs@gmail.com",
    "Parliamentarian":"philoparliamentarian.nbs@gmail.com",
    "Sergeant-at-Arms":"philosergeantatarms.nbs@gmail.com",
}
POSITION_QUESTIONS={
 "President":["What major business did you oversee this month?","What decisions or follow-ups are still outstanding?","Which committees or officers need support?","What should the membership know for next month?"],
 "Vice President":["Which committees or projects did you support?","What progress was made?","What remains outstanding?","What support is needed next month?"],
 "Recording Secretary":["Which meetings/minutes were completed?","What correspondence or records were handled?","What items require follow-up?","What is due next month?"],
 "Financial Secretary":["What funds or payments were received?","What records/reconciliations were completed?","What items are outstanding?","What financial follow-up is needed next month?"],
 "Treasurer":["What was the beginning balance?","What income was received?","What expenses were paid?","What is the ending balance and are any reimbursements outstanding?"],
 "Historian":["What chapter activities were documented?","What photos, programs, or records were collected?","What historical items still need to be gathered?","What will you document next month?"],
 "Chaplain":["What spiritual/wellness support was provided?","Were there member concerns requiring follow-up?","What observances or acknowledgements are upcoming?","What assistance is needed?"],
 "Parliamentarian":["What parliamentary questions arose?","Were any bylaws/SOP items identified for review?","Were motions or procedures needing clarification documented?","What governance work is planned next month?"],
 "Sergeant-at-Arms":["What meeting/order responsibilities were completed?","Were there attendance, access, or room concerns?","What materials or logistics require follow-up?","What is needed for the next meeting?"],
}

def safe_rows(name, **eq):
    try:return rows(name,**eq)
    except:return []

def member_offices(mid):
    try:
        return table('officer_assignments').select('*').eq('member_id',mid).eq('active',True).execute().data or []
    except:return []

def vacant_offices():
    assigned={x.get('position') for x in safe_rows('officer_assignments',active=True)}
    return [p for p in OFFICER_POSITIONS if p not in assigned]

def is_president(mid):
    return any(x.get('position')=='President' for x in member_offices(mid))

def profile_photo_url(mid):
    try:
        r=table('member_profiles').select('photo_path').eq('member_id',mid).execute().data or []
        return signed_url(r[0].get('photo_path',''),3600) if r and r[0].get('photo_path') else ''
    except:return ''

def report_pdf(title, subtitle, sections, signature=""):
    bio=BytesIO()
    doc=SimpleDocTemplate(bio,pagesize=letter,rightMargin=54,leftMargin=54,topMargin=52,bottomMargin=50)
    styles=getSampleStyleSheet()
    gold=colors.HexColor("#C99A20")
    title_style=ParagraphStyle('NBS_Title',parent=styles['Title'],fontName='Times-Bold',fontSize=20,leading=24,textColor=gold,spaceAfter=8)
    sub_style=ParagraphStyle('NBS_Sub',parent=styles['Normal'],fontName='Times-Roman',fontSize=10,alignment=1,spaceAfter=18)
    h=ParagraphStyle('NBS_H',parent=styles['Heading2'],fontName='Times-Bold',fontSize=12,textColor=colors.black,spaceBefore=10,spaceAfter=5)
    body=ParagraphStyle('NBS_Body',parent=styles['BodyText'],fontName='Times-Roman',fontSize=11,leading=16)
    story=[Paragraph("NU BETA SIGMA ALUMNAE CHAPTER",title_style),Paragraph("PHILO AFFILIATE",sub_style),
           Table([[""]],colWidths=[7*inch],rowHeights=[1],style=TableStyle([('BACKGROUND',(0,0),(-1,-1),gold)])),
           Spacer(1,12),Paragraph(title,title_style),Paragraph(subtitle,sub_style)]
    for head,bodytxt in sections:
        safe_head=html.escape(str(head or ''))
        safe_body=html.escape(str(bodytxt or "No information reported.")).replace("\n","<br/>")
        story += [Paragraph(safe_head,h),Paragraph(safe_body,body)]
    if signature:
        story += [Spacer(1,20),Paragraph(f"<b>Submitted/Signed:</b> {html.escape(str(signature))}",body)]
    story += [Spacer(1,24),Paragraph("Nu Beta Sigma Philo Affiliate",sub_style)]
    doc.build(story); bio.seek(0); return bio.getvalue()

def governance_pdf(doc_title, body_text, effective_date="", adopted_at=""):
    sections=[("Current Adopted Text",body_text)]
    if effective_date: sections.append(("Effective Date",effective_date))
    if adopted_at: sections.append(("Adopted At",adopted_at))
    return report_pdf(doc_title,"Official Member Print Version",sections)

def position_email(position):
    try:
        custom=setting("position_email_"+position.lower().replace(" ","_").replace("-","_"),"")
        return custom or POSITION_EMAILS.get(position,"")
    except:return POSITION_EMAILS.get(position,"")

def set_page(p):
    st.session_state['page']=p
    st.rerun()

def render_top_nav(is_admin, offices):
    st.markdown("<div class='nav-note'>Choose where you want to go:</div>",unsafe_allow_html=True)
    nav=[("🏠 Home","🏠 Dashboard"),("📅 Calendar","📅 Calendar & Events"),
         ("✅ Events & Tasks","✅ Events & Tasks"),("🧾 Forms","🧾 Forms"),
         ("📚 Governance","📚 Governance"),("👥 Members","👥 Members"),
         ("🌱 Intake","🌱 Interest & Intake"),("📸 Media & History","📸 Historian & Communications")]
    if chat_enabled():
        nav.append(("💬 Messages","💬 Messages"))
    if offices or is_admin:
        nav.append(("🏅 Officer","🏅 Officer Dashboard"))
    if is_admin or is_president(member_id) or chaired_committees(member_id):
        nav.append(("📋 Committee","📋 Committee Dashboard"))
    if is_admin:
        nav.append(("⚙️ Admin","⚙️ Admin Center"))
    cols=st.columns(5)
    for i,(label,p) in enumerate(nav):
        if cols[i%5].button(label,key=f"nav_{i}_{p}",use_container_width=True):
            set_page(p)
    st.divider()


def finance_rows():
    try:
        return table('finance_transactions').select('*').order('transaction_date',desc=True).execute().data or []
    except Exception as ex:
        raise RuntimeError(f'Financial ledger is unavailable; totals were not calculated: {ex}')

def month_finance_summary(ym):
    rows_=finance_rows()
    month_rows=[r for r in rows_ if str(r.get('transaction_date',''))[:7]==ym]
    income=sum(float(r.get('amount') or 0) for r in month_rows if r.get('direction')=='Income')
    expenses=sum(float(r.get('amount') or 0) for r in month_rows if r.get('direction')=='Expense')
    return month_rows,income,expenses,income-expenses

def member_dues_status(year_label):
    try:
        payments=table('dues_payments').select('*').eq('fiscal_year',year_label).execute().data or []
    except:payments=[]
    members_=philo_members()
    out=[]
    for m in members_:
        p=[x for x in payments if x.get('member_id')==m['id']]
        paid=sum(float(x.get('amount') or 0) for x in p)
        required=float(setting('annual_local_dues','0') or 0)
        status='Paid' if required and paid>=required else ('Partial' if paid>0 else 'Not Paid')
        out.append({'Member':m['full_name'],'Paid':paid,'Required':required,'Balance':max(0,required-paid),'Status':status})
    return out

def editable_financial_report_text(ym,position):
    month_rows,income,expenses,net=month_finance_summary(ym)
    income_lines=[f"- {r.get('transaction_date')}: {r.get('category') or 'Income'} — ${float(r.get('amount') or 0):,.2f} ({r.get('payer_payee') or ''})" for r in month_rows if r.get('direction')=='Income']
    expense_lines=[f"- {r.get('transaction_date')}: {r.get('category') or 'Expense'} — ${float(r.get('amount') or 0):,.2f} ({r.get('payer_payee') or ''})" for r in month_rows if r.get('direction')=='Expense']
    return f"""NBS PHILO FINANCIAL REPORT
Reporting Period: {ym}
Prepared for: {position}

INCOMING FUNDS
Total Income: ${income:,.2f}
{chr(10).join(income_lines) if income_lines else '- No incoming funds recorded.'}

EXPENSES / DISBURSEMENTS
Total Expenses: ${expenses:,.2f}
{chr(10).join(expense_lines) if expense_lines else '- No expenses recorded.'}

NET ACTIVITY
Income less expenses: ${net:,.2f}

VOUCHERS / REIMBURSEMENTS
Review the pending and paid voucher section for detailed disbursement activity.

NOTES / EXPLANATION
"""

def money_value(value):
    try:return Decimal(str(value or 0)).quantize(Decimal('0.01'),rounding=ROUND_HALF_UP)
    except (InvalidOperation,ValueError,TypeError):raise ValueError(f'Invalid currency amount: {value!r}')


def assessment_total(count_assessment,count_fundraiser,count_gala,count_late,count_delegate):
    assessment=Decimal(str(count_assessment))*Decimal('125.00')
    fundraiser=Decimal(str(count_fundraiser))*Decimal('5.00')
    gala=Decimal(str(count_gala))*Decimal('150.00')
    late=Decimal(str(count_late))*Decimal('10.00')
    delegate=Decimal(str(count_delegate))*Decimal('50.00')
    total=(assessment+fundraiser+gala+late+delegate).quantize(Decimal('0.01'),rounding=ROUND_HALF_UP)
    paypal=(total*Decimal('1.0299')+Decimal('1.00')).quantize(Decimal('0.01'),rounding=ROUND_HALF_UP) if total>0 else Decimal('0.00')
    return tuple(float(x) for x in (assessment,fundraiser,gala,late,delegate,total,paypal))

def ner_assessment_pdf(data, roster):
    sections=[
      ("Affiliate Information",f"""Date Prepared: {data['date_prepared']}
Affiliate Name: {data['affiliate_name']}
Affiliate Address: {data['affiliate_address']}
City, State, Zip: {data['city_state_zip']}

President: {data['president_name']} | {data['president_phone']} | {data['president_email']}
Financial Secretary: {data['fs_name']} | {data['fs_phone']} | {data['fs_email']}"""),
      ("Chapter Assessments",f"""Assessment Fee: ${data['assessment_total']:,.2f}
Fundraiser Fee: ${data['fundraiser_total']:,.2f}
Pageant/Gala Fee: ${data['gala_total']:,.2f}
Late Fee: ${data['late_total']:,.2f}
Delegate Fine: ${data['delegate_total']:,.2f}
TOTAL FEES: ${data['total_fees']:,.2f}
PayPal Amount Submitted (2.99% + $1): ${data['paypal_total']:,.2f}
E-check Amount Submitted (no processing fee): ${data['total_fees']:,.2f}"""),
      ("Financial Secretary Use Only",f"""Date Received: {data.get('date_received','')}
Amount Received: ${float(data.get('amount_received') or 0):,.2f}
Balance Due: ${float(data.get('balance_due') or 0):,.2f}
Credit Due: ${float(data.get('credit_due') or 0):,.2f}"""),
      ("Affiliate Roster","\n".join([f"{r.get('name','')} | {r.get('address','')} | {r.get('city_state_zip','')} | {r.get('phone','')} | {r.get('title_email','')}" for r in roster]) or "No roster entries.")
    ]
    return report_pdf("SY NER Philo Assessment Form","Remittance of Funds",sections,data.get('fs_name',''))


def position_finance_summary(position):
    rows_=finance_rows()
    income=sum(float(r.get('amount') or 0) for r in rows_ if r.get('direction')=='Income' and (r.get('position') or '')==position)
    spent=sum(float(r.get('amount') or 0) for r in rows_ if r.get('direction')=='Expense' and (r.get('position') or '')==position)
    vouchers=table('reimbursements').select('*').execute().data or []
    committed=sum(float(r.get('amount') or 0) for r in vouchers if (r.get('budget_position') or '')==position and r.get('status')=='Approved')
    budget=safe_rows('position_budgets',position=position)
    budget_amount=float(budget[0].get('budget_amount') or 0) if budget else 0.0
    remaining=budget_amount+income-spent-committed
    return {'budget':budget_amount,'income':income,'spent':spent,'committed':committed,'remaining':remaining}

def ensure_paid_voucher_transaction(voucher):
    existing=table('finance_transactions').select('id').eq('source_type','Voucher').eq('source_id',voucher['id']).execute().data or []
    if existing:return existing[0].get('id')
    inserted=table('finance_transactions').insert({
        'transaction_date':date.today().isoformat(),'fiscal_year':finance_fy(),'direction':'Expense',
        'category':'Reimbursement / Voucher','position':voucher.get('budget_position') or 'General Chapter','budget_line_id':voucher.get('budget_line_id'),
        'payer_payee':voucher.get('payee_name') or voucher.get('submitted_by_name') or '',
        'amount':float(voucher.get('amount_approved') if voucher.get('amount_approved') is not None else (voucher.get('amount') or 0)),
        'payment_method':voucher.get('payment_method') or '',
        'reference_number':'','notes':voucher.get('description') or '',
        'source_type':'Voucher','source_id':voucher['id'],
        'entered_by_member_id':member_id,'entered_by_name':member_name
    }).execute()
    row=(inserted.data or [{}])[0]
    if not row.get('id'): raise RuntimeError('Voucher ledger transaction was not created.')
    return row.get('id')

def csv_safe_cell(value):
    s=str(value or '')
    # Spreadsheet programs may execute cells beginning with formula/control prefixes.
    if s.startswith(('=','+','-','@','\t','\r','\n')):
        return "'"+s
    return s

def quickbooks_csv_bytes(rows_):
    out=io.StringIO()
    w=csv.writer(out)
    w.writerow(['Date','Type','Account/Category','Class/Position','Name','Memo','Amount','Payment Method','Reference No.'])
    for r in sorted(rows_,key=lambda x:str(x.get('transaction_date',''))):
        amount=float(r.get('amount') or 0)
        signed=amount if r.get('direction')=='Income' else -amount
        w.writerow([
            r.get('transaction_date',''),
            'Deposit' if r.get('direction')=='Income' else 'Expense',
            csv_safe_cell(r.get('category','')),
            csv_safe_cell(r.get('position','')),
            csv_safe_cell(r.get('payer_payee','')),
            csv_safe_cell(r.get('notes','')),
            f"{signed:.2f}",
            csv_safe_cell(r.get('payment_method','')),
            csv_safe_cell(r.get('reference_number',''))
        ])
    return out.getvalue().encode('utf-8')

def eoy_member_summary(mid,name,year_label,position=None):
    tasks=[x for x in safe_rows('event_tasks') if x.get('assigned_member_id')==mid]
    complete=sum(1 for x in tasks if x.get('status')=='Complete')
    sessions=[x for x in safe_rows('service_sessions',member_id=mid) if x.get('check_out')]
    hours=sum(float(x.get('hours') or 0) for x in sessions)
    reports=[x for x in safe_rows('monthly_reports') if x.get('member_id')==mid]
    recs=[x for x in safe_rows('recommendations') if x.get('member_id')==mid]
    body=f"""Reporting Year: {year_label}

MEMBER / OFFICER
{name}
{position or 'Philo Member'}

SERVICE
Community Service Hours Recorded: {hours:.2f}

EVENT RESPONSIBILITIES
Tasks Assigned: {len(tasks)}
Tasks Completed: {complete}
Tasks Outstanding: {max(0,len(tasks)-complete)}

MONTHLY REPORTING
Reports Generated/Submitted: {len(reports)}

RECOMMENDATIONS
Recommendations Submitted: {len(recs)}

YEAR-END REFLECTION
Add accomplishments, challenges, lessons learned, and recommendations for the next year below.
"""
    return body

def advisor_year_summary(year_label):
    events=table('events').select('*').execute().data or []
    tasks=safe_rows('event_tasks')
    service=safe_rows('service_sessions')
    reports=safe_rows('monthly_reports')
    recs=table('recommendations').select('*').execute().data or []
    finance=finance_rows()
    income=sum(float(r.get('amount') or 0) for r in finance if r.get('direction')=='Income')
    expenses=sum(float(r.get('amount') or 0) for r in finance if r.get('direction')=='Expense')
    hours=sum(float(r.get('hours') or 0) for r in service if r.get('check_out'))
    applied=sum(1 for r in recs if r.get('status')=='Applied to Governing Document')
    return f"""NBS PHILO AFFILIATE — YEAR IN REVIEW
Reporting Year: {year_label}

MEMBERSHIP & LEADERSHIP
Active Philos: {len(philo_members())}
Executive Board Positions Filled: {len(safe_rows('officer_assignments',active=True))}
Vacant Executive Board Positions: {len(vacant_offices())}

PROGRAMS & EVENTS
Events Recorded: {len(events)}
Event Tasks Assigned: {len(tasks)}
Event Tasks Completed: {sum(1 for x in tasks if x.get('status')=='Complete')}

COMMUNITY SERVICE
Total Service Hours Recorded: {hours:.2f}

FINANCIAL ACTIVITY
Total Income Recorded: ${income:,.2f}
Total Expenses Recorded: ${expenses:,.2f}
Net Recorded Activity: ${income-expenses:,.2f}

REPORTING
Monthly Reports Recorded: {len(reports)}

GOVERNANCE
Recommendations Submitted: {len(recs)}
Recommendations Applied to Governing Documents: {applied}

ADVISOR NARRATIVE
Add highlights, challenges, accomplishments, membership notes, program impact, and goals for the next Sorority Year below.
"""

def budget_lines(active_only=True):
    try:
        q=table('budget_lines').select('*')
        if active_only:q=q.eq('active',True)
        return q.order('name').execute().data or []
    except:return []

def budget_line_name(line_id):
    try:
        r=table('budget_lines').select('name').eq('id',line_id).single().execute().data
        return r.get('name','') if r else ''
    except:return ''

def budget_line_summary(line):
    line_id=line.get('id')
    rows_=finance_rows()
    income=sum(float(r.get('amount') or 0) for r in rows_ if r.get('budget_line_id')==line_id and r.get('direction')=='Income')
    spent=sum(float(r.get('amount') or 0) for r in rows_ if r.get('budget_line_id')==line_id and r.get('direction')=='Expense')
    vouchers=table('reimbursements').select('*').execute().data or []
    committed=sum(float(v.get('amount') or 0) for v in vouchers if v.get('budget_line_id')==line_id and v.get('status')=='Approved')
    budget=float(line.get('budget_amount') or 0)
    available=budget+income-spent-committed
    return {'budget':budget,'income':income,'spent':spent,'committed':committed,'available':available}

def all_budget_balances():
    out=[]
    for line in budget_lines():
        s=budget_line_summary(line)
        out.append({
            'Budget Line':line.get('name'),
            'Assigned Position':line.get('owner_position') or 'General Chapter',
            'Budget':s['budget'],'Income':s['income'],'Spent':s['spent'],
            'Committed':s['committed'],'Available':s['available']
        })
    return out

def position_budget_totals(position):
    lines=[x for x in budget_lines() if (x.get('owner_position') or 'General Chapter')==position]
    sums={'budget':0,'income':0,'spent':0,'committed':0,'available':0}
    for line in lines:
        s=budget_line_summary(line)
        for k in sums:sums[k]+=s[k]
    return sums,lines


def chat_enabled():
    return setting('chat_enabled','false').strip().lower()=='true'

def set_chat_enabled(enabled):
    save_setting('chat_enabled','true' if enabled else 'false')


def display_member_name(member_row, profile=None):
    profile=profile or {}
    prefix=(profile.get('prefix') or '').strip()
    first=(profile.get('first_name') or '').strip()
    last=(profile.get('last_name') or '').strip()
    if first or last:
        base=" ".join([x for x in [first,last] if x])
    else:
        base=member_row.get('full_name','')
    return f"{prefix} {base}".strip()

PROFILE_HOBBY_OPTIONS=[
    "Reading","Watching movies / TV","Puzzles","Gaming","Swimming",
    "Trying new restaurants / foods","Traveling","Cooking / Baking",
    "Photography","Crafting / DIY","Music","Dancing","Fitness / Working Out",
    "Gardening","Volunteering","Spending Time with Family","Writing",
    "Fashion / Styling","Other"
]

PROFILE_SERVICE_REASON_OPTIONS=[
    "Giving back to my community",
    "Supporting youth and the next generation",
    "Helping families and neighbors",
    "Building a stronger, more connected community",
    "Continuing a personal or family tradition of service",
    "Honoring people who helped me along the way",
    "Supporting causes that matter to me",
    "Other"
]

PROFILE_PHILO_VALUE_OPTIONS=[
    "Sisterhood and lifelong friendships",
    "Friendship and lasting connections",
    "Supporting Sigma Gamma Rho programs and service",
    "Opportunities to serve the community",
    "Personal and professional growth",
    "Building confidence and leadership skills",
    "Mentorship and encouragement",
    "Being part of a supportive Philo Affiliate",
    "Other"
]

PROFILE_STRENGTH_OPTIONS=[
    "Organization and Planning","Public Speaking","Technology / Google Workspace",
    "Canva / Graphic Design","Photography","Social Media Management",
    "Writing / Communication","Budgeting / Finance","Event Planning",
    "Mentorship / Coaching","Problem Solving","Adaptability","Teamwork",
    "Hospitality","Other"
]

PROFILE_FUN_FACT_OPTIONS=[
    "A hidden talent",
    "A place I have traveled",
    "A favorite hobby or fandom",
    "A personal goal I am working toward",
    "Something people are surprised to learn about me",
    "A memorable or funny story",
    "Other"
]

PROFILE_FUTURE_GOAL_OPTIONS=[
    "Help grow the Philo Affiliate's reach",
    "Support a new community-service project",
    "Mentor and encourage youth affiliates",
    "Take on a Philo leadership position",
    "Strengthen community partnerships",
    "Develop a new skill",
    "Support more affiliate programs and events",
    "Other"
]

def _profile_join(items):
    vals=[str(x).strip() for x in (items or []) if str(x).strip() and str(x).strip()!="Other"]
    if not vals:return ""
    if len(vals)==1:return vals[0]
    if len(vals)==2:return f"{vals[0]} and {vals[1]}"
    return ", ".join(vals[:-1])+f", and {vals[-1]}"

def _profile_detail(text):
    s=str(text or "").strip()
    if not s:return ""
    return s if s.endswith((".","!","?")) else s+"."

def generate_profile_bio_from_dropdowns(
    hobbies=None,hobbies_detail="",
    service_reasons=None,service_detail="",
    philo_values=None,philo_detail="",
    strengths=None,strengths_detail="",
    fun_fact_category="",fun_fact_detail="",
    future_goals=None,future_detail=""
):
    """Build an About Me draft in the Philo's own first-person voice."""
    parts=[]

    joined=_profile_join(hobbies)
    if joined:
        parts.append(f"In my free time, I enjoy {joined.lower()}.")
    if _profile_detail(hobbies_detail):
        parts.append(_profile_detail(hobbies_detail))

    joined=_profile_join(service_reasons)
    if joined:
        parts.append(f"Community service matters to me because I value {joined.lower()}.")
    if _profile_detail(service_detail):
        parts.append(_profile_detail(service_detail))

    joined=_profile_join(philo_values)
    if joined:
        parts.append(f"Being a Philo is meaningful to me because of {joined.lower()}.")
    if _profile_detail(philo_detail):
        parts.append(_profile_detail(philo_detail))

    joined=_profile_join(strengths)
    if joined:
        parts.append(f"The strengths I bring include {joined.lower()}.")
    if _profile_detail(strengths_detail):
        parts.append(_profile_detail(strengths_detail))

    if fun_fact_category and fun_fact_category!="Other" and not str(fun_fact_detail or "").strip():
        parts.append(f"A fun fact about me relates to {fun_fact_category.lower()}.")
    if _profile_detail(fun_fact_detail):
        parts.append(f"A fun fact about me is {_profile_detail(fun_fact_detail)[0].lower()+_profile_detail(fun_fact_detail)[1:]}")

    joined=_profile_join(future_goals)
    if joined:
        parts.append(f"Looking ahead, I hope to {joined[0].lower()+joined[1:]}.")
    if _profile_detail(future_detail):
        parts.append(_profile_detail(future_detail))

    return " ".join(parts).strip()


def _bio_fragment(value, leadins=()):
    """Clean a questionnaire answer so a generated sentence does not duplicate its lead-in."""
    s=str(value or '').strip().rstrip('.')
    for lead in leadins:
        if s.lower().startswith(lead.lower()):
            s=s[len(lead):].lstrip(" :-,")
            break
    return s

def bio_from_questionnaire(q):
    intro=[]
    interests=_bio_fragment(q.get('interests'),('I enjoy','She enjoys'))
    community=_bio_fragment(q.get('community'),('Community service is meaningful to me because','Community service is meaningful to her because'))
    philo=_bio_fragment(q.get('philo'),('As a Philo, I value','As a Philo, she values','I value','She values'))
    skills=_bio_fragment(q.get('skills'),('I bring strengths in','She brings strengths in','My strengths include'))
    fun=_bio_fragment(q.get('fun'),('A fun fact about me is','A fun fact about her is'))
    goal=_bio_fragment(q.get('goal'),('I hope to','She hopes to'))

    if interests: intro.append(f"She enjoys {interests}.")
    if community: intro.append(f"Community service is meaningful to her because {community}.")
    if philo: intro.append(f"As a Philo, she values {philo}.")
    if skills: intro.append(f"She brings strengths in {skills}.")
    if fun: intro.append(f"A fun fact about her is {fun}.")
    if goal: intro.append(f"She hopes to {goal}.")
    return " ".join(intro)

def birthday_today(profile):
    b=profile.get('birthday')
    if not b:return False
    try:
        d=pd.to_datetime(b).date()
        today=date.today()
        return d.month==today.month and d.day==today.day
    except:return False



def spreadsheet_id_from_url(value):
    raw=str(value or '').strip()
    if not raw:return ''
    m=re.search(r'/spreadsheets/d/([A-Za-z0-9_-]+)',raw)
    if m:return m.group(1)
    return raw if re.fullmatch(r'[A-Za-z0-9_-]{20,}',raw) else ''

def officer_contact(position):
    assignment=next((x for x in safe_rows('officer_assignments',active=True) if x.get('position')==position),None)
    if not assignment:
        return {'name':'','phone':'','email':position_email(position)}
    mid=assignment.get('member_id')
    member_row=next((m for m in active_members() if m.get('id')==mid),{})
    profiles=safe_rows('member_profiles',member_id=mid) if mid else []
    p=profiles[0] if profiles else {}
    return {
        'name':assignment.get('member_name') or member_row.get('full_name') or '',
        'phone':member_row.get('phone') or p.get('phone') or '',
        'email':position_email(position) or member_row.get('email') or ''
    }

def assessment_roster_from_hub():
    rows=[]
    for m in philo_members():
        ps=safe_rows('member_profiles',member_id=m['id'])
        p=ps[0] if ps else {}
        offices_=member_offices(m['id'])
        title=', '.join(x.get('position') for x in offices_ if x.get('position')) or 'Philo'
        pos_email=''
        if offices_:
            pos_email=position_email(offices_[0].get('position'))
        email=pos_email or m.get('email') or ''
        address=p.get('address') or ''
        city=p.get('city_state_zip') or ''
        rows.append({
            'name':m.get('full_name') or '',
            'address':address,
            'city_state_zip':city,
            'phone':m.get('phone') or '',
            'title_email':f"{title} / {email}" if email else title
        })
    return rows

def assessment_autofill_from_hub():
    president=officer_contact('President')
    fs=officer_contact('Financial Secretary')
    treasurer=officer_contact('Treasurer')
    roster=assessment_roster_from_hub()
    return {
        'date_prepared':date.today().isoformat(),
        'affiliate_name':setting('affiliate_name','Nu Beta Sigma Philo Affiliate'),
        'affiliate_address':setting('affiliate_address','P.O. Box 280651'),
        'city_state_zip':setting('affiliate_city_state_zip','Queens Village, NY 11428'),
        'president_name':president['name'],
        'president_phone':president['phone'],
        'president_email':president['email'],
        'fs_name':fs['name'],
        'fs_phone':fs['phone'],
        'fs_email':fs['email'],
        'treasurer_name':treasurer['name'],
        'treasurer_email':treasurer['email'],
        'advisor_name':setting('philo_advisor_name','Chennel Chapman'),
        'advisor_contact':setting('philo_advisor_contact',''),
        'assessment_count':1,
        'fundraiser_count':len(roster),
        'gala_count':0,
        'late_count':0,
        'delegate_count':0,
        'roster':roster
    }

def remittance_sheet_values(form_row):
    fd=form_row.get('form_data') or {}
    roster=(fd.get('roster') or assessment_roster_from_hub())[:21]
    values={
        # Header / affiliate info. Coordinates match the supplied NER template.
        'D2':fd.get('date_prepared',''),
        'D3':fd.get('affiliate_name',''),
        'D4':fd.get('affiliate_address',''),
        'D5':fd.get('city_state_zip',''),
        'F3':fd.get('president_name',''),
        'H3':fd.get('president_phone',''),
        'J3':fd.get('president_email',''),
        'F5':fd.get('fs_name',''),
        'H5':fd.get('fs_phone',''),
        'J5':fd.get('fs_email',''),
        # Assessment quantities and totals.
        'D9':fd.get('assessment_count',0),
        'D10':fd.get('fundraiser_count',0),
        'D11':fd.get('gala_count',0),
        'D12':fd.get('late_count',0),
        'D13':fd.get('delegate_count',0),
        'H9':fd.get('assessment_total',0),
        'H10':fd.get('fundraiser_total',0),
        'H11':fd.get('gala_total',0),
        'H12':fd.get('late_total',0),
        'H13':fd.get('delegate_total',0),
        'H14':fd.get('total_fees',0),
        'H15':max(float(fd.get('paypal_total') or 0)-float(fd.get('total_fees') or 0),0),
        'H17':fd.get('paypal_total',0),
        'H19':fd.get('echeck_total',fd.get('total_fees',0)),
        # Treasurer / Advisor lines.
        'D21':fd.get('treasurer_name',''),
        'D22':fd.get('treasurer_email',''),
        'H21':fd.get('advisor_name',''),
        'H22':fd.get('advisor_contact',''),
        # Financial Secretary use only.
        'B28':fd.get('date_received',''),
        'D28':fd.get('amount_received',0),
        'H28':fd.get('balance_due',0),
        'J28':fd.get('credit_due',0),
        # Affiliate roster heading.
        'D32':fd.get('affiliate_name','')
    }
    start_row=36
    for idx,r in enumerate(roster):
        row=start_row+idx
        values[f'B{row}']=r.get('name','')
        values[f'D{row}']=r.get('address','')
        values[f'F{row}']=r.get('city_state_zip','')
        values[f'H{row}']=r.get('phone','')
        values[f'J{row}']=r.get('title_email','')
    return values

def create_completed_remittance_sheet(form_row):
    if st.session_state.get('demo_admin_mode'):
        raise RuntimeError('ADMIN DEMO MODE: Google/Drive actions are disabled in the sandbox.')
    if test_preview_active():
        raise RuntimeError('TEST MODE: external Google/Drive actions are disabled.')
    if not google_connected():
        raise RuntimeError('Connect the Advisor Google account first.')
    template_id=spreadsheet_id_from_url(setting('ner_remittance_template_url',''))
    if not template_id:
        raise RuntimeError('Add the NER Remittance Google Sheet Template URL in Admin Settings first.')
    drive=drive_service();sheets=sheets_service()
    if not drive or not sheets:
        raise RuntimeError('Google Drive/Sheets connection is unavailable.')
    fd=form_row.get('form_data') or {}
    prepared=str(fd.get('date_prepared') or date.today().isoformat())
    name=f"NBS NER Remittance - {prepared} - Form {form_row.get('id')}"
    body={'name':name}
    folder_id=setting('ner_remittance_completed_folder_id','').strip()
    if folder_id:
        body['parents']=[folder_id]
    copied=drive.files().copy(fileId=template_id,body=body,fields='id,name,webViewLink').execute()
    spreadsheet_id=copied['id']
    sheet_name=setting('ner_remittance_sheet_name','RemittanceSummary').strip() or 'RemittanceSummary'
    data=[]
    for cell,value in remittance_sheet_values(form_row).items():
        data.append({'range':f"'{sheet_name}'!{cell}",'values':[[value]]})
    sheets.spreadsheets().values().batchUpdate(
        spreadsheetId=spreadsheet_id,
        body={'valueInputOption':'USER_ENTERED','data':data}
    ).execute()
    url=copied.get('webViewLink') or f'https://docs.google.com/spreadsheets/d/{spreadsheet_id}/edit'
    return {'id':spreadsheet_id,'url':url,'name':copied.get('name') or name}

def assessment_forms_for_status(statuses=None):
    try:
        q=table('ner_assessment_forms').select('*').order('created_at',desc=True)
        rows_=q.execute().data or []
        if statuses:
            rows_=[r for r in rows_ if r.get('status') in statuses]
        return rows_
    except:
        return []

def assessment_form_label(row):
    fd=row.get('form_data') or {}
    return f"#{row.get('id')} • {fd.get('affiliate_name','Assessment')} • ${float(row.get('total_fees') or 0):,.2f} • {row.get('status','Draft')}"


TASK_TEMPLATES=[
    "Contact venue",
    "Confirm venue reservation",
    "Confirm event date/time",
    "Create flyer",
    "Send flyer for approval",
    "Post event announcement",
    "Confirm volunteers",
    "Create volunteer sign-up sheet",
    "Order food",
    "Pick up food",
    "Bring plates",
    "Bring cups",
    "Bring napkins",
    "Bring utensils",
    "Bring tablecloths",
    "Purchase supplies",
    "Prepare event materials",
    "Prepare sign-in sheet",
    "Manage check-in table",
    "Set up event space",
    "Decorate event space",
    "Clean up event space",
    "Take photos",
    "Collect photos from members",
    "Prepare social media recap",
    "Contact community partner",
    "Confirm guest speaker",
    "Confirm transportation",
    "Prepare program / agenda",
    "Print handouts",
    "Collect receipts",
    "Submit receipts",
    "Prepare reimbursement / voucher",
    "Complete event report",
    "Follow up with venue",
    "Follow up with participants",
    "Other / Write-in"
]

POSITION_TOOLKIT={
 "President":{
   "title":"President Toolkit",
   "items":["Review vacant offices and acting access","Review chapter account balances","Review paperwork requiring President approval","Check upcoming events and officer deadlines","Prepare meeting agenda / New Business items","Review recommendations and governance items"]
 },
 "Vice President":{
   "title":"Vice President Toolkit",
   "items":["Review membership and reclamation follow-up","Check committee/event progress","Review upcoming chapter activities","Track delegated assignments","Prepare monthly officer update"]
 },
 "Recording Secretary":{
   "title":"Recording Secretary Toolkit",
   "items":["Prepare meeting agenda materials","Track minutes and corrections","Maintain membership/meeting records","Track correspondence and action items","Prepare records for upcoming meeting"]
 },
 "Financial Secretary":{
   "title":"Financial Secretary Toolkit",
   "items":["Record incoming monies","Update dues payments","Prepare assessment/remittance forms","Review returned forms from Treasurer","Prepare monthly financial secretary report"]
 },
 "Treasurer":{
   "title":"Treasurer Toolkit",
   "items":["Review incoming funds posted by Financial Secretary","Manage budget lines","Review assessment/remittance forms","Review vouchers and expenses","Prepare payments and financial reports","Export transactions for QuickBooks if needed"]
 },
 "Historian":{
   "title":"Historian Toolkit",
   "items":["Review upcoming events that need documentation","Track photos and event records","Prepare publicity/history notes","Maintain program documentation","Prepare monthly historian report"]
 },
 "Chaplain":{
   "title":"Chaplain Toolkit",
   "items":["Review prayer requests","Prepare devotion/meditation","Track illness/bereavement follow-up","Prepare cards/acknowledgements","Prepare monthly chaplain update"]
 },
 "Parliamentarian":{
   "title":"Parliamentarian Toolkit",
   "items":["Review parliamentary questions","Review bylaws/SOP recommendations","Prepare governance items for meeting","Check Robert's Rules / local procedure questions","Track proposed amendments"]
 },
 "Sergeant-at-Arms":{
   "title":"Sergeant-at-Arms Toolkit",
   "items":["Prepare meeting access/check-in","Review attendance responsibilities","Prepare room/materials","Track meeting order/logistics","Prepare any required attendance or late-arrival notes"]
 }
}

def is_philo_member(mid):
    try:
        r=table('member_private').select('is_philo').eq('member_id',mid).execute().data or []
    except Exception:
        return False
    if r and r[0].get('is_philo') is not None:
        return bool(r[0].get('is_philo'))
    try:
        m=table('members').select('role').eq('id',mid).single().execute().data or {}
    except Exception:
        return False
    return str(m.get('role') or '').strip().lower()=='member'

def philo_members():
    return [m for m in active_members() if is_philo_member(m['id'])]


def advisor_only_sigma_officer_titles(level='all'):
    """Return Sigma Gamma Rho officer titles only inside a verified Advisor/Admin session."""
    if not bool(st.session_state.get('logged_in')) or not bool(st.session_state.get('is_admin')):
        return []
    level=str(level or 'all').strip().lower()
    if level in {'ner','regional','region'}:
        return list(_SIGMA_NER_OFFICER_TITLES)
    if level in {'international','national'}:
        return list(_SIGMA_INTERNATIONAL_OFFICER_TITLES)
    return list(_SIGMA_NER_OFFICER_TITLES)+list(_SIGMA_INTERNATIONAL_OFFICER_TITLES)


def advisor_only_sigma_people():
    """Non-Philo/Sigma-side names are enumerable only by Advisor/Admin."""
    if not bool(st.session_state.get('logged_in')) or not bool(st.session_state.get('is_admin')):
        return []
    return [m for m in active_members() if not is_philo_member(m['id'])]


def philo_dropdown_people():
    """Single source of truth for member-facing person dropdowns."""
    return philo_members()

def relevant_voucher_for_position(voucher, position):
    if position in ['President','Financial Secretary','Treasurer']:
        return True
    # Officer whose own budget line/position is being charged can see/sign it.
    return (voucher.get('budget_position') or '')==position

def pending_paperwork_for_position(position):
    vouchers=table('reimbursements').select('*').execute().data or []
    return [r for r in vouchers if r.get('status')=='Pending' and relevant_voucher_for_position(r,position)]

def render_position_toolkit(position):
    tk=POSITION_TOOLKIT.get(position)
    if not tk:return
    with st.container(border=True):
        st.markdown(f"### {tk['title']}")
        for item in tk['items']:
            st.write(f"• {item}")

def prayer_requests_for_chaplain():
    try:return table('prayer_requests').select('*').order('created_at',desc=True).execute().data or []
    except:return []


COMMITTEE_NAMES=[
    "Membership",
    "Budget and Finance",
    "Achievements and Awards",
    "Bylaws, Policy and Procedures",
    "Scholarship",
    "Public Relations",
    "Spotlight",
    "Nominations",
    "Elections",
    "Recommendations",
    "Programs",
    "Protocol",
    "Community Service",
    "Fundraising",
    "Other / Local Committee",
]

MOTION_RANKING=[
    {"rank":1,"motion":"Fix the Time to Which to Adjourn","second":"Yes","debatable":"No","amendable":"Yes","vote":"Majority"},
    {"rank":2,"motion":"Adjourn","second":"Yes","debatable":"No","amendable":"No","vote":"Majority"},
    {"rank":3,"motion":"Recess","second":"Yes","debatable":"No","amendable":"Yes","vote":"Majority"},
    {"rank":4,"motion":"Raise a Question of Privilege","second":"No","debatable":"No","amendable":"No","vote":"Chair rules"},
    {"rank":5,"motion":"Call for the Orders of the Day","second":"No","debatable":"No","amendable":"No","vote":"No vote unless set aside"},
    {"rank":6,"motion":"Lay on the Table","second":"Yes","debatable":"No","amendable":"No","vote":"Majority"},
    {"rank":7,"motion":"Previous Question (Close Debate)","second":"Yes","debatable":"No","amendable":"No","vote":"Two-thirds"},
    {"rank":8,"motion":"Limit or Extend Limits of Debate","second":"Yes","debatable":"No","amendable":"Yes","vote":"Two-thirds"},
    {"rank":9,"motion":"Postpone to a Certain Time","second":"Yes","debatable":"Yes","amendable":"Yes","vote":"Majority"},
    {"rank":10,"motion":"Commit or Refer","second":"Yes","debatable":"Yes","amendable":"Yes","vote":"Majority"},
    {"rank":11,"motion":"Amend","second":"Yes","debatable":"When main motion is debatable","amendable":"Yes","vote":"Majority"},
    {"rank":12,"motion":"Postpone Indefinitely","second":"Yes","debatable":"Yes","amendable":"No","vote":"Majority"},
    {"rank":13,"motion":"Main Motion","second":"Yes","debatable":"Yes","amendable":"Yes","vote":"Majority"}
]

AGENDA_STANDARD_ITEMS=[
    "Prayer",
    "Call to Order",
    "Roll Call / Attendance",
    "Adoption of Previous Meeting Minutes",
    "Adoption of Agenda",
    "Pledge",
    "President Report",
    "Vice President Report",
    "Recording Secretary Report",
    "Financial Secretary Report",
    "Treasurer Report",
    "Historian Report",
    "Chaplain Report",
    "Parliamentarian Report",
    "Sergeant-at-Arms Report",
    "Committee Reports",
    "Unfinished Business",
    "New Business",
    "Announcements",
    "Advisor Update",
    "Adjournment",
    "Hymn"
]

def local_asset_bytes(path):
    p=Path(__file__).resolve().parent/path
    try:return p.read_bytes()
    except:return b""

def show_pdf_asset(path,height=720):
    data=local_asset_bytes(path)
    if not data:
        st.error("Document file is not available in this deployment.")
        return
    b64=base64.b64encode(data).decode()
    html=f'<iframe src="data:application/pdf;base64,{b64}" width="100%" height="{height}" type="application/pdf"></iframe>'
    components.html(html,height=height+20,scrolling=True)

def fee_schedule_rows():
    try:return table('fee_schedule').select('*').eq('active',True).order('sort_order').execute().data or []
    except:return []

def seed_default_fee_schedule():
    try:
        existing=table('fee_schedule').select('id').limit(1).execute().data or []
        if existing:return
        defaults=[
            ("Individual","Returning","National Sigma - Returning",25.00,10),
            ("Individual","Inductee","National Sigma - Inductee",75.00,20),
            ("Individual","Any","National Sigma - Late Fee",10.00,30),
            ("Individual","Returning","National Philo - Returning",10.00,40),
            ("Individual","Inductee","National Philo - Inductee",15.00,50),
            ("Individual","Any","National Philo - Late Fee",10.00,60),
            ("Individual","Reactivating","Reactivation Fee",5.00,70),
            ("Individual","Any","Regional Philo - Fundraiser Fee",5.00,80),
            ("Individual","Returning","Local NBS Dues",122.74,90),
            ("Affiliate","Any","National Sigma - Assessment",25.00,110),
            ("Affiliate","Any","National Sigma - Late Fee",25.00,120),
            ("Affiliate","Any","National Philo - Assessment",25.00,130),
            ("Affiliate","Any","National Philo - Late Fee",25.00,140),
            ("Affiliate","Any","Regional Philo - Assessment",125.00,150),
            ("Affiliate","Any","First Lady / Gala Fee",150.00,160),
            ("Affiliate","Any","Regional Late Fee",10.00,170),
            ("Affiliate","Any","Regional Sigma",10.00,180),
            ("Affiliate","Any","Regional Sigma - Late Fee",10.00,190),
            ("Affiliate","Any","Service Fee",0.35,200),
        ]
        for scope,mtype,name,amount,order in defaults:
            table('fee_schedule').insert({
                'scope':scope,'member_type':mtype,'fee_name':name,'amount':amount,
                'sort_order':order,'active':True
            }).execute()
    except:
        pass

def member_type(mid):
    try:
        r=table('member_private').select('member_type').eq('member_id',mid).execute().data or []
        return (r[0].get('member_type') if r else '') or 'Returning'
    except:return 'Returning'

def member_financial_summary(mid, fiscal_year):
    seed_default_fee_schedule()
    mtype=member_type(mid)
    schedule=[r for r in fee_schedule_rows() if r.get('scope')=='Individual' and r.get('member_type') in [mtype,'Any']]
    required=sum(float(r.get('amount') or 0) for r in schedule)
    if mtype=='Returning':
        required=200.00
    try:
        pays=table('dues_payments').select('*').eq('member_id',mid).eq('fiscal_year',fiscal_year).execute().data or []
    except:pays=[]
    paid=sum(float(x.get('amount') or 0) for x in pays)
    balance=max(0,required-paid)
    return {'member_type':mtype,'required':required,'paid':paid,'balance':balance,'status':'Financial' if required>0 and balance<=0.005 else ('Partial' if paid>0 else 'Not Financial')}

def committee_chair_positions(mid):
    try:return table('committee_assignments').select('*').eq('chair_member_id',mid).eq('active',True).execute().data or []
    except:return []

def can_create_agenda(mid,is_admin_flag=False):
    return is_admin_flag or is_president(mid) or bool(committee_chair_positions(mid))

def agenda_pdf(agenda,items):
    title=agenda.get('title') or 'Meeting Agenda'
    sub=f"{agenda.get('meeting_date','')} • {agenda.get('meeting_time','')} • {agenda.get('location','')}"
    body=[]
    for idx,it in enumerate(items,1):
        roman=["I","II","III","IV","V","VI","VII","VIII","IX","X","XI","XII","XIII","XIV","XV","XVI","XVII","XVIII","XIX","XX","XXI","XXII","XXIII","XXIV","XXV"]
        marker=roman[idx-1] if idx<=len(roman) else str(idx)
        line=f"{marker}. {it.get('item_text','')}"
        if it.get('notes'): line+=f"\n    {it.get('notes')}"
        body.append(line)
    return report_pdf(title,sub,[("Order of Business","\n".join(body))],agenda.get('created_by_name',''))

def agenda_rows(published_only=False):
    try:
        q=table('meeting_agendas').select('*').order('meeting_date',desc=True)
        rows_=q.execute().data or []
        return [r for r in rows_ if (r.get('status')=='Published')] if published_only else rows_
    except:return []

def agenda_items(agenda_id):
    try:return table('agenda_items').select('*').eq('agenda_id',agenda_id).order('sort_order').execute().data or []
    except:return []

def meeting_agenda_creator_label(mid):
    if is_president(mid): return "President"
    chairs=committee_chair_positions(mid)
    if chairs:return chairs[0].get('committee_name')+" Committee Chair"
    return "Advisor/Admin"


def voucher_approvals(voucher_id):
    try:
        return table('document_approvals').select('*').eq('document_type','reimbursement').eq('document_id',voucher_id).order('signed_at').execute().data or []
    except:
        return []

def officer_holder(position):
    try:
        r=table('officer_assignments').select('*').eq('position',position).eq('active',True).execute().data or []
        return r[0] if r else None
    except:
        return None

def voucher_required_positions(voucher):
    submitter_id=voucher.get('submitted_by_member_id')
    submitter_positions=[x.get('position') for x in member_offices(submitter_id)] if submitter_id else []

    # Standard local routing: Financial Secretary -> President -> Treasurer.
    # Prevent a submitter from authorizing/certifying her own voucher.
    required=['Financial Secretary','President','Treasurer']
    if 'President' in submitter_positions:
        required=[p for p in required if p!='President']
    if 'Treasurer' in submitter_positions:
        required=[p for p in required if p!='Treasurer']
    if 'Financial Secretary' in submitter_positions:
        required=[p for p in required if p!='Financial Secretary']
    return required

def voucher_signature_status(voucher):
    approvals=voucher_approvals(voucher['id'])
    out=[]
    for pos in voucher_required_positions(voucher):
        holder=officer_holder(pos)
        actions=[a for a in approvals if a.get('position')==pos]
        action=actions[-1] if actions else None
        out.append({
            'Position':pos,
            'Officer':holder.get('member_name') if holder else 'VACANT',
            'Decision':action.get('decision') if action else 'Waiting',
            'Reason / Note':action.get('reason_note') if action else '',
            'Signed By':action.get('signature_name') if action else '',
            'Signed At':action.get('signed_at') if action else ''
        })
    return out

def voucher_can_delete(voucher):
    # Permanent delete only before any officer has acted.
    return len(voucher_approvals(voucher['id']))==0 and voucher.get('status') in ['Pending','Cancelled']

def voucher_can_cancel(voucher):
    # Once paid, cancellation is no longer appropriate.
    return voucher.get('status') not in ['Paid','Cancelled','Deleted']


def toolkit_records(position=None, member_id_filter=None, record_type=None):
    try:
        q=table('officer_toolkit_records').select('*')
        if position:q=q.eq('position',position)
        if member_id_filter:q=q.eq('member_id',member_id_filter)
        if record_type:q=q.eq('record_type',record_type)
        return q.order('created_at',desc=True).execute().data or []
    except:return []

def save_toolkit_record(position,record_type,title,data):
    table('officer_toolkit_records').insert({
        'position':position,
        'member_id':member_id,
        'member_name':member_name,
        'record_type':record_type,
        'title':title,
        'data':data,
        'created_at':datetime.now(timezone.utc).isoformat(),
        'updated_at':datetime.now(timezone.utc).isoformat()
    }).execute()

def position_event_activity(mid):
    tasks=[x for x in safe_rows('event_tasks') if x.get('assigned_member_id')==mid]
    complete=[x for x in tasks if x.get('status')=='Complete']
    inprog=[x for x in tasks if x.get('status')=='In Progress']
    return tasks,complete,inprog

def polished_position_report(position,mid,name,period,extra_notes=""):
    records=toolkit_records(position=position,member_id_filter=mid)
    tasks,complete,inprog=position_event_activity(mid)
    events=upcoming_events(200)

    sections=[]
    opening=f"During {period}, {name} served as {position} and carried out responsibilities associated with the office while supporting the programs, operations, and activities of the NBS Philo Affiliate."
    sections.append(("Overview",opening))

    if tasks:
        task_names=", ".join([x.get('task','') for x in complete[:6] if x.get('task')])
        txt=f"{name} was assigned {len(tasks)} event or operational task(s) during the reporting period. {len(complete)} task(s) were marked complete"
        if inprog: txt+=f", with {len(inprog)} still in progress"
        txt+="."
        if task_names: txt+=f" Completed responsibilities included {task_names}."
        sections.append(("Assignments and Event Support",txt))

    # Generic toolkit record summary
    if records:
        grouped={}
        for r in records:
            grouped.setdefault(r.get('record_type','Activity'),[]).append(r)
        for rtype,items in grouped.items():
            snippets=[]
            for item in items[:8]:
                data=item.get('data') or {}
                if isinstance(data,str):
                    try:data=json.loads(data)
                    except:data={'details':data}
                detail_bits=[]
                for k,v in data.items():
                    if v not in ['',None,False,[]]:
                        detail_bits.append(f"{k.replace('_',' ').title()}: {v}")
                snippets.append(f"{item.get('title')}: " + "; ".join(detail_bits[:4]))
            if snippets:
                sections.append((rtype.replace('_',' ').title(), " ".join(snippets)))

    # Position-specific auto summaries
    if position=='Recording Secretary':
        mins=[r for r in records if r.get('record_type')=='meeting_minutes']
        corr=[r for r in records if r.get('record_type')=='correspondence']
        act=[r for r in records if r.get('record_type')=='action_item']
        txt=f"Meeting and records management remained a key focus. {len(mins)} meeting record(s) were created or updated, {len(corr)} correspondence item(s) were logged, and {len(act)} follow-up action item(s) were tracked."
        sections.append(("Records, Minutes, and Correspondence",txt))
    elif position=='President':
        dec=[r for r in records if r.get('record_type')=='decision']
        chk=[r for r in records if r.get('record_type')=='officer_checkin']
        sections.append(("Leadership and Oversight",f"The President documented {len(dec)} leadership decision(s) and {len(chk)} officer/committee check-in(s) during the period, while maintaining oversight of chapter business, vacant offices, and upcoming priorities."))
    elif position=='Vice President':
        mem=[r for r in records if r.get('record_type')=='membership_followup']
        prog=[r for r in records if r.get('record_type')=='committee_progress']
        sections.append(("Membership and Committee Support",f"The Vice President tracked {len(mem)} membership follow-up item(s) and {len(prog)} committee/program progress update(s), supporting continuity and follow-through on chapter initiatives."))
    elif position=='Historian':
        hist=[r for r in records if r.get('record_type') in ['history_entry','media_log']]
        photos=safe_rows('historian_photos');communications=safe_rows('communications_requests');published=[r for r in communications if r.get('status')=='Published'];nominations=safe_rows('philo_month_nominations')
        sections.append(("History and Documentation",f"The Historian documented {len(hist)} history log item(s), maintained {len(photos)} photo submission(s), tracked {len(communications)} communications request(s), published {len(published)} communications item(s), and monitored {len(nominations)} Philo of the Month nomination(s)."))
    elif position=='Chaplain':
        prayers=prayer_requests_for_chaplain()
        courtesy=[r for r in records if r.get('record_type')=='courtesy']
        sections.append(("Chaplain and Member Care",f"The Chaplain monitored {len(prayers)} prayer request(s) in the Hub and documented {len(courtesy)} courtesy, condolence, recognition, or member-care follow-up item(s)."))
    elif position=='Parliamentarian':
        parl=[r for r in records if r.get('record_type')=='parliamentary_ruling']
        gov=[r for r in records if r.get('record_type')=='governance_review']
        sections.append(("Parliamentary and Governance Support",f"The Parliamentarian documented {len(parl)} parliamentary question/ruling item(s) and {len(gov)} bylaws, policy, procedure, or governance review item(s)."))
    elif position=='Sergeant-at-Arms':
        att=[r for r in records if r.get('record_type')=='meeting_order']
        sections.append(("Meeting Order and Logistics",f"The Sergeant-at-Arms documented {len(att)} meeting-order, attendance, access, or logistics item(s) during the reporting period."))
    elif position=='Financial Secretary':
        rows_=finance_rows()
        inc=[r for r in rows_ if r.get('direction')=='Income' and r.get('entered_by_member_id')==mid]
        sections.append(("Incoming Funds and Financial Records",f"The Financial Secretary recorded {len(inc)} incoming financial transaction(s) in the Hub and maintained dues, assessment, receipt, and remittance records for Treasurer review."))
    elif position=='Treasurer':
        rows_=finance_rows()
        expenses=[r for r in rows_ if r.get('direction')=='Expense']
        balances=all_budget_balances()
        sections.append(("Treasury and Budget Management",f"The Treasurer maintained the financial ledger, monitored {len(balances)} active budget line(s), and recorded or reviewed {len(expenses)} expense transaction(s) in support of accurate budget-to-actual reporting."))

    if extra_notes.strip():
        sections.append(("Additional Highlights",extra_notes.strip()))

    closing="The information above reflects activity recorded in the NBS Philo Hub and may be edited to add context, accomplishments, challenges, recommendations, or details that were completed outside of the app."
    sections.append(("Summary and Next Steps",closing))

    text_parts=[]
    for head,body in sections:
        text_parts.append(head.upper()+"\n"+body)
    return "\n\n".join(text_parts)

def render_actual_toolkit(position):
    st.markdown(f"## {position} Work Toolkit")

    # shared tabs
    if position=='Recording Secretary':
        tabs=st.tabs(['Minutes Builder','Action Items','Correspondence Log','Records Checklist'])
        with tabs[0]:
            st.subheader('Meeting Minutes Builder')
            with st.form('rs_minutes',clear_on_submit=True):
                mdate=st.date_input('Meeting Date')
                mtype=st.text_input('Meeting Type / Title')
                attendees=st.text_area('Attendance / Roll Call')
                business=st.text_area('Business Discussed',height=140)
                motions=st.text_area('Motions / Votes / Decisions',height=120)
                adj=st.text_input('Adjournment Time')
                save=st.form_submit_button('Save Meeting Record')
            if save:
                save_toolkit_record(position,'meeting_minutes',mtype or str(mdate),{
                    'meeting_date':mdate.isoformat(),'attendance':attendees,'business':business,
                    'motions_votes_decisions':motions,'adjournment_time':adj
                });st.success('Meeting record saved.');st.rerun()

        with tabs[1]:
            with st.form('rs_action',clear_on_submit=True):
                item=st.text_input('Action Item')
                assigned=st.text_input('Assigned To')
                due=st.date_input('Due Date')
                status=st.selectbox('Status',['Open','In Progress','Complete'])
                save=st.form_submit_button('Save Action Item')
            if save and item.strip():
                save_toolkit_record(position,'action_item',item,{
                    'assigned_to':assigned,'due_date':due.isoformat(),'status':status
                });st.rerun()
            rows_=toolkit_records(position=position,member_id_filter=member_id,record_type='action_item')
            if rows_:st.dataframe(pd.DataFrame([{'Action':r['title'],**(r.get('data') or {})} for r in rows_]),hide_index=True,use_container_width=True)

        with tabs[2]:
            with st.form('rs_corr',clear_on_submit=True):
                cdate=st.date_input('Date')
                direction=st.selectbox('Type',['Received','Sent'])
                party=st.text_input('From / To')
                subject=st.text_input('Subject')
                followup=st.text_area('Follow-up / Notes')
                save=st.form_submit_button('Save Correspondence')
            if save:
                save_toolkit_record(position,'correspondence',subject or 'Correspondence',{
                    'date':cdate.isoformat(),'type':direction,'from_to':party,'follow_up':followup
                });st.rerun()

        with tabs[3]:
            checklist=['Previous meeting minutes','Current membership roster','Committee reports','Correspondence','Action items','Legal/governing records','Upcoming meeting documents']
            for c in checklist:
                st.checkbox(c,key=f"rscheck_{c}")

    elif position=='President':
        tabs=st.tabs(['Leadership Decisions','Officer Check-ins','Quick Meeting Prep','Oversight'])
        with tabs[0]:
            with st.form('pres_decision',clear_on_submit=True):
                title=st.text_input('Decision / Issue')
                detail=st.text_area('Decision or Direction')
                follow=st.text_area('Follow-up Required')
                save=st.form_submit_button('Save Leadership Note')
            if save and title.strip():
                save_toolkit_record(position,'decision',title,{'decision':detail,'follow_up':follow});st.rerun()
        with tabs[1]:
            with st.form('pres_checkin',clear_on_submit=True):
                officer=st.selectbox('Officer',OFFICER_POSITIONS)
                update=st.text_area('Update / Concern / Support Needed')
                due=st.date_input('Follow-up Date')
                save=st.form_submit_button('Save Check-in')
            if save:
                save_toolkit_record(position,'officer_checkin',officer,{'update':update,'follow_up_date':due.isoformat()});st.rerun()
        with tabs[2]:
            st.info('Use Meeting Center to create/publish the agenda. Use this area to collect items before building the agenda.')
            item=st.text_area('New Business / Agenda Prep Notes')
            if st.button('Save Meeting Prep Note') and item.strip():
                save_toolkit_record(position,'meeting_prep','Meeting Prep',{'notes':item});st.rerun()
        with tabs[3]:
            balances=all_budget_balances()
            if balances:st.dataframe(pd.DataFrame(balances),hide_index=True,use_container_width=True)
            st.write(f"Vacant offices: {', '.join(vacant_offices()) if vacant_offices() else 'None'}")

    elif position=='Vice President':
        tabs=st.tabs(['Membership Follow-up','Committee Progress','Delegated Work'])
        with tabs[0]:
            with st.form('vp_member_follow',clear_on_submit=True):
                person=st.text_input('Member / Prospect / Reactivation')
                reason=st.selectbox('Follow-up Type',['Recruitment','Retention','Reactivation','Member Support','Other'])
                notes=st.text_area('Notes / Next Step')
                due=st.date_input('Follow-up Date')
                save=st.form_submit_button('Save Follow-up')
            if save:
                save_toolkit_record(position,'membership_followup',person or reason,{'type':reason,'notes':notes,'follow_up_date':due.isoformat()});st.rerun()
        with tabs[1]:
            with st.form('vp_progress',clear_on_submit=True):
                committee=st.selectbox('Committee / Area',COMMITTEE_NAMES)
                update=st.text_area('Progress Update')
                concern=st.text_area('Concern / Support Needed')
                save=st.form_submit_button('Save Progress Update')
            if save:
                save_toolkit_record(position,'committee_progress',committee,{'progress':update,'support_needed':concern});st.rerun()
        with tabs[2]:
            mytasks=[x for x in safe_rows('event_tasks') if x.get('assigned_member_id')==member_id]
            if mytasks:st.dataframe(pd.DataFrame(mytasks)[['task','due_date','status','completion_notes']],hide_index=True,use_container_width=True)

    elif position=='Historian':
        tabs=st.tabs(['History Log','Photo Archive','Communications Queue','Celebrations & Nominations','Archive Checklist'])
        with tabs[0]:
            with st.form('hist_entry',clear_on_submit=True):
                event=st.text_input('Event / Activity');d=st.date_input('Date');summary=st.text_area('What happened?',height=120);significance=st.text_area('Why should this be preserved?');save=st.form_submit_button('Save History Entry',use_container_width=True)
            if save:save_toolkit_record(position,'history_entry',event or str(d),{'date':d.isoformat(),'summary':summary,'significance':significance});st.rerun()
        with tabs[1]:
            photos=table('historian_photos').select('*').order('created_at',desc=True).execute().data or [];mm=historian_member_choice_map()
            if not photos:st.info('No photo submissions yet.')
            for p in photos[:60]:
                with st.expander(f"{p.get('event_name')} • {p.get('event_date')} • {p.get('status')}"):
                    u=signed_url(p.get('photo_path'),3600) if p.get('photo_path') else None
                    if u:st.image(u,use_container_width=True)
                    st.write(f"**Submitted by:** {p.get('submitted_by_name')}");st.write(p.get('caption') or '')
                    statuses=['Submitted','Historian Review','Needs Information','Approved for Archive','Archived'];cur=p.get('status') or 'Submitted';status=st.selectbox('Status',statuses,index=statuses.index(cur) if cur in statuses else 0,key=f"hpst{p['id']}");note=st.text_area('Historian Note',value=p.get('historian_notes') or '',key=f"hpno{p['id']}")
                    if st.button('Save Photo Review',key=f"hpsv{p['id']}",use_container_width=True):table('historian_photos').update({'status':status,'historian_notes':note,'reviewed_by_member_id':member_id,'reviewed_at':datetime.now(timezone.utc).isoformat()}).eq('id',p['id']).execute();st.rerun()
                    tags=historian_photo_tags(p['id'])
                    if tags:st.dataframe(pd.DataFrame([{'Category':x.get('person_category'),'Name':public_historian_name(x.get('person_category'),x.get('linked_member_name') or x.get('manual_name'),x.get('preferred_public_title') or '')} for x in tags]),hide_index=True,use_container_width=True)
                    cat=st.selectbox('Person Category',HISTORIAN_PERSON_CATEGORIES,key=f"hpcat{p['id']}");opts=[0]+list(mm);person=st.selectbox('Person',opts,format_func=lambda x:'Not Listed — Enter Name' if x==0 else mm[x],key=f"hpper{p['id']}");manual=st.text_input('Manual Name',key=f"hpman{p['id']}");title=st.text_input('Preferred Public Title',key=f"hptit{p['id']}")
                    if st.button('Add Person Tag',key=f"hpadd{p['id']}",use_container_width=True):
                        linked=mm.get(person,'') if person else '';manual_name=normalize_historian_name(manual) if not person else ''
                        if not linked and not manual_name:st.warning('Choose or enter a name.')
                        else:table('historian_photo_people').insert({'photo_id':p['id'],'person_category':cat,'linked_member_id':person or None,'linked_member_name':linked,'manual_name':manual_name,'preferred_public_title':title.strip()}).execute();st.rerun()
        with tabs[2]:
            reqs=table('communications_requests').select('*').order('created_at').execute().data or []
            if not reqs:st.info('No communications requests.')
            for r in reqs[:80]:
                with st.expander(f"{r.get('request_type')} • {r.get('title')} • {r.get('status')}"):
                    st.write(f"**Requested by:** {r.get('requested_by_name')}");st.write(r.get('details') or '')
                    if r.get('attachment_path'):
                        u=signed_url(r.get('attachment_path'),3600)
                        if u:st.link_button('Open Attachment',u,use_container_width=True)
                    cur=r.get('status') or 'Submitted';status=st.selectbox('Workflow Status',COMMUNICATION_STATUSES,index=COMMUNICATION_STATUSES.index(cur) if cur in COMMUNICATION_STATUSES else 0,key=f"hcst{r['id']}");note=st.text_area('Historian Notes',value=r.get('historian_notes') or '',key=f"hcno{r['id']}");sd=st.date_input('Scheduled Date',value=date.today(),key=f"hcsd{r['id']}");stm=st.time_input('Scheduled Time',key=f"hcstm{r['id']}")
                    if st.button('Save Communications Update',key=f"hcsv{r['id']}",use_container_width=True):
                        up={'status':status,'historian_notes':note,'reviewed_by_member_id':member_id,'reviewed_at':datetime.now(timezone.utc).isoformat()}
                        if status in ['Scheduled','Published']:up['scheduled_for']=datetime.combine(sd,stm).replace(tzinfo=timezone.utc).isoformat()
                        if status=='Published':up['published_at']=datetime.now(timezone.utc).isoformat()
                        table('communications_requests').update(up).eq('id',r['id']).execute();st.rerun()
        with tabs[3]:
            b=next_birthdays(90)
            if b:st.dataframe(pd.DataFrame(b),hide_index=True,use_container_width=True)
            noms=table('philo_month_nominations').select('*').order('created_at',desc=True).execute().data or []
            for n in noms[:50]:
                with st.expander(f"{n.get('nominee_name')} • {n.get('nomination_month')} • {n.get('status')}"):
                    st.write(n.get('reason') or '');opts=['Submitted','Under Review','Selected','Not Selected','Archived'];status=st.selectbox('Status',opts,index=opts.index(n.get('status')) if n.get('status') in opts else 0,key=f"hnst{n['id']}")
                    if st.button('Save Nomination',key=f"hnsv{n['id']}"):table('philo_month_nominations').update({'status':status,'reviewed_by_member_id':member_id,'reviewed_at':datetime.now(timezone.utc).isoformat()}).eq('id',n['id']).execute();st.rerun()
        with tabs[4]:
            for c in ['Event photos collected','People in photos identified','Programs/flyers saved','Awards documented','Press/publicity saved','Celebrations reviewed','Annual history updated','Files organized for year-end']:st.checkbox(c,key=f"histcheck_{c}")

    elif position=='Chaplain':
        tabs=st.tabs(['Prayer Requests','Courtesy / Condolence Log','Devotion Prep'])
        with tabs[0]:
            prayers=prayer_requests_for_chaplain()
            if not prayers:st.caption('No prayer requests have been submitted.')
            for pr in prayers:
                with st.container(border=True):
                    st.markdown(f"**{pr.get('member_name')}**")
                    st.write(pr.get('request_text'))
                    status=st.selectbox('Status',['New','Acknowledged','Follow-up Needed','Completed'],index=(['New','Acknowledged','Follow-up Needed','Completed'].index(pr.get('status','New')) if pr.get('status','New') in ['New','Acknowledged','Follow-up Needed','Completed'] else 0),key=f"chappr_{pr['id']}")
                    note=st.text_area('Chaplain Note',value=pr.get('chaplain_note') or '',key=f"chapnote_{pr['id']}")
                    if st.button('Save',key=f"chapsave_{pr['id']}"):
                        table('prayer_requests').update({'status':status,'chaplain_note':note,'updated_at':datetime.now(timezone.utc).isoformat()}).eq('id',pr['id']).execute();st.rerun()
        with tabs[1]:
            with st.form('courtesy_log',clear_on_submit=True):
                person=st.text_input('Member / Family')
                occasion=st.selectbox('Type',['Illness','Hospitalization','Bereavement','Birthday','Congratulations','Other'])
                action=st.text_area('Card / Call / Flowers / Follow-up')
                save=st.form_submit_button('Save Courtesy Record')
            if save:
                save_toolkit_record(position,'courtesy',person or occasion,{'type':occasion,'action':action});st.rerun()
        with tabs[2]:
            theme=st.text_input('Devotion / Meditation Theme')
            notes=st.text_area('Prayer / Scripture / Notes')
            if st.button('Save Devotion Prep') and (theme.strip() or notes.strip()):
                save_toolkit_record(position,'devotion',theme or 'Devotion Prep',{'notes':notes});st.rerun()

    elif position=='Parliamentarian':
        tabs=st.tabs(['Question / Ruling Log','Governance Review Notes','Motion Quick Guide'])
        with tabs[0]:
            with st.form('parl_ruling',clear_on_submit=True):
                q=st.text_area('Parliamentary Question')
                source=st.selectbox('Primary Source',GOVERNANCE_SOURCE_OPTIONS)
                ruling=st.text_area('Ruling / Guidance')
                save=st.form_submit_button('Save Parliamentary Record')
            if save:
                save_toolkit_record(position,'parliamentary_ruling',q[:80] or 'Parliamentary Question',{'source':source,'ruling_guidance':ruling});st.rerun()
        with tabs[1]:
            with st.form('gov_review',clear_on_submit=True):
                item=st.text_input('Bylaw / Procedure / Recommendation')
                concern=st.text_area('Concern / Alignment Issue')
                recommendation=st.text_area('Recommended Follow-up')
                save=st.form_submit_button('Save Governance Review Note')
            if save:
                save_toolkit_record(position,'governance_review',item,{'concern':concern,'recommended_follow_up':recommendation});st.rerun()
        with tabs[2]:
            st.dataframe(pd.DataFrame(MOTION_RANKING).rename(columns={'rank':'Rank','motion':'Motion','second':'Second?','debatable':'Debatable?','amendable':'Amendable?','vote':'Vote'}),hide_index=True,use_container_width=True)

    elif position=='Sergeant-at-Arms':
        tabs=st.tabs(['Meeting Order Log','Attendance / Access','Meeting Prep'])
        with tabs[0]:
            with st.form('saa_order',clear_on_submit=True):
                d=st.date_input('Meeting Date')
                issue=st.text_input('Meeting Order / Logistics Item')
                action=st.text_area('Action Taken / Notes')
                save=st.form_submit_button('Save Meeting Order Record')
            if save:
                save_toolkit_record(position,'meeting_order',issue or str(d),{'date':d.isoformat(),'action_notes':action});st.rerun()
        with tabs[1]:
            st.info('Use this area to record meeting attendance/access notes that need to be retained.')
            note=st.text_area('Attendance / Access Note')
            if st.button('Save Attendance Note') and note.strip():
                save_toolkit_record(position,'meeting_order','Attendance / Access',{'notes':note});st.rerun()
        with tabs[2]:
            for c in ['Room/access ready','Sign-in ready','Materials ready','Order/protocol reminders ready','Special guests/entrance instructions reviewed']:
                st.checkbox(c,key=f"saacheck_{c}")

    elif position in ['Financial Secretary','Treasurer']:
        st.info('Your finance tabs below are your primary working toolkit. Use them for dues, incoming funds, remittances, budgets, expenses, vouchers, and financial reporting.')

    else:
        st.info('Use Events & Tasks, Reports, and this dashboard to manage work assigned to your office.')


def meeting_link_allowed_positions():
    raw=setting('meeting_link_positions','Recording Secretary')
    vals=[x.strip() for x in raw.split(',') if x.strip()]
    return vals or ['Recording Secretary']

def can_manage_meeting_links(mid,is_admin_flag=False):
    if is_admin_flag:return True
    positions=[x.get('position') for x in member_offices(mid)]
    return any(p in meeting_link_allowed_positions() for p in positions)

def active_meeting_links_now():
    now=datetime.now(timezone.utc)
    try:
        rows_=table('meeting_links').select('*').eq('active',True).order('meeting_start').execute().data or []
    except:
        return []
    active=[]
    for r in rows_:
        try:
            visible_from=pd.to_datetime(r.get('visible_from'),utc=True).to_pydatetime()
            visible_until=pd.to_datetime(r.get('visible_until'),utc=True).to_pydatetime()
            if not (visible_from <= now <= visible_until):
                continue
            audience=r.get('audience_type') or 'All Philos'
            if audience=='All Philos':
                if is_philo_member(member_id):
                    active.append(r)
            else:
                cid=r.get('committee_id')
                if cid and any(m.get('member_id')==member_id for m in committee_members(cid)):
                    active.append(r)
        except:
            continue
    return active

def all_meeting_links():
    try:return table('meeting_links').select('*').order('meeting_start',desc=True).execute().data or []
    except:return []


def active_committees():
    try:
        return table('committees').select('*').eq('active',True).order('committee_name').execute().data or []
    except:
        return []

def committee_members(committee_id):
    try:
        return table('committee_members').select('*').eq('committee_id',committee_id).eq('active',True).order('member_name').execute().data or []
    except:
        return []

def committees_for_member(mid):
    out=[]
    try:
        memberships=table('committee_members').select('*').eq('member_id',mid).eq('active',True).execute().data or []
        for m in memberships:
            try:
                c=table('committees').select('*').eq('id',m.get('committee_id')).single().execute().data
                if c and c.get('active'):
                    out.append(c)
            except:
                pass
    except:
        pass
    return out

def chaired_committees(mid):
    return [c for c in active_committees() if c.get('chair_member_id')==mid]

def visible_published_agendas(mid):
    visible=[]
    for a in agenda_rows(True):
        audience=a.get('audience_type') or 'All Philos'
        if audience=='All Philos':
            if is_philo_member(mid):
                visible.append(a)
        else:
            cid=a.get('committee_id')
            if cid and any(m.get('member_id')==mid for m in committee_members(cid)):
                visible.append(a)
    return visible

def published_reports():
    try:
        return table('published_reports').select('*').eq('active',True).order('published_at',desc=True).execute().data or []
    except:
        return []

def visible_published_reports(mid):
    visible=[]
    for r in published_reports():
        audience=r.get('audience_type') or 'All Philos'
        if audience=='All Philos':
            if is_philo_member(mid):
                visible.append(r)
        else:
            cid=r.get('committee_id')
            if cid and any(m.get('member_id')==mid for m in committee_members(cid)):
                visible.append(r)
    return visible

def can_use_management_centers(mid,is_admin_flag=False):
    return is_admin_flag or bool(member_offices(mid)) or bool(chaired_committees(mid))

def committee_report_text(committee,period,extra=""):
    members_=committee_members(committee['id'])
    member_ids=[m.get('member_id') for m in members_]
    tasks=[x for x in safe_rows('event_tasks') if x.get('assigned_member_id') in member_ids]
    completed=[x for x in tasks if x.get('status')=='Complete']
    in_progress=[x for x in tasks if x.get('status')=='In Progress']

    names=", ".join([m.get('member_name','') for m in members_]) or "No members recorded"
    task_summary=", ".join([x.get('task','') for x in completed[:8] if x.get('task')]) or "No completed tasks recorded in the Hub."

    text=f"""OVERVIEW
During {period}, the {committee.get('committee_name','')} Committee continued its work in support of the NBS Philo Affiliate.

COMMITTEE LEADERSHIP AND MEMBERSHIP
Chair: {committee.get('chair_member_name','')}
Members: {names}

ACTIVITIES AND ASSIGNMENTS
Committee member tasks recorded in the Hub: {len(tasks)}
Completed: {len(completed)}
In Progress: {len(in_progress)}
Completed work included: {task_summary}

ACCOMPLISHMENTS AND HIGHLIGHTS
Add the committee's major accomplishments, programs, decisions, service, or outcomes here.

UNFINISHED BUSINESS / CHALLENGES
Add unresolved items, concerns, or follow-up needs here.

NEXT STEPS
Add upcoming priorities, recommendations, deadlines, and assistance needed here."""
    if extra.strip():
        text += "\n\nADDITIONAL INFORMATION\n"+extra.strip()
    return text


def create_member_account(first_name,last_name,email,phone,is_philo=True,member_type='Returning',role='Member'):
    full_name=" ".join([x.strip() for x in [first_name,last_name] if x.strip()])
    if not full_name:
        raise ValueError("First and last name are required.")
    if not email.strip():
        raise ValueError("Email is required.")

    existing=table('members').select('*').eq('email',email.strip().lower()).execute().data or []
    if existing:
        raise ValueError("A member with that email already exists.")

    result=table('members').insert({
        'full_name':full_name,
        'email':email.strip().lower(),
        'phone':phone.strip(),
        'role':role,
        'active':True
    }).execute()

    member_row=(result.data or [{}])[0]
    mid=member_row.get('id')
    if not mid:
        # fallback lookup
        row=table('members').select('*').eq('email',email.strip().lower()).single().execute().data
        mid=row.get('id')

    table('member_profiles').upsert({
        'member_id':mid,
        'first_name':first_name.strip(),
        'last_name':last_name.strip(),
        'updated_at':datetime.now(timezone.utc).isoformat()
    }).execute()

    table('member_private').upsert({
        'member_id':mid,
        'is_philo':bool(is_philo),
        'member_type':member_type,
        'advisor_notes':''
    }).execute()

    return mid


INITIAL_MEMBER_CODE='1943'


def _pin_pepper():
    a=auth_cfg() or {}
    secret=str(a.get('admin_password') or '').encode('utf-8')
    if not secret:
        raise RuntimeError('Server authentication secret is not configured.')
    return hashlib.sha256(b'nbs-philo-pin-pepper:'+secret).digest()


def hash_pin(pin):
    salt=hashlib.sha256(f"{member_id}:{_pin_pepper().hex()}".encode('utf-8')).digest()[:16]
    digest=hashlib.pbkdf2_hmac('sha256',str(pin).encode('utf-8'),salt,200000)
    return 'pbkdf2_sha256$200000$'+digest.hex()


def _legacy_pin_hash(pin):
    return hashlib.sha256(str(pin).encode('utf-8')).hexdigest()

def member_pin_record(mid):
    try:
        r=table('member_access').select('*').eq('member_id',mid).execute().data or []
        return r[0] if r else None
    except:
        return None

def member_has_personal_pin(mid):
    r=member_pin_record(mid)
    return bool(r and r.get('pin_hash'))

def set_member_pin(mid,pin):
    if not (str(pin).isdigit() and len(str(pin))==4):
        raise ValueError('PIN must be exactly 4 digits.')
    salt=hashlib.sha256(f"{mid}:{_pin_pepper().hex()}".encode('utf-8')).digest()[:16]
    digest=hashlib.pbkdf2_hmac('sha256',str(pin).encode('utf-8'),salt,200000).hex()
    table('member_access').upsert({
        'member_id':mid,
        'pin_hash':'pbkdf2_sha256$200000$'+digest,
        'pin_set_at':datetime.now(timezone.utc).isoformat(),
        'reset_required':False
    }).execute()

def check_member_pin(mid,pin):
    r=member_pin_record(mid)
    if not r or not r.get('pin_hash') or r.get('reset_required'):
        return False
    stored=str(r.get('pin_hash') or '')
    # hash_pin uses the logged-in candidate member_id global; temporarily derive directly for mid.
    if stored.startswith('pbkdf2_sha256$'):
        salt=hashlib.sha256(f"{mid}:{_pin_pepper().hex()}".encode('utf-8')).digest()[:16]
        digest=hashlib.pbkdf2_hmac('sha256',str(pin).encode('utf-8'),salt,200000).hex()
        ok=hmac.compare_digest(stored,'pbkdf2_sha256$200000$'+digest)
    else:
        ok=hmac.compare_digest(stored,_legacy_pin_hash(pin))
        if ok:
            # One-time transparent migration from the legacy unsalted SHA-256 hash.
            salt=hashlib.sha256(f"{mid}:{_pin_pepper().hex()}".encode('utf-8')).digest()[:16]
            digest=hashlib.pbkdf2_hmac('sha256',str(pin).encode('utf-8'),salt,200000).hex()
            table('member_access').update({'pin_hash':'pbkdf2_sha256$200000$'+digest}).eq('member_id',mid).execute()
    return ok

def reset_member_pin(mid):
    table('member_access').upsert({
        'member_id':mid,
        'pin_hash':'',
        'pin_set_at':None,
        'reset_required':True
    }).execute()


BULK_MEMBER_COLUMNS=[
    'First Name','Last Name','Email','Phone','Prefix','Pronouns',
    'Address','City State Zip','Birthday','Membership Type',
    'Counts as Philo','Philo ID','Account Type'
]

def normalize_bulk_header(value):
    return re.sub(r'[^a-z0-9]+','',str(value or '').strip().lower())

def bulk_column_map(columns):
    aliases={
        'firstname':'First Name','first':'First Name','givenname':'First Name',
        'lastname':'Last Name','last':'Last Name','surname':'Last Name','familyname':'Last Name',
        'email':'Email','emailaddress':'Email',
        'phone':'Phone','phonenumber':'Phone','number':'Phone',
        'prefix':'Prefix','title':'Prefix',
        'pronouns':'Pronouns','pronoun':'Pronouns',
        'address':'Address','streetaddress':'Address',
        'citystatezip':'City State Zip','citystatezipcode':'City State Zip','citystatepostalcode':'City State Zip',
        'birthday':'Birthday','birthdate':'Birthday','dateofbirth':'Birthday','dob':'Birthday',
        'membershiptype':'Membership Type','membertype':'Membership Type',
        'countsasphilo':'Counts as Philo','isphilo':'Counts as Philo','philo':'Counts as Philo',
        'philoid':'Philo ID','philoadid':'Philo ID','memberid':'Philo ID',
        'accounttype':'Account Type','role':'Account Type'
    }
    result={}
    for col in columns:
        canon=aliases.get(normalize_bulk_header(col))
        if canon:
            result[col]=canon
    return result

def parse_yes_no(value,default=True):
    if pd.isna(value) or str(value).strip()=='':
        return default
    v=str(value).strip().lower()
    if v in ['yes','y','true','1','philo','member']:
        return True
    if v in ['no','n','false','0','advisor','admin']:
        return False
    return default

def clean_cell(value):
    if pd.isna(value):
        return ''
    return str(value).strip()

def read_bulk_member_file(uploaded):
    name=(uploaded.name or '').lower()
    raw=bytes(uploaded.getbuffer())
    if len(raw)>5*1024*1024:
        raise ValueError('Bulk member file is too large. Maximum size is 5 MB.')
    bio=BytesIO(raw)
    if name.endswith('.csv'):
        df=pd.read_csv(bio,dtype=str).fillna('')
    elif name.endswith('.xlsx'):
        df=pd.read_excel(bio,dtype=str,engine='openpyxl').fillna('')
    else:
        raise ValueError('Upload a .xlsx or .csv file.')
    if len(df)>2000:
        raise ValueError('Bulk member file has too many rows. Maximum is 2,000 rows per import.')
    mapping=bulk_column_map(df.columns)
    df=df.rename(columns=mapping)
    missing=[c for c in ['First Name','Last Name','Email'] if c not in df.columns]
    if missing:
        raise ValueError('Missing required column(s): '+', '.join(missing))
    # Keep only recognized fields, adding optional columns as blanks.
    for c in BULK_MEMBER_COLUMNS:
        if c not in df.columns:
            df[c]=''
    return df[BULK_MEMBER_COLUMNS].copy()

def import_bulk_members(df):
    added=[]
    skipped=[]
    errors=[]

    existing_people=table('members').select('id,email').execute().data or []
    existing_emails={str(x.get('email') or '').strip().lower() for x in existing_people if x.get('email')}

    for idx,row in df.iterrows():
        row_num=int(idx)+2
        first=clean_cell(row.get('First Name'))
        last=clean_cell(row.get('Last Name'))
        email=clean_cell(row.get('Email')).lower()
        phone=clean_cell(row.get('Phone'))

        if not first or not last or not email:
            errors.append({'Row':row_num,'Name':f"{first} {last}".strip(),'Email':email,'Reason':'Missing first name, last name, or email'})
            continue

        if email in existing_emails:
            skipped.append({'Row':row_num,'Name':f"{first} {last}",'Email':email,'Reason':'Email already exists'})
            continue

        mtype=clean_cell(row.get('Membership Type')).title() or 'Returning'
        if mtype not in ['Returning','Inductee','Reactivating']:
            mtype='Returning'

        requested_role=clean_cell(row.get('Account Type')).title() or 'Member'
        is_philo=parse_yes_no(row.get('Counts as Philo'),True)
        # Bulk roster files are for Philo accounts. Advisor/Admin creation is deliberately manual.
        role='Member'
        if requested_role=='Admin':
            errors.append({'Row':row_num,'Name':f"{first} {last}",'Email':email,'Reason':'Admin accounts cannot be created by bulk import; create the Advisor/Admin account manually.'})
            continue

        try:
            mid=create_member_account(
                first,last,email,phone,
                is_philo=is_philo,
                member_type=mtype,
                role=role
            )

            profile_updates={
                'member_id':mid,
                'first_name':first,
                'last_name':last,
                'prefix':clean_cell(row.get('Prefix')),
                'pronouns':clean_cell(row.get('Pronouns')),
                'address':clean_cell(row.get('Address')),
                'city_state_zip':clean_cell(row.get('City State Zip')),
                'updated_at':datetime.now(timezone.utc).isoformat()
            }
            birthday=clean_cell(row.get('Birthday'))
            if birthday:
                try:
                    profile_updates['birthday']=pd.to_datetime(birthday).date().isoformat()
                except:
                    pass
            table('member_profiles').upsert(profile_updates).execute()

            private_updates={
                'member_id':mid,
                'is_philo':is_philo,
                'member_type':mtype
            }
            philo_id=clean_cell(row.get('Philo ID'))
            if philo_id:
                private_updates['philo_id']=philo_id
            table('member_private').upsert(private_updates).execute()

            existing_emails.add(email)
            added.append({'Row':row_num,'Name':f"{first} {last}",'Email':email})
        except Exception as ex:
            errors.append({'Row':row_num,'Name':f"{first} {last}",'Email':email,'Reason':str(ex)})

    return added,skipped,errors


# ===== Treasurer Finance Center v3 =====
FINANCE_OPEN_STATUSES=['Pending','Submitted','Under Review','Needs Information','Needs Correction','Returned','Approved']

def finance_fy():
    return setting('current_fiscal_year','2026-2027').strip() or '2026-2027'

def finance_settings():
    fy=finance_fy()
    try:
        r=table('treasurer_financial_settings').select('*').eq('fiscal_year',fy).execute().data or []
        if r:return r[0]
    except:pass
    return {'fiscal_year':fy,'opening_balance':0,'protected_minimum':500,'notes':''}

def finance_year_rows():
    fy=finance_fy()
    return [r for r in finance_rows() if (r.get('fiscal_year') or fy)==fy]

def normalized_voucher_status(v):
    s=(v.get('status') or 'Pending').strip()
    return {'Returned':'Needs Information','Needs Correction':'Needs Information'}.get(s,s)

def voucher_commitment(v):
    if normalized_voucher_status(v)!='Approved' or v.get('posted_to_ledger'):return 0.0
    return float(v.get('amount_approved') if v.get('amount_approved') is not None else (v.get('amount') or 0))

def finance_budget_rows():
    fy=finance_fy(); tx=finance_year_rows(); vouchers=safe_rows('reimbursements'); out=[]
    for b in budget_lines():
        if (b.get('fiscal_year') or fy)!=fy:continue
        bid=b['id']
        income=sum(float(r.get('amount') or 0) for r in tx if r.get('budget_line_id')==bid and r.get('direction')=='Income')
        spent=sum(float(r.get('amount') or 0) for r in tx if r.get('budget_line_id')==bid and r.get('direction')=='Expense')
        committed=sum(voucher_commitment(v) for v in vouchers if v.get('budget_line_id')==bid)
        budget=float(b.get('budget_amount') or 0)
        out.append({'Budget ID':bid,'Budget':b.get('name'),'Assigned To':b.get('owner_position') or 'General Chapter',
                    'Budgeted':budget,'Income':income,'Spent':spent,'Committed':committed,
                    'Available':budget+income-spent-committed})
    return out

def finance_snapshot_v3():
    s=finance_settings(); tx=finance_year_rows()
    deposits=sum(float(r.get('amount') or 0) for r in tx if r.get('direction')=='Income')
    withdrawals=sum(float(r.get('amount') or 0) for r in tx if r.get('direction')=='Expense')
    opening=float(s.get('opening_balance') or 0); protected=float(s.get('protected_minimum') or 0)
    current=opening+deposits-withdrawals
    budgets_=finance_budget_rows()
    committed=sum(float(x.get('Committed') or 0) for x in budgets_)
    reserved=sum(max(float(x.get('Available') or 0),0) for x in budgets_)
    operating=current-reserved
    return {'fiscal_year':finance_fy(),'opening_balance':opening,'current_balance':current,
            'deposits':deposits,'withdrawals':withdrawals,'committed':committed,'reserved':reserved,
            'operating':operating,'protected':protected,'above_protected':max(operating-protected,0)}

def monthly_finance_v3():
    grouped={}
    for r in finance_year_rows():
        m=str(r.get('transaction_date') or '')[:7]
        if not m:continue
        g=grouped.setdefault(m,{'Month':m,'Deposits':0.0,'Withdrawals':0.0,'Net':0.0})
        a=float(r.get('amount') or 0)
        if r.get('direction')=='Income':g['Deposits']+=a;g['Net']+=a
        else:g['Withdrawals']+=a;g['Net']-=a
    return [grouped[k] for k in sorted(grouped)]

def finance_audit(action,entity_type,entity_id='',details=None):
    try:
        table('finance_audit_log').insert({'member_id':member_id,'member_name':member_name,'action':action,
            'entity_type':entity_type,'entity_id':str(entity_id or ''),'details':details or {},
            'created_at':datetime.now(timezone.utc).isoformat()}).execute()
    except:pass

def prior_voucher_approvals(v):
    submitter_positions=[x.get('position') for x in member_offices(v.get('submitted_by_member_id'))] if v.get('submitted_by_member_id') else []
    required=[p for p in ['Financial Secretary','President'] if p not in submitter_positions]
    approvals=voucher_approvals(v['id']); out=[]
    for p in required:
        acts=[a for a in approvals if a.get('position')==p]; a=acts[-1] if acts else None
        out.append({'Position':p,'Decision':a.get('decision') if a else 'Waiting',
                    'Signed By':a.get('signature_name') if a else '',
                    'Reason / Note':a.get('reason_note') if a else ''})
    return out

def voucher_ready_for_treasurer(v):
    p=prior_voucher_approvals(v)
    return all(x['Decision']=='Approved' for x in p) if p else True

def save_voucher_decision(v,position,decision,signature,note=''):
    try:
        old=table('document_approvals').select('*').eq('document_type','reimbursement').eq('document_id',v['id']).eq('position',position).execute().data or []
        for x in old:table('document_approvals').delete().eq('id',x['id']).execute()
    except:pass
    table('document_approvals').insert({'document_type':'reimbursement','document_id':v['id'],'position':position,
        'member_id':member_id,'decision':decision,'reason_note':note.strip(),'signature_name':signature.strip(),
        'signed_at':datetime.now(timezone.utc).isoformat()}).execute()
    if decision=='Denied':
        table('reimbursements').update({'status':'Denied','status_reason_note':note.strip()}).eq('id',v['id']).execute()
    elif decision=='Needs Information':
        table('reimbursements').update({'status':'Needs Information','status_reason_note':note.strip()}).eq('id',v['id']).execute()
    finance_audit('VOUCHER_'+decision.upper().replace(' ','_'),'voucher',v['id'],{'position':position,'note':note})

def treasurer_approve_voucher(v,approved_amount,note,signature):
    if not voucher_ready_for_treasurer(v):raise RuntimeError('Financial Secretary and President authorization must be complete first.')
    save_voucher_decision(v,'Treasurer','Approved',signature,note)
    table('reimbursements').update({'status':'Approved','amount_approved':approved_amount,'treasurer_notes':note.strip(),
        'reviewed_by_member_id':member_id,'reviewed_by_name':member_name,'approved_at':datetime.now(timezone.utc).isoformat(),
        'status_reason_note':note.strip()}).eq('id',v['id']).execute()

def mark_paid_v3(v,payment_date,payment_method,reference='',note=''):
    fresh=table('reimbursements').select('*').eq('id',v['id']).single().execute().data
    if normalized_voucher_status(fresh)!='Approved':raise RuntimeError('Voucher must be Approved before payment.')
    if fresh.get('posted_to_ledger'):raise RuntimeError('Voucher was already posted to the ledger.')
    existing=table('finance_transactions').select('id').eq('source_type','Voucher').eq('source_id',fresh['id']).execute().data or []
    if existing:raise RuntimeError('Duplicate voucher payment blocked: ledger entry already exists.')
    amount=float(fresh.get('amount_approved') if fresh.get('amount_approved') is not None else (fresh.get('amount') or 0))
    table('finance_transactions').insert({'transaction_date':payment_date.isoformat(),'fiscal_year':finance_fy(),
        'direction':'Expense','category':'Reimbursement / Voucher','position':fresh.get('budget_position') or 'General Chapter',
        'budget_line_id':fresh.get('budget_line_id'),'payer_payee':fresh.get('payee_name') or fresh.get('submitted_by_name') or '',
        'amount':amount,'payment_method':payment_method,'reference_number':reference.strip(),'notes':note.strip() or fresh.get('description') or '',
        'source_type':'Voucher','source_id':fresh['id'],'entered_by_member_id':member_id,'entered_by_name':member_name}).execute()
    table('reimbursements').update({'status':'Paid','payment_method':payment_method,'payment_reference':reference.strip(),
        'paid_at':datetime.now(timezone.utc).isoformat(),'posted_to_ledger':True,'status_reason_note':note.strip() or 'Payment completed.'}).eq('id',fresh['id']).execute()
    finance_audit('MARK_VOUCHER_PAID','voucher',fresh['id'],{'amount':amount,'method':payment_method,'reference':reference})

def treasurer_handoff_v3():
    s=finance_snapshot_v3(); b=finance_budget_rows()
    openv=[v for v in safe_rows('reimbursements') if normalized_voucher_status(v) in FINANCE_OPEN_STATUSES]
    return f"""NBS PHILO TREASURER HANDOFF SUMMARY
Fiscal Year: {s['fiscal_year']}

FINANCIAL POSITION
Opening Balance: ${s['opening_balance']:,.2f}
Current Balance: ${s['current_balance']:,.2f}
YTD Deposits: ${s['deposits']:,.2f}
YTD Withdrawals: ${s['withdrawals']:,.2f}
Approved / Not Paid: ${s['committed']:,.2f}
Operating Funds: ${s['operating']:,.2f}
Protected Minimum: ${s['protected']:,.2f}

BUDGET STATUS
{chr(10).join([f"- {x['Budget']}: budget ${x['Budgeted']:,.2f}; spent ${x['Spent']:,.2f}; committed ${x['Committed']:,.2f}; available ${x['Available']:,.2f}" for x in b]) or '- No active budgets.'}

OPEN VOUCHERS
{chr(10).join([f"- #{v.get('id')} {v.get('submitted_by_name')} — ${float(v.get('amount') or 0):,.2f} — {v.get('status')}" for v in openv]) or '- No open vouchers.'}

HANDOFF NOTES
Add reconciliation notes, pending deposits/payments, outstanding documentation, deadlines, and other transition details.
"""

def treasurer_report_v3(period):
    s=finance_snapshot_v3(); m=next((x for x in monthly_finance_v3() if x['Month']==period),{'Deposits':0,'Withdrawals':0,'Net':0})
    b=finance_budget_rows()
    return f"""NBS PHILO TREASURER FINANCIAL REPORT
Reporting Period: {period}
Fiscal Year: {s['fiscal_year']}

MONTHLY ACTIVITY
Deposits: ${m['Deposits']:,.2f}
Withdrawals: ${m['Withdrawals']:,.2f}
Net: ${m['Net']:,.2f}

YEAR-TO-DATE
Opening Balance: ${s['opening_balance']:,.2f}
YTD Deposits: ${s['deposits']:,.2f}
YTD Withdrawals: ${s['withdrawals']:,.2f}
Current Balance: ${s['current_balance']:,.2f}
Approved / Not Paid: ${s['committed']:,.2f}
Operating Funds: ${s['operating']:,.2f}
Protected Minimum: ${s['protected']:,.2f}

BUDGETS
{chr(10).join([f"- {x['Budget']}: budget ${x['Budgeted']:,.2f}; spent ${x['Spent']:,.2f}; committed ${x['Committed']:,.2f}; available ${x['Available']:,.2f}" for x in b]) or '- No active budgets.'}

TREASURER NOTES / EXPLANATION
Add reconciliation notes, unusual activity, outstanding payments, expected deposits, and recommendations.
"""

def normalize_login_name(value):
    return " ".join(str(value or "").strip().casefold().split())

def member_login_names(member_row):
    vals=[]
    ps=safe_rows('member_profiles',member_id=member_row['id']);p=ps[0] if ps else {}
    first=(p.get('first_name') or '').strip();last=(p.get('last_name') or '').strip()
    if first and last:vals.append((normalize_login_name(first),normalize_login_name(last)))
    parts=(member_row.get('full_name') or '').strip().split()
    if len(parts)>=2:vals.append((normalize_login_name(parts[0]),normalize_login_name(' '.join(parts[1:]))))
    return list(dict.fromkeys(vals))

def find_member_by_typed_name(first_name,last_name):
    f=normalize_login_name(first_name);l=normalize_login_name(last_name)
    if not f or not l:return None
    matches=[]
    for m in active_members():
        if any(mf==f and ml==l for mf,ml in member_login_names(m)):
            matches.append(m)
    if len(matches)>1:
        raise ValueError('More than one active account has that name. Ask the Advisor to update the duplicate account names before logging in.')
    return matches[0] if matches else None

HISTORIAN_PERSON_CATEGORIES=['Philo','Soror (Sigma Member)','Rhoer','Rhosebud','Philo Advisor','Speaker / Presenter','Community Partner','Guest / Community Participant','Vendor','Elected Official','Other']
HISTORIAN_EVENT_TYPES=['Affiliate Meeting','Mixer','Community Service','Social','Fundraiser','Membership Event','Recruitment Event','Workshop','Conference','Regional Event','National Event','Induction','Celebration','Partner Event','Other']
COMMUNICATION_STATUSES=['Submitted','Historian Review','Needs Information','Draft Ready','Sent to Chapter','Chapter Review','Revision Requested','Approved','Scheduled','Published','Archived']

def normalize_historian_name(name):return " ".join(str(name or '').split())
def historian_photo_tags(photo_id):return safe_rows('historian_photo_people',photo_id=photo_id)
def historian_member_choice_map():return {m['id']:m['full_name'] for m in philo_dropdown_people()}
def public_historian_name(category,name,preferred_title=''):
    name=normalize_historian_name(name)
    if not name:return ''
    if category in ['Philo','Rhoer','Rhosebud'] and not name.casefold().startswith(category.casefold()+' '):return f'{category} {name}'
    if preferred_title and preferred_title!='No Title' and not name.casefold().startswith(preferred_title.casefold()+' '):return f'{preferred_title} {name}'
    defaults={'Advisor':'Advisor','Speaker / Presenter':'Speaker','Community Partner':'Community Partner','Vendor':'Vendor','Elected Official':'Elected Official'}
    title=defaults.get(category)
    if title and not name.casefold().startswith(title.casefold()+' '):return f'{title} {name}'
    if category=='Guest / Community Participant':return f'{name}, Community Guest'
    return name

def next_birthdays(days=60):
    today=date.today();out=[]
    for m in philo_dropdown_people():
        ps=safe_rows('member_profiles',member_id=m['id']);p=ps[0] if ps else {}
        if not p.get('birthday'):continue
        try:
            bd=pd.to_datetime(p['birthday']).date();candidate=bd.replace(year=today.year)
            if candidate<today:candidate=bd.replace(year=today.year+1)
            delta=(candidate-today).days
            if 0<=delta<=days:out.append({'Name':display_member_name(m,p),'Date':candidate.strftime('%B %d'),'Days Away':delta})
        except:pass
    return sorted(out,key=lambda x:x['Days Away'])

def historian_resume_text(mid,name):
    ps=safe_rows('member_profiles',member_id=mid);p=ps[0] if ps else {}
    offices=member_offices(mid);committees=committees_for_member(mid)
    sessions=[x for x in safe_rows('service_sessions',member_id=mid) if x.get('check_out')];hours=sum(float(x.get('hours') or 0) for x in sessions)
    tasks=[x for x in safe_rows('event_tasks') if x.get('assigned_member_id')==mid];done=[x for x in tasks if x.get('status')=='Complete']
    return f"""PHILO SERVICE RESUME\n\n{name}\n\nPROFILE\n{p.get('bio') or 'Add a short service-focused biography here.'}\n\nPHILO SERVICE\nInduction Date: {p.get('induction_date') or 'Add if applicable'}\nCommunity Service Hours Recorded: {hours:.2f}\n\nLEADERSHIP / OFFICES\n{chr(10).join([f"- {x.get('position')}" for x in offices]) or '- No officer positions recorded.'}\n\nCOMMITTEE SERVICE\n{chr(10).join([f"- {x.get('committee_name')}" for x in committees]) or '- No committee assignments recorded.'}\n\nEVENT / PROJECT RESPONSIBILITIES\nTasks Completed: {len(done)} of {len(tasks)}\n{chr(10).join([f"- {x.get('task')}" for x in done[:12]]) or '- No completed tasks recorded yet.'}\n\nHIGHLIGHTS\nAdd awards, special projects, recognitions, conferences, programs, and other accomplishments here."""


D9_ORGANIZATIONS=[
    'Alpha Kappa Alpha Sorority, Incorporated',
    'Alpha Phi Alpha Fraternity, Incorporated',
    'Delta Sigma Theta Sorority, Incorporated',
    'Iota Phi Theta Fraternity, Incorporated',
    'Kappa Alpha Psi Fraternity, Incorporated',
    'Omega Psi Phi Fraternity, Incorporated',
    'Phi Beta Sigma Fraternity, Incorporated',
    'Sigma Gamma Rho Sorority, Incorporated',
    'Zeta Phi Beta Sorority, Incorporated',
    'Other — Type In'
]
EVENT_AUDIENCE_OPTIONS=['Philos','Sorors','Philos & Sorors','Other Organization']
EVENT_TYPE_OPTIONS=[
    'Community Service','Social','Fundraiser','Affiliate Meeting',
    'Membership Event','Workshop / Program','Philo Regional / National Event','Other'
]

def event_type_label(event):
    return event.get('event_type') or 'Other'

def service_sessions_for_member(mid):
    return [
        x for x in safe_rows('service_sessions',member_id=mid)
        if x.get('check_out')
    ]

def member_service_hours_total(mid):
    return sum(float(x.get('hours') or 0) for x in service_sessions_for_member(mid))

def open_service_session(mid,event_id):
    try:
        rows_=table('service_sessions').select('*').eq(
            'member_id',mid
        ).eq('event_id',event_id).is_('check_out','null').execute().data or []
        return rows_[0] if rows_ else None
    except:
        return None

def clock_in_service_event(event_id):
    current=open_service_session(member_id,event_id)
    if current:
        raise RuntimeError('You are already clocked in to this event.')
    table('service_sessions').insert({
        'event_id':event_id,
        'member_id':member_id,
        'member_name':member_name,
        'check_in':datetime.now(timezone.utc).isoformat(),
        'verified':False
    }).execute()

def clock_out_service_event(event_id):
    current=open_service_session(member_id,event_id)
    if not current:
        raise RuntimeError('No open clock-in was found for this event.')
    end=datetime.now(timezone.utc)
    start=pd.to_datetime(current['check_in'],utc=True).to_pydatetime()
    hours=max(0,(end-start).total_seconds()/3600)
    table('service_sessions').update({
        'check_out':end.isoformat(),
        'hours':round(hours,2)
    }).eq('id',current['id']).execute()
    return round(hours,2)

def service_history_for_event(event_id):
    return [
        x for x in safe_rows('service_sessions',event_id=event_id)
        if x.get('check_out')
    ]


def all_members_admin():
    try:
        return table('members').select('*').order('full_name').execute().data or []
    except:
        return []

def record_member_status_change(mid,old_active,new_active,old_role,new_role,reason=''):
    try:
        target=table('members').select('*').eq('id',mid).single().execute().data or {}
        table('member_status_history').insert({
            'member_id':mid,
            'member_name':target.get('full_name') or '',
            'old_active':bool(old_active),
            'new_active':bool(new_active),
            'old_account_type':old_role or '',
            'new_account_type':new_role or '',
            'reason':reason.strip(),
            'changed_by_member_id':member_id if 'member_id' in globals() else None,
            'changed_by_name':member_name if 'member_name' in globals() else '',
            'changed_at':datetime.now(timezone.utc).isoformat()
        }).execute()
    except:
        pass

def start_or_touch_login_activity(mid,name,is_admin_flag=False):
    now=datetime.now(timezone.utc)
    session_id=st.session_state.get('login_activity_id')
    if session_id:
        try:
            table('login_activity').update({
                'last_activity_at':now.isoformat(),
                'last_page':st.session_state.get('page','🏠 Dashboard')
            }).eq('id',session_id).execute()
            return session_id
        except:
            st.session_state.pop('login_activity_id',None)

    try:
        result=table('login_activity').insert({
            'member_id':mid,
            'member_name':name,
            'account_type':'Admin' if is_admin_flag else 'Member',
            'login_at':now.isoformat(),
            'last_activity_at':now.isoformat(),
            'last_page':st.session_state.get('page','🏠 Dashboard'),
            'signed_out':False
        }).execute()
        session_id=(result.data or [{}])[0].get('id')
        if session_id:
            st.session_state['login_activity_id']=session_id
        return session_id
    except:
        return None

def close_login_activity():
    session_id=st.session_state.get('login_activity_id')
    if not session_id:return
    now=datetime.now(timezone.utc)
    try:
        table('login_activity').update({
            'logout_at':now.isoformat(),
            'last_activity_at':now.isoformat(),
            'signed_out':True
        }).eq('id',session_id).execute()
    except:
        pass

def login_activity_rows(limit=500):
    try:
        return table('login_activity').select('*').order('login_at',desc=True).limit(limit).execute().data or []
    except:
        return []

def activity_duration_label(row):
    try:
        start=pd.to_datetime(row.get('login_at'),utc=True)
        end=pd.to_datetime(row.get('logout_at') or row.get('last_activity_at'),utc=True)
        minutes=max(int((end-start).total_seconds()/60),0)
        if minutes<60:return f'{minutes} min'
        hours=minutes//60;rem=minutes%60
        return f'{hours} hr {rem} min'
    except:
        return ''

def activity_status(row):
    if row.get('signed_out') or row.get('logout_at'):return 'Signed Out'
    try:
        last=pd.to_datetime(row.get('last_activity_at'),utc=True)
        age=(pd.Timestamp.now(tz='UTC')-last).total_seconds()/60
        return 'Active' if age<=15 else 'Last Seen'
    except:
        return 'Last Seen'

def extract_google_calendar_id(value):
    raw=str(value or '').strip()
    if not raw:return ''
    if '@' in raw and '://' not in raw:
        return raw
    try:
        parsed=urllib.parse.urlparse(raw)
        qs=urllib.parse.parse_qs(parsed.query)
        if qs.get('src'):
            return urllib.parse.unquote(qs['src'][0])
        if qs.get('cid'):
            cid=qs['cid'][0]
            try:
                padded=cid + '='*((4-len(cid)%4)%4)
                decoded=base64.urlsafe_b64decode(padded.encode()).decode()
                if '@' in decoded:return decoded
            except:
                pass
        parts=[urllib.parse.unquote(x) for x in parsed.path.split('/') if x]
        if 'ical' in parts:
            idx=parts.index('ical')
            if idx+1<len(parts):return parts[idx+1]
        for part in parts:
            if '@' in part:return part
    except:
        pass
    return raw

def current_next_month_window():
    now=pd.Timestamp.now(tz='America/New_York')
    start=now.normalize().replace(day=1)
    next_month=(start+pd.DateOffset(months=1)).replace(day=1)
    end=(start+pd.DateOffset(months=2)).replace(day=1)
    return start,next_month,end

def current_next_month_calendar_events():
    start,_,end=current_next_month_window()
    calendar_id=configured_calendar_id() or setting('google_calendar_id','')
    public_rows=public_google_calendar_events(calendar_id,start.isoformat(),end.isoformat()) if calendar_id else []
    synced=[]
    try:
        synced=[
            e for e in table('events').select('*').eq('active',True).order('start_at').execute().data or []
            if start.date().isoformat() <= str(e.get('start_at') or '')[:10] < end.date().isoformat()
        ]
    except:
        pass
    if not public_rows:
        return dedupe_event_rows(synced)
    return dedupe_event_rows(list(public_rows)+list(synced))

def ensure_current_calendar_events_in_board():
    count=0
    for e in current_next_month_calendar_events():
        gid=e.get('google_event_id') or ''
        event_key=e.get('public_event_key') or f"{gid}|{e.get('start_at')}|{e.get('title')}"
        existing=[]
        if gid:
            try:existing=table('events').select('*').eq('google_event_id',gid).execute().data or []
            except:existing=[]
        if not existing:
            try:existing=table('events').select('*').eq('calendar_event_key',event_key).execute().data or []
            except:existing=[]
        if not existing:
            try:
                candidates=table('events').select('*').eq('active',True).execute().data or []
                existing=[r for r in candidates if event_rows_equivalent(r,e)][:1]
            except:
                existing=[]
        payload={
            'google_event_id':gid,
            'calendar_event_key':event_key,
            'title':e.get('title') or 'Untitled Event',
            'start_at':e.get('start_at'),
            'end_at':e.get('end_at') or e.get('start_at'),
            'location':e.get('location') or '',
            'description':e.get('description') or '',
            'event_type':(existing[0].get('event_type') if existing else None) or 'Other',
            'source':'google',
            'active':True
        }
        try:
            if existing:
                table('events').update(payload).eq('id',existing[0]['id']).execute()
            else:
                table('events').insert(payload).execute()
            count+=1
        except:
            pass
    return count

def event_org_display(event):
    audience=event.get('event_audience') or 'Not Classified'
    if audience!='Other Organization':return audience
    return event.get('external_organization') or 'Other Organization'

def chapter_report_source_rows(period=''):
    published=safe_rows('published_reports')
    monthly=safe_rows('monthly_reports')
    if period:
        published=[r for r in published if period.casefold() in str(r.get('report_period') or '').casefold()]
        monthly=[r for r in monthly if period.casefold() in str(r.get('report_month') or '').casefold()]
    return published,monthly

def advisor_chapter_report_v32(period,report_style='Standard Chapter Report'):
    published,monthly=chapter_report_source_rows(period)
    officers=safe_rows('officer_assignments',active=True)
    committees=active_committees()
    tasks=safe_rows('event_tasks')
    service=safe_rows('service_sessions')
    history=safe_rows('historian_photos')
    comms=safe_rows('communications_requests')
    recs=safe_rows('recommendations')
    snap=finance_snapshot_v3()
    events=current_next_month_calendar_events()
    hours=sum(float(x.get('hours') or 0) for x in service if x.get('check_out'))

    source_summaries=[]
    for r in published:
        body=(r.get('report_text') or '').strip()
        source_summaries.append(
            f"{r.get('title') or r.get('position') or 'Report'} — "
            f"{r.get('published_by_name') or r.get('member_name') or ''}\n"
            f"{body[:1200] if body else '[Uploaded completed report on file]'}"
        )
    for r in monthly:
        body=(r.get('report_text') or '').strip()
        source_summaries.append(
            f"{r.get('position') or 'Officer'} — {r.get('member_name') or ''}\n{body[:1200]}"
        )

    detail = report_style=='Detailed Affiliate Report'
    quick = report_style=='Quick Summary'

    report=f"""NU BETA SIGMA PHILO AFFILIATE
{report_style.upper()}
Reporting Period: {period}

AFFILIATE SNAPSHOT
Active Philos: {len(philo_members())}
Executive Board Positions Filled: {len(officers)}
Active Committees: {len(committees)}
Community Service Hours Recorded: {hours:.2f}
Event/Project Tasks Completed: {sum(1 for x in tasks if x.get('status')=='Complete')} of {len(tasks)}

FINANCIAL SUMMARY
Current Balance: ${snap['current_balance']:,.2f}
Operating Funds: ${snap['operating']:,.2f}
Approved / Not Yet Paid: ${snap['committed']:,.2f}
YTD Deposits: ${snap['deposits']:,.2f}
YTD Withdrawals: ${snap['withdrawals']:,.2f}

HISTORIAN / COMMUNICATIONS
Photo Submissions on File: {len(history)}
Communications Requests: {len(comms)}
Published Communications: {sum(1 for x in comms if x.get('status')=='Published')}

GOVERNANCE
Recommendations on File: {len(recs)}

UPCOMING EVENTS — CURRENT & NEXT MONTH
{chr(10).join([f"- {e.get('title')} — {calendar_event_date_time(e)[0]} at {calendar_event_date_time(e)[1]}" for e in events[:20]]) or '- No upcoming events listed.'}
"""
    if not quick:
        report+=f"""
OFFICER / COMMITTEE REPORT SOURCE SUMMARY
{chr(10)+chr(10).join(source_summaries) if source_summaries else 'No officer/committee reports matched the selected reporting period.'}

CHAPTER HIGHLIGHTS
Use the source reports and dashboard data above to summarize accomplishments, programs, service, membership activity, collaborations, and notable outcomes.

ITEMS NEEDING CHAPTER ATTENTION
Add pending approvals, incomplete reports, upcoming deadlines, concerns, and matters requiring follow-up.
"""
    if detail:
        report+=f"""
DETAILED OPERATIONS
Open Event/Project Tasks: {sum(1 for x in tasks if x.get('status')!='Complete')}
Vacant Executive Board Positions: {len(vacant_offices())}
Communications Awaiting Completion: {sum(1 for x in comms if x.get('status') not in ['Published','Archived'])}
Reports Included in Source Set: {len(published)+len(monthly)}

ADVISOR ANALYSIS / RECOMMENDATIONS
Add context, recommendations, chapter-facing explanations, goals, and next steps here.
"""
    report+="\n\nPrepared from NBS Philo Hub officer, committee, finance, event, service, Historian, governance, and reporting records."
    return report


TEST_WORDS=['test','testing','demo','sample','dummy']

def is_testing_record(record):
    if not record:return False
    fields=[
        record.get('title'),record.get('report_title'),record.get('name'),
        record.get('description'),record.get('form_type'),record.get('event_name'),
        record.get('request_type')
    ]
    text=" ".join(str(x or '') for x in fields).casefold()
    return any(re.search(rf'\b{re.escape(word)}\b',text) for word in TEST_WORDS)

def visible_non_test_rows(rows_):
    return [r for r in (rows_ or []) if not is_testing_record(r)]

def member_payment_links():
    return {
        'card':setting('member_card_payment_url','').strip(),
        'paypal':setting('member_paypal_payment_url','').strip()
    }

def render_member_financial_status(mid,compact=False,key_prefix='financial'):
    if not is_philo_member(mid):
        return
    fy=setting('current_fiscal_year','2026-2027')
    fsum=member_financial_summary(mid,fy)
    st.markdown('### My Financial Status')
    with st.container(border=True):
        if compact:
            c1,c2,c3=st.columns(3)
            c1.metric('Required',f"${fsum['required']:,.2f}")
            c2.metric('Paid',f"${fsum['paid']:,.2f}")
            c3.metric('Balance',f"${fsum['balance']:,.2f}")
        else:
            c1,c2,c3,c4=st.columns(4)
            c1.metric('Member Type',fsum['member_type'])
            c2.metric('Required',f"${fsum['required']:,.2f}")
            c3.metric('Paid',f"${fsum['paid']:,.2f}")
            c4.metric('Balance',f"${fsum['balance']:,.2f}")

        if fsum['status']=='Financial':
            st.success('Financial Status: Financial')
        else:
            st.warning(f"Financial Status: {fsum['status']}")

        if fsum['balance']>0:
            links=member_payment_links()
            st.markdown('#### Pay My Balance')
            p1,p2=st.columns(2)
            if links['card']:
                p1.link_button(
                    f"💳 Pay ${fsum['balance']:,.2f} by Card",
                    links['card'],
                    use_container_width=True
                )
            else:
                p1.button(
                    '💳 Pay by Card — Not Connected Yet',
                    disabled=True,
                    use_container_width=True,
                    key=f'{key_prefix}_card_disabled'
                )
            if links['paypal']:
                p2.link_button(
                    f"Pay ${fsum['balance']:,.2f} with PayPal",
                    links['paypal'],
                    use_container_width=True
                )
            else:
                p2.button(
                    'PayPal — Not Connected Yet',
                    disabled=True,
                    use_container_width=True,
                    key=f'{key_prefix}_paypal_disabled'
                )
            st.caption(
                'Payments open in the chapter’s secure hosted payment page. '
                'Card or PayPal account information is not stored in the Philo Hub.'
            )
        return fsum

def cleanup_testing_records():
    candidates=[
        ('published_reports',['title','report_text','file_name']),
        ('monthly_reports',['report_title','report_text','position']),
        ('communications_requests',['title','details','request_type']),
        ('historian_photos',['event_name','caption']),
        ('recommendations',['title','recommendation_text']),
        ('events',['title','description']),
    ]
    removed=[]
    for table_name,fields in candidates:
        try:
            rows_=table(table_name).select('*').execute().data or []
            for r in rows_:
                text=" ".join(str(r.get(f) or '') for f in fields).casefold()
                if any(re.search(rf'\b{re.escape(w)}\b',text) for w in TEST_WORDS):
                    table(table_name).delete().eq('id',r['id']).execute()
                    removed.append(f"{table_name} #{r['id']}")
        except:
            continue
    return removed


def embeddable_governance_url(url):
    raw=str(url or '').strip()
    if not raw:
        return ''
    try:
        m=re.search(r'drive\.google\.com/file/d/([^/]+)',raw)
        if m:
            return f"https://drive.google.com/file/d/{m.group(1)}/preview"

        m=re.search(r'docs\.google\.com/(document|spreadsheets|presentation)/d/([^/]+)',raw)
        if m:
            kind=m.group(1)
            doc_id=m.group(2)
            if kind=='document':
                return f"https://docs.google.com/document/d/{doc_id}/preview"
            if kind=='spreadsheets':
                return f"https://docs.google.com/spreadsheets/d/{doc_id}/preview"
            return f"https://docs.google.com/presentation/d/{doc_id}/preview"

        return raw
    except:
        return raw

def show_governance_url_in_app(url,height=850):
    preview=embeddable_governance_url(url)
    if not preview:
        st.warning('This governing document has not been added yet.')
        return
    iframe_html=(
        f'<div style="width:100%;height:{height}px;border:1px solid #d8cfb8;'
        f'border-radius:12px;overflow:hidden;background:#fff;">'
        f'<iframe src="{html.escape(preview,quote=True)}" '
        f'style="width:100%;height:100%;border:0;" allow="fullscreen" loading="lazy"></iframe>'
        f'</div>'
    )
    components.html(
        iframe_html,
        height=height+10,
        scrolling=False
    )


# ============================================================
# v3.6.18 — Interest & Intake Center
# ============================================================

INTEREST_STATUSES=[
    'Interest Submitted',
    'Under Officer Review',
    'Meet & Greet / Conversation',
    'Selected for Member Vote',
    'Approved to Begin Intake',
    'Not Moving Forward',
    'Intake in Progress',
    'Intake Complete',
    'Archived'
]

INTAKE_TASK_STATUSES=['Not Started','In Progress','Submitted','Complete','Needs Attention']

RECRUITMENT_STARTER_LEADS=[
    {
        'school':'Queensborough Community College',
        'office':'Office of Student Activities',
        'contact':'Gisela Rivera / Raymond Volel',
        'email':'grivera@qcc.cuny.edu / rvolel@qcc.cuny.edu',
        'phone':'718-631-6233',
        'location':'Bayside, Queens',
        'approach':'Ask about tabling/special-event scheduling, student organization fairs, heritage-month programming, and approved off-campus/community resources.'
    },
    {
        'school':'LaGuardia Community College',
        'office':'Office of Campus Life',
        'contact':'Campus Life',
        'email':'CampusLife@lagcc.cuny.edu',
        'phone':'718-482-5190',
        'location':'Long Island City, Queens',
        'approach':'Ask about approved community organizations participating in Campus Life events, resource fairs, or tabling opportunities.'
    },
    {
        'school':'Kingsborough Community College',
        'office':'Office of Student Life',
        'contact':'Melissa Merced, Director / Student Life Team',
        'email':'KCC_StudentLife@kbcc.cuny.edu',
        'phone':'718-368-5597',
        'location':'Brooklyn',
        'approach':'Ask about community-partner tabling, student-life events, service/leadership programming, and women-centered campus outreach.'
    },
    {
        'school':'Nassau Community College',
        'office':'Student Activities Office',
        'contact':'Student Activities',
        'email':'studentactivities@ncc.edu',
        'phone':'516-572-7148',
        'location':'Garden City, Nassau County',
        'approach':'Ask about outside/community organization tabling, cultural programs, club fairs, campus calendar opportunities, and event procedures.'
    }
]

def interest_candidate_by_email(email):
    try:
        rows_=table('interest_candidates').select('*').ilike('email',str(email or '').strip()).execute().data or []
        return rows_[0] if rows_ else None
    except:
        return None

def interest_pin_digest(candidate_id,pin):
    salt=hashlib.sha256(f"interest:{candidate_id}:{_pin_pepper().hex()}".encode('utf-8')).digest()[:16]
    digest=hashlib.pbkdf2_hmac('sha256',str(pin).encode('utf-8'),salt,200000).hex()
    return 'pbkdf2_sha256$200000$'+digest

def set_interest_pin(candidate_id,pin):
    if not(str(pin).isdigit() and len(str(pin))==4):
        raise ValueError('PIN must be exactly 4 digits.')
    table('interest_candidates').update({
        'pin_hash':interest_pin_digest(candidate_id,pin),
        'updated_at':datetime.now(timezone.utc).isoformat()
    }).eq('id',candidate_id).execute()

def check_interest_pin(candidate,pin):
    stored=str(candidate.get('pin_hash') or '')
    return bool(stored) and hmac.compare_digest(stored,interest_pin_digest(candidate['id'],pin))

def interest_bio_text(c):
    first=c.get('first_name') or 'She'
    occupation=c.get('occupation') or ''
    education=c.get('education') or ''
    community=c.get('community_involvement') or ''
    interests=c.get('interests') or ''
    why=c.get('why_philo') or ''
    strengths=c.get('strengths') or ''
    parts=[f"{first} {c.get('last_name') or ''}".strip()]
    intro=[]
    if occupation:intro.append(f"works in {occupation}")
    if education:intro.append(f"has a background in {education}")
    if intro:parts.append("She "+ " and ".join(intro)+".")
    if community:parts.append(f"Her community involvement includes {community.rstrip('.')}.")
    if interests:parts.append(f"Her interests include {interests.rstrip('.')}.")
    if strengths:parts.append(f"She describes her strengths as {strengths.rstrip('.')}.")
    if why:parts.append(f"She is interested in becoming a Philo because {why.rstrip('.')}.")
    return " ".join(parts)

def candidate_public_status(candidate):
    return candidate.get('status') or 'Interest Submitted'

def candidate_photo_url(candidate):
    path=candidate.get('photo_path')
    return signed_url(path,3600) if path else ''

def candidate_tasks(candidate_id):
    try:
        return table('intake_candidate_tasks').select('*').eq('candidate_id',candidate_id).order('due_date').execute().data or []
    except:return []

def candidate_documents(candidate_id):
    try:
        return table('intake_candidate_documents').select('*').eq('candidate_id',candidate_id).order('uploaded_at',desc=True).execute().data or []
    except:return []

def intake_candidate_pool():
    try:
        return table('interest_candidates').select('*').order('created_at',desc=True).execute().data or []
    except:return []

def active_officer_positions_for_member(mid):
    return [x.get('position') for x in member_offices(mid)]

def can_manage_intake(mid,is_admin_flag=False):
    if is_admin_flag:return True
    positions=active_officer_positions_for_member(mid)
    if 'Vice President' in positions or 'President' in positions:return True
    try:
        chaired=chaired_committees(mid)
        return any('membership' in str(c.get('committee_name') or c.get('name') or c).lower() for c in chaired)
    except:return False

def can_view_interest_pool(mid,is_admin_flag=False):
    return is_admin_flag or bool(active_officer_positions_for_member(mid))

def candidate_vote_window():
    enabled=str(setting('intake_vote_open','false')).strip().lower() in {'true','1','yes','on'}
    start_raw=setting('intake_vote_start','').strip()
    end_raw=setting('intake_vote_end','').strip()
    now=datetime.now().astimezone()
    def parse_local(v):
        if not v:return None
        try:
            dt=datetime.fromisoformat(v)
            if dt.tzinfo is None:
                dt=dt.replace(tzinfo=now.tzinfo)
            return dt.astimezone(now.tzinfo)
        except:return None
    start=parse_local(start_raw);end=parse_local(end_raw)
    active=enabled and (start is None or now>=start) and (end is None or now<=end)
    return {'enabled':enabled,'start':start,'end':end,'now':now,'active':active}

def candidate_vote_open():
    return candidate_vote_window()['active']

def fmt_vote_time(dt):
    return dt.strftime('%m/%d/%Y at %I:%M %p') if dt else 'Not set'

def finalist_candidates():
    try:
        return table('interest_candidates').select('*').eq('finalist',True).order('last_name').execute().data or []
    except:return []

def member_candidate_vote(candidate_id,mid):
    try:
        rows_=table('intake_member_votes').select('*').eq('candidate_id',candidate_id).eq('member_id',mid).execute().data or []
        return rows_[0] if rows_ else None
    except:return None

def elevator_speech_text(what_philo,why_joined,chapter_work,ideal_person,personal_touch):
    return (
        f"Philos are {what_philo.strip().rstrip('.')}. "
        f"I became involved because {why_joined.strip().rstrip('.')}. "
        f"Through our affiliate, we {chapter_work.strip().rstrip('.')}. "
        f"If you are a woman who {ideal_person.strip().rstrip('.')}, this may be a sisterhood and service opportunity worth learning more about. "
        f"{personal_touch.strip()}".strip()
    )

def render_interest_candidate_portal():
    st.markdown(f"""<div class='philo-header'><img class='brand-logo' src='{PHILO_LOGO_URI}'><div><h1>Philo Interest Portal</h1><p>Learn, connect, and complete your interest profile.</p><div class='gold-rule'></div></div><img class='brand-crest' src='{PHILO_CREST_URI}'></div>""",unsafe_allow_html=True)

    if not st.session_state.get('interest_candidate_id'):
        mode=st.radio('Choose an option',['I already have an Interest Profile','Create My Interest Profile'],horizontal=True,key='interest_mode_choice')
        if mode=='I already have an Interest Profile':
            with st.form('interest_login'):
                email=st.text_input('Email Address')
                pin=st.text_input('4-digit PIN',type='password',max_chars=4)
                go=st.form_submit_button('Enter Interest Portal',use_container_width=True)
            if go:
                c=interest_candidate_by_email(email)
                if not c or not check_interest_pin(c,pin):
                    st.error('Email or PIN did not match an Interest Profile.')
                else:
                    st.session_state['interest_candidate_id']=c['id'];st.rerun()
        else:
            st.info('Creating an Interest Profile does not guarantee selection or membership. It begins the affiliate’s interest-review process.')
            with st.form('interest_register',clear_on_submit=False):
                a,b=st.columns(2)
                first=a.text_input('First Name')
                last=b.text_input('Last Name')
                email=a.text_input('Email Address')
                phone=b.text_input('Phone Number')
                pin=a.text_input('Create a 4-digit PIN',type='password',max_chars=4)
                pin2=b.text_input('Confirm PIN',type='password',max_chars=4)
                consent=st.checkbox('I understand this information will be reviewed by authorized Nu Beta Sigma Philo Affiliate leadership.')
                create=st.form_submit_button('Create Interest Profile',use_container_width=True)
            if create:
                if not(first.strip() and last.strip() and email.strip()):
                    st.warning('First name, last name, and email are required.')
                elif interest_candidate_by_email(email):
                    st.warning('An Interest Profile already exists for this email. Choose the login option instead.')
                elif not(pin.isdigit() and len(pin)==4 and pin==pin2):
                    st.warning('Create matching 4-digit PINs.')
                elif not consent:
                    st.warning('Please acknowledge the profile-review notice.')
                else:
                    result=table('interest_candidates').insert({
                        'first_name':first.strip(),'last_name':last.strip(),'email':email.strip().lower(),
                        'phone':phone.strip(),'status':'Interest Submitted','finalist':False,
                        'created_at':datetime.now(timezone.utc).isoformat(),'updated_at':datetime.now(timezone.utc).isoformat()
                    }).execute()
                    cid=(result.data or [{}])[0].get('id')
                    if cid:
                        set_interest_pin(cid,pin);st.session_state['interest_candidate_id']=cid;st.rerun()
        if st.button('← Return to Member Login',use_container_width=True,key='interest_return_member'):
            st.session_state.pop('interest_portal_mode',None);st.rerun()
        st.stop()

    cid=st.session_state['interest_candidate_id']
    rows_=table('interest_candidates').select('*').eq('id',cid).execute().data or []
    if not rows_:
        st.session_state.pop('interest_candidate_id',None);st.rerun()
    c=rows_[0]
    top1,top2=st.columns([5,1])
    top1.markdown(f"## Welcome, {c.get('first_name')}")
    if top2.button('Sign Out',use_container_width=True,key='interest_logout'):
        st.session_state.pop('interest_candidate_id',None);st.rerun()

    st.info(f"**Current Status:** {candidate_public_status(c)}")
    if c.get('public_next_step'):
        st.success(f"**What happens next:** {c.get('public_next_step')}")

    tabs=st.tabs(['My Profile & Bio','My Journey','Tasks','Documents','Dues & What to Expect'])
    with tabs[0]:
        photo_url=candidate_photo_url(c)
        if photo_url:st.image(photo_url,width=160)
        photo=st.file_uploader('Upload / Change Profile Photo',type=['png','jpg','jpeg'],key='interest_photo')
        with st.form('interest_profile_edit'):
            a,b=st.columns(2)
            first=a.text_input('First Name',value=c.get('first_name') or '')
            last=b.text_input('Last Name',value=c.get('last_name') or '')
            phone=a.text_input('Phone',value=c.get('phone') or '')
            location=b.text_input('City / Borough',value=c.get('location') or '')
            occupation=st.text_input('Occupation / Professional Area',value=c.get('occupation') or '')
            education=st.text_input('Education / Training',value=c.get('education') or '')
            community=st.text_area('Community Involvement / Service',value=c.get('community_involvement') or '')
            interests=st.text_area('Interests / Causes You Care About',value=c.get('interests') or '')
            strengths=st.text_area('Strengths You Would Bring',value=c.get('strengths') or '')
            why=st.text_area('Why are you interested in becoming a Philo?',value=c.get('why_philo') or '')
            availability=st.text_area('Availability / Scheduling Notes',value=c.get('availability') or '')
            save=st.form_submit_button('Save Profile & Generate Bio',use_container_width=True)
        if save:
            payload={'first_name':first.strip(),'last_name':last.strip(),'phone':phone.strip(),'location':location.strip(),
                     'occupation':occupation.strip(),'education':education.strip(),'community_involvement':community.strip(),
                     'interests':interests.strip(),'strengths':strengths.strip(),'why_philo':why.strip(),'availability':availability.strip()}
            payload['generated_bio']=interest_bio_text(payload);payload['updated_at']=datetime.now(timezone.utc).isoformat()
            table('interest_candidates').update(payload).eq('id',cid).execute();st.success('Profile and bio updated.');st.rerun()
        if photo:
            safe=re.sub(r'[^A-Za-z0-9._-]+','_',photo.name)
            path=upload_private(bytes(photo.getbuffer()),f"intake/candidates/{cid}/profile_{safe}",photo.type or 'image/jpeg')
            if path:table('interest_candidates').update({'photo_path':path}).eq('id',cid).execute();st.success('Photo saved.');st.rerun()
        st.markdown('### My Bio')
        st.write(c.get('generated_bio') or 'Complete the questionnaire above to generate your bio.')

    with tabs[1]:
        st.markdown('### Your Philo Interest Journey')
        steps=[
            ('1','Interest Profile','Submit your interest profile and introductory information.'),
            ('2','Leadership Review','Authorized officers review the interest pool.'),
            ('3','Conversation / Meet & Greet','You may be invited to learn more and ask questions.'),
            ('4','Affiliate Decision Process','Qualified finalists may move through the affiliate’s internal decision process.'),
            ('5','Intake Preparation','If approved, you receive authorized next steps, deadlines, documents, and financial information.'),
            ('6','Formal Intake','Ceremony-specific information is provided only by authorized leadership at the appropriate time.')
        ]
        for num,title,body in steps:st.markdown(f"**{num}. {title}** — {body}")
        st.caption('National, regional, and local governing requirements control the official process. This portal intentionally does not disclose ceremony details.')

    with tabs[2]:
        tasks_=candidate_tasks(cid)
        if not tasks_:st.info('No intake tasks have been assigned yet.')
        for task in tasks_:
            with st.container(border=True):
                st.markdown(f"**{task.get('title')}**")
                st.caption(f"Due: {task.get('due_date') or 'No date'}")
                if task.get('instructions'):st.write(task.get('instructions'))
                status=st.selectbox('Status',INTAKE_TASK_STATUSES,index=INTAKE_TASK_STATUSES.index(task.get('status')) if task.get('status') in INTAKE_TASK_STATUSES else 0,key=f"cand_task_{task['id']}")
                note=st.text_area('Update / Note',value=task.get('candidate_note') or '',key=f"cand_task_note_{task['id']}")
                if st.button('Save Task Update',key=f"cand_task_save_{task['id']}",use_container_width=True):
                    table('intake_candidate_tasks').update({'status':status,'candidate_note':note.strip(),'updated_at':datetime.now(timezone.utc).isoformat()}).eq('id',task['id']).execute();st.rerun()

    with tabs[3]:
        with st.form('candidate_document_upload',clear_on_submit=True):
            dtype=st.selectbox('Document Type',['Requested Intake Document','Identification / Verification','Application / Questionnaire','Payment / Receipt','Other'])
            doc=st.file_uploader('Upload Document',type=['pdf','doc','docx','png','jpg','jpeg'])
            note=st.text_input('Document Note')
            up=st.form_submit_button('Upload Document',use_container_width=True)
        if up and doc:
            safe=re.sub(r'[^A-Za-z0-9._-]+','_',doc.name)
            path=upload_private(bytes(doc.getbuffer()),f"intake/candidates/{cid}/documents/{datetime.now().strftime('%Y%m%d%H%M%S')}_{safe}",doc.type or 'application/octet-stream')
            if path:
                table('intake_candidate_documents').insert({'candidate_id':cid,'document_type':dtype,'file_path':path,'file_name':doc.name,'candidate_note':note.strip(),'status':'Submitted','uploaded_at':datetime.now(timezone.utc).isoformat()}).execute();st.success('Document uploaded.');st.rerun()
        docs=candidate_documents(cid)
        if docs:
            st.dataframe(pd.DataFrame([{'Document':d.get('document_type'),'File':d.get('file_name'),'Status':d.get('status'),'Reviewer Note':d.get('reviewer_note') or ''} for d in docs]),hide_index=True,use_container_width=True)

    with tabs[4]:
        dues=setting('intake_initial_dues','').strip()
        if dues:st.warning(f"**Current initial dues estimate/information:** {dues}")
        else:st.info('Initial dues information will be provided by authorized affiliate leadership before payment is required.')
        msg=setting('intake_public_expectations','').strip()
        if msg:st.write(msg)
        else:
            st.write('If you are selected to continue, leadership will explain deadlines, required documents, financial obligations, meetings, and the next authorized step. Ceremony-specific information is not released through the general interest portal.')

    st.stop()

def render_interest_intake_center(member_id,is_admin):
    st.title('Interest & Intake Center')
    officer_view=can_view_interest_pool(member_id,is_admin)
    manage=can_manage_intake(member_id,is_admin)
    finalists=finalist_candidates()

    if not officer_view:
        st.subheader('Final Candidate Vote')
        window=candidate_vote_window()
        if window['start'] or window['end']:
            st.caption(f"Voting window: {fmt_vote_time(window['start'])} through {fmt_vote_time(window['end'])}")
        if not window['active']:
            if window['enabled'] and window['start'] and window['now']<window['start']:
                st.info(f"Candidate voting is scheduled to open {fmt_vote_time(window['start'])}.")
            elif window['end'] and window['now']>window['end']:
                st.info(f"Candidate voting closed {fmt_vote_time(window['end'])}.")
            else:
                st.info('There is no candidate vote open at this time.')
            return
        if not finalists:
            st.info('No candidates are currently in the final voting pool.')
            return
        for c in finalists:
            with st.container(border=True):
                purl=candidate_photo_url(c)
                if purl:st.image(purl,width=120)
                st.markdown(f"### {c.get('first_name')} {c.get('last_name')}")
                st.write(c.get('generated_bio') or 'Candidate bio is being prepared.')
                prior=member_candidate_vote(c['id'],member_id)
                choice=st.radio('Should this candidate move forward into intake?',['Move Forward','Do Not Move Forward'],index=0 if not prior or prior.get('vote')=='Move Forward' else 1,key=f"candidate_vote_{c['id']}")
                if st.button('Submit / Update My Vote',key=f"candidate_vote_save_{c['id']}",use_container_width=True):
                    table('intake_member_votes').upsert({'candidate_id':c['id'],'member_id':member_id,'member_name':member_name,'vote':choice,'voted_at':datetime.now(timezone.utc).isoformat()},on_conflict='candidate_id,member_id').execute();st.success('Vote saved.');st.rerun()
        return

    tabs=st.tabs(['Interest Pool','Reviews & Selection','Member Vote','Candidate Tasks & Documents','Intake Guide'])
    pool=intake_candidate_pool()

    with tabs[0]:
        st.caption('All officers may view the Interest Pool. Private review notes are limited to authorized intake leadership.')
        if not pool:st.info('No interested candidates have created profiles yet.')
        for c in pool:
            with st.expander(f"{c.get('first_name')} {c.get('last_name')} • {c.get('status')}"):
                purl=candidate_photo_url(c)
                if purl:st.image(purl,width=140)
                st.write(c.get('generated_bio') or 'Bio not yet completed.')
                st.write(f"**Email:** {c.get('email')}  |  **Phone:** {c.get('phone') or ''}")
                st.write(f"**Location:** {c.get('location') or ''}")
                st.write(f"**Why interested:** {c.get('why_philo') or ''}")
                st.write(f"**Community involvement:** {c.get('community_involvement') or ''}")

    with tabs[1]:
        if not manage:
            st.info('Review notes and finalist selection are managed by the Vice President/Membership leadership, President, or Advisor/Admin.')
        else:
            for c in pool:
                with st.expander(f"{c.get('first_name')} {c.get('last_name')} • Review"):
                    reviews=table('intake_candidate_reviews').select('*').eq('candidate_id',c['id']).order('created_at',desc=True).execute().data or []
                    if reviews:st.dataframe(pd.DataFrame([{'Reviewer':r.get('reviewer_name'),'Recommendation':r.get('recommendation'),'Notes':r.get('notes'),'Date':fmt_dt(r.get('created_at'))} for r in reviews]),hide_index=True,use_container_width=True)
                    rec=st.selectbox('Recommendation',['Continue Review','Recommend Move Forward','Do Not Recommend'],key=f"intake_rec_{c['id']}")
                    notes=st.text_area('Private Membership/Intake Notes',key=f"intake_notes_{c['id']}",placeholder='Document strengths, concerns, follow-up needed, or why the candidate should/should not move forward.')
                    if st.button('Save Review',key=f"save_intake_review_{c['id']}",use_container_width=True):
                        table('intake_candidate_reviews').insert({'candidate_id':c['id'],'reviewer_member_id':member_id,'reviewer_name':member_name,'recommendation':rec,'notes':notes.strip(),'created_at':datetime.now(timezone.utc).isoformat()}).execute();st.success('Review saved.');st.rerun()
                    finalist=st.checkbox('Place in Final Member Voting Pool',value=bool(c.get('finalist')),key=f"finalist_{c['id']}")
                    status=st.selectbox('Candidate Status',INTEREST_STATUSES,index=INTEREST_STATUSES.index(c.get('status')) if c.get('status') in INTEREST_STATUSES else 0,key=f"cand_status_{c['id']}")
                    public_next=st.text_area('Candidate-Facing Next Step Message',value=c.get('public_next_step') or '',key=f"cand_next_{c['id']}")
                    if st.button('Save Selection / Status',key=f"save_cand_status_{c['id']}",use_container_width=True):
                        table('interest_candidates').update({'finalist':finalist,'status':status,'public_next_step':public_next.strip(),'updated_at':datetime.now(timezone.utc).isoformat()}).eq('id',c['id']).execute();st.success('Candidate status updated.');st.rerun()

    with tabs[2]:
        window=candidate_vote_window()
        if manage:
            st.markdown('### Candidate Voting Control')
            st.caption('The Vice President/Membership leadership can schedule the exact voting period. Times display in 12-hour AM/PM format.')
            default_start=window['start'] or datetime.now().astimezone()
            default_end=window['end'] or (datetime.now().astimezone()+timedelta(days=3))
            c1,c2=st.columns(2)
            vote_start_date=c1.date_input('Voting Opens — Date',value=default_start.date(),key='vote_start_date')
            vote_start_time=c1.time_input('Voting Opens — Time',value=default_start.time().replace(second=0,microsecond=0),key='vote_start_time')
            vote_end_date=c2.date_input('Voting Closes — Date',value=default_end.date(),key='vote_end_date')
            vote_end_time=c2.time_input('Voting Closes — Time',value=default_end.time().replace(second=0,microsecond=0),key='vote_end_time')
            duration_choice=st.selectbox('Quick Voting Time Frame',[
                'Use dates/times above','24 Hours','48 Hours','72 Hours','5 Days','7 Days'
            ],index=0,key='candidate_vote_duration')
            if duration_choice!='Use dates/times above':
                hours={'24 Hours':24,'48 Hours':48,'72 Hours':72,'5 Days':120,'7 Days':168}[duration_choice]
                st.caption(f"If activated with this option, voting will close automatically {duration_choice.lower()} after activation.")
            b1,b2=st.columns(2)
            if b1.button('🗳️ Activate / Schedule Candidate Voting',use_container_width=True,key='activate_candidate_vote'):
                start_dt=datetime.combine(vote_start_date,vote_start_time).astimezone()
                if duration_choice!='Use dates/times above':
                    # If selected start is already in the past, activation begins now.
                    now_=datetime.now().astimezone()
                    if start_dt<now_:start_dt=now_
                    hours={'24 Hours':24,'48 Hours':48,'72 Hours':72,'5 Days':120,'7 Days':168}[duration_choice]
                    end_dt=start_dt+timedelta(hours=hours)
                else:
                    end_dt=datetime.combine(vote_end_date,vote_end_time).astimezone()
                if end_dt<=start_dt:
                    st.error('Voting close time must be after the opening time.')
                elif not finalist_candidates():
                    st.error('Select at least one finalist before activating voting.')
                else:
                    save_setting('intake_vote_start',start_dt.isoformat())
                    save_setting('intake_vote_end',end_dt.isoformat())
                    save_setting('intake_vote_open','true')
                    st.success(f"Candidate voting scheduled from {fmt_vote_time(start_dt)} through {fmt_vote_time(end_dt)}.")
                    st.rerun()
            if b2.button('🔒 Close Voting Now',use_container_width=True,key='close_candidate_vote'):
                save_setting('intake_vote_open','false')
                save_setting('intake_vote_end',datetime.now().astimezone().isoformat())
                st.success('Candidate voting is closed.')
                st.rerun()

            window=candidate_vote_window()
            if window['active']:
                st.success(f"Voting is OPEN now and closes {fmt_vote_time(window['end'])}.")
            elif window['enabled'] and window['start'] and window['now']<window['start']:
                st.info(f"Voting is SCHEDULED to open {fmt_vote_time(window['start'])} and close {fmt_vote_time(window['end'])}.")
            elif window['end']:
                st.info(f"Voting is CLOSED. Last close time: {fmt_vote_time(window['end'])}.")
            else:
                st.info('Voting has not been scheduled.')

        st.markdown('### Candidate Voting Results')
        if finalists:
            for c in finalists:
                votes=table('intake_member_votes').select('*').eq('candidate_id',c['id']).execute().data or []
                yes=sum(1 for v in votes if v.get('vote')=='Move Forward')
                no=sum(1 for v in votes if v.get('vote')=='Do Not Move Forward')
                total=len(votes)
                eligible=len(philo_members())
                pct=(yes/total*100) if total else 0
                with st.container(border=True):
                    st.markdown(f"**{c.get('first_name')} {c.get('last_name')}**")
                    r1,r2,r3,r4=st.columns(4)
                    r1.metric('Move Forward',yes)
                    r2.metric('Do Not Move Forward',no)
                    r3.metric('Votes Cast',total)
                    r4.metric('Participation',f"{total}/{eligible}")
                    st.progress(min(1.0,pct/100) if total else 0.0,text=f"{pct:.0f}% of votes cast: Move Forward")
                    if manage:
                        show_names=st.checkbox('Show individual member votes',key=f"show_candidate_votes_{c['id']}")
                        if show_names and votes:
                            st.dataframe(pd.DataFrame([{
                                'Member':v.get('member_name'),
                                'Vote':v.get('vote'),
                                'Voted':fmt_dt(v.get('voted_at'))
                            } for v in votes]),hide_index=True,use_container_width=True)
        else:
            st.info('No candidates have been placed in the final voting pool.')

    with tabs[3]:
        if not manage:
            st.info('Candidate task/document management is limited to intake leadership.')
        else:
            cmap={c['id']:f"{c.get('first_name')} {c.get('last_name')}" for c in pool}
            if cmap:
                cid=st.selectbox('Candidate',list(cmap),format_func=lambda x:cmap[x],key='intake_manage_candidate')
                st.markdown('### Assign Task')
                with st.form('assign_intake_task',clear_on_submit=True):
                    title=st.text_input('Task')
                    instructions=st.text_area('Instructions')
                    due=st.date_input('Due Date',value=date.today()+timedelta(days=7))
                    add=st.form_submit_button('Assign Task',use_container_width=True)
                if add and title.strip():
                    table('intake_candidate_tasks').insert({'candidate_id':cid,'title':title.strip(),'instructions':instructions.strip(),'due_date':due.isoformat(),'status':'Not Started','assigned_by_member_id':member_id,'assigned_by_name':member_name,'created_at':datetime.now(timezone.utc).isoformat()}).execute();st.success('Task assigned.');st.rerun()
                tasks_=candidate_tasks(cid)
                if tasks_:st.dataframe(pd.DataFrame([{'Task':x.get('title'),'Due':x.get('due_date'),'Status':x.get('status'),'Candidate Note':x.get('candidate_note') or ''} for x in tasks_]),hide_index=True,use_container_width=True)
                st.markdown('### Candidate Documents')
                docs=candidate_documents(cid)
                for d in docs:
                    with st.expander(f"{d.get('document_type')} • {d.get('file_name')} • {d.get('status')}"):
                        url=signed_url(d.get('file_path'),1800) if d.get('file_path') else ''
                        if url:st.link_button('Open Document',url,use_container_width=True)
                        status=st.selectbox('Review Status',['Submitted','Accepted','Needs Replacement','Reviewed'],index=['Submitted','Accepted','Needs Replacement','Reviewed'].index(d.get('status')) if d.get('status') in ['Submitted','Accepted','Needs Replacement','Reviewed'] else 0,key=f"docstat_{d['id']}")
                        note=st.text_area('Reviewer Note',value=d.get('reviewer_note') or '',key=f"docnote_{d['id']}")
                        if st.button('Save Document Review',key=f"docsave_{d['id']}",use_container_width=True):
                            table('intake_candidate_documents').update({'status':status,'reviewer_note':note.strip(),'reviewed_by_member_id':member_id,'reviewed_at':datetime.now(timezone.utc).isoformat()}).eq('id',d['id']).execute();st.rerun()

    with tabs[4]:
        st.markdown('### Intake Process Guide')
        st.markdown("""
1. **Interest Pool** — Interested women create a profile, questionnaire-based bio, and photo.
2. **Officer Review** — Officers can view the pool; authorized Membership/Intake leadership records private evaluation notes.
3. **Conversation / Meet & Greet** — Leadership may invite women to learn more, ask questions, and understand expectations.
4. **Finalist Selection** — The VP/Membership leadership selects the final pool for affiliate consideration.
5. **Member Vote** — Financial/eligible members vote on whether each finalist should move forward.
6. **Notification** — Leadership communicates the decision privately and provides the next authorized step.
7. **Intake Preparation** — Approved candidates receive deadlines, requested documents, initial financial obligations, meetings, and portal tasks.
8. **Formal Intake** — Ceremony or ritual details are provided only by authorized leadership at the appropriate stage; they are not placed in the general Interest Portal.
9. **Completion / Transition** — After completion, leadership updates the candidate record and creates the appropriate member account.
        """)
        st.warning('Use National, Northeastern Region, and local governing documents as the controlling authority. The Hub guide is administrative workflow support, not a replacement for official intake rules.')
        if manage:
            dues=st.text_input('Initial Dues / Financial Information Shown to Approved Candidates',value=setting('intake_initial_dues',''),key='intake_dues_setting')
            expectations=st.text_area('Candidate-Facing “What to Expect Next” Message',value=setting('intake_public_expectations',''),height=180,key='intake_expectations_setting')
            if st.button('Save Candidate Guidance',use_container_width=True):
                save_setting('intake_initial_dues',dues.strip());save_setting('intake_public_expectations',expectations.strip());st.success('Candidate guidance saved.')

def render_elevator_speech_builder(mid,name):
    st.markdown('### 🎤 My Philo Elevator Speech')
    with st.expander('Build / Update My Talking Points',expanded=False):
        existing=[]
        try:existing=table('member_elevator_speeches').select('*').eq('member_id',mid).execute().data or []
        except:pass
        e=existing[0] if existing else {}
        what=st.text_input('In your own words, what are Philos?',value=e.get('what_philo') or 'women connected by sisterhood, service, and support of Sigma Gamma Rho Sorority, Incorporated and our communities',key='elev_what')
        why=st.text_input('Why did you become involved?',value=e.get('why_joined') or '',key='elev_why')
        work=st.text_input('What does your affiliate do?',value=e.get('chapter_work') or 'serve our community, support programs, build relationships, and create opportunities for women to get involved',key='elev_work')
        ideal=st.text_input('What type of woman might enjoy being a Philo?',value=e.get('ideal_person') or 'enjoys service, sisterhood, community involvement, and meeting women with similar values',key='elev_ideal')
        personal=st.text_input('Optional personal closing line',value=e.get('personal_touch') or 'I would be happy to tell you more about what we do and invite you to an upcoming interest event.',key='elev_personal')
        if st.button('Generate My Elevator Speech',use_container_width=True,key='elev_generate'):
            speech=elevator_speech_text(what,why,work,ideal,personal)
            st.session_state['elevator_speech']=speech
            table('member_elevator_speeches').upsert({'member_id':mid,'member_name':name,'what_philo':what,'why_joined':why,'chapter_work':work,'ideal_person':ideal,'personal_touch':personal,'speech_text':speech,'updated_at':datetime.now(timezone.utc).isoformat()},on_conflict='member_id').execute()
        speech=st.text_area('My Saved Speech',value=st.session_state.get('elevator_speech',e.get('speech_text') or ''),height=180,key='elev_speech_text')
        if speech.strip() and st.button('Save Edited Speech',use_container_width=True,key='elev_save_edit'):
            table('member_elevator_speeches').upsert({'member_id':mid,'member_name':name,'what_philo':what,'why_joined':why,'chapter_work':work,'ideal_person':ideal,'personal_touch':personal,'speech_text':speech.strip(),'updated_at':datetime.now(timezone.utc).isoformat()},on_conflict='member_id').execute();st.success('Elevator speech saved.')


def candidate_test_profile():
    return st.session_state.setdefault('candidate_test_profile',{
        'first_name':'Test',
        'last_name':'Candidate',
        'email':'test.candidate@example.com',
        'phone':'(555) 555-0192',
        'location':'Queens, NY',
        'occupation':'Community Professional',
        'education':'College / Professional Training',
        'community_involvement':'Volunteers with local community programs and enjoys service projects.',
        'interests':'Community service, professional growth, mentorship, and sisterhood.',
        'strengths':'Dependable, organized, friendly, and willing to serve.',
        'why_philo':'she wants to connect with service-minded women and make a positive impact in the community.',
        'availability':'Evenings and weekends',
        'generated_bio':'Test Candidate is a community-minded professional who values service, personal growth, and meaningful connections. She enjoys volunteering, supporting community programs, and bringing dependable organization to group efforts. She is interested in becoming a Philo to connect with service-minded women and make a positive impact in the community.',
        'status':'Interest Submitted',
        'public_next_step':'Your interest profile has been received. Leadership will contact you if you are selected for the next step.'
    })

def render_candidate_test_preview():
    c=candidate_test_profile()
    st.markdown("## 🧪 Candidate Portal Preview")
    st.warning('TEST MODE — You are viewing a simulated candidate. Nothing entered here is saved, uploaded, voted, sent, or counted.')
    st.info(f"**Current Status:** {c.get('status')}")
    if c.get('public_next_step'):st.success(f"**What happens next:** {c.get('public_next_step')}")
    tabs=st.tabs(['My Profile & Bio','My Journey','Tasks','Documents','Dues & What to Expect'])
    with tabs[0]:
        st.markdown('### Profile Photo')
        photo=st.file_uploader('Upload / Change Profile Photo',type=['png','jpg','jpeg'],key='preview_candidate_photo')
        if photo:st.success('Preview only: photo selected successfully. It was NOT uploaded.')
        a,b=st.columns(2)
        c['first_name']=a.text_input('First Name',value=c.get('first_name',''),key='preview_c_first')
        c['last_name']=b.text_input('Last Name',value=c.get('last_name',''),key='preview_c_last')
        c['phone']=a.text_input('Phone',value=c.get('phone',''),key='preview_c_phone')
        c['location']=b.text_input('City / Borough',value=c.get('location',''),key='preview_c_loc')
        c['occupation']=st.text_input('Occupation / Professional Area',value=c.get('occupation',''),key='preview_c_occ')
        c['education']=st.text_input('Education / Training',value=c.get('education',''),key='preview_c_edu')
        c['community_involvement']=st.text_area('Community Involvement / Service',value=c.get('community_involvement',''),key='preview_c_comm')
        c['interests']=st.text_area('Interests / Causes You Care About',value=c.get('interests',''),key='preview_c_int')
        c['strengths']=st.text_area('Strengths You Would Bring',value=c.get('strengths',''),key='preview_c_strength')
        c['why_philo']=st.text_area('Why are you interested in becoming a Philo?',value=c.get('why_philo',''),key='preview_c_why')
        if st.button('Preview Generated Bio',use_container_width=True,key='preview_generate_bio'):
            c['generated_bio']=interest_bio_text(c)
        st.markdown('### My Bio')
        st.write(c.get('generated_bio'))
    with tabs[1]:
        for title,body in [
            ('Interest Profile','Your profile and introductory information are received.'),
            ('Leadership Review','Authorized leadership privately reviews the interest pool.'),
            ('Conversation / Meet & Greet','Selected women may be invited to learn more and ask questions.'),
            ('Decision','You are only told whether you are selected to continue; internal voting/review details are not shown.'),
            ('Intake Preparation','If selected, additional tasks, documents, deadlines, and approved financial information appear here.'),
            ('Formal Intake','Authorized leadership provides ceremony-specific information only at the appropriate time.')
        ]:st.markdown(f"**{title}** — {body}")
    with tabs[2]:
        demo_tasks=st.session_state.setdefault('candidate_preview_tasks',[
            {'title':'Review Welcome Information','due':'09/05/2026','status':'Not Started'},
            {'title':'Complete Requested Intake Questionnaire','due':'09/10/2026','status':'Not Started'}
        ])
        for i,task in enumerate(demo_tasks):
            with st.container(border=True):
                st.markdown(f"**{task['title']}**")
                st.caption(f"Due: {task['due']}")
                task['status']=st.selectbox('Status',INTAKE_TASK_STATUSES,index=INTAKE_TASK_STATUSES.index(task['status']) if task['status'] in INTAKE_TASK_STATUSES else 0,key=f"preview_task_{i}")
                st.text_area('Update / Note',key=f"preview_task_note_{i}")
        st.caption('Changes above exist only in this browser preview session.')
    with tabs[3]:
        dtype=st.selectbox('Document Type',['Requested Intake Document','Identification / Verification','Application / Questionnaire','Payment / Receipt','Other'],key='preview_doc_type')
        doc=st.file_uploader('Upload Document',type=['pdf','doc','docx','png','jpg','jpeg'],key='preview_candidate_doc')
        if doc:st.success(f"Preview only: {doc.name} selected. It was NOT uploaded.")
    with tabs[4]:
        dues=setting('intake_initial_dues','').strip()
        st.warning(f"**Initial dues information:** {dues or 'Leadership will provide the approved amount before payment is required.'}")
        st.write(setting('intake_public_expectations','').strip() or 'If selected to continue, leadership explains deadlines, requested documents, financial obligations, meetings, and the next authorized step.')
        st.caption('A candidate never sees officer recommendations, member voting, vote counts, or internal selection notes.')


ADVISOR_HELP_PRIORITIES=['Routine','Soon','Urgent']
GRIEVANCE_CATEGORIES=['Interpersonal Conflict','Officer / Role Concern','Meeting / Procedure Concern','Financial Concern','Harassment / Bullying / Disrespect','Discrimination / Bias Concern','Safety / Threat Concern','Privacy / Confidentiality Concern','Membership / Intake Concern','Other']
GRIEVANCE_STATUSES=['Submitted','Under Review','More Information Needed','Referred','Resolved','Closed']

def grievance_guidance(category,details=''):
    c=str(category or ''); d=str(details or '').lower()
    if c=='Financial Concern':
        guidance=['Preserve receipts, approvals, and transaction records.','Separate a possible error from suspected misconduct until records are reviewed.','Use the financial approval/reconciliation chain.'];roles=['Advisor','President','Treasurer','Financial Secretary']
    elif c=='Meeting / Procedure Concern':
        guidance=['Identify the exact procedural action being questioned.','Check local bylaws/SOPs first, then parliamentary authority.','Keep personal disagreements separate from procedural rulings.'];roles=['President','Parliamentarian','Advisor']
    elif c=='Officer / Role Concern':
        guidance=['Review officer duties and governing documents.','Clarify whether the issue is authority, performance, communication, or conduct.','Document expectations and corrective steps.'];roles=['Advisor','President','Parliamentarian if governing rules are involved']
    elif c in ['Harassment / Bullying / Disrespect','Discrimination / Bias Concern']:
        guidance=['Preserve dates, messages, witnesses, and specific conduct.','Limit unnecessary sharing during review.','Do not require informal mediation if the reporting member feels unsafe or intimidated.'];roles=['Advisor','President']
    elif c=='Safety / Threat Concern':
        guidance=['Prioritize immediate physical safety over internal process.','Preserve evidence and do not require confrontation.','If danger is immediate, use venue/security or emergency resources first.'];roles=['Advisor','President','Venue/Security or emergency services if there is immediate danger']
    elif c=='Privacy / Confidentiality Concern':
        guidance=['Identify what information was disclosed and to whom.','Restrict further access while the concern is reviewed.','Correct permissions/records if needed and document the response.'];roles=['Advisor','President','Relevant records officer']
    elif c=='Membership / Intake Concern':
        guidance=['Limit candidate/member information to authorized membership/intake leadership.','Review governing membership/intake rules before acting.','Do not disclose private deliberations, votes, or ceremony details.'];roles=['Advisor','Vice President / Membership Lead','President']
    elif c=='Interpersonal Conflict':
        guidance=['Document the concern factually.','Consider a private conversation or facilitated discussion if safe.','Set a clear expectation and follow-up date.'];roles=['Advisor','President']
    else:
        guidance=['Document the facts and requested outcome.','Identify the officer role with actual authority.','Use the least escalated appropriate process while preserving confidentiality.'];roles=['Advisor','President']
    if any(x in d for x in ['weapon','threat','assault','stalking','hurt me','kill']):
        if 'Venue/Security or emergency services if there is immediate danger' not in roles:roles.append('Venue/Security or emergency services if there is immediate danger')
        guidance.insert(0,'The description may involve immediate safety; do not rely only on the internal grievance process.')
    return {'guidance':guidance,'roles':roles,'summary':' '.join([f"{i+1}. {x}" for i,x in enumerate(guidance)])+' Suggested roles: '+', '.join(roles)+'.'}

def advisor_help_rows():
    try:return table('advisor_help_requests').select('*').order('created_at',desc=True).execute().data or []
    except:return []
def grievance_rows():
    try:return table('grievances').select('*').order('created_at',desc=True).execute().data or []
    except:return []
def event_rsvp_rows(event_id):
    try:return table('event_attendance_rsvps').select('*').eq('event_id',event_id).order('member_name').execute().data or []
    except:return []
def current_member_rsvp(event_id,mid):return next((r for r in event_rsvp_rows(event_id) if r.get('member_id')==mid),None)
def save_member_rsvp(event_id,mid,name,status):
    table('event_attendance_rsvps').upsert({'event_id':event_id,'member_id':mid,'member_name':name,'status':status,'updated_at':datetime.now(timezone.utc).isoformat()},on_conflict='event_id,member_id').execute()
def event_attendee_names(event_id):return [r.get('member_name') for r in event_rsvp_rows(event_id) if r.get('status')=='Attending']
def event_geo_settings(event_id):
    try:
        r=table('event_geo_settings').select('*').eq('event_id',event_id).execute().data or []
        return r[0] if r else None
    except:return None

FLYER_REQUIRED_CHECKS=[
    ("affiliate_name_present","Affiliate name present","Add the full Philo Affiliate name so viewers know who is hosting."),
    ("event_title_present","Event title is clear","Make the event title larger and easier to identify at a glance."),
    ("date_present","Date is present and readable","Add the complete event date in a prominent location."),
    ("time_present","Time is present when applicable","Add the event time or clearly mark the event as all-day when appropriate."),
    ("location_present","Location / access details are complete","Add the full venue/address or virtual access instructions."),
    ("cta_present","Call to action / contact information is included","Add a clear next step such as Register, RSVP, Donate, Attend, or Contact."),
    ("philo_colors_ok","Philo branding uses Gold & White appropriately","Use Philo Gold & White as the primary palette; black/gray may be neutral."),
    ("contrast_ok","Text/background contrast is readable","Increase contrast or simplify the background behind text."),
    ("font_count_ok","Typography is readable and not overly mixed","Use no more than about two font families and avoid script for body copy."),
    ("image_quality_ok","Images/logos are clear and not distorted","Replace blurry images and keep graphics proportional."),
    ("logo_integrity_ok","Approved marks/logos are not stretched or altered","Use approved marks at their correct proportions and do not cover registration marks."),
    ("copy_clear","Copy is concise and easy to scan","Shorten long paragraphs and keep only information needed for action."),
    ("proofread_ok","Spelling, grammar, dates, and details were proofread","Proofread names, dates, times, links, and event details before submission."),
    ("accessibility_ok","Accessibility considerations were checked","Confirm readable font size, contrast, plain language, and alt-text readiness."),
]

FLYER_REVIEW_ACCESS_ROLES={
    "President","Vice President","Recording Secretary","Financial Secretary",
    "Treasurer","Historian","Chaplain","Parliamentarian","Sergeant-at-Arms"
}

FLYER_STATUS_FLOW=[
    "Draft",
    "Ready for Branding Review",
    "Needs Changes",
    "Ready to Submit to Chapter",
    "Waiting Approval from Chapter",
    "Approved by Chapter",
    "Not Approved / Revise",
    "Published / Final",
    "Archived",
    "Cancelled",
]

def can_use_flyer_workflow(mid,is_admin_flag=False):
    return bool(is_admin_flag or (current_position_names(mid) & FLYER_REVIEW_ACCESS_ROLES))

def flyer_brand_review_summary(values):
    passed,total,pct=flyer_score_from_checks(values)
    missing=[label for key,label,_ in FLYER_REQUIRED_CHECKS if not values.get(key)]
    if not missing:
        overall="PASS — Ready to Submit to Chapter"
    elif len(missing)<=3:
        overall="NEEDS ATTENTION — Fix the remaining items"
    else:
        overall="MUST FIX BEFORE SUBMISSION"
    return passed,total,pct,overall,missing

def flyer_feedback(values):
    feedback=[]
    for key,label,fix in FLYER_REQUIRED_CHECKS:
        if values.get(key):
            feedback.append({"Result":"✓ PASS","Check":label,"Recommendation":"No change required."})
        else:
            feedback.append({"Result":"✕ FIX REQUIRED","Check":label,"Recommendation":fix})
    return feedback

def flyer_file_validate(uploaded):
    if uploaded is None:
        raise ValueError("Choose the completed flyer file.")
    raw=bytes(uploaded.getbuffer())
    if not raw:
        raise ValueError("The selected flyer file is empty.")
    if len(raw)>20*1024*1024:
        raise ValueError("Flyer file is too large. Maximum size is 20 MB.")
    ext=Path(uploaded.name or "").suffix.lower()
    if ext not in {".pdf",".png",".jpg",".jpeg"}:
        raise ValueError("Flyer upload supports PDF, PNG, JPG, or JPEG.")
    return raw,ext

def flyer_progress(status):
    order={
        "Draft":1,
        "Ready for Branding Review":2,
        "Needs Changes":3,
        "Ready to Submit to Chapter":4,
        "Waiting Approval from Chapter":5,
        "Approved by Chapter":6,
        "Not Approved / Revise":3,
        "Published / Final":6,
        "Archived":6,
        "Cancelled":0,
    }
    return order.get(status,1)

def flyer_rows(**filters):
    q=table("philo_flyer_submissions").select("*")
    for k,v in filters.items():
        q=q.eq(k,v)
    return q.order("updated_at",desc=True).execute().data or []

def flyer_score_from_checks(values):
    total=len(FLYER_REQUIRED_CHECKS)
    passed=sum(1 for key,_,_ in FLYER_REQUIRED_CHECKS if values.get(key))
    pct=round((passed/total)*100) if total else 0
    return passed,total,pct

def flyer_status_banner(status):
    icon={
        "Draft":"📝",
        "Ready for Branding Review":"🔎",
        "Needs Changes":"⚠️",
        "Ready to Submit to Chapter":"✅",
        "Waiting Approval from Chapter":"⏳",
        "Approved by Chapter":"✅",
        "Not Approved / Revise":"✏️",
        "Published / Final":"📣",
        "Archived":"🗂️",
        "Cancelled":"🚫",
    }.get(status,"ℹ️")
    st.markdown(f"### {icon} STATUS: {status.upper()}")

def historian_bulk_template_bytes():
    cols=[
        "sorority_year","date_known","event_date","record_type","title","location",
        "description","people_involved","verification_status","source_note",
        "has_photos","has_flyer_program","needs_identification","source_conflict"
    ]
    return pd.DataFrame(columns=cols).to_csv(index=False).encode("utf-8")

def historian_change_log(action,record_type,record_id,details,mid):
    try:
        table("historian_change_log").insert({
            "action":action,
            "record_type":record_type,
            "record_id":str(record_id or ""),
            "details":str(details or ""),
            "changed_by_member_id":mid,
            "changed_by_name":member_name,
            "created_at":datetime.now(timezone.utc).isoformat(),
        }).execute()
    except Exception:
        pass

def historian_delegate_rows(mid=None):
    try:
        q=table("historian_delegate_access").select("*").eq("active",True)
        if mid is not None:
            q=q.eq("member_id",mid)
        return q.execute().data or []
    except Exception:
        return []

def historian_event_rows(include_archived=False):
    q=table("historian_event_memory").select("*")
    if not include_archived:
        q=q.eq("is_archived",False)
    return q.order("event_date",desc=True).execute().data or []

def historian_file_validate(uploaded,max_mb=20):
    if uploaded is None:
        raise ValueError("Choose a source document.")
    raw=bytes(uploaded.getbuffer())
    if not raw:
        raise ValueError("The selected file is empty.")
    if len(raw)>max_mb*1024*1024:
        raise ValueError(f"File is too large. Maximum size is {max_mb} MB.")
    ext=Path(uploaded.name or "").suffix.lower()
    if ext not in {".pdf",".docx",".xlsx",".xls",".csv",".png",".jpg",".jpeg"}:
        raise ValueError("Supported source files: PDF, DOCX, XLSX/XLS, CSV, PNG, JPG/JPEG.")
    return raw,ext

def historian_preference(mid,key,default=False):
    try:
        rows_=table("historian_preferences").select("preference_value").eq(
            "member_id",mid
        ).eq("preference_key",key).execute().data or []
        if rows_:
            return str(rows_[0].get("preference_value") or "").casefold()=="true"
    except Exception:
        pass
    return bool(default)

def render_philo_flyer_approval_workflow(mid,is_admin_flag=False):
    if not can_use_flyer_workflow(mid,is_admin_flag):
        return

    if "flyer_help_guide" not in st.session_state:
        st.session_state["flyer_help_guide"]=True
    if "flyer_large_controls" not in st.session_state:
        st.session_state["flyer_large_controls"]=False

    st.markdown("## Philo Flyer Review & Chapter Approval")
    st.caption(
        "Upload the finished flyer, run the required branding check, fix any issues, "
        "then submit to Chapter approval. The Chapter button stays locked until the review passes."
    )

    hc1,hc2=st.columns(2)
    with hc1:
        st.toggle("Show Flyer Help Guide",key="flyer_help_guide")
    with hc2:
        st.toggle("Large / Easy Controls",key="flyer_large_controls")

    if st.session_state["flyer_large_controls"]:
        st.markdown(
            "<style>"
            "div.stButton > button {min-height:54px;font-size:18px;font-weight:700;}"
            "div[data-baseweb='input'] input, textarea {font-size:18px !important;}"
            "</style>",
            unsafe_allow_html=True
        )

    if st.session_state["flyer_help_guide"]:
        st.info(
            "**How this works:** 1 Upload/Save Draft → 2 Run Branding Review → "
            "3 Fix anything flagged → 4 Submit to Chapter → 5 Record the Chapter decision."
        )

    tabs=st.tabs([
        "1. Upload / Draft",
        "2. Branding Review",
        "3. Submit to Chapter",
        "4. Chapter Decision",
        "5. Notifications",
        "6. Archive / Versions",
    ])

    # -----------------------------------------------------
    # 1. Upload / Draft / Replace / Delete / Cancel
    # -----------------------------------------------------
    with tabs[0]:
        if st.session_state["flyer_help_guide"]:
            st.info(
                "Upload the finished design here. **Save Draft** stores it without sending anything. "
                "You can replace the flyer with a corrected version later."
            )

        with st.form("philo_flyer_submit_form",clear_on_submit=False):
            title=st.text_input("Flyer / Event Title",max_chars=180)
            committee=st.text_input("Officer / Committee Responsible")
            requested_by=st.text_input("Requested By",value=member_name)
            event_date=st.date_input("Event Date",value=date.today())
            event_time=st.text_input("Event Time",placeholder="Example: 2:00 PM or All Day")
            location=st.text_input("Location / Access Details")
            cta=st.text_input("CTA / Contact",placeholder="Example: RSVP by..., Register at..., Contact...")
            description=st.text_area("Purpose / Notes",max_chars=1500)
            flyer_file=st.file_uploader("Completed Flyer",type=["pdf","png","jpg","jpeg"],key="philo_flyer_upload")
            b1,b2=st.columns(2)
            save_draft=b1.form_submit_button("Save Draft",use_container_width=True)
            save_review=b2.form_submit_button("Save & Go to Branding Review",use_container_width=True)

        if save_draft or save_review:
            try:
                if not title.strip():
                    raise ValueError("Enter the flyer/event title.")
                if not committee.strip():
                    raise ValueError("Enter the responsible officer or committee.")
                raw,ext=flyer_file_validate(flyer_file)
                digest=hashlib.sha256(raw).hexdigest()
                dup=table("philo_flyer_submissions").select("id,title,status").eq("file_sha256",digest).execute().data or []
                if dup:
                    raise ValueError(
                        f"This exact flyer is already stored as '{dup[0].get('title')}' "
                        f"with status {dup[0].get('status')}."
                    )
                safe_name=re.sub(r"[^A-Za-z0-9._-]+","_",Path(flyer_file.name or f"flyer{ext}").name)
                path=f"philo-flyers/{event_date.year}/{digest[:12]}_{safe_name}"
                stored=upload_private(raw,path,flyer_file.type or "application/octet-stream")
                if not stored:
                    raise RuntimeError("Flyer storage upload failed.")
                status="Ready for Branding Review" if save_review else "Draft"
                inserted=table("philo_flyer_submissions").insert({
                    "title":title.strip(),
                    "committee_office":committee.strip(),
                    "requested_by":requested_by.strip(),
                    "event_date":event_date.isoformat(),
                    "event_time":event_time.strip(),
                    "location_details":location.strip(),
                    "cta_contact":cta.strip(),
                    "description":description.strip(),
                    "file_path":stored,
                    "file_name":flyer_file.name or safe_name,
                    "file_sha256":digest,
                    "version_number":1,
                    "status":status,
                    "branding_passed":False,
                    "submitted_by_member_id":mid,
                    "submitted_by_name":member_name,
                    "created_at":datetime.now(timezone.utc).isoformat(),
                    "updated_at":datetime.now(timezone.utc).isoformat(),
                }).execute().data or []
                st.success(f"Flyer saved. Status: {status}.")
                st.rerun()
            except Exception as ex:
                st.error(f"Flyer was not saved: {ex}")

        rows_=flyer_rows()
        editable=[r for r in rows_ if r.get("status") in {"Draft","Ready for Branding Review","Needs Changes","Not Approved / Revise"}]
        if editable:
            st.markdown("### Manage Existing Draft / Revision")
            fmap={r["id"]:f"{r.get('title')} • v{r.get('version_number') or 1} • {r.get('status')}" for r in editable}
            fid=st.selectbox("Choose Flyer",list(fmap.keys()),format_func=lambda x:fmap[x],key="manage_flyer_id")
            row=next(r for r in editable if r["id"]==fid)
            flyer_status_banner(row.get("status"))

            replacement=st.file_uploader("Replace Flyer with Corrected Version",type=["pdf","png","jpg","jpeg"],key=f"replace_flyer_{fid}")
            c1,c2,c3=st.columns(3)
            if c1.button("Replace Flyer",use_container_width=True,key=f"replace_btn_{fid}"):
                try:
                    raw,ext=flyer_file_validate(replacement)
                    digest=hashlib.sha256(raw).hexdigest()
                    dup=table("philo_flyer_versions").select("id").eq("file_sha256",digest).execute().data or []
                    if dup:
                        raise ValueError("This exact file version is already stored.")
                    safe_name=re.sub(r"[^A-Za-z0-9._-]+","_",Path(replacement.name or f"flyer{ext}").name)
                    new_ver=int(row.get("version_number") or 1)+1
                    path=f"philo-flyers/{event_date.year}/{digest[:12]}_v{new_ver}_{safe_name}"
                    stored=upload_private(raw,path,replacement.type or "application/octet-stream")
                    table("philo_flyer_versions").insert({
                        "flyer_id":fid,
                        "version_number":new_ver,
                        "file_path":stored,
                        "file_name":replacement.name or safe_name,
                        "file_sha256":digest,
                        "replaced_by_member_id":mid,
                        "replaced_by_name":member_name,
                        "created_at":datetime.now(timezone.utc).isoformat(),
                    }).execute()
                    table("philo_flyer_submissions").update({
                        "file_path":stored,
                        "file_name":replacement.name or safe_name,
                        "file_sha256":digest,
                        "version_number":new_ver,
                        "status":"Ready for Branding Review",
                        "branding_passed":False,
                        "updated_at":datetime.now(timezone.utc).isoformat(),
                    }).eq("id",fid).execute()
                    st.success(f"Corrected flyer saved as version {new_ver}. Branding review must be rerun.")
                    st.rerun()
                except Exception as ex:
                    st.error(f"Replacement was not saved: {ex}")

            if c2.button("Cancel Submission",use_container_width=True,key=f"cancel_flyer_{fid}"):
                table("philo_flyer_submissions").update({
                    "status":"Cancelled",
                    "updated_at":datetime.now(timezone.utc).isoformat(),
                }).eq("id",fid).execute()
                st.success("Submission cancelled.")
                st.rerun()

            if c3.button("Delete Draft",use_container_width=True,key=f"delete_flyer_{fid}"):
                st.session_state[f"confirm_delete_flyer_{fid}"]=True

            if st.session_state.get(f"confirm_delete_flyer_{fid}"):
                st.warning("This permanently deletes the draft workflow record. Are you sure?")
                d1,d2=st.columns(2)
                if d1.button("Yes, Delete Draft",use_container_width=True,key=f"confirm_delete_{fid}"):
                    table("philo_flyer_submissions").delete().eq("id",fid).execute()
                    st.session_state.pop(f"confirm_delete_flyer_{fid}",None)
                    st.success("Draft deleted.")
                    st.rerun()
                if d2.button("No, Keep It",use_container_width=True,key=f"keep_flyer_{fid}"):
                    st.session_state.pop(f"confirm_delete_flyer_{fid}",None)
                    st.rerun()

    # -----------------------------------------------------
    # 2. Required branding review with feedback
    # -----------------------------------------------------
    with tabs[1]:
        reviewable=[r for r in flyer_rows() if r.get("status") in {
            "Draft","Ready for Branding Review","Needs Changes","Not Approved / Revise"
        }]
        if not reviewable:
            st.info("No flyers are available for branding review.")
        else:
            fmap={r["id"]:f"{r.get('title')} • v{r.get('version_number') or 1}" for r in reviewable}
            fid=st.selectbox("Flyer to Review",list(fmap.keys()),format_func=lambda x:fmap[x],key="brand_review_flyer_id")
            row=next(r for r in reviewable if r["id"]==fid)
            flyer_status_banner(row.get("status"))
            st.progress(min(flyer_progress(row.get("status"))/6,1.0),text="Workflow progress")

            if st.session_state["flyer_help_guide"]:
                st.info(
                    "The checklist uses Philo-appropriate Gold & White branding plus general design-quality principles: "
                    "clear imagery, minimal clutter, white space, readable hierarchy, limited fonts, and strong contrast."
                )

            # Metadata pre-checks
            metadata_feedback=[]
            if not row.get("event_time"):
                metadata_feedback.append("Event time is missing from the workflow record.")
            if not row.get("location_details"):
                metadata_feedback.append("Location/access details are missing from the workflow record.")
            if not row.get("cta_contact"):
                metadata_feedback.append("CTA/contact information is missing from the workflow record.")
            event_d = row.get("event_date")
            try:
                if event_d and pd.to_datetime(event_d).date() < date.today():
                    metadata_feedback.append("The event date has already passed. Confirm this is an archival/revision workflow.")
            except Exception:
                pass

            if metadata_feedback:
                st.warning("**Pre-check findings:**\n\n- " + "\n- ".join(metadata_feedback))

            st.markdown("### Required Branding Checklist")
            vals={}
            for key,label,fix in FLYER_REQUIRED_CHECKS:
                vals[key]=st.checkbox(label,key=f"brand_check_{fid}_{key}")

            notes=st.text_area(
                "Reviewer Notes",
                placeholder="Add specific observations or recommended changes.",
                key=f"brand_review_notes_{fid}"
            )

            if st.button("Run Branding Review",use_container_width=True,key=f"run_brand_review_{fid}"):
                passed,total,pct,overall,missing=flyer_brand_review_summary(vals)
                feedback=flyer_feedback(vals)
                table("philo_flyer_reviews").insert({
                    "flyer_id":fid,
                    **vals,
                    "score_percent":pct,
                    "review_outcome":overall,
                    "review_notes":notes.strip(),
                    "reviewed_by_member_id":mid,
                    "reviewed_by_name":member_name,
                    "created_at":datetime.now(timezone.utc).isoformat(),
                }).execute()

                if missing or metadata_feedback:
                    table("philo_flyer_submissions").update({
                        "status":"Needs Changes",
                        "branding_passed":False,
                        "branding_score":pct,
                        "last_branding_review_at":datetime.now(timezone.utc).isoformat(),
                        "updated_at":datetime.now(timezone.utc).isoformat(),
                    }).eq("id",fid).execute()
                    save_flyer_notification(
                        fid,row.get("committee_office"),
                        f"Flyer '{row.get('title')}' needs changes before Chapter submission.",
                        "Branding Review",mid
                    )
                else:
                    table("philo_flyer_submissions").update({
                        "status":"Ready to Submit to Chapter",
                        "branding_passed":True,
                        "branding_score":pct,
                        "last_branding_review_at":datetime.now(timezone.utc).isoformat(),
                        "updated_at":datetime.now(timezone.utc).isoformat(),
                    }).eq("id",fid).execute()
                    save_flyer_notification(
                        fid,row.get("committee_office"),
                        f"Flyer '{row.get('title')}' passed branding review and is ready to submit to Chapter.",
                        "Branding Review",mid
                    )

                st.session_state[f"brand_results_{fid}"]={
                    "pct":pct,"overall":overall,"feedback":feedback,"metadata":metadata_feedback
                }

            result=st.session_state.get(f"brand_results_{fid}")
            if result:
                st.markdown("### Branding Review Results")
                if result["overall"].startswith("PASS") and not result["metadata"]:
                    st.success(f"✓ PASS — {result['pct']}% — Ready to Submit to Chapter")
                elif result["overall"].startswith("NEEDS"):
                    st.warning(f"⚠ NEEDS ATTENTION — {result['pct']}%")
                else:
                    st.error(f"✕ MUST FIX BEFORE SUBMISSION — {result['pct']}%")
                st.dataframe(pd.DataFrame(result["feedback"]),hide_index=True,use_container_width=True)

                failed=[x for x in result["feedback"] if x["Result"]!="✓ PASS"]
                if failed or result["metadata"]:
                    st.markdown("### Recommended Improvements")
                    for item in failed:
                        st.write(f"- **{item['Check']}:** {item['Recommendation']}")
                    for item in result["metadata"]:
                        st.write(f"- **Workflow detail:** {item}")

    # -----------------------------------------------------
    # 3. Submit to Chapter — locked until branding passes
    # -----------------------------------------------------
    with tabs[2]:
        ready=[r for r in flyer_rows() if r.get("status")=="Ready to Submit to Chapter" and r.get("branding_passed")]
        if not ready:
            st.info(
                "No flyer is currently eligible for Chapter submission. "
                "Run the Branding Review and fix all required items first."
            )
        else:
            fmap={r["id"]:f"{r.get('title')} • v{r.get('version_number') or 1} • Branding {r.get('branding_score') or 0}%" for r in ready}
            fid=st.selectbox("Flyer Ready for Chapter",list(fmap.keys()),format_func=lambda x:fmap[x],key="chapter_ready_flyer_id")
            row=next(r for r in ready if r["id"]==fid)
            flyer_status_banner(row.get("status"))
            st.success("✓ Branding gate passed.")
            st.write(f"**Responsible Officer/Committee:** {row.get('committee_office')}")
            st.write(f"**Event Date:** {row.get('event_date')}")
            st.write(f"**Branding Score:** {row.get('branding_score')}%")

            confirm=st.checkbox(
                "I confirm this is the final reviewed flyer and it is ready to be sent to Chapter for approval.",
                key=f"chapter_submit_confirm_{fid}"
            )
            if st.button(
                "Submit to Chapter for Approval",
                use_container_width=True,
                disabled=not confirm,
                key=f"submit_to_chapter_{fid}"
            ):
                table("philo_flyer_submissions").update({
                    "status":"Waiting Approval from Chapter",
                    "sent_to_chapter_at":datetime.now(timezone.utc).isoformat(),
                    "sent_to_chapter_by":member_name,
                    "updated_at":datetime.now(timezone.utc).isoformat(),
                }).eq("id",fid).execute()
                save_flyer_notification(
                    fid,row.get("committee_office"),
                    f"Flyer '{row.get('title')}' is now Waiting Approval from Chapter.",
                    "Status Update",mid
                )
                st.success("Status changed to Waiting Approval from Chapter.")
                st.rerun()

    # -----------------------------------------------------
    # 4. Record Chapter decision + evidence
    # -----------------------------------------------------
    with tabs[3]:
        waiting=flyer_rows(status="Waiting Approval from Chapter")
        if not waiting:
            st.info("There are no pending Chapter approval decisions.")
        else:
            fmap={r["id"]:f"{r.get('title')} • {r.get('committee_office')}" for r in waiting}
            fid=st.selectbox("Flyer",list(fmap.keys()),format_func=lambda x:fmap[x],key="flyer_decision_id")
            row=next(r for r in waiting if r["id"]==fid)
            flyer_status_banner(row.get("status"))
            decision=st.selectbox("Chapter Decision",["Approved by Chapter","Not Approved / Revise"])
            decision_date=st.date_input("Decision Date",value=date.today())
            approver=st.text_input("Chapter Approver / Approval Source")
            notes=st.text_area("Decision Notes / Required Revisions")
            evidence=st.file_uploader(
                "Optional Approval Evidence (email PDF/screenshot)",
                type=["pdf","png","jpg","jpeg"],
                key=f"approval_evidence_{fid}"
            )
            if st.button("Save Chapter Decision",use_container_width=True):
                if not approver.strip():
                    st.error("Enter the Chapter approver or approval source.")
                elif decision=="Not Approved / Revise" and not notes.strip():
                    st.error("Enter the revisions requested by Chapter.")
                else:
                    evidence_path=""
                    if evidence:
                        raw,ext=flyer_file_validate(evidence)
                        digest=hashlib.sha256(raw).hexdigest()
                        safe_name=re.sub(r"[^A-Za-z0-9._-]+","_",Path(evidence.name).name)
                        path=f"philo-flyers/approval-evidence/{digest[:12]}_{safe_name}"
                        evidence_path=upload_private(raw,path,evidence.type or "application/octet-stream") or ""
                    table("philo_flyer_submissions").update({
                        "status":decision,
                        "chapter_decision_date":decision_date.isoformat(),
                        "chapter_approver_source":approver.strip(),
                        "chapter_decision_notes":notes.strip(),
                        "chapter_approval_evidence_path":evidence_path,
                        "chapter_decision_recorded_by":member_name,
                        "updated_at":datetime.now(timezone.utc).isoformat(),
                    }).eq("id",fid).execute()

                    if decision=="Approved by Chapter":
                        msg=f"Flyer '{row.get('title')}' has been Approved by Chapter. You may proceed with the approved final flyer."
                    else:
                        msg=f"Flyer '{row.get('title')}' was not approved as submitted. Please revise before resubmission. {notes.strip()}"
                    save_flyer_notification(fid,row.get("committee_office"),msg,"Chapter Decision",mid)
                    st.success("Chapter decision saved and notification created.")
                    st.rerun()

    # -----------------------------------------------------
    # 5. Action-oriented notifications
    # -----------------------------------------------------
    with tabs[4]:
        notices=table("philo_flyer_notifications").select("*").order("created_at",desc=True).execute().data or []
        if not notices:
            st.info("No flyer notifications yet.")
        else:
            for n in notices[:30]:
                kind=str(n.get("notification_type") or "")
                msg=str(n.get("message") or "")
                if "needs changes" in msg.casefold() or "not approved" in msg.casefold():
                    st.warning(f"**ACTION REQUIRED** — {msg}")
                elif "approved by chapter" in msg.casefold():
                    st.success(f"**APPROVED** — {msg}")
                else:
                    st.info(f"**{kind or 'UPDATE'}** — {msg}")

    # -----------------------------------------------------
    # 6. Archive and version history
    # -----------------------------------------------------
    with tabs[5]:
        rows_=flyer_rows()
        if not rows_:
            st.info("No flyer workflow records yet.")
        else:
            st.dataframe(pd.DataFrame([{
                "Flyer":r.get("title"),
                "Version":r.get("version_number"),
                "Responsible":r.get("committee_office"),
                "Event Date":r.get("event_date"),
                "Status":r.get("status"),
                "Branding Score":r.get("branding_score"),
                "Chapter Decision Date":r.get("chapter_decision_date"),
                "Approval Source":r.get("chapter_approver_source"),
                "File":r.get("file_name"),
            } for r in rows_]),hide_index=True,use_container_width=True)

            versions=table("philo_flyer_versions").select("*").order("created_at",desc=True).execute().data or []
            if versions:
                st.markdown("### Flyer Version History")
                st.dataframe(pd.DataFrame(versions),hide_index=True,use_container_width=True)

def save_flyer_notification(flyer_id,audience,message,notification_type,mid):
    table("philo_flyer_notifications").insert({
        "flyer_id":flyer_id,
        "audience":audience or "Submitting Officer/Committee",
        "message":message,
        "notification_type":notification_type,
        "created_by_member_id":mid,
        "created_by_name":member_name,
        "created_at":datetime.now(timezone.utc).isoformat(),
    }).execute()

def save_historian_preference(mid,key,value):
    try:
        table("historian_preferences").upsert({
            "member_id":mid,
            "preference_key":key,
            "preference_value":"true" if bool(value) else "false",
            "updated_at":datetime.now(timezone.utc).isoformat(),
        },on_conflict="member_id,preference_key").execute()
    except Exception:
        pass


# ============================================================
# v3.6.23 — Leadership Updates, Chapter Votes, Official Reports
# ============================================================

LEADERSHIP_UPDATE_TYPES=[
    'Information',
    'Action Needed',
    'Report Update',
    'Vote / Decision Needed',
    'Deadline',
    'Guidance'
]
LEADERSHIP_UPDATE_PRIORITIES=['Routine','Important','Urgent']

def leadership_updates_for_member(mid,position=''):
    try:
        all_rows=table('leadership_updates').select('*').order('created_at',desc=True).execute().data or []
    except:
        return []
    chaired_names={str(c.get('committee_name') or '') for c in chaired_committees(mid)}
    out=[]
    for r in all_rows:
        rt=str(r.get('recipient_type') or '')
        rv=str(r.get('recipient_value') or '')
        if rt=='Member' and rv==str(mid):
            out.append(r)
        elif rt=='Officer Position' and position and rv==position:
            out.append(r)
        elif rt=='All Officers' and position:
            out.append(r)
        elif rt=='Committee' and rv in chaired_names:
            out.append(r)
    return out

def leadership_report_queue(mid):
    try:
        return table('leadership_report_queue').select('*').eq('member_id',mid).eq('active',True).order('created_at').execute().data or []
    except:
        return []

def leadership_update_text(update_id):
    try:
        rows_=table('leadership_updates').select('*').eq('id',update_id).execute().data or []
        if not rows_:return ''
        r=rows_[0]
        parts=[]
        if r.get('subject'):parts.append(str(r.get('subject')))
        if r.get('message'):parts.append(str(r.get('message')))
        if r.get('suggested_action'):parts.append('Suggested action: '+str(r.get('suggested_action')))
        return " — ".join(parts)
    except:return ''

def queue_leadership_update_for_report(update_id,mid,name):
    table('leadership_report_queue').upsert({
        'update_id':update_id,'member_id':mid,'member_name':name,'active':True,
        'created_at':datetime.now(timezone.utc).isoformat()
    },on_conflict='update_id,member_id').execute()
    table('leadership_updates').update({
        'status':'Added to Report',
        'updated_at':datetime.now(timezone.utc).isoformat()
    }).eq('id',update_id).execute()

def remove_leadership_update_from_report(update_id,mid):
    try:
        table('leadership_report_queue').update({'active':False}).eq('update_id',update_id).eq('member_id',mid).execute()
    except:pass

def open_chapter_votes():
    try:
        rows_=table('chapter_votes').select('*').eq('active',True).order('opens_at').execute().data or []
    except:return []
    now=datetime.now(timezone.utc)
    out=[]
    for v in rows_:
        try:
            start=pd.to_datetime(v.get('opens_at'),utc=True).to_pydatetime() if v.get('opens_at') else None
            end=pd.to_datetime(v.get('closes_at'),utc=True).to_pydatetime() if v.get('closes_at') else None
            if (start is None or now>=start) and (end is None or now<=end):
                out.append(v)
        except:
            out.append(v)
    return out

def vote_response(vote_id,mid):
    try:
        rows_=table('chapter_vote_responses').select('*').eq('vote_id',vote_id).eq('member_id',mid).execute().data or []
        return rows_[0] if rows_ else None
    except:return None

def chapter_vote_results(vote_id):
    try:return table('chapter_vote_responses').select('*').eq('vote_id',vote_id).execute().data or []
    except:return []

def render_open_chapter_votes(mid,name):
    votes=open_chapter_votes()
    if not votes:return
    st.markdown('### 🗳️ Open Chapter Votes')
    with st.container(border=True):
        for v in votes:
            st.markdown(f"**{v.get('title') or 'Chapter Vote'}**")
            if v.get('question'):st.write(v.get('question'))
            if v.get('closes_at'):st.caption(f"Voting closes: {fmt_dt(v.get('closes_at'))}")
            prior=vote_response(v['id'],mid)
            options=['Yes','No','Abstain']
            choice=st.radio(
                'Your Vote',
                options,
                index=options.index(prior.get('vote')) if prior and prior.get('vote') in options else 0,
                horizontal=True,
                key=f"chapter_vote_choice_{v['id']}"
            )
            if st.button('Submit / Update Vote',key=f"chapter_vote_save_{v['id']}",use_container_width=True):
                table('chapter_vote_responses').upsert({
                    'vote_id':v['id'],'member_id':mid,'member_name':name,'vote':choice,
                    'voted_at':datetime.now(timezone.utc).isoformat()
                },on_conflict='vote_id,member_id').execute()
                st.success('Vote recorded.')
                st.rerun()

def report_author_options(current_mid,is_admin_flag):
    if not is_admin_flag:
        return [current_mid]
    try:
        return [m['id'] for m in active_members() if m.get('id') is not None]
    except:return [current_mid]

def report_author_info(mid):
    try:
        rows_=table('members').select('*').eq('id',mid).execute().data or []
        if rows_:
            r=rows_[0]
            return {'id':mid,'name':r.get('full_name') or 'Philo Member','email':r.get('email') or ''}
    except:pass
    return {'id':mid,'name':'Philo Member','email':''}

def _wrap_report_lines(text,max_chars=78):
    lines=[]
    for raw in str(text or '').splitlines():
        raw=raw.strip().lstrip('•-').strip()
        if not raw:continue
        wrapped=textwrap.wrap(raw,width=max_chars) or ['']
        lines.extend(wrapped)
    return lines

def official_chapter_report_pdf(
    office_or_committee,
    author_name,
    author_email,
    meeting_date,
    accomplishments='',
    updates='',
    upcoming='',
    reminders='',
    action_items='',
    full_report=''
):
    bg=ASSET_DIR/'chapter_report_template.png'
    if not bg.exists():
        raise RuntimeError('Official chapter report template image is missing from assets.')

    bio=BytesIO()
    c=canvas.Canvas(bio,pagesize=letter)
    W,H=letter
    img=ImageReader(str(bg))

    # template was rendered at 1275 x 1650, same aspect ratio as US Letter
    sx=W/1275.0
    sy=H/1650.0

    def x(px):return px*sx
    def y_top(py):return H-(py*sy)
    def white_rect(px,py,pw,ph):
        c.setFillColor(colors.white)
        c.setStrokeColor(colors.white)
        c.rect(x(px),y_top(py+ph),x(pw),ph*sy,fill=1,stroke=0)

    def draw_bullets(text,px,py,pw,ph,font_size=9.2,max_items=8):
        c.setFillColor(colors.black)
        c.setFont('Helvetica',font_size)
        max_chars=max(22,int(pw/13))
        lines=_wrap_report_lines(text,max_chars=max_chars)
        yy=y_top(py)
        bottom=y_top(py+ph)
        leading=font_size+2.7
        count=0
        for line in lines:
            if yy-leading < bottom:break
            if count==0 or (count>0 and (not line.startswith(' '))):
                c.drawString(x(px),yy,u'•')
                c.drawString(x(px+18),yy,line[:max_chars])
            else:
                c.drawString(x(px+18),yy,line[:max_chars])
            yy-=leading
            count+=1
            if count>=max_items:break

    # First page: exact uploaded template as background.
    c.drawImage(img,0,0,width=W,height=H,mask='auto')

    # Replace gray example values while preserving black labels.
    white_rect(390,255,650,103)
    c.setFillColor(colors.black)
    c.setFont('Helvetica',9.2)
    header_x=x(392)
    c.drawString(header_x,y_top(273),str(office_or_committee or 'Philo Affiliate'))
    c.drawString(header_x,y_top(298),str(author_name or 'Philo Member'))
    c.drawString(header_x,y_top(323),str(author_email or ''))
    c.drawString(header_x,y_top(348),str(meeting_date or ''))

    # Greeting placeholder dots.
    white_rect(105,395,930,86)
    c.setFont('Helvetica',9.4)
    c.drawString(x(108),y_top(414),'Dear Sorors,')
    c.setFont('Helvetica',8.7)
    intro=f"I respectfully submit the {office_or_committee or 'Philo Affiliate'} report for the current reporting period."
    for i,line in enumerate(textwrap.wrap(intro,95)[:3]):
        c.drawString(x(108),y_top(440+i*18),line)

    # White out only the inside of the five report boxes (keep border + headings).
    white_rect(64,550,970,128)
    white_rect(64,777,970,128)
    white_rect(64,1005,970,124)
    white_rect(64,1248,470,125)
    white_rect(565,1248,468,125)

    draw_bullets(accomplishments,87,570,905,95,font_size=9.0,max_items=7)
    draw_bullets(updates,87,797,905,95,font_size=9.0,max_items=7)
    draw_bullets(upcoming,87,1024,905,91,font_size=9.0,max_items=7)
    draw_bullets(reminders,90,1267,405,92,font_size=8.7,max_items=6)
    draw_bullets(action_items,590,1267,405,92,font_size=8.7,max_items=6)

    # Full narrative continues on an official continuation page if present.
    if str(full_report or '').strip():
        c.showPage()
        c.drawImage(img,0,0,width=W,height=H,mask='auto')
        # Keep official header, clear everything below horizontal rule.
        c.setFillColor(colors.white)
        c.rect(0,0,W,y_top(245),fill=1,stroke=0)
        c.setFillColor(colors.black)
        c.setFont('Helvetica-Bold',13)
        c.drawCentredString(W/2,H-150,'Report Continuation')
        c.setFont('Helvetica',9.5)
        tx=c.beginText(54,H-180)
        tx.setLeading(13)
        for para in str(full_report).splitlines():
            if not para.strip():
                tx.textLine('')
                continue
            for line in textwrap.wrap(para.strip(),95):
                if tx.getY()<55:
                    c.drawText(tx)
                    c.showPage()
                    c.drawImage(img,0,0,width=W,height=H,mask='auto')
                    c.setFillColor(colors.white)
                    c.rect(0,0,W,y_top(245),fill=1,stroke=0)
                    c.setFillColor(colors.black)
                    c.setFont('Helvetica-Bold',13)
                    c.drawCentredString(W/2,H-150,'Report Continuation')
                    c.setFont('Helvetica',9.5)
                    tx=c.beginText(54,H-180)
                    tx.setLeading(13)
                tx.textLine(line)
        c.drawText(tx)

    c.save()
    bio.seek(0)
    return bio.getvalue()

def leadership_update_recipient_label(r):
    rt=r.get('recipient_type') or ''
    rv=r.get('recipient_value') or ''
    return f"{rt}: {rv}" if rv else rt


# ============================================================
# v3.6.24 — External Admin Demo Sandbox
# ============================================================

def demo_admin_enabled():
    return str(setting('demo_admin_enabled','true')).strip().lower() in {'1','true','yes','on'}

def demo_admin_code():
    # Default demo code can be changed by real Admin in App Settings.
    return str(setting('demo_admin_code','1922')).strip() or '1922'

def demo_admin_display_name():
    return str(setting('demo_admin_display_name','Demo Admin')).strip() or 'Demo Admin'

def demo_admin_login_allowed(name,code):
    if not demo_admin_enabled():
        return False
    return (
        normalize_login_name(name)==normalize_login_name(demo_admin_display_name())
        and str(code or '').strip()==demo_admin_code()
    )

def demo_admin_fake_member():
    return {
        'id':-424242,
        'full_name':demo_admin_display_name(),
        'email':'demo@preview.local',
        'phone':'',
        'role':'Admin',
        'active':True,
        'is_philo':False,
        'membership_type':'Demo Soror',
        'account_type':'Admin Demo'
    }

def demo_admin_notice():
    st.error(
        '🧪 ADMIN DEMO MODE — This is a protected sandbox. '
        'Nothing you submit, approve, delete, upload, vote on, edit, or change is saved to the live Philo Hub.'
    )
    st.caption(
        'Demo Admin can explore all dashboards and workflows, but live member records, documents, '
        'Google Drive files, settings, votes, payments, reports, and Philo data cannot be changed.'
    )

OFFICER_DOCUMENT_TYPES=[
    "Treasurer Report","Financial Secretary Report","President Report",
    "Vice President Report","Recording Secretary Report","Historian Report",
    "Committee Report","Meeting Minutes","Annual / End-of-Year Report",
    "Budget / Financial Report","Event Report","Other Officer Record"
]

def officer_doc_validate(uploaded,max_mb=25):
    if uploaded is None:
        raise ValueError("Choose a document to upload.")
    raw=bytes(uploaded.getbuffer())
    if not raw:
        raise ValueError("The selected document is empty.")
    if len(raw)>max_mb*1024*1024:
        raise ValueError(f"File is too large. Maximum size is {max_mb} MB.")
    name=Path(uploaded.name or "document").name
    ext=Path(name).suffix.lower()
    if ext not in {".pdf",".docx",".xlsx",".xls",".csv",".txt"}:
        raise ValueError("Supported files: PDF, DOCX, XLSX/XLS, CSV, TXT.")
    return raw,name,ext

def officer_doc_extract(uploaded):
    raw,name,ext=officer_doc_validate(uploaded)
    parts=[]; warnings=[]
    if ext==".pdf":
        if PdfReader is None:
            warnings.append("PDF parser unavailable in this deployment.")
        else:
            reader=PdfReader(BytesIO(raw))
            for i,p in enumerate(reader.pages[:100]):
                try:
                    parts.append(p.extract_text() or "")
                except Exception:
                    warnings.append(f"Could not read PDF page {i+1}.")
    elif ext==".docx":
        if DocxDocument is None:
            warnings.append("DOCX parser unavailable in this deployment.")
        else:
            d=DocxDocument(BytesIO(raw))
            parts.extend([p.text for p in d.paragraphs if p.text.strip()])
            for t in d.tables:
                for row in t.rows:
                    parts.append(" | ".join(c.text.strip() for c in row.cells))
    elif ext in {".xlsx",".xls"}:
        xl=pd.ExcelFile(BytesIO(raw))
        for sheet in xl.sheet_names[:25]:
            df=pd.read_excel(BytesIO(raw),sheet_name=sheet,header=None).dropna(how="all")
            if df.empty: continue
            parts.append(f"[Sheet: {sheet}]")
            for _,r in df.head(1000).iterrows():
                vals=[str(x).strip() for x in r.tolist() if str(x).strip() not in {"","nan","None"}]
                if vals: parts.append(" | ".join(vals))
    elif ext==".csv":
        df=pd.read_csv(BytesIO(raw))
        for _,r in df.head(2000).iterrows():
            vals=[str(x).strip() for x in r.tolist() if str(x).strip() not in {"","nan","None"}]
            if vals: parts.append(" | ".join(vals))
    else:
        parts.append(raw.decode("utf-8",errors="replace"))
    return {
        "raw":raw,"name":name,"ext":ext,"text":"\n".join(parts),
        "sha256":hashlib.sha256(raw).hexdigest(),
        "mime":str(getattr(uploaded,"type","") or "application/octet-stream"),
        "warnings":warnings
    }

def _extract_year(text):
    m=re.search(r"\b(20\d{2})\s*[-–—/]\s*(20\d{2})\b",text)
    return f"{m.group(1)}-{m.group(2)}" if m else ""

def _extract_date(text):
    pats=[
        r"\b(20\d{2}-\d{1,2}-\d{1,2})\b",
        r"\b(\d{1,2}/\d{1,2}/20\d{2})\b",
        r"\b((?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)\s+\d{1,2},?\s+20\d{2})\b",
    ]
    for pat in pats:
        m=re.search(pat,text,re.I)
        if m:
            try: return pd.to_datetime(m.group(1)).date().isoformat()
            except Exception: pass
    return None

def _extract_money(label,text):
    for pat in [rf"{label}\s*[:\-]?\s*\$?\s*([0-9][0-9,]*\.?[0-9]{{0,2}})",
                rf"{label}.*?\$?\s*([0-9][0-9,]*\.?[0-9]{{0,2}})"]:
        m=re.search(pat,text,re.I|re.S)
        if m:
            try: return float(m.group(1).replace(",",""))
            except Exception: pass
    return None

def officer_doc_detect_office(doc_type,text,position_hint=""):
    combined=f"{doc_type} {position_hint} {text[:5000]}".casefold()
    for office,token in [
        ("Financial Secretary","financial secretary"),("Treasurer","treasurer"),
        ("Historian","historian"),("Recording Secretary","recording secretary"),
        ("Vice President","vice president"),("President","president"),
        ("Chaplain","chaplain"),("Parliamentarian","parliamentarian"),
        ("Sergeant-at-Arms","sergeant")
    ]:
        if token in combined: return office
    return position_hint or ""

def officer_doc_route(office,doc_type):
    if office in {"Treasurer","Financial Secretary"} or "Financial" in doc_type or "Budget" in doc_type:
        return "Finance"
    if office=="Historian" or doc_type in {"Historian Report","Event Report"}:
        return "Historian"
    if office=="Recording Secretary" or doc_type=="Meeting Minutes":
        return "Minutes / Secretary"
    return "Officer Report"

def officer_doc_common(parsed,doc_type,position_hint=""):
    txt=parsed["text"]
    office=officer_doc_detect_office(doc_type,txt,position_hint)
    lines=[x.strip() for x in txt.splitlines() if x.strip()]
    title=next((x for x in lines[:40] if len(x)<=180 and any(k in x.casefold() for k in ["report","minutes","budget","historian","treasurer","financial secretary"])), "")
    return {
        "office":office,"document_type":doc_type,
        "title":title or f"{doc_type} — {parsed['name']}",
        "sorority_year":_extract_year(txt),"report_date":_extract_date(txt),
        "summary":txt[:5000].strip()
    }

def officer_doc_finance(parsed):
    txt=parsed["text"]
    return {
        "opening_balance":_extract_money(r"(?:opening balance|beginning balance|starting balance)",txt),
        "deposits":_extract_money(r"(?:total deposits|deposits|total income|income)",txt),
        "withdrawals":_extract_money(r"(?:total withdrawals|withdrawals|total expenses|expenses)",txt),
        "closing_balance":_extract_money(r"(?:closing balance|ending balance|bank balance|account balance)",txt),
        "dues":_extract_money(r"(?:dues)",txt),
        "donations":_extract_money(r"(?:donations?)",txt),
        "fundraising":_extract_money(r"(?:fundraising|fundraiser)",txt),
        "zeffy":_extract_money(r"(?:zeffy)",txt),
        "assessments":_extract_money(r"(?:assessments?)",txt),
        "reimbursements":_extract_money(r"(?:reimbursements?|club expenses?)",txt),
        "gifts":_extract_money(r"(?:gifts?)",txt),
    }

def officer_doc_historian(parsed):
    lines=[x.strip() for x in parsed["text"].splitlines() if x.strip()]
    markers=["walk","meeting","induction","award","luncheon","service","event","boule",
             "photoshoot","parade","spotlight","reception","training","tech day"]
    found=[]
    for line in lines:
        if len(line)<=300 and any(k in line.casefold() for k in markers):
            found.append({"event_date":_extract_date(line),"title":line[:180],"description":line[:500]})
        if len(found)>=40: break
    return found

def save_officer_intake(parsed,common,mid):
    dup=table("officer_document_intake").select("id,status").eq("file_sha256",parsed["sha256"]).execute().data or []
    if dup: return dup[0]["id"],True
    safe=re.sub(r"[^A-Za-z0-9._-]+","_",parsed["name"])
    path=f"officer-document-intake/{common.get('sorority_year') or 'unknown-year'}/{parsed['sha256'][:12]}_{safe}"
    stored=upload_private(parsed["raw"],path,parsed["mime"])
    if not stored: raise RuntimeError("Source document could not be stored.")
    row=table("officer_document_intake").insert({
        "file_name":parsed["name"],"file_path":stored,"file_sha256":parsed["sha256"],
        "document_type":common["document_type"],"detected_office":common["office"],
        "sorority_year":common.get("sorority_year") or "","report_date":common.get("report_date"),
        "detected_title":common.get("title") or "","extracted_text_preview":common.get("summary") or "",
        "route_destination":officer_doc_route(common["office"],common["document_type"]),
        "status":"Extracted — Review Required","uploaded_by_member_id":mid,
        "uploaded_by_name":member_name,"created_at":datetime.now(timezone.utc).isoformat(),
        "updated_at":datetime.now(timezone.utc).isoformat()
    }).execute().data or []
    if not row: raise RuntimeError("Document intake record was not created.")
    return row[0]["id"],False

def render_universal_document_intake(mid,position,is_admin_flag=False):
    st.markdown("## Upload Officer Document")
    st.caption(
        "Upload the report once. The app extracts what it can and prepares reviewable drafts. "
        "Existing forms remain available only as backup/manual correction."
    )
    with st.expander("📄 Upload & Extract Officer Document",expanded=False):
        doc_type=st.selectbox("Document Type",OFFICER_DOCUMENT_TYPES,key=f"u_doc_type_{position}_{mid}")
        upload=st.file_uploader("Officer Document",type=["pdf","docx","xlsx","xls","csv","txt"],key=f"u_doc_{position}_{mid}")
        if st.button("Extract Information from Document",use_container_width=True,key=f"u_extract_{position}_{mid}"):
            try:
                parsed=officer_doc_extract(upload)
                common=officer_doc_common(parsed,doc_type,position)
                result={"parsed":parsed,"common":common,"finance":officer_doc_finance(parsed),"events":officer_doc_historian(parsed)}
                st.session_state[f"u_result_{position}_{mid}"]=result
            except Exception as ex:
                st.error(f"Document could not be extracted: {ex}")

        result=st.session_state.get(f"u_result_{position}_{mid}")
        if result:
            parsed=result["parsed"]; common=result["common"]
            for w in parsed.get("warnings") or []: st.warning(w)
            st.markdown("### Extracted Information — Review Before Saving")
            title=st.text_input("Title",value=common["title"],key=f"u_title_{position}_{mid}")
            sy=st.text_input("Sorority Year",value=common["sorority_year"],placeholder="2025-2026",key=f"u_sy_{position}_{mid}")
            rdate=st.text_input("Report Date",value=common["report_date"] or "",placeholder="YYYY-MM-DD",key=f"u_date_{position}_{mid}")
            office=st.text_input("Officer / Office",value=common["office"] or position,key=f"u_office_{position}_{mid}")
            summary=st.text_area("Extracted Summary / Notes",value=common["summary"],height=220,key=f"u_summary_{position}_{mid}")
            route=officer_doc_route(office,doc_type)
            st.info(f"**Suggested destination:** {route}")

            finance_edit=None
            if route=="Finance":
                st.markdown("#### Financial Values Found")
                finance_edit={}
                for key,label in [
                    ("opening_balance","Opening Balance"),("deposits","Deposits / Income"),
                    ("withdrawals","Withdrawals / Expenses"),("closing_balance","Closing / Ending Balance"),
                    ("dues","Dues"),("donations","Donations"),("fundraising","Fundraising"),
                    ("zeffy","Zeffy / Online Giving"),("assessments","Assessments"),
                    ("reimbursements","Reimbursements / Club Expenses"),("gifts","Gifts")
                ]:
                    finance_edit[key]=st.number_input(label,value=float(result["finance"].get(key) or 0),
                        step=0.01,format="%.2f",key=f"u_fin_{position}_{mid}_{key}")
            elif route=="Historian":
                st.markdown("#### Possible Historical Events Found")
                if result["events"]:
                    st.dataframe(pd.DataFrame(result["events"]),hide_index=True,use_container_width=True)
                    st.caption("These remain drafts until the Historian reviews them.")
                else:
                    st.info("No event candidates were confidently identified.")

            c1,c2=st.columns(2)
            if c1.button("Save Extracted Drafts",use_container_width=True,key=f"u_save_{position}_{mid}"):
                try:
                    common2=dict(common)
                    common2.update({"title":title.strip(),"sorority_year":sy.strip(),"report_date":rdate.strip() or None,
                                    "office":office.strip(),"summary":summary.strip()})
                    intake_id,duplicate=save_officer_intake(parsed,common2,mid)
                    if duplicate: st.warning("This exact document was already uploaded; the existing intake record was reused.")
                    if route=="Finance":
                        f=finance_edit or result["finance"]
                        table("officer_finance_extraction_drafts").upsert({
                            "intake_id":intake_id,"sorority_year":sy.strip(),"report_date":rdate.strip() or None,
                            **f,"status":"Draft from Uploaded Document",
                            "created_by_member_id":mid,"created_by_name":member_name,
                            "created_at":datetime.now(timezone.utc).isoformat()
                        },on_conflict="intake_id").execute()
                        msg="Financial extraction draft saved for review."
                    elif route=="Historian":
                        created=0
                        for ev in result["events"]:
                            title_ev=str(ev.get("title") or "").strip()
                            if not title_ev: continue
                            table("officer_historian_extraction_drafts").insert({
                                "intake_id":intake_id,"sorority_year":sy.strip(),"event_date":ev.get("event_date"),
                                "title":title_ev,"description":ev.get("description") or "",
                                "verification_status":"Previously Documented — Needs Source",
                                "status":"Draft from Uploaded Document","created_by_member_id":mid,
                                "created_by_name":member_name,"created_at":datetime.now(timezone.utc).isoformat()
                            }).execute()
                            created+=1
                        msg=f"Historian extraction saved with {created} candidate event draft(s)."
                    else:
                        table("officer_report_extraction_drafts").upsert({
                            "intake_id":intake_id,"office":office.strip(),"document_type":doc_type,
                            "sorority_year":sy.strip(),"report_date":rdate.strip() or None,
                            "title":title.strip(),"summary":summary.strip(),"status":"Draft from Uploaded Document",
                            "created_by_member_id":mid,"created_by_name":member_name,
                            "created_at":datetime.now(timezone.utc).isoformat()
                        },on_conflict="intake_id").execute()
                        st.session_state["officer_doc_report_prefill"]={
                            "office":office.strip(),
                            "document_type":doc_type,
                            "sorority_year":sy.strip(),
                            "report_date":rdate.strip() or "",
                            "title":title.strip(),
                            "summary":summary.strip(),
                            "intake_id":intake_id
                        }
                        msg="Officer report extraction draft saved and is ready for the Report Center."
                    table("officer_document_intake").update({"status":"Drafts Created","updated_at":datetime.now(timezone.utc).isoformat()}).eq("id",intake_id).execute()
                    st.success(msg); st.rerun()
                except Exception as ex:
                    st.error(f"Extracted drafts were not saved: {ex}")
            if c2.button("Cancel Extraction",use_container_width=True,key=f"u_cancel_{position}_{mid}"):
                st.session_state.pop(f"u_result_{position}_{mid}",None); st.rerun()

    report_prefill=st.session_state.get("officer_doc_report_prefill")
    if report_prefill and officer_doc_route(report_prefill.get("office",""),report_prefill.get("document_type",""))=="Officer Report":
        st.success("Your reviewed extraction is ready. You do not need to type the report again.")
        if st.button("📝 Open Reviewed Draft in Report Center",use_container_width=True,key=f"u_open_report_{position}_{mid}"):
            st.session_state["managed_report_draft"]=report_prefill.get("summary") or ""
            st.session_state["managed_report_title"]=report_prefill.get("title") or f"{report_prefill.get('office') or position} Report"
            st.session_state["managed_report_period"]=(report_prefill.get("report_date") or report_prefill.get("sorority_year") or "")
            st.session_state["report_prefill_source"]=f"Officer: {report_prefill.get('office') or position}"
            st.session_state["report_prefill_updates"]=report_prefill.get("summary") or ""
            set_page("📝 Reports")

    try:
        recent=table("officer_document_intake").select(
            "file_name,document_type,detected_office,sorority_year,report_date,route_destination,status,created_at"
        ).order("created_at",desc=True).limit(20).execute().data or []
        if recent:
            st.markdown("### Recent Officer Document Uploads")
            st.dataframe(pd.DataFrame(recent),hide_index=True,use_container_width=True)
    except Exception:
        pass

def auth_cfg():
    a=sec('auth')
    return {'member_code':str(a.get('member_code','')),'admin_password':str(a.get('admin_password',''))} if a else None

def active_members():
    data=table('members').select('*').order('full_name').execute().data or []
    out=[]
    for m in data:
        active=m.get('active',True)
        if active is True or str(active).strip().lower() in {'true','1','yes','y'}:
            out.append(m)
    return out

def supabase_project_host():
    c=cfg()
    if not c:return 'Not configured'
    return c['url'].replace('https://','').replace('http://','').rstrip('/')

def require_login():
    if not cfg():st.error('Supabase is not configured. Follow SETUP_GUIDE.md.');st.stop()
    if st.session_state.get('logged_in'):return
    st.markdown(f"""<div class='philo-header'><img class='brand-logo' src='{PHILO_LOGO_URI}'><div><h1>NBS Philo Hub</h1><p>Members, service, forms, files, and meeting preparation in one place.</p><div class='gold-rule'></div></div><img class='brand-crest' src='{PHILO_CREST_URI}'></div>""",unsafe_allow_html=True)
    members=active_members()
    if not members:
        st.warning('No active members were returned to the app.')
        st.caption(f"Connected Supabase project: {supabase_project_host()}")
        try:
            count=len(table('members').select('id').execute().data or [])
            st.caption(f"Rows visible to the app in members table: {count}")
        except Exception as e:
            st.error(f"Member table check failed: {e}")
        if st.button('Refresh Member List',use_container_width=True):
            st.cache_resource.clear()
            st.rerun()
        st.stop()
    st.markdown('<div class="pearl-divider"></div>',unsafe_allow_html=True)
    if st.session_state.get('interest_portal_mode'):
        render_interest_candidate_portal()
    if st.button('🌱 Interested in Becoming a Philo?',use_container_width=True,key='open_interest_portal'):
        st.session_state['interest_portal_mode']=True
        st.rerun()
    st.caption('Interested women can create a separate Interest Profile. This does not provide access to the member Hub.')

    if demo_admin_enabled():
        with st.expander('🧪 Admin Demo — Explore the Full Hub',expanded=False):
            st.caption(
                'For invited Sorors who want to explore the app. Demo Admin can try every dashboard and workflow, '
                'but nothing in the live Hub can be changed.'
            )
            with st.form('demo_admin_login_form'):
                demo_name=st.text_input('Demo Login Name',value='')
                demo_code=st.text_input('Demo Access Code',type='password')
                enter_demo=st.form_submit_button('Enter Admin Demo Sandbox',use_container_width=True)
            if enter_demo:
                if demo_admin_login_allowed(demo_name,demo_code):
                    st.session_state['demo_admin_mode']=True
                    st.session_state['member_id']=-424242
                    st.session_state['is_admin']=True
                    st.session_state['page']='🏠 Dashboard'
                    st.rerun()
                else:
                    st.error('Demo Admin name or access code did not match.')
    if not st.session_state.get('login_lookup_member_id'):
        with st.form('typed_name_lookup'):
            first_login=st.text_input('First Name',autocomplete='given-name')
            last_login=st.text_input('Last Name',autocomplete='family-name')
            admin_lookup=st.checkbox('Administrator')
            go=st.form_submit_button('Continue',use_container_width=True)
        if go:
            try:
                found=find_member_by_typed_name(first_login,last_login)
            except ValueError as ex:
                st.error(str(ex)); found=None
            if not found:st.error('We could not match that first and last name to an active account.')
            else:
                # Admin mode may be requested only for an account explicitly marked Admin.
                requested_admin=bool(admin_lookup)
                account_is_admin=(str(found.get('role') or '').strip().lower()=='admin')
                if requested_admin and not account_is_admin:
                    st.error('Administrator sign-in is restricted to the designated Advisor/Admin account.')
                else:
                    st.session_state['login_lookup_member_id']=found['id'];st.session_state['login_lookup_is_admin']=requested_admin;st.rerun()
        st.stop()
    mid=st.session_state['login_lookup_member_id'];admin=bool(st.session_state.get('login_lookup_is_admin'))
    if st.button('← Use a Different Name',key='login_change_name'):
        st.session_state.pop('login_lookup_member_id',None);st.session_state.pop('login_lookup_is_admin',None);st.rerun()
    pin_record=member_pin_record(mid);needs_setup=(not pin_record) or (not pin_record.get('pin_hash')) or bool(pin_record.get('reset_required'))
    if admin:
        with st.form('admin_login'):
            admin_pw=st.text_input('Admin password',type='password');submit_admin=st.form_submit_button('Enter Philo Hub',use_container_width=True)
        if submit_admin:
            a=auth_cfg() or {'admin_password':''}
            admin_member=next((m for m in active_members() if m.get('id')==mid),None)
            if not admin_member or str(admin_member.get('role') or '').strip().lower()!='admin':
                st.error('This account is not authorized for Advisor/Admin access.')
            elif not a.get('admin_password'):
                st.error('The Advisor/Admin password is not configured.')
            elif not hmac.compare_digest(str(admin_pw),str(a['admin_password'])):
                st.error('Admin password is incorrect.')
            else:
                st.session_state.update({'logged_in':True,'member_id':mid,'is_admin':True,'login_failures':0})
                st.session_state.pop('login_lookup_member_id',None);st.session_state.pop('login_lookup_is_admin',None);st.rerun()
    elif needs_setup:
        st.info('First-time setup: use the affiliate first-time code once, then create your private 4-digit PIN.')
        with st.form('first_time_pin_setup'):
            initial_code=st.text_input('Affiliate first-time access code',type='password',max_chars=4)
            new_pin=st.text_input('Create a 4-digit PIN',type='password',max_chars=4)
            confirm_pin=st.text_input('Confirm your 4-digit PIN',type='password',max_chars=4)
            setup=st.form_submit_button('Create PIN & Enter Philo Hub',use_container_width=True)
        if setup:
            if not hmac.compare_digest(str(initial_code),INITIAL_MEMBER_CODE):
                st.error('Initial access code is incorrect.')
            elif not(new_pin.isdigit() and len(new_pin)==4):
                st.error('Your PIN must be exactly 4 digits.')
            elif new_pin!=confirm_pin:
                st.error('The PINs do not match.')
            else:
                set_member_pin(mid,new_pin)
                st.session_state.update({'logged_in':True,'member_id':mid,'is_admin':False,'login_failures':0})
                st.session_state.pop('login_lookup_member_id',None)
                st.session_state.pop('login_lookup_is_admin',None)
                st.rerun()
    else:
        with st.form('member_pin_login'):
            code=st.text_input('4-digit PIN',type='password',max_chars=4);submit=st.form_submit_button('Enter Philo Hub',use_container_width=True)
        if submit:
            failures=int(st.session_state.get('login_failures',0) or 0)
            if failures>=5:
                st.error('Too many unsuccessful PIN attempts in this session. Use “Use a Different Name” or contact the Advisor for a PIN reset.')
            elif not check_member_pin(mid,code):
                st.session_state['login_failures']=failures+1
                st.error('PIN is incorrect.')
            else:
                st.session_state.update({'logged_in':True,'member_id':mid,'is_admin':False,'login_failures':0})
                st.session_state.pop('login_lookup_member_id',None);st.session_state.pop('login_lookup_is_admin',None);st.rerun()
    st.stop()

components.html("""<script>try{const d=window.parent.document;if(!d.querySelector('link[data-nbs-manifest]')){let m=d.createElement('link');m.rel='manifest';m.href='/app/static/manifest.json';m.setAttribute('data-nbs-manifest','1');d.head.appendChild(m);let a=d.createElement('link');a.rel='apple-touch-icon';a.href='/app/static/apple-touch-icon.png';d.head.appendChild(a);let c=d.createElement('meta');c.name='theme-color';c.content='#c99a20';d.head.appendChild(c);}}catch(e){}</script>""",height=0)

GOOGLE_SCOPES=['https://www.googleapis.com/auth/calendar.readonly','https://www.googleapis.com/auth/drive','https://www.googleapis.com/auth/spreadsheets']
def google_oauth_cfg():
    s=sec('google_oauth')
    if not s:return None
    return {'client_id':str(s.get('client_id','')).strip(),'client_secret':str(s.get('client_secret','')).strip(),'redirect_uri':str(s.get('redirect_uri','')).strip(),'calendar_id':str(s.get('calendar_id','')).strip()}
def google_flow():
    c=google_oauth_cfg()
    if not c or not c['client_id'] or not c['client_secret'] or not c['redirect_uri']:return None
    conf={'web':{'client_id':c['client_id'],'client_secret':c['client_secret'],'auth_uri':'https://accounts.google.com/o/oauth2/auth','token_uri':'https://oauth2.googleapis.com/token','redirect_uris':[c['redirect_uri']]}}
    # This is a confidential server-side OAuth client. Do not auto-generate a
    # PKCE verifier because Streamlit may create a fresh Python Flow object
    # after Google's redirect; that new object would not have the original
    # verifier and Google would reject the token exchange as "Missing code verifier."
    return Flow.from_client_config(
        conf,
        scopes=GOOGLE_SCOPES,
        redirect_uri=c['redirect_uri'],
        autogenerate_code_verifier=False
    )
def google_connected():return bool(st.session_state.get('google_token'))
def google_credentials():
    x=st.session_state.get('google_token')
    if not x:return None
    return Credentials(token=x.get('token'),refresh_token=x.get('refresh_token'),token_uri='https://oauth2.googleapis.com/token',client_id=x.get('client_id'),client_secret=x.get('client_secret'),scopes=x.get('scopes') or GOOGLE_SCOPES)
def google_oauth_state_secret():
    cfg=google_oauth_cfg() or {}
    secret=str(cfg.get('client_secret') or '').strip()
    if not secret:
        raise RuntimeError('Google OAuth client secret is required to sign OAuth state.')
    return secret.encode('utf-8')

def make_google_oauth_state(mid,is_admin_flag):
    payload={
        'member_id':int(mid),
        'is_admin':bool(is_admin_flag),
        'issued_at':int(datetime.now(timezone.utc).timestamp())
    }
    raw=json.dumps(payload,separators=(',',':')).encode('utf-8')
    body=base64.urlsafe_b64encode(raw).decode('ascii').rstrip('=')
    sig=hmac.new(
        google_oauth_state_secret(),
        body.encode('utf-8'),
        hashlib.sha256
    ).hexdigest()
    return f"{body}.{sig}"

def verify_google_oauth_state(state):
    try:
        body,sig=str(state or '').split('.',1)
        expected=hmac.new(
            google_oauth_state_secret(),
            body.encode('utf-8'),
            hashlib.sha256
        ).hexdigest()
        if not hmac.compare_digest(sig,expected):
            return None
        padded=body + '='*((4-len(body)%4)%4)
        payload=json.loads(base64.urlsafe_b64decode(padded.encode('ascii')).decode('utf-8'))
        issued=int(payload.get('issued_at') or 0)
        age=int(datetime.now(timezone.utc).timestamp())-issued
        if age<0 or age>900:
            return None
        return payload
    except:
        return None

def google_connect_url():
    if st.session_state.get('demo_admin_mode'):
        raise RuntimeError('ADMIN DEMO MODE: Google/Drive actions are disabled in the sandbox.')
    if test_preview_active():
        raise RuntimeError('TEST MODE: external Google/Drive actions are disabled.')
    if not bool(st.session_state.get('logged_in')) or not bool(st.session_state.get('is_admin')):
        return ''
    f=google_flow()
    if not f:return ''
    mid=st.session_state.get('member_id')
    if not mid:return ''
    signed_state=make_google_oauth_state(
        mid,
        st.session_state.get('is_admin',False)
    )
    u,_=f.authorization_url(
        access_type='offline',
        include_granted_scopes='true',
        prompt='consent',
        state=signed_state
    )
    return u

def process_google_oauth_callback():
    try:
        code=st.query_params.get('code')
        state=st.query_params.get('state')
        oauth_error=st.query_params.get('error')
    except:
        return False

    if oauth_error:
        st.query_params.clear()
        st.error(f'Google connection was not completed: {oauth_error}')
        return False

    if not code:
        return False

    payload=verify_google_oauth_state(state)
    if not payload:
        st.query_params.clear()
        st.error(
            'Google sign-in return could not be verified or it expired. '
            'Log in to the Hub and press Connect Advisor Google Account again.'
        )
        return False

    f=google_flow()
    if not f:
        st.query_params.clear()
        st.error('Google OAuth settings are incomplete.')
        return False

    try:
        f.fetch_token(code=code)
        c=f.credentials

        # Restore the Hub login that initiated the Google connection.
        restored_member_id=int(payload['member_id'])
        restored=table('members').select('*').eq(
            'id',restored_member_id
        ).single().execute().data
        if not restored or not restored.get('active',True):
            raise RuntimeError('The Hub account that started this connection is no longer active.')
        if str(restored.get('role') or '').strip().lower()!='admin':
            raise PermissionError('This account is no longer authorized for Advisor/Admin Google access.')
        if not bool(payload.get('is_admin')):
            raise PermissionError('The Google authorization was not started from an Advisor/Admin session.')

        st.session_state['logged_in']=True
        st.session_state['member_id']=restored_member_id
        st.session_state['is_admin']=True
        st.session_state['google_token']={
            'token':c.token,
            'refresh_token':c.refresh_token,
            'client_id':c.client_id,
            'client_secret':c.client_secret,
            'scopes':list(c.scopes or GOOGLE_SCOPES)
        }
        st.session_state['google_connected_notice']=True

        # Critical: remove Google's one-time code immediately so it cannot be reused.
        st.query_params.clear()
        st.rerun()
        return True
    except Exception as ex:
        st.query_params.clear()
        st.error(f'Google sign-in failed: {ex}')
        st.info(
            'This Google authorization attempt has been cleared. '
            'Return to Google & App Settings and press Connect Advisor Google Account '
            'to begin a completely new connection.'
        )
        return False
def google_disconnect():
    st.session_state.pop('google_token',None)
def calendar_service():
    c=google_credentials();return build('calendar','v3',credentials=c,cache_discovery=False) if c else None
def drive_service():
    c=google_credentials();return build('drive','v3',credentials=c,cache_discovery=False) if c else None
def sheets_service():
    c=google_credentials();return build('sheets','v4',credentials=c,cache_discovery=False) if c else None
def configured_calendar_id():
    c=google_oauth_cfg() or {};return setting('google_calendar_id','') or c.get('calendar_id','')
def _normalize_event_text(value):
    s=str(value or '').casefold().strip()
    s=re.sub(r'[^a-z0-9]+',' ',s)
    return re.sub(r'\s+',' ',s).strip()

def canonical_event_fingerprint(event):
    """Match the same real-world event across Google API, ICS, and local rows."""
    title=_normalize_event_text(event.get('title') or event.get('summary') or '')
    location=_normalize_event_text(event.get('location') or '')
    raw_start=str(event.get('start_at') or event.get('start') or '').strip()

    try:
        ts=pd.to_datetime(raw_start)
        if getattr(ts,'tzinfo',None) is not None:
            ts=ts.tz_convert('America/New_York')
        # All-day feeds may represent the same event as YYYY-MM-DD or midnight.
        if 'T' not in raw_start or ts.strftime('%H:%M')=='00:00':
            start_key=ts.strftime('%Y-%m-%d')
        else:
            start_key=ts.strftime('%Y-%m-%dT%H:%M')
    except Exception:
        start_key=raw_start[:16]

    # Location is retained for collision resistance, but blank location must not create a duplicate
    # when one source simply omits it.
    return (title,start_key,location)

def event_rows_equivalent(a,b):
    fa=canonical_event_fingerprint(a)
    fb=canonical_event_fingerprint(b)
    if fa==fb:
        return True

    # Same normalized title + same normalized start should be treated as one event even if
    # one feed omits location.
    if fa[0] and fa[0]==fb[0] and fa[1]==fb[1]:
        if not fa[2] or not fb[2] or fa[2]==fb[2]:
            return True

    # Google/ICS can disagree by a few minutes because of timezone/rendering conversions.
    # Only relax this when titles match.
    if fa[0] and fa[0]==fb[0]:
        try:
            ta=pd.to_datetime(str(a.get('start_at') or a.get('start') or ''))
            tb=pd.to_datetime(str(b.get('start_at') or b.get('start') or ''))
            if getattr(ta,'tzinfo',None) is not None:
                ta=ta.tz_convert('America/New_York')
            if getattr(tb,'tzinfo',None) is not None:
                tb=tb.tz_convert('America/New_York')
            if abs((ta-tb).total_seconds())<=300:
                return True
        except Exception:
            pass
    return False



def dedupe_event_rows(rows_):
    """Collapse equivalent calendar rows without deleting linked database records."""
    chosen=[]

    def score(r):
        s=0
        if str(r.get('source') or '').lower()=='manual': s+=50
        if r.get('event_type') and r.get('event_type')!='Other': s+=20
        if r.get('classification_note'): s+=10
        if r.get('event_audience'): s+=8
        if r.get('calendar_event_key'): s+=3
        if r.get('google_event_id'): s+=2
        return s

    for row in rows_ or []:
        match_index=None
        for i,current in enumerate(chosen):
            if event_rows_equivalent(row,current):
                match_index=i
                break

        if match_index is None:
            chosen.append(row)
        elif score(row)>score(chosen[match_index]):
            chosen[match_index]=row

    return sorted(chosen,key=lambda x:str(x.get('start_at') or ''))


def sync_google_calendar():
    if not bool(st.session_state.get('logged_in')) or not bool(st.session_state.get('is_admin')):
        raise PermissionError('Only Advisor/Admin may sync the Google calendar.')
    service=calendar_service();cal_id=configured_calendar_id()
    if not service:raise RuntimeError('Connect Google first.')
    if not cal_id:raise RuntimeError('Calendar ID is not configured.')

    items=service.events().list(
        calendarId=cal_id,
        timeMin=datetime.now(timezone.utc).isoformat(),
        singleEvents=True,
        orderBy='startTime',
        maxResults=250
    ).execute().get('items',[])

    existing_rows=table('events').select('*').eq('active',True).execute().data or []
    by_google_id={
        str(r.get('google_event_id') or ''):r
        for r in existing_rows if r.get('google_event_id')
    }
    by_fingerprint={}
    for r in existing_rows:
        by_fingerprint.setdefault(canonical_event_fingerprint(r),r)

    changed=0
    for e in items:
        s=e.get('start',{});en=e.get('end',{})
        payload={
            'google_event_id':e.get('id',''),
            'title':e.get('summary','Untitled Event'),
            'start_at':s.get('dateTime') or s.get('date'),
            'end_at':en.get('dateTime') or en.get('date'),
            'location':e.get('location',''),
            'description':e.get('description',''),
            'source':'google',
            'active':True
        }
        gid=str(payload.get('google_event_id') or '')
        existing=by_google_id.get(gid) if gid else None
        if not existing:
            existing=by_fingerprint.get(canonical_event_fingerprint(payload))
        if not existing:
            existing=next((r for r in existing_rows if event_rows_equivalent(r,payload)),None)

        if existing:
            # Preserve local classification fields on an already-existing event.
            update_payload=dict(payload)
            if existing.get('source')=='manual':
                update_payload['source']='manual'
            table('events').update(update_payload).eq('id',existing['id']).execute()
            by_google_id[gid]=existing
        else:
            inserted=table('events').insert(payload).execute().data or []
            if inserted:
                row=inserted[0]
                by_google_id[gid]=row
                by_fingerprint[canonical_event_fingerprint(row)]=row
        changed+=1
    return changed

def upcoming_events(limit=50):
    data=table('events').select('*').eq('active',True).order('start_at').execute().data or []
    today=date.today().isoformat()
    future=[r for r in data if str(r.get('start_at',''))[:10]>=today]
    return dedupe_event_rows(future)[:limit]
@st.cache_data(ttl=600,show_spinner=False)
def public_google_calendar_events(calendar_id,start_iso,end_iso):
    if not calendar_id:
        return []
    try:
        encoded=urllib.parse.quote(calendar_id,safe='')
        ics_url=f"https://calendar.google.com/calendar/ical/{encoded}/public/basic.ics"
        response=requests.get(ics_url,timeout=12)
        response.raise_for_status()

        cal=Calendar.from_ical(response.content)
        start_dt=pd.Timestamp(start_iso).to_pydatetime()
        end_dt=pd.Timestamp(end_iso).to_pydatetime()

        # Treat the requested window as America/New_York where needed.
        if start_dt.tzinfo is None:
            start_dt=pd.Timestamp(start_dt,tz='America/New_York').to_pydatetime()
        if end_dt.tzinfo is None:
            end_dt=pd.Timestamp(end_dt,tz='America/New_York').to_pydatetime()

        expanded=recurring_ical_events.of(cal).between(start_dt,end_dt)
        out=[]

        for component in expanded:
            try:
                start_value=component.decoded('DTSTART')
                end_value=component.decoded('DTEND') if component.get('DTEND') else None

                if isinstance(start_value,date) and not isinstance(start_value,datetime):
                    start_ts=pd.Timestamp(start_value)
                    all_day=True
                else:
                    start_ts=pd.Timestamp(start_value)
                    all_day=False

                if end_value is not None:
                    end_ts=pd.Timestamp(end_value)
                else:
                    end_ts=None

                uid=str(component.get('UID') or '')
                recurrence_id=str(component.get('RECURRENCE-ID') or '')
                event_key=f"{uid}|{recurrence_id}|{start_ts.isoformat()}"

                out.append({
                    'public_event_key':event_key,
                    'google_event_id':uid,
                    'title':str(component.get('SUMMARY') or 'Untitled Event'),
                    'start_at':start_ts.isoformat(),
                    'end_at':end_ts.isoformat() if end_ts is not None else '',
                    'location':str(component.get('LOCATION') or ''),
                    'description':str(component.get('DESCRIPTION') or ''),
                    'all_day':all_day,
                    'source':'google_public'
                })
            except:
                continue

        out.sort(key=lambda x:str(x.get('start_at') or ''))
        return out
    except:
        return []

def refresh_calendar_cache_if_possible():
    # The embedded Google calendar can display events because the browser is
    # signed into Google even when the public ICS feed is unavailable.
    # When Advisor/Admin has connected Google, refresh those same events into
    # Supabase so every member dashboard can use the cached event list.
    if not google_connected():
        return 0

    now=datetime.now(timezone.utc)
    last=st.session_state.get('last_calendar_auto_sync')
    if last:
        try:
            last_dt=pd.to_datetime(last,utc=True).to_pydatetime()
            if (now-last_dt).total_seconds() < 600:
                return 0
        except:
            pass

    try:
        count=sync_google_calendar()
        st.session_state['last_calendar_auto_sync']=now.isoformat()
        return count
    except:
        return 0

def public_events_next_two_months():
    start=pd.Timestamp(date.today())
    end=start+pd.DateOffset(months=2)
    calendar_id=configured_calendar_id() or '47c2d5153540fd42e274906f301e4fd65a6501c716056066a41341106cedc064@group.calendar.google.com'
    return public_google_calendar_events(
        calendar_id,
        start.isoformat(),
        end.isoformat()
    )

def two_month_events_for_members():
    public_rows=public_events_next_two_months()
    synced_rows=events_next_two_months()

    # Prefer public rows when available, but merge in synced Hub events that
    # are missing so the dashboard never shows zero simply because the public
    # calendar feed is unavailable.
    if not public_rows:
        return synced_rows

    return dedupe_event_rows(list(public_rows)+list(synced_rows))

def calendar_event_date_time(event):
    try:
        start=pd.to_datetime(event.get('start_at'))
        if event.get('all_day'):
            return start.strftime('%A, %B %d, %Y'), 'All Day'
        return start.strftime('%A, %B %d, %Y'), start.strftime('%I:%M %p').lstrip('0')
    except:
        return str(event.get('start_at') or ''), ''

def events_next_two_months():
    start=pd.Timestamp(date.today())
    end=start+pd.DateOffset(months=2)
    out=[]
    for r in upcoming_events(250):
        try:
            d=pd.to_datetime(str(r.get("start_at",""))[:10])
            if start <= d <= end: out.append(r)
        except: pass
    return out

def volunteer_link_for_event(event_id):
    return setting(f"volunteer_sheet_{int(event_id)}","").strip()

def save_volunteer_link_for_event(event_id,url):
    save_setting(f"volunteer_sheet_{int(event_id)}",url.strip())

def drive_upload_bytes(data,filename,mime_type,folder_id):
    if st.session_state.get('demo_admin_mode'):
        raise RuntimeError('ADMIN DEMO MODE: Google/Drive actions are disabled in the sandbox.')
    service=drive_service()
    if not service:
        raise RuntimeError('Connect Google first.')
    if not folder_id:
        raise RuntimeError('The destination Drive folder is not configured.')
    media=MediaIoBaseUpload(BytesIO(data),mimetype=mime_type or 'application/octet-stream',resumable=False)
    metadata={'name':filename,'parents':[folder_id]}
    created=service.files().create(
        body=metadata,
        media_body=media,
        fields='id,name,webViewLink'
    ).execute()
    return created

def officer_reports_drive_folder_id():
    return setting('officer_reports_drive_folder_id','1k0KcbEbDF87YP76OXJG7vwthKQVxOWlR').strip()

def committee_reports_drive_folder_id():
    return setting('committee_reports_drive_folder_id','15IfzQd916c8xbN4cYSXsS3qMYWJ4F8LA').strip()

def report_drive_folder_for_source(source_label):
    return committee_reports_drive_folder_id() if str(source_label).startswith('Committee:') else officer_reports_drive_folder_id()

def storage_download(path):
    if not path:return b''
    try:
        return sb().storage.from_(cfg()['bucket']).download(path)
    except:
        return b''

def drive_files():
    service=drive_service();folder_id=setting('google_drive_folder_id','')
    if not service:raise RuntimeError('Connect Google first.')
    if not folder_id:return []
    q=f"'{folder_id}' in parents and trashed=false";out=[];token=None
    while True:
        r=service.files().list(q=q,fields='nextPageToken, files(id,name,mimeType,modifiedTime,webViewLink)',orderBy='name',pageToken=token,pageSize=100).execute();out.extend(r.get('files',[]));token=r.get('nextPageToken')
        if not token:break
    return out
# Process the Google return BEFORE the Hub login gate. OAuth reloads the
# Streamlit page, so doing this later would discard the initiating Hub session
# and could cause Google's one-time code to be reused.
process_google_oauth_callback()

require_login()
real_member_id=st.session_state['member_id']
real_is_admin=st.session_state.get('is_admin',False)

if st.session_state.get('demo_admin_mode'):
    demo_admin_notice()
    dmember=demo_admin_fake_member()
    member_id=dmember['id']
    is_admin=True
    member=dmember
    member_name=dmember['full_name']
else:
    member_id=real_member_id
    is_admin=real_is_admin

if st.session_state.get('test_preview_kind')=='candidate':
    if not real_is_admin:
        st.session_state.pop('test_preview_kind',None)
    else:
        b1,b2=st.columns([6,1])
        b1.error('🧪 TEST MODE — Candidate Preview. Nothing in this preview is saved or affects the live Hub.')
        if b2.button('Exit Test Mode',use_container_width=True,key='exit_candidate_preview_top'):
            for k in ['test_preview_kind','test_preview_member_id','candidate_test_profile','candidate_preview_tasks']:
                st.session_state.pop(k,None)
            st.rerun()
        render_candidate_test_preview()
        st.stop()

if not st.session_state.get('demo_admin_mode'):
    member_id=real_member_id
    is_admin=real_is_admin
if st.session_state.get('test_preview_kind')=='member':
    if not real_is_admin:
        st.session_state.pop('test_preview_kind',None)
    else:
        preview_mid=st.session_state.get('test_preview_member_id')
        prow=sb().table('members').select('*').eq('id',preview_mid).execute().data or []
        if prow:
            member_id=preview_mid
            is_admin=False
        else:
            st.session_state.pop('test_preview_kind',None)
            st.session_state.pop('test_preview_member_id',None)

if not st.session_state.get('demo_admin_mode'):
    member=table('members').select('*').eq('id',member_id).single().execute().data
    member_name=member['full_name']
    start_or_touch_login_activity(member_id,member_name,is_admin)
else:
    member=demo_admin_fake_member()
    member_name=member['full_name']

if st.session_state.get('test_preview_kind')=='member':
    b1,b2=st.columns([6,1])
    b1.error(f"🧪 TEST MODE — Viewing the Hub as {member_name}. All database/storage changes are blocked.")
    if b2.button('Exit Test Mode',use_container_width=True,key='exit_member_preview_top'):
        st.session_state.pop('test_preview_kind',None)
        st.session_state.pop('test_preview_member_id',None)
        st.session_state.pop('test_preview_last_action',None)
        st.rerun()
    if st.session_state.get('test_preview_last_action'):
        st.info(st.session_state.pop('test_preview_last_action'))

if st.session_state.get('demo_admin_mode') and not st.session_state.get('admin_demo_welcome_acknowledged'):
    st.markdown('## Welcome to the NBS Philo Hub — Admin Demo')
    with st.container(border=True):
        st.markdown('### 🧪 This is a protected demonstration space')
        st.write(
            'This Admin Demo is designed so invited Sorors can experience the Philo Hub from an administrator’s point of view. '
            'You are encouraged to explore every corner of the app, open every dashboard, try the forms and workflows, '
            'review officer tools, test reports, votes, intake, Historian features, Treasurer features, and other available areas.'
        )
        st.success(
            'Anything you do while using Admin Demo will NOT affect the Philos’ live information or documents in any way.'
        )
        st.write(
            'Your demo actions are not saved to live member records, Philo documents, Google Drive files, votes, tasks, '
            'forms, messages, settings, approvals, uploads, financial information, or other live Hub data.'
        )
        st.info(
            'Feel free to click, type, test, approve, reject, vote, upload, edit, and explore. '
            'The purpose of this account is to let you see how the full Hub works without changing the real app.'
        )
        if st.button('I Understand — Enter the Admin Demo',use_container_width=True,key='admin_demo_welcome_continue'):
            st.session_state['admin_demo_welcome_acknowledged']=True
            st.rerun()
    st.stop()

if st.session_state.get('demo_admin_mode'):
    c_demo1,c_demo2=st.columns([6,1])
    c_demo1.warning('🧪 Admin Demo Sandbox is active.')
    if c_demo2.button('Exit Demo',use_container_width=True,key='exit_demo_admin_top'):
        st.session_state.clear()
        st.rerun()

if st.session_state.pop('google_connected_notice',False):
    st.success('Google account connected successfully.')

def fmt_dt(v):
    if not v:return ''
    try:
        d=pd.to_datetime(v);return d.strftime('%b %d, %Y') if len(str(v))<=10 else d.strftime('%b %d, %Y • %I:%M %p')
    except:return str(v)
def reimbursement_pdf(r):
    bio=BytesIO();c=canvas.Canvas(bio,pagesize=letter);y=10.3*inch;c.setFont('Helvetica-Bold',16);c.drawString(.7*inch,y,'NBS Philo Reimbursement / Voucher');y-=.35*inch;c.setFont('Helvetica',9)
    fields=[('Submission #',r.get('id','')),('Submitted By',r.get('submitted_by_name','')),('Form Type',r.get('form_type','')),('Payee',r.get('payee_name','')),('Date',r.get('expense_date','')),('Amount',f"${float(r.get('amount') or 0):,.2f}"),('Category',r.get('category','')),('Event',r.get('event_name','')),('Payment Method',r.get('payment_method','')),('Payment Email',r.get('payment_email','')),('Status',r.get('status','Pending'))]
    for label,val in fields:c.setFont('Helvetica-Bold',9);c.drawString(.7*inch,y,f'{label}:');c.setFont('Helvetica',9);c.drawString(2.2*inch,y,str(val or ''));y-=.22*inch
    y-=.1*inch;c.setFont('Helvetica-Bold',9);c.drawString(.7*inch,y,'Description / Business Purpose:');y-=.2*inch;tx=c.beginText(.7*inch,y);tx.setFont('Helvetica',9)
    for line in textwrap.wrap(str(r.get('description','')),95):tx.textLine(line)
    c.drawText(tx);y=tx.getY()-.25*inch;c.drawString(.7*inch,y,f"Typed Signature: {r.get('signature_name','')}");c.save();bio.seek(0);return bio.getvalue()

st.markdown(
    f"""<div class='philo-header'>
    <img class='brand-logo' src='{PHILO_LOGO_URI}'>
    <div>
      <h1>NBS Philo Hub</h1>
      <p>Members, service, forms, files, and meeting preparation in one place.</p>
      <div class='gold-rule'></div>
    </div>
    <img class='brand-crest' src='{PHILO_CREST_URI}'>
    </div>""",
    unsafe_allow_html=True
)

offices=member_offices(member_id)
president=is_president(member_id)
page=st.session_state.get('page','🏠 Dashboard')
render_top_nav(is_admin,offices)

if page=='🏠 Dashboard':
    st.markdown('<div class="pearl-divider"></div>',unsafe_allow_html=True)
    myprof=safe_rows('member_profiles',member_id=member_id)
    myprof=myprof[0] if myprof else {}
    if birthday_today(myprof):
        st.success(f"🎉 Happy Birthday, {display_member_name(member,myprof)}! Wishing you a wonderful day from your NBS Philo family.")
        if not st.session_state.get('birthday_balloons_shown'):
            st.balloons()
            st.session_state['birthday_balloons_shown']=True
    if is_admin:
        with st.container(border=True):
            st.markdown("### Advisor Controls")
            current_chat=chat_enabled()
            col_a,col_b=st.columns([3,1])
            with col_a:
                st.write("Philo Chat is currently **ON**." if current_chat else "Philo Chat is currently **OFF**.")
                st.caption("Turning chat off hides Messages from members but keeps existing messages stored.")
            with col_b:
                if st.button("Turn Chat Off" if current_chat else "Enable Philo Chat",use_container_width=True):
                    set_chat_enabled(not current_chat)
                    st.rerun()
    photo=profile_photo_url(member_id)
    greet_img,greet_text=st.columns([1,6])
    with greet_img:
        if photo:
            st.image(photo,width=105)
        else:
            st.markdown("### 🌼")
    with greet_text:
        st.markdown(f"## Welcome, {member_name}")
        roles=[x.get('position') for x in offices]
        if roles:
            st.markdown(" ".join([f"<span class='role-pill'>{r}</span>" for r in roles]),unsafe_allow_html=True)
        if president and vacant_offices():
            st.caption("President acting access: "+", ".join(vacant_offices()))
    st.markdown('### My Portal')
    h1,h2,h3,h4=st.columns(4)

    with h1:
        with st.container(border=True):
            st.markdown('#### 👤 My Profile')
            st.caption('Update your profile, photo, privacy choices, and personal information.')
            if st.button('Open My Profile',key='home_my_profile_v35',use_container_width=True):
                st.session_state['page']='👥 Members'
                st.session_state['member_portal_view']='profile'
                st.rerun()

    with h2:
        with st.container(border=True):
            st.markdown('#### 📄 Monthly Reports')
            st.caption('View the monthly reports that have been published for members.')
            if st.button('Open Monthly Reports',key='home_monthly_reports_v35',use_container_width=True):
                st.session_state['page']='👥 Members'
                st.session_state['member_portal_view']='reports'
                st.rerun()

    with h3:
        with st.container(border=True):
            st.markdown('#### 🙏 Prayer Request')
            st.caption('Send a private prayer request to the Chaplain.')
            if st.button('Send Prayer Request',key='home_prayer_request_v35',use_container_width=True):
                st.session_state['page']='👥 Members'
                st.session_state['member_portal_view']='prayer'
                st.rerun()

    with h4:
        with st.container(border=True):
            st.markdown('#### 🧾 My Forms')
            st.caption('Open vouchers, submitted forms, volunteer sign-ups, and service records.')
            if st.button('Open My Forms',key='home_my_forms_v35',use_container_width=True):
                st.session_state['page']='🧾 Forms'
                st.session_state['open_member_forms_home']=True
                st.rerun()

    render_elevator_speech_builder(member_id,member_name)
    st.divider()

    if is_philo_member(member_id):
        render_member_financial_status(member_id,compact=True,key_prefix='home_financial_v35')
        st.divider()

    finished_agendas=visible_published_agendas(member_id)
    finished_reports=visible_non_test_rows(visible_published_reports(member_id))

    if finished_agendas or finished_reports:
        st.markdown('### Published for You')

    if finished_agendas:
        st.markdown('#### Meeting Agendas')
        for ag in finished_agendas[:8]:
            with st.container(border=True):
                st.markdown(f"**{ag.get('title')}**")
                st.caption(f"{ag.get('meeting_date') or ''} • {ag.get('meeting_time') or ''}")
                if ag.get('committee_name'):
                    st.write(f"Committee: {ag.get('committee_name')}")
                with st.expander('View Finished Agenda'):
                    items=agenda_items(ag['id'])
                    for idx,it in enumerate(items,1):
                        st.write(f"{idx}. {it.get('item_text')}")
                    pdf=agenda_pdf(ag,items)
                    st.download_button(
                        'Print / Download Agenda',pdf,
                        file_name=f"NBS_Agenda_{ag['id']}.pdf",
                        mime='application/pdf',
                        key=f"member_agenda_pdf_{ag['id']}",
                        use_container_width=True
                    )

    if finished_reports:
        st.markdown('#### Reports')
        for rr in finished_reports[:10]:
            with st.container(border=True):
                st.markdown(f"**{rr.get('title')}**")
                st.caption(f"{rr.get('report_period') or ''} • Published by {rr.get('published_by_name') or ''}")
                if rr.get('committee_name'):
                    st.write(f"Committee: {rr.get('committee_name')}")
                with st.expander('View Finished Report'):
                    if rr.get('report_text'):
                        st.write(rr.get('report_text') or '')

                    if rr.get('file_path'):
                        furl=signed_url(rr.get('file_path'),3600)
                        fname=rr.get('file_name') or 'Completed Report'
                        st.write(f"**Uploaded file:** {fname}")
                        if furl:
                            st.link_button('Open / Download Uploaded Report',furl,use_container_width=True)
                        else:
                            st.warning('The uploaded report could not be opened right now.')
                    elif rr.get('report_text'):
                        pdf=report_pdf(
                            rr.get('title') or 'Published Report',
                            rr.get('report_period') or '',
                            [('Report',rr.get('report_text') or '')],
                            rr.get('published_by_name') or ''
                        )
                        st.download_button(
                            'Print / Download Report',pdf,
                            file_name=f"NBS_Published_Report_{rr['id']}.pdf",
                            mime='application/pdf',
                            key=f"member_report_pdf_{rr['id']}",
                            use_container_width=True
                        )

    active_links=active_meeting_links_now()
    if active_links:
        for ml in active_links:
            with st.container(border=True):
                c1,c2=st.columns([3,1])
                with c1:
                    st.markdown(f"### {ml.get('title') or 'Philo Meeting'}")
                    try:
                        start_label=pd.to_datetime(ml.get('meeting_start')).strftime('%B %d, %Y • %I:%M %p')
                    except:
                        start_label=ml.get('meeting_start') or ''
                    st.write(f"Meeting time: {start_label}")
                    if ml.get('note'):st.caption(ml.get('note'))
                with c2:
                    st.link_button('Join Meeting',ml.get('meeting_url'),use_container_width=True)

    if is_admin and google_connected():
        refresh_calendar_cache_if_possible()

    two_month=current_next_month_calendar_events()
    posted=sum(1 for e in upcoming_events(100) if volunteer_link_for_event(e['id']))
    mytasks=safe_rows('event_tasks',assigned_member_id=member_id)
    open_tasks=sum(1 for x in mytasks if x.get('status')!='Complete')

    st.markdown('### Your Quick View')
    q1,q2,q3=st.columns(3)
    with q1:
        with st.container(border=True):
            st.markdown('### 📅 Current & Next Month')
            st.markdown(f"## {len(two_month)}")
            st.caption('scheduled events')
            if st.button('View Calendar & Events',key='home_calendar_shortcut',use_container_width=True):
                st.session_state['page']='📅 Calendar & Events'
                st.rerun()
    with q2:
        with st.container(border=True):
            st.markdown('### ✅ My Open Tasks')
            st.markdown(f"## {open_tasks}")
            st.caption('event responsibilities')
            if st.button('Event Responsibilities',key='home_tasks_shortcut',use_container_width=True):
                st.session_state['page']='✅ Events & Tasks'
                st.session_state['show_my_responsibilities']=True
                st.rerun()
    with q3:
        with st.container(border=True):
            st.markdown('### 🙋 Volunteer Sign-Ups')
            st.markdown(f"## {posted}")
            st.caption('posted opportunities')
            if st.button('View Volunteer Sign-Ups',key='home_volunteer_shortcut',use_container_width=True):
                st.session_state['page']='🧾 Forms'
                st.session_state['open_forms_tab']='Volunteer Sign-Up'
                st.rerun()

    st.subheader('Current & Next Month at a Glance')
    if not two_month:
        st.info('No events are currently scheduled during the next two months.')
        if is_admin:
            if google_connected():
                st.caption('Google is connected. Open Calendar → Admin Calendar Tools and press Refresh NBS Calendar Now if events were just added.')
            else:
                st.caption('Connect the Philo Advisor Google account once under Calendar → Admin Calendar Tools so the Hub can cache shared-calendar events for member dashboards.')
    for e in two_month[:10]:
        with st.container(border=True):
            event_date,event_time=calendar_event_date_time(e)
            st.markdown(f"**{e.get('title','Untitled Event')}**")
            st.write(f"📅 {event_date}  •  🕒 {event_time}")
            if e.get('location'):
                st.caption(f"📍 {e.get('location')}")
            # Volunteer links remain available when this event also exists in the Hub.
            if e.get('id'):
                link=volunteer_link_for_event(e['id'])
                if link:
                    st.link_button('Volunteer',link,use_container_width=True)
    render_open_chapter_votes(member_id,member_name)
    st.divider()


elif page=='📅 Calendar & Events':
    st.title('Calendar & Events')
    st.caption('This is the shared NBS Philo calendar. Members do not connect or sync personal calendars.')

    if is_admin and google_connected():
        refresh_calendar_cache_if_possible()

    calendar_id=setting('google_calendar_id','') or '47c2d5153540fd42e274906f301e4fd65a6501c716056066a41341106cedc064@group.calendar.google.com'
    encoded_calendar=urllib.parse.quote(calendar_id, safe='')
    embed_url=(
        'https://calendar.google.com/calendar/embed'
        '?src='+encoded_calendar+
        '&ctz=America%2FNew_York'
        '&mode=MONTH'
        '&showTitle=0'
        '&showNav=1'
        '&showDate=1'
        '&showPrint=0'
        '&showTabs=1'
        '&showCalendars=0'
        '&showTz=0'
    )

    st.markdown('### Live NBS Calendar')
    components.iframe(embed_url,height=720,scrolling=True)

    public_calendar_rows=public_events_next_two_months()
    if not public_calendar_rows:
        st.info('If the calendar appears blank or asks for Google access, the shared NBS calendar still needs to be made viewable to members. The app itself is not connecting anyone\'s personal calendar.')

    if is_admin:
        with st.expander('Admin Calendar Tools',expanded=False):
            st.caption('Paste the shared Google Calendar link or Calendar ID once. The Hub will use it for the live calendar and the current/next-month event list.')
            calendar_link_value=st.text_input(
                'Shared Google Calendar Link or Calendar ID',
                value=setting('google_calendar_link','') or setting('google_calendar_id',''),
                key='admin_calendar_link_input'
            )
            if st.button('Save Calendar Link',use_container_width=True,key='save_calendar_link_v32'):
                parsed_id=extract_google_calendar_id(calendar_link_value)
                save_setting('google_calendar_link',calendar_link_value.strip())
                save_setting('google_calendar_id',parsed_id.strip())
                st.success('Calendar saved. The event list will use this calendar.')
                st.cache_data.clear()
                st.rerun()
            if google_connected():
                if st.button('🔄 Refresh NBS Calendar Now',use_container_width=True):
                    try:
                        n=sync_google_calendar()
                        st.success(f'Synced {n} upcoming NBS events.')
                        st.rerun()
                    except Exception as ex:
                        st.error(f'Calendar sync failed: {ex}')
            else:
                u=google_connect_url() if google_oauth_cfg() else ''
                if u:
                    st.link_button('Connect Philo Advisor Google Account',u,use_container_width=True)

    st.markdown('### Current & Next Month Events')
    two_month=current_next_month_calendar_events()
    if not two_month:
        st.caption('No events are currently scheduled for the current or next month.')
    else:
        for e in two_month:
            with st.container(border=True):
                event_date,event_time=calendar_event_date_time(e)
                st.markdown(f"**{e.get('title','Untitled Event')}**")
                st.write(f"📅 {event_date}  •  🕒 {event_time}")
                if e.get('location'):
                    st.caption(f"📍 {e.get('location')}")
                if e.get('id'):
                    link=volunteer_link_for_event(e['id'])
                    if link:
                        st.link_button('Volunteer',link,use_container_width=True)

elif page=='🗓️ Meeting Center':
    if not can_use_management_centers(member_id,is_admin):
        st.info('Meeting Center is available through Officer and Committee dashboards.')
        if st.button('Return Home',use_container_width=True):
            set_page('🏠 Dashboard')
        st.stop()

    st.title('Meeting Center')
    st.caption('Create and manage working meeting materials here. Members only see the finished published agenda they are authorized to view.')

    managed_committees=active_committees() if (is_admin or president) else chaired_committees(member_id)
    audience_options=['All Philos']+[c.get('committee_name') for c in managed_committees]

    tabs=st.tabs(['Agenda Manager','Meeting Links'])

    with tabs[0]:
        with st.form('agenda_create_v28',clear_on_submit=True):
            audience=st.selectbox('Who should see the finished agenda?',audience_options)
            cobj=next((c for c in managed_committees if c.get('committee_name')==audience),None)
            title=st.text_input(
                'Meeting / Agenda Title',
                value='Philo Affiliate Meeting' if audience=='All Philos' else f'{audience} Committee Meeting'
            )
            meeting_date=st.date_input('Date')
            meeting_time=st.time_input('Time')
            location=st.text_input('Location')
            subject=st.text_input('Subject / Theme')
            selected_items=st.multiselect(
                'Standard Agenda Items',
                AGENDA_STANDARD_ITEMS,
                default=[
                    'Prayer','Call to Order','Roll Call / Attendance',
                    'Adoption of Previous Meeting Minutes','Adoption of Agenda',
                    'Pledge','Committee Reports','Unfinished Business','New Business',
                    'Announcements','Advisor Update','Adjournment','Hymn'
                ]
            )
            extra=st.text_area('Additional Agenda Items (one per line)')
            create_agenda=st.form_submit_button('Create Agenda Draft',use_container_width=True)

        if create_agenda:
            result=table('meeting_agendas').insert({
                'title':title,
                'committee_name':'' if audience=='All Philos' else audience,
                'committee_id':None if audience=='All Philos' else (cobj.get('id') if cobj else None),
                'audience_type':'All Philos' if audience=='All Philos' else 'Committee',
                'meeting_date':meeting_date.isoformat(),
                'meeting_time':meeting_time.strftime('%I:%M %p'),
                'location':location,
                'subject':subject,
                'status':'Draft',
                'created_by_member_id':member_id,
                'created_by_name':member_name,
                'created_at':datetime.now(timezone.utc).isoformat()
            }).execute()
            aid=(result.data or [{}])[0].get('id')
            all_items=list(selected_items)+[x.strip() for x in extra.splitlines() if x.strip()]
            for order,item in enumerate(all_items,1):
                table('agenda_items').insert({
                    'agenda_id':aid,'sort_order':order,'item_text':item,'notes':''
                }).execute()
            st.success('Agenda draft created.')
            st.rerun()

        st.markdown('### Drafts and Published Agendas')
        available=[]
        for a in agenda_rows(False):
            if is_admin or president or a.get('created_by_member_id')==member_id:
                available.append(a)
            elif a.get('committee_id') and any(c.get('id')==a.get('committee_id') for c in chaired_committees(member_id)):
                available.append(a)

        for a in available:
            current_audience='All Philos' if (a.get('audience_type') or 'All Philos')=='All Philos' else a.get('committee_name')
            with st.expander(f"{a.get('title')} • {current_audience} • {a.get('meeting_date')} • {a.get('status')}"):
                items=agenda_items(a['id'])
                editable_audiences=['All Philos']+[c.get('committee_name') for c in managed_committees]
                if current_audience not in editable_audiences:
                    editable_audiences.append(current_audience)
                new_audience=st.selectbox(
                    'Audience',
                    editable_audiences,
                    index=editable_audiences.index(current_audience),
                    key=f"agenda_audience28_{a['id']}"
                )

                for it in items:
                    c1,c2=st.columns([3,2])
                    txt=c1.text_input('Agenda Item',value=it.get('item_text') or '',key=f"agenda_item28_{it['id']}")
                    note=c2.text_input('Notes / Presenter',value=it.get('notes') or '',key=f"agenda_note28_{it['id']}")
                    if st.button('Save Item',key=f"agenda_save_item28_{it['id']}"):
                        table('agenda_items').update({'item_text':txt,'notes':note}).eq('id',it['id']).execute()
                        st.rerun()

                add_item=st.text_input('Add Another Agenda Item',key=f"agenda_add28_{a['id']}")
                if st.button('Add Item',key=f"agenda_add_btn28_{a['id']}") and add_item.strip():
                    nxt=max([int(x.get('sort_order') or 0) for x in items] or [0])+1
                    table('agenda_items').insert({
                        'agenda_id':a['id'],'sort_order':nxt,'item_text':add_item.strip(),'notes':''
                    }).execute()
                    st.rerun()

                c1,c2,c3=st.columns(3)
                if c1.button('Save Audience',key=f"agenda_save_audience28_{a['id']}",use_container_width=True):
                    chosen=next((c for c in managed_committees if c.get('committee_name')==new_audience),None)
                    table('meeting_agendas').update({
                        'audience_type':'All Philos' if new_audience=='All Philos' else 'Committee',
                        'committee_id':None if new_audience=='All Philos' else (chosen.get('id') if chosen else None),
                        'committee_name':'' if new_audience=='All Philos' else new_audience
                    }).eq('id',a['id']).execute()
                    st.rerun()

                if c2.button('Publish Finished Agenda',key=f"agenda_publish28_{a['id']}",use_container_width=True):
                    table('meeting_agendas').update({
                        'status':'Published',
                        'published_at':datetime.now(timezone.utc).isoformat()
                    }).eq('id',a['id']).execute()
                    st.success('Finished agenda published to the selected audience.')
                    st.rerun()

                pdf=agenda_pdf(a,items)
                c3.download_button(
                    'Print Agenda',pdf,
                    file_name=f"NBS_Agenda_{a['id']}.pdf",
                    mime='application/pdf',
                    key=f"agenda_print28_{a['id']}",
                    use_container_width=True
                )

    with tabs[1]:
        if not can_manage_meeting_links(member_id,is_admin):
            st.info('You do not currently have permission to manage meeting links.')
        else:
            st.caption('Schedule a Join Meeting button for All Philos or for one committee.')

            with st.form('meeting_link_create_v28',clear_on_submit=True):
                audience=st.selectbox('Who should see the Join Meeting button?',audience_options,key='meeting_link_audience28')
                link_cobj=next((c for c in managed_committees if c.get('committee_name')==audience),None)
                title=st.text_input('Meeting Title',value='Philo Affiliate Meeting')
                url=st.text_input('Zoom / Meeting Link')
                c1,c2=st.columns(2)
                with c1:
                    meeting_date=st.date_input('Meeting Date')
                    meeting_time=st.time_input('Meeting Time')
                with c2:
                    visible_date=st.date_input('Button Appears On',value=date.today())
                    visible_time=st.time_input('Button Appears At')
                c3,c4=st.columns(2)
                with c3:
                    hide_date=st.date_input('Button Disappears On')
                    hide_time=st.time_input('Button Disappears At')
                with c4:
                    note=st.text_area('Optional Member Note')
                create_link=st.form_submit_button('Schedule Join Meeting Button',use_container_width=True)

            if create_link:
                visible_dt=datetime.combine(visible_date,visible_time).replace(tzinfo=timezone.utc)
                hide_dt=datetime.combine(hide_date,hide_time).replace(tzinfo=timezone.utc)
                meeting_dt=datetime.combine(meeting_date,meeting_time).replace(tzinfo=timezone.utc)
                if not url.strip():
                    st.warning('Enter a meeting link.')
                elif hide_dt<=visible_dt:
                    st.error('The disappear time must be after the button appears.')
                else:
                    table('meeting_links').insert({
                        'title':title.strip() or 'Philo Affiliate Meeting',
                        'meeting_url':url.strip(),
                        'meeting_start':meeting_dt.isoformat(),
                        'visible_from':visible_dt.isoformat(),
                        'visible_until':hide_dt.isoformat(),
                        'note':note.strip(),
                        'active':True,
                        'audience_type':'All Philos' if audience=='All Philos' else 'Committee',
                        'committee_id':None if audience=='All Philos' else (link_cobj.get('id') if link_cobj else None),
                        'committee_name':'' if audience=='All Philos' else audience,
                        'created_by_member_id':member_id,
                        'created_by_name':member_name,
                        'created_at':datetime.now(timezone.utc).isoformat()
                    }).execute()
                    st.success('Meeting button scheduled.')
                    st.rerun()

            st.markdown('### Scheduled Meeting Buttons')
            for ml in all_meeting_links():
                aud='All Philos' if (ml.get('audience_type') or 'All Philos')=='All Philos' else ml.get('committee_name')
                with st.expander(f"{ml.get('title')} • {aud} • {'Active' if ml.get('active') else 'Off'}"):
                    st.write(f"Visible: {ml.get('visible_from')} through {ml.get('visible_until')}")
                    enabled=st.checkbox('Enabled',value=bool(ml.get('active',True)),key=f"meeting_link_enabled28_{ml['id']}")
                    c1,c2=st.columns(2)
                    if c1.button('Save Status',key=f"meeting_link_save28_{ml['id']}"):
                        table('meeting_links').update({'active':enabled}).eq('id',ml['id']).execute()
                        st.rerun()
                    confirm=st.checkbox('Confirm delete',key=f"meeting_link_delete_confirm28_{ml['id']}")
                    if c2.button('Delete',key=f"meeting_link_delete28_{ml['id']}",disabled=not confirm):
                        table('meeting_links').delete().eq('id',ml['id']).execute()
                        st.rerun()

            if is_admin:
                st.divider()
                st.subheader('Meeting Link Permissions')
                current=meeting_link_allowed_positions()
                allowed_positions=st.multiselect(
                    'Officer positions that may manage meeting links',
                    OFFICER_POSITIONS,
                    default=[x for x in current if x in OFFICER_POSITIONS]
                )
                if st.button('Save Meeting Link Permissions',use_container_width=True):
                    save_setting('meeting_link_positions',",".join(allowed_positions))
                    st.rerun()
elif page=='📝 Reports':
    if not can_use_management_centers(member_id,is_admin):
        st.info('Report creation is available through Officer and Committee dashboards.')
        if st.button('Return Home',use_container_width=True):
            set_page('🏠 Dashboard')
        st.stop()

    st.title('Report Center')
    my_positions=[x.get('position') for x in offices]
    if president:
        my_positions += [x for x in vacant_offices() if x not in my_positions]

    my_chairs=chaired_committees(member_id)
    report_sources=[f"Officer: {p}" for p in my_positions]
    report_sources += [f"Committee: {c.get('committee_name')}" for c in my_chairs]

    if is_admin:
        for p in OFFICER_POSITIONS:
            if f"Officer: {p}" not in report_sources:
                report_sources.append(f"Officer: {p}")
        for c in active_committees():
            label=f"Committee: {c.get('committee_name')}"
            if label not in report_sources:
                report_sources.append(label)

    if not report_sources:
        st.info('No officer or committee report area is assigned to you.')
        st.stop()

    tabs=st.tabs(['Create Full Report','End-of-Year','Published Report Manager'])

    with tabs[0]:
        prefill_source=st.session_state.get('report_prefill_source','')
        source_index=report_sources.index(prefill_source) if prefill_source in report_sources else 0
        source=st.selectbox('Report For',report_sources,index=source_index)

        author_ids=report_author_options(member_id,is_admin)
        author_map={aid:report_author_info(aid) for aid in author_ids}
        default_author=member_id if member_id in author_ids else author_ids[0]
        report_author_id=st.selectbox(
            'Official Report Name / Letterhead Person',
            author_ids,
            index=author_ids.index(default_author),
            format_func=lambda x:author_map[x]['name'],
            help='Admin may generate a report on behalf of another Philo. The selected person’s name and email appear on the official chapter template.'
        )
        report_author=author_map[report_author_id]
        chapter_meeting_date=st.date_input('Chapter Meeting Date',value=date.today(),key='official_chapter_meeting_date')
        if st.session_state.get('report_prefill_updates'):
            st.success('Reviewed officer-document extraction loaded. Edit anything you want; you do not need to retype the uploaded report.')

        month=st.selectbox(
            'Reporting Month',
            [f"{datetime.now().year}-{m:02d}" for m in range(1,13)],
            index=datetime.now().month-1
        )

        accomplishments=st.text_area('Accomplishments',height=100)
        challenges=st.text_area('Updates / concerns / unfinished business',value=st.session_state.get('report_prefill_updates',''),height=100)
        upcoming=st.text_area('Upcoming Events / plans',height=90)
        recommendations=st.text_area('Reminders / recommendations',height=90)

        queued_updates=leadership_report_queue(report_author_id)
        queued_texts=[leadership_update_text(q.get('update_id')) for q in queued_updates if leadership_update_text(q.get('update_id'))]
        action_default='\n'.join(queued_texts)
        action_items=st.text_area(
            'Action Items',
            value=action_default,
            height=100,
            help='Advisor/leadership updates that were added to the report queue appear here automatically. Edit the wording however you want.'
        )

        if st.button('Generate Full Report Draft',use_container_width=True):
            extra_parts=[]
            if accomplishments.strip():extra_parts.append("Accomplishments: "+accomplishments.strip())
            if challenges.strip():extra_parts.append("Challenges/Unfinished Business: "+challenges.strip())
            if upcoming.strip():extra_parts.append("Next Month Priorities: "+upcoming.strip())
            if recommendations.strip():extra_parts.append("Recommendations/Assistance Needed: "+recommendations.strip())
            if action_items.strip():extra_parts.append("Action Items: "+action_items.strip())
            extra="\n".join(extra_parts)

            if source.startswith('Officer: '):
                position=source.replace('Officer: ','',1)
                draft=polished_position_report(position,member_id,member_name,month,extra)
                title=f"{position} Monthly Report"
                committee_id=None
                committee_name=''
            else:
                cname=source.replace('Committee: ','',1)
                committee=next((c for c in active_committees() if c.get('committee_name')==cname),None)
                draft=committee_report_text(committee,month,extra) if committee else extra
                title=f"{cname} Committee Report"
                committee_id=committee.get('id') if committee else None
                committee_name=cname

            st.session_state['managed_report_draft']=draft
            st.session_state['managed_report_title']=title
            st.session_state['managed_report_committee_id']=committee_id
            st.session_state['managed_report_committee_name']=committee_name
            st.session_state['managed_report_period']=month

        draft=st.text_area(
            'Editable Full Report Draft',
            value=st.session_state.get('managed_report_draft',''),
            height=620
        )

        if draft.strip():
            title=st.session_state.get('managed_report_title','NBS Report')
            period=st.session_state.get('managed_report_period',month)
            report_committee_id=st.session_state.get('managed_report_committee_id')
            report_committee_name=st.session_state.get('managed_report_committee_name','')

            office_label=source.replace('Officer: ','').replace('Committee: ','')
            official_pdf=official_chapter_report_pdf(
                office_or_committee=office_label,
                author_name=report_author['name'],
                author_email=report_author['email'],
                meeting_date=chapter_meeting_date.strftime('%B %-d, %Y') if os.name!='nt' else chapter_meeting_date.strftime('%B %d, %Y').replace(' 0',' '),
                accomplishments=accomplishments,
                updates=challenges,
                upcoming=upcoming,
                reminders=recommendations,
                action_items=action_items,
                full_report=draft
            )
            st.download_button(
                '📄 Download / Print Official Chapter Report',
                official_pdf,
                file_name=f"{title.replace(' ','_')}_{period}_Official.pdf",
                mime='application/pdf',
                use_container_width=True
            )
            with st.expander('Plain Hub Draft PDF'):
                plain_pdf=report_pdf(title,f"Reporting Period: {period} • Prepared by {report_author['name']}",[('Report',draft)],report_author['name'])
                st.download_button(
                    'Download Plain Draft PDF',
                    plain_pdf,
                    file_name=f"{title.replace(' ','_')}_{period}_Plain.pdf",
                    mime='application/pdf',
                    use_container_width=True,
                    key='plain_report_pdf_623'
                )

            st.markdown('### Publish Finished Report')
            audience_candidates=['All Philos']
            if is_admin or president:
                audience_candidates += [c.get('committee_name') for c in active_committees()]
            else:
                audience_candidates += [c.get('committee_name') for c in my_chairs]

            if report_committee_name and report_committee_name not in audience_candidates:
                audience_candidates.append(report_committee_name)

            audience=st.selectbox('Who should see the finished report?',audience_candidates,key='publish_report_audience')
            generated_quick_drive=st.checkbox('Quick Add this finished report to the proper Google Drive folder',value=False,key='generated_quick_drive')
            if st.button('Publish Finished Report',use_container_width=True):
                target=next((c for c in active_committees() if c.get('committee_name')==audience),None)
                table('published_reports').insert({
                    'title':title,
                    'report_period':period,
                    'report_text':draft,
                    'source_type':'Generated',
                    'audience_type':'All Philos' if audience=='All Philos' else 'Committee',
                    'committee_id':None if audience=='All Philos' else (target.get('id') if target else None),
                    'committee_name':'' if audience=='All Philos' else audience,
                    'published_by_member_id':member_id,
                    'published_by_name':member_name,
                    'published_at':datetime.now(timezone.utc).isoformat(),
                    'active':True
                }).execute()
                drive_message=''
                if generated_quick_drive:
                    try:
                        folder_id=report_drive_folder_for_source(source)
                        drive_name=f"{title.replace(' ','_')}_{period}.pdf"
                        created_drive=drive_upload_bytes(official_pdf,drive_name,'application/pdf',folder_id)
                        drive_message=f" Added to Google Drive as {created_drive.get('name',drive_name)}."
                    except Exception as ex:
                        drive_message=f" Published in the Hub, but Drive upload needs attention: {ex}"
                st.success('Finished report published to the selected audience.'+drive_message)
                st.rerun()

    with tabs[1]:
        year_label=st.text_input('Reporting Year / Sorority Year',value=setting('current_fiscal_year','2026-2027'))
        source=st.selectbox('Year-End Report For',report_sources,key='eoy_source28')
        extra=st.text_area('Year-end highlights, challenges, recommendations, or other information',height=150)

        if st.button('Generate Full Year-End Draft',use_container_width=True):
            if source.startswith('Officer: '):
                position=source.replace('Officer: ','',1)
                draft=polished_position_report(position,member_id,member_name,year_label,extra)
                title=f"{position} End-of-Year Report"
                cname=''
            else:
                cname=source.replace('Committee: ','',1)
                committee=next((c for c in active_committees() if c.get('committee_name')==cname),None)
                draft=committee_report_text(committee,year_label,extra) if committee else extra
                title=f"{cname} Committee End-of-Year Report"
            st.session_state['managed_eoy_draft']=draft
            st.session_state['managed_eoy_title']=title
            st.session_state['managed_eoy_period']=year_label
            st.session_state['managed_eoy_committee_name']=cname

        edraft=st.text_area(
            'Editable Year-End Draft',
            value=st.session_state.get('managed_eoy_draft',''),
            height=650
        )
        if edraft.strip():
            etitle=st.session_state.get('managed_eoy_title','End-of-Year Report')
            eperiod=st.session_state.get('managed_eoy_period',year_label)
            pdf=report_pdf(etitle,eperiod,[('Year in Review',edraft)],member_name)
            st.download_button(
                'Download / Print Year-End Report',pdf,
                file_name=f"{etitle.replace(' ','_')}_{eperiod}.pdf",
                mime='application/pdf',
                use_container_width=True
            )

            audience_candidates=['All Philos']
            if is_admin or president:
                audience_candidates += [c.get('committee_name') for c in active_committees()]
            else:
                audience_candidates += [c.get('committee_name') for c in my_chairs]
            ecname=st.session_state.get('managed_eoy_committee_name','')
            if ecname and ecname not in audience_candidates:
                audience_candidates.append(ecname)

            audience=st.selectbox('Who should see the finished year-end report?',audience_candidates,key='publish_eoy_audience28')
            eoy_quick_drive=st.checkbox('Quick Add this year-end report to the proper Google Drive folder',value=False,key='eoy_quick_drive')
            if st.button('Publish Finished Year-End Report',use_container_width=True):
                target=next((c for c in active_committees() if c.get('committee_name')==audience),None)
                table('published_reports').insert({
                    'title':etitle,
                    'report_period':eperiod,
                    'report_text':edraft,
                    'source_type':'Generated',
                    'audience_type':'All Philos' if audience=='All Philos' else 'Committee',
                    'committee_id':None if audience=='All Philos' else (target.get('id') if target else None),
                    'committee_name':'' if audience=='All Philos' else audience,
                    'published_by_member_id':member_id,
                    'published_by_name':member_name,
                    'published_at':datetime.now(timezone.utc).isoformat(),
                    'active':True
                }).execute()
                drive_message=''
                if eoy_quick_drive:
                    try:
                        folder_id=report_drive_folder_for_source(source)
                        drive_name=f"{etitle.replace(' ','_')}_{eperiod}.pdf"
                        created_drive=drive_upload_bytes(pdf,drive_name,'application/pdf',folder_id)
                        drive_message=f" Added to Google Drive as {created_drive.get('name',drive_name)}."
                    except Exception as ex:
                        drive_message=f" Published in the Hub, but Drive upload needs attention: {ex}"
                st.success('Year-end report published.'+drive_message)
                st.rerun()

        if is_admin:
            st.divider()
            st.subheader('Advisor / Chapter Year-in-Review')
            advisor_extra=st.text_area('Advisor highlights / context',height=130)
            if st.button('Generate Advisor Year-in-Review Draft',use_container_width=True):
                adraft=advisor_year_summary(year_label)
                if advisor_extra.strip():
                    adraft+="\n\nADVISOR HIGHLIGHTS\n"+advisor_extra.strip()
                st.session_state['advisor_eoy_draft']=adraft
            adraft=st.text_area('Editable Advisor Draft',value=st.session_state.get('advisor_eoy_draft',''),height=650)
            if adraft.strip():
                apdf=report_pdf('Advisor / Chapter End-of-Year Report',year_label,[('Chapter Year in Review',adraft)],member_name)
                st.download_button('Download Advisor Report',apdf,file_name=f"NBS_Advisor_End_of_Year_{year_label}.pdf",mime='application/pdf',use_container_width=True)

    with tabs[2]:
        st.subheader('Upload & Publish a Completed Report')
        st.caption('Use this when the final report was completed outside the Hub. Publish the uploaded PDF or Word file to All Philos or to one committee.')

        upload_audiences=['All Philos']
        if is_admin or president:
            upload_audiences += [c.get('committee_name') for c in active_committees()]
        else:
            upload_audiences += [c.get('committee_name') for c in my_chairs]

        with st.form('upload_completed_report_v284',clear_on_submit=True):
            uploaded_report=st.file_uploader('Completed Report',type=['pdf','docx'])
            upload_source=st.selectbox('Report Belongs To',report_sources,key='uploaded_report_source')
            upload_title=st.text_input('Report Title',placeholder='Example: Recording Secretary Monthly Report')
            upload_period=st.text_input('Reporting Period',placeholder='Example: August 2026 or 2026–2027')
            upload_audience=st.selectbox('Who should see this finished report?',upload_audiences)
            quick_drive=st.checkbox('Quick Add to the proper Google Drive reports folder',value=False)
            publish_upload=st.form_submit_button('Upload & Publish Finished Report',use_container_width=True)

        if publish_upload:
            if not uploaded_report:
                st.warning('Choose a PDF or Word report to upload.')
            elif not upload_title.strip():
                st.warning('Enter a report title.')
            else:
                target=next((c for c in active_committees() if c.get('committee_name')==upload_audience),None)
                safe_name=re.sub(r'[^A-Za-z0-9._-]+','_',uploaded_report.name)
                path=upload_private(
                    bytes(uploaded_report.getbuffer()),
                    f"published_reports/{member_id}/{datetime.now().strftime('%Y%m%d%H%M%S')}_{safe_name}",
                    uploaded_report.type or 'application/octet-stream'
                )
                if path:
                    table('published_reports').insert({
                        'title':upload_title.strip(),
                        'report_period':upload_period.strip(),
                        'report_text':'',
                        'file_path':path,
                        'file_name':uploaded_report.name,
                        'mime_type':uploaded_report.type or 'application/octet-stream',
                        'source_type':'Uploaded',
                        'audience_type':'All Philos' if upload_audience=='All Philos' else 'Committee',
                        'committee_id':None if upload_audience=='All Philos' else (target.get('id') if target else None),
                        'committee_name':'' if upload_audience=='All Philos' else upload_audience,
                        'published_by_member_id':member_id,
                        'published_by_name':member_name,
                        'published_at':datetime.now(timezone.utc).isoformat(),
                        'active':True
                    }).execute()

                    drive_message=''
                    if quick_drive:
                        try:
                            folder_id=report_drive_folder_for_source(upload_source)
                            created_drive=drive_upload_bytes(
                                bytes(uploaded_report.getbuffer()),
                                uploaded_report.name,
                                uploaded_report.type or 'application/octet-stream',
                                folder_id
                            )
                            drive_message=f" Added to Google Drive as {created_drive.get('name',uploaded_report.name)}."
                        except Exception as ex:
                            drive_message=f" Report published in the Hub, but Drive upload needs attention: {ex}"

                    st.success('Completed report uploaded and published to the selected audience.'+drive_message)
                    st.rerun()

        st.divider()
        st.subheader('Published Reports')
        managed=[]
        for rr in published_reports():
            if is_admin or president or rr.get('published_by_member_id')==member_id:
                managed.append(rr)
            elif rr.get('committee_id') and any(c.get('id')==rr.get('committee_id') for c in my_chairs):
                managed.append(rr)

        if not managed:
            st.info('No published reports to manage.')

        for rr in managed:
            aud='All Philos' if (rr.get('audience_type') or 'All Philos')=='All Philos' else rr.get('committee_name')
            with st.expander(f"{rr.get('title')} • {aud} • {rr.get('report_period')}"):
                if rr.get('report_text'):
                    st.write(rr.get('report_text') or '')
                if rr.get('file_path'):
                    st.write(f"**Uploaded file:** {rr.get('file_name') or 'Completed Report'}")
                    manager_url=signed_url(rr.get('file_path'),3600)
                    if manager_url:
                        st.link_button('Open Uploaded Report',manager_url,use_container_width=True)
                active=st.checkbox('Published / Visible',value=bool(rr.get('active',True)),key=f"pub_report_active28_{rr['id']}")
                c1,c2=st.columns(2)
                if c1.button('Save Visibility',key=f"pub_report_save28_{rr['id']}"):
                    table('published_reports').update({'active':active}).eq('id',rr['id']).execute()
                    st.rerun()
                confirm=st.checkbox('Confirm delete',key=f"pub_report_delete_confirm28_{rr['id']}")
                if c2.button('Delete Published Copy',key=f"pub_report_delete28_{rr['id']}",disabled=not confirm):
                    table('published_reports').delete().eq('id',rr['id']).execute()
                    st.rerun()
elif page=='✅ Events & Tasks':
    st.title('Events & Task Board')
    can_manage=is_admin or bool(offices)
    recording_secretary=any(x.get('position')=='Recording Secretary' for x in offices)
    can_classify_calendar=is_admin or recording_secretary

    my_responsibilities=[
        x for x in safe_rows('event_tasks',assigned_member_id=member_id)
        if x.get('status')!='Complete'
    ]
    if st.session_state.pop('show_my_responsibilities',False) or my_responsibilities:
        with st.expander('✅ My Event Responsibilities',expanded=True):
            if not my_responsibilities:
                st.success('You have no open event responsibilities.')
            for task in my_responsibilities:
                event_rows=safe_rows('events',id=task.get('event_id')) if task.get('event_id') else []
                task_event=event_rows[0] if event_rows else {}
                with st.container(border=True):
                    st.markdown(f"**{task.get('task')}**")
                    if task_event:
                        st.caption(f"{task_event.get('title') or 'Event'} • Due: {task.get('due_date') or 'No due date'}")
                    else:
                        st.caption(f"Due: {task.get('due_date') or 'No due date'}")
                    if task.get('notes'):
                        st.write(f"**Instructions:** {task.get('notes')}")
                    current_status=task.get('status') or 'Not Started'
                    statuses=['Not Started','In Progress','Complete']
                    personal_status=st.selectbox(
                        'Status',
                        statuses,
                        index=statuses.index(current_status) if current_status in statuses else 0,
                        key=f"myresp_status_{task['id']}"
                    )
                    personal_note=st.text_area(
                        'Progress Notes',
                        value=task.get('completion_notes') or '',
                        key=f"myresp_note_{task['id']}"
                    )
                    st.caption(
                        'Examples: “Called venue — waiting for confirmation.” • '
                        '“Purchased plates; receipt uploaded.” • '
                        '“Spoke with chair and confirmed delivery time.” • '
                        '“Completed and delivered to the event chair.”'
                    )
                    if st.button('Save My Task Update',key=f"myresp_save_{task['id']}",use_container_width=True):
                        table('event_tasks').update({
                            'status':personal_status,
                            'completion_notes':personal_note.strip(),
                            'completed_at':datetime.now(timezone.utc).isoformat() if personal_status=='Complete' else None,
                            'updated_by_member_id':member_id,
                            'updated_by_name':member_name,
                            'updated_at':datetime.now(timezone.utc).isoformat()
                        }).eq('id',task['id']).execute()
                        st.success('Task updated.')
                        st.rerun()

    if can_classify_calendar:
        with st.expander('📅 Calendar Event Classification',expanded=False):
            st.caption('Current and next month Google Calendar events are added to the Event Board. Classify who the event is for before assigning tasks.')
            if st.button('Refresh Calendar Events into Event Board',use_container_width=True,key='refresh_calendar_into_events_v32'):
                if is_admin and google_connected():
                    try:sync_google_calendar()
                    except:pass
                n=ensure_current_calendar_events_in_board()
                st.success(f'Refreshed {n} current/next-month calendar event(s).')
                st.rerun()

            cal_events=current_next_month_calendar_events()
            if not cal_events:
                st.info('No current/next-month calendar events found.')
            for ce in cal_events:
                gid=ce.get('google_event_id') or ''
                candidates=[]
                if gid:
                    candidates=safe_rows('events',google_event_id=gid)
                if not candidates and ce.get('public_event_key'):
                    candidates=safe_rows('events',calendar_event_key=ce.get('public_event_key'))
                event_row=candidates[0] if candidates else None
                with st.container(border=True):
                    d_,tm_=calendar_event_date_time(ce)
                    st.markdown(f"**{ce.get('title','Event')}**")
                    st.caption(f"{d_} • {tm_}")
                    if not event_row:
                        st.caption('Press Refresh Calendar Events into Event Board to classify this event.')
                        continue
                    current_type=event_row.get('event_type') or 'Other'
                    event_type=st.selectbox(
                        'Event Type',
                        EVENT_TYPE_OPTIONS,
                        index=EVENT_TYPE_OPTIONS.index(current_type) if current_type in EVENT_TYPE_OPTIONS else len(EVENT_TYPE_OPTIONS)-1,
                        key=f"event_type_v34_{event_row['id']}"
                    )

                    current_audience=event_row.get('event_audience') or 'Philos & Sorors'
                    audience=st.selectbox(
                        'Event Group',
                        EVENT_AUDIENCE_OPTIONS,
                        index=EVENT_AUDIENCE_OPTIONS.index(current_audience) if current_audience in EVENT_AUDIENCE_OPTIONS else 2,
                        key=f"event_audience_v32_{event_row['id']}"
                    )
                    org=''
                    if audience=='Other Organization':
                        current_org=event_row.get('external_organization') or D9_ORGANIZATIONS[0]
                        org_choice=st.selectbox(
                            'Organization',
                            D9_ORGANIZATIONS,
                            index=D9_ORGANIZATIONS.index(current_org) if current_org in D9_ORGANIZATIONS else len(D9_ORGANIZATIONS)-1,
                            key=f"event_org_v32_{event_row['id']}"
                        )
                        if org_choice=='Other — Type In':
                            org=st.text_input(
                                'Organization Name',
                                value='' if current_org in D9_ORGANIZATIONS else current_org,
                                key=f"event_org_other_v32_{event_row['id']}"
                            ).strip()
                        else:
                            org=org_choice
                    notes=st.text_input(
                        'Classification Note',
                        value=event_row.get('classification_note') or '',
                        key=f"event_class_note_v32_{event_row['id']}"
                    )
                    if st.button('Save Event Classification',key=f"save_event_class_v32_{event_row['id']}",use_container_width=True):
                        table('events').update({
                            'event_type':event_type,
                            'event_audience':audience,
                            'external_organization':org if audience=='Other Organization' else '',
                            'classification_note':notes.strip(),
                            'classified_by_member_id':member_id,
                            'classified_by_name':member_name,
                            'classified_at':datetime.now(timezone.utc).isoformat()
                        }).eq('id',event_row['id']).execute()
                        st.success('Event classification saved.')
                        st.rerun()

    if can_manage:
        with st.expander('➕ Create an Event',expanded=False):
            with st.form('officer_create_event',clear_on_submit=True):
                title=st.text_input('Event Name')
                event_date=st.date_input('Event Date',value=date.today())
                start_time=st.time_input('Start Time',value=datetime.now().replace(hour=10,minute=0,second=0,microsecond=0).time())
                end_time=st.time_input('End Time',value=datetime.now().replace(hour=12,minute=0,second=0,microsecond=0).time())
                location=st.text_input('Location')
                description=st.text_area('Event Description / Notes')
                manual_event_type=st.selectbox(
                    'Event Type',
                    EVENT_TYPE_OPTIONS,
                    index=0
                )
                manual_audience=st.selectbox('Event Group',EVENT_AUDIENCE_OPTIONS,index=2)
                manual_org=''
                if manual_audience=='Other Organization':
                    manual_org_choice=st.selectbox('Organization',D9_ORGANIZATIONS)
                    manual_org=st.text_input('Other Organization Name') if manual_org_choice=='Other — Type In' else manual_org_choice
                create=st.form_submit_button('Create Event',use_container_width=True)
            if create and title.strip():
                start_dt=datetime.combine(event_date,start_time)
                end_dt=datetime.combine(event_date,end_time)
                table('events').insert({
                    'title':title.strip(),
                    'start_at':start_dt.isoformat(),
                    'end_at':end_dt.isoformat(),
                    'location':location.strip(),
                    'description':description.strip(),
                    'event_type':manual_event_type,
                    'event_audience':manual_audience,
                    'external_organization':manual_org.strip() if isinstance(manual_org,str) else '',
                    'source':'manual',
                    'active':True,
                    'created_by_member_id':member_id,
                    'created_by_name':member_name
                }).execute()
                st.success('Event created.')
                st.rerun()

    events=upcoming_events(150)
    if not events:
        st.info('No upcoming events.')
    else:
        em={e['id']:f"{fmt_dt(e.get('start_at'))} — {e['title']}" for e in events}
        eid=st.selectbox('Event',list(em),format_func=lambda x:em[x])
        selected_event=next((e for e in events if e['id']==eid),None)
        if selected_event:
            st.caption(
                f"Event Type: {event_type_label(selected_event)} • "
                f"Event Group: {event_org_display(selected_event)}"
            )
            if selected_event.get('classification_note'):
                st.caption(selected_event.get('classification_note'))

            st.markdown('### 🙋 Attendance / RSVP')
            current_rsvp=current_member_rsvp(eid,member_id)
            current_status=current_rsvp.get('status') if current_rsvp else 'No Response'
            st.caption(f"Your current response: **{current_status}**")
            a1,a2,a3=st.columns(3)
            if a1.button('✅ I’m Attending',key=f"rsvp_yes_{eid}",use_container_width=True):save_member_rsvp(eid,member_id,member_name,'Attending');st.rerun()
            if a2.button('🤔 Maybe',key=f"rsvp_maybe_{eid}",use_container_width=True):save_member_rsvp(eid,member_id,member_name,'Maybe');st.rerun()
            if a3.button('❌ Not Attending',key=f"rsvp_no_{eid}",use_container_width=True):save_member_rsvp(eid,member_id,member_name,'Not Attending');st.rerun()
            attendees=event_attendee_names(eid)
            if attendees:st.markdown(f"**Attending ({len(attendees)}):** "+", ".join(attendees))
            else:st.caption('No one has marked Attending yet.')

            if event_type_label(selected_event)=='Community Service':
                st.markdown('### ⏱️ Community Service Clock')
                current_service=open_service_session(member_id,eid)

                if current_service:
                    try:
                        start_label=pd.to_datetime(
                            current_service.get('check_in'),
                            utc=True
                        ).tz_convert('America/New_York').strftime('%I:%M %p')
                    except:
                        start_label=fmt_dt(current_service.get('check_in'))
                    st.success(f'You are currently clocked in. Start time: {start_label}')

                clock1,clock2=st.columns(2)
                if clock1.button(
                    '🟢 Clock In',
                    disabled=current_service is not None,
                    key=f"service_clock_in_{eid}",
                    use_container_width=True
                ):
                    try:
                        clock_in_service_event(eid)
                        st.success('Clocked in for community service.')
                        st.rerun()
                    except Exception as ex:
                        st.error(str(ex))

                if clock2.button(
                    '🔴 Clock Out',
                    disabled=current_service is None,
                    key=f"service_clock_out_{eid}",
                    use_container_width=True
                ):
                    try:
                        hours=clock_out_service_event(eid)
                        st.success(f'Clocked out. {hours:.2f} service hour(s) recorded.')
                        st.rerun()
                    except Exception as ex:
                        st.error(str(ex))

                my_event_hours=sum(
                    float(x.get('hours') or 0)
                    for x in service_history_for_event(eid)
                    if x.get('member_id')==member_id
                )
                my_total_hours=member_service_hours_total(member_id)
                h1,h2=st.columns(2)
                h1.metric('My Hours for This Event',f"{my_event_hours:.2f}")
                h2.metric('My Total Service Hours',f"{my_total_hours:.2f}")

        if selected_event and event_type_label(selected_event)=='Community Service':
            geo=event_geo_settings(eid)
            if geo and geo.get('active'):
                st.markdown('### 📍 Location-Assisted Clock-In')
                st.caption(f"Venue: {geo.get('location_label') or selected_event.get('location') or 'Configured location'} • Radius: {int(geo.get('radius_meters') or 250)} meters")
                st.info('Background automatic tracking is not enabled. The location/radius is stored so a phone-browser location verification button can be connected safely.')

        if can_manage and selected_event:
            with st.expander('📍 Configure Community Service Location',expanded=False):
                geo=event_geo_settings(eid) or {}
                with st.form(f'event_geo_{eid}'):
                    glabel=st.text_input('Venue / Location Label',value=geo.get('location_label') or selected_event.get('location') or '')
                    lat=st.number_input('Latitude',value=float(geo.get('latitude') or 0.0),format='%.6f')
                    lon=st.number_input('Longitude',value=float(geo.get('longitude') or 0.0),format='%.6f')
                    radius=st.number_input('Allowed Radius (meters)',min_value=50,max_value=5000,value=int(geo.get('radius_meters') or 250),step=50)
                    active=st.checkbox('Enable Location-Assisted Check-In',value=bool(geo.get('active')))
                    save_geo=st.form_submit_button('Save Location Settings',use_container_width=True)
                if save_geo:
                    table('event_geo_settings').upsert({'event_id':eid,'location_label':glabel.strip(),'latitude':lat,'longitude':lon,'radius_meters':radius,'active':active,'updated_by_member_id':member_id,'updated_at':datetime.now(timezone.utc).isoformat()},on_conflict='event_id').execute()
                    st.success('Location settings saved.');st.rerun()

        if can_manage and selected_event:
            owns_event=(selected_event.get('created_by_member_id')==member_id)
            can_edit_event=is_admin or president or owns_event
            if can_edit_event:
                with st.expander('✏️ Edit / Delete Event',expanded=False):
                    current_start=pd.to_datetime(selected_event.get('start_at')).to_pydatetime() if selected_event.get('start_at') else datetime.now()
                    current_end=pd.to_datetime(selected_event.get('end_at')).to_pydatetime() if selected_event.get('end_at') else current_start
                    edit_title=st.text_input('Event Name',value=selected_event.get('title') or '',key=f"edit_title_{eid}")
                    edit_date=st.date_input('Event Date',value=current_start.date(),key=f"edit_date_{eid}")
                    edit_start=st.time_input('Start Time',value=current_start.time().replace(tzinfo=None),key=f"edit_start_{eid}")
                    edit_end=st.time_input('End Time',value=current_end.time().replace(tzinfo=None),key=f"edit_end_{eid}")
                    edit_location=st.text_input('Location',value=selected_event.get('location') or '',key=f"edit_loc_{eid}")
                    current_edit_type=selected_event.get('event_type') or 'Other'
                    edit_type=st.selectbox(
                        'Event Type',
                        EVENT_TYPE_OPTIONS,
                        index=EVENT_TYPE_OPTIONS.index(current_edit_type) if current_edit_type in EVENT_TYPE_OPTIONS else len(EVENT_TYPE_OPTIONS)-1,
                        key=f"edit_type_{eid}"
                    )
                    edit_desc=st.text_area('Event Description / Notes',value=selected_event.get('description') or '',key=f"edit_desc_{eid}")

                    e1,e2=st.columns(2)
                    if e1.button('Save Event Changes',key=f"save_event_{eid}",use_container_width=True):
                        new_start=datetime.combine(edit_date,edit_start)
                        new_end=datetime.combine(edit_date,edit_end)
                        table('events').update({
                            'title':edit_title.strip(),
                            'start_at':new_start.isoformat(),
                            'end_at':new_end.isoformat(),
                            'location':edit_location.strip(),
                            'event_type':edit_type,
                            'description':edit_desc.strip()
                        }).eq('id',eid).execute()
                        st.success('Event updated.')
                        st.rerun()

                    related_tasks=safe_rows('event_tasks',event_id=eid)
                    related_service=safe_rows('service_sessions',event_id=eid)
                    confirm_delete=st.checkbox(
                        'I understand deleting this event also removes or disconnects related event records.',
                        key=f"delete_event_confirm_{eid}"
                    )
                    if related_tasks or related_service:
                        st.warning(f"This event currently has {len(related_tasks)} task(s) and {len(related_service)} service record(s).")
                    if e2.button('Delete Event',key=f"delete_event_{eid}",disabled=not confirm_delete,use_container_width=True):
                        table('events').delete().eq('id',eid).execute()
                        st.success('Event deleted.')
                        st.rerun()

        if can_manage:
            with st.expander('➕ Assign an Event Task',expanded=False):
                members=philo_members()
                mm={m['id']:m['full_name'] for m in members}
                with st.form('taskadd_v25',clear_on_submit=True):
                    mid=st.selectbox('Assign to',list(mm),format_func=lambda x:mm[x])
                    task_choice=st.selectbox('Task',TASK_TEMPLATES)
                    custom_task=''
                    if task_choice=='Other / Write-in':
                        custom_task=st.text_input('Write in the task')
                    due=st.date_input('Due date')
                    notes=st.text_area('Instructions / Notes')
                    save=st.form_submit_button('Assign Task')
                task_text=custom_task.strip() if task_choice=='Other / Write-in' else task_choice
                if save:
                    if not task_text:
                        st.warning('Enter a task.')
                    else:
                        table('event_tasks').insert({
                            'event_id':eid,
                            'assigned_member_id':mid,
                            'assigned_member_name':mm[mid],
                            'task':task_text,
                            'due_date':due.isoformat(),
                            'status':'Not Started',
                            'notes':notes.strip(),
                            'created_by_member_id':member_id,
                            'created_by_name':member_name,
                            'created_at':datetime.now(timezone.utc).isoformat()
                        }).execute()
                        st.success('Task assigned.')
                        st.rerun()

        tasks=safe_rows('event_tasks',event_id=eid)
        if not tasks:
            st.caption('No tasks have been assigned to this event yet.')

        for x in tasks:
            is_assignee=(x.get('assigned_member_id')==member_id)
            is_assigner=(x.get('created_by_member_id')==member_id)
            can_update_task=is_admin or is_assignee or is_assigner
            can_delete_task=is_admin or is_assigner

            if not (can_update_task or can_manage):
                continue

            with st.container(border=True):
                st.markdown(f"**{x.get('task')}** — {x.get('assigned_member_name')}")
                st.caption(
                    f"Due: {x.get('due_date') or 'No date'}"
                    + (f" • Assigned by: {x.get('created_by_name')}" if x.get('created_by_name') else '')
                )
                if x.get('notes'):
                    st.write(f"**Instructions:** {x.get('notes')}")

                if can_update_task:
                    statuses=['Not Started','In Progress','Complete']
                    status=st.selectbox(
                        'Status',
                        statuses,
                        index=statuses.index(x.get('status','Not Started')) if x.get('status','Not Started') in statuses else 0,
                        key=f"ts{x['id']}"
                    )
                    note=st.text_area(
                        'Progress notes',
                        value=x.get('completion_notes') or '',
                        key=f"tn{x['id']}"
                    )
                    st.caption(
                        'Examples: “Called venue — waiting for confirmation.” • '
                        '“Picked up supplies; receipt submitted.” • '
                        '“Waiting for the chair to approve the final count.” • '
                        '“Completed and delivered.”'
                    )
                    if st.button('Save Task Update',key=f"tb{x['id']}",use_container_width=True):
                        table('event_tasks').update({
                            'status':status,
                            'completion_notes':note.strip(),
                            'completed_at':datetime.now(timezone.utc).isoformat() if status=='Complete' else None,
                            'updated_by_member_id':member_id,
                            'updated_by_name':member_name,
                            'updated_at':datetime.now(timezone.utc).isoformat()
                        }).eq('id',x['id']).execute()
                        st.success('Task updated.')
                        st.rerun()
                else:
                    st.caption(f"Status: {x.get('status') or 'Not Started'}")
                    if x.get('completion_notes'):
                        st.write(f"**Progress:** {x.get('completion_notes')}")

                if can_delete_task:
                    delete_reason=st.text_input(
                        'Reason for deleting task',
                        key=f"task_delete_reason_{x['id']}",
                        placeholder='Example: Duplicate task or no longer needed'
                    )
                    confirm_task_delete=st.checkbox(
                        'Confirm delete task',
                        key=f"task_delete_confirm_{x['id']}"
                    )
                    if st.button(
                        'Delete Task',
                        key=f"task_delete_{x['id']}",
                        disabled=not confirm_task_delete,
                        use_container_width=True
                    ):
                        if not delete_reason.strip():
                            st.warning('Enter a reason before deleting the task.')
                        else:
                            try:
                                table('task_deletion_log').insert({
                                    'task_id':x['id'],
                                    'event_id':x.get('event_id'),
                                    'task_text':x.get('task') or '',
                                    'assigned_member_id':x.get('assigned_member_id'),
                                    'assigned_member_name':x.get('assigned_member_name') or '',
                                    'deleted_by_member_id':member_id,
                                    'deleted_by_name':member_name,
                                    'reason':delete_reason.strip(),
                                    'deleted_at':datetime.now(timezone.utc).isoformat()
                                }).execute()
                            except:
                                pass
                            table('event_tasks').delete().eq('id',x['id']).execute()
                            st.success('Task deleted.')
                            st.rerun()

        if can_manage and selected_event and event_type_label(selected_event)=='Community Service':
            st.markdown('### Community Service Attendance')
            service_rows=service_history_for_event(eid)
            if service_rows:
                service_df=pd.DataFrame([{
                    'Member':x.get('member_name'),
                    'Clock In':fmt_dt(x.get('check_in')),
                    'Clock Out':fmt_dt(x.get('check_out')),
                    'Hours':float(x.get('hours') or 0),
                    'Verified':'Yes' if x.get('verified') else 'No'
                } for x in service_rows])
                st.dataframe(service_df,hide_index=True,use_container_width=True)
                st.metric(
                    'Total Recorded Event Service Hours',
                    f"{sum(float(x.get('hours') or 0) for x in service_rows):.2f}"
                )
            else:
                st.caption('No completed service sessions recorded for this event yet.')

        if can_manage and tasks:
            body="\n".join([
                f"{x.get('assigned_member_name')}: {x.get('task')} — {x.get('status')}. Notes: {x.get('completion_notes') or ''}"
                for x in tasks
            ])
            pdf=report_pdf('Event Task Report',em[eid],[('Assignments & Progress',body)],member_name)
            st.download_button('⬇️ Generate Event Task Report',pdf,file_name='NBS_Event_Task_Report.pdf',mime='application/pdf')
elif page=='🧾 Forms':
    st.title('Forms & Signatures')

    if st.session_state.pop('open_member_forms_home',False):
        st.markdown('## My Forms')
        myv=visible_non_test_rows(
            table('reimbursements').select('*').eq(
                'submitted_by_member_id',member_id
            ).order('submitted_at',desc=True).execute().data or []
        )
        st.markdown('### My Submitted Forms')
        if myv:
            st.dataframe(pd.DataFrame([{
                'Form':'Voucher / Reimbursement',
                'Date':v.get('submitted_at'),
                'Reason':v.get('description'),
                'Amount':float(v.get('amount') or 0),
                'Status':v.get('status')
            } for v in myv]),hide_index=True,use_container_width=True)
        else:
            st.caption('You do not have any submitted forms yet.')

        formc1,formc2=st.columns(2)
        with formc1:
            st.info('Use **Reimbursement / Voucher** below to submit a new financial request.')
        with formc2:
            st.info('Use **Community Service** below to view or record eligible service hours.')
        st.divider()

    if st.session_state.pop('open_forms_tab',None)=='Volunteer Sign-Up':
        st.markdown('### 🙋 Volunteer Sign-Ups')
        volunteer_events=[
            e for e in upcoming_events(100)
            if volunteer_link_for_event(e['id'])
        ]
        if not volunteer_events:
            st.info('No volunteer sign-up sheets are currently posted.')
        else:
            for ve in volunteer_events:
                with st.container(border=True):
                    st.markdown(f"**{ve.get('title')}**")
                    st.caption(fmt_dt(ve.get('start_at')))
                    st.link_button(
                        'Open Volunteer Sign-Up',
                        volunteer_link_for_event(ve['id']),
                        use_container_width=True
                    )
        st.divider()

    tabs=st.tabs(['Reimbursement / Voucher','My Vouchers','Grievance / Concern','Volunteer Sign-Up','Community Service'])
    with tabs[0]:
        st.subheader('Voucher / Reimbursement / Check Disbursement Request')
        with st.form('voucher_v2'):
            request_date=st.date_input('Date',value=date.today());reason=st.text_area('Reason for Request')
            blines_v=budget_lines()
            if blines_v:
                budget_line_id=st.selectbox('Committee / Office or Budget Line Item',[x['id'] for x in blines_v],format_func=lambda x:next((b['name'] for b in blines_v if b['id']==x),str(x)))
                line_obj=next((b for b in blines_v if b['id']==budget_line_id),{})
                committee=line_obj.get('name','')
                budget_position=line_obj.get('owner_position') or 'General Chapter'
            else:
                st.warning('No budget lines are available yet. Ask the Treasurer to create one first.')
                budget_line_id=None
                committee=''
                budget_position='General Chapter'
            amount=st.number_input('Amount Requested',min_value=0.0,step=1.0)
            payment=st.radio('Payment Method',['E-check','Reimbursement eVoucher','Mail Check'])
            payee=st.text_input('Check / Payment Payable To');pay_email=st.text_input('Email Address for E-check');mail=st.text_area('Mailing Address (if check)')
            receipt=st.file_uploader('Supporting documentation / receipts',type=['pdf','png','jpg','jpeg'])
            signature=st.text_input('Type your full legal name as your signature')
            certify=st.checkbox('I certify that this typed signature represents my signature for this request.')
            submit=st.form_submit_button('Sign & Submit Request',use_container_width=True)
        if submit:
            if not reason.strip() or amount<=0 or not signature.strip() or not certify:st.error('Complete the reason, amount, signature, and certification.')
            else:
                rp=''
                if receipt:rp=upload_private(bytes(receipt.getbuffer()),f"reimbursements/{member_id}/{datetime.now().strftime('%Y%m%d%H%M%S')}_{receipt.name}",receipt.type)
                table('reimbursements').insert({'submitted_by_member_id':member_id,'submitted_by_name':member_name,'form_type':'Voucher/Reimbursement','expense_date':request_date.isoformat(),'amount':amount,'category':committee,'budget_position':budget_position,'budget_line_id':budget_line_id,'description':reason,'payment_method':payment,'payee_name':payee,'payment_email':pay_email,'receipt_path':rp,'signature_name':signature,'status':'Pending','submitted_at':datetime.now(timezone.utc).isoformat()}).execute();st.success('Signed request submitted.')
    with tabs[1]:
        st.subheader('My Submitted Vouchers')
        st.caption('Track each request from submission through final payment.')

        my_vouchers=visible_non_test_rows(
            table('reimbursements').select('*').eq('submitted_by_member_id',member_id).order('submitted_at',desc=True).execute().data or []
        )

        if not my_vouchers:
            st.info('You have not submitted any voucher or reimbursement requests yet.')

        for v in my_vouchers:
            status=v.get('status') or 'Pending'
            budget_name=budget_line_name(v.get('budget_line_id')) if v.get('budget_line_id') else (v.get('category') or 'Not assigned')
            with st.expander(f"Voucher #{v['id']} • ${float(v.get('amount') or 0):,.2f} • {status}",expanded=False):
                c1,c2=st.columns(2)
                with c1:
                    st.write(f"**Submitted:** {v.get('submitted_at') or ''}")
                    st.write(f"**Reason:** {v.get('description') or ''}")
                    st.write(f"**Budget Line:** {budget_name}")
                    st.write(f"**Payment Method:** {v.get('payment_method') or ''}")
                with c2:
                    st.write(f"**Current Status:** {status}")
                    if v.get('status_reason_note'):
                        st.write(f"**Status Reason / Note:** {v.get('status_reason_note')}")
                    st.write(f"**Your Signature:** {v.get('signature_name') or 'Not recorded'}")
                    if v.get('payee_name'):
                        st.write(f"**Payable To:** {v.get('payee_name')}")

                st.markdown('#### Approval / Signature Trail')
                sigrows=voucher_signature_status(v)

                # Submitter signature is shown separately because it is part of the original request.
                submitter_sig={
                    'Position':'Submitter',
                    'Officer':v.get('submitted_by_name') or member_name,
                    'Decision':'Submitted',
                    'Reason / Note':v.get('status_reason_note') or '',
                    'Signed By':v.get('signature_name') or '',
                    'Signed At':v.get('submitted_at') or ''
                }
                display_rows=[submitter_sig]+sigrows
                st.dataframe(pd.DataFrame(display_rows),hide_index=True,use_container_width=True)

                denied=[r for r in sigrows if r.get('Decision')=='Denied']
                approved=[r for r in sigrows if r.get('Decision')=='Approved']
                waiting=[r for r in sigrows if r.get('Decision') in ['Waiting','Pending']]

                if denied:
                    st.error('This voucher has been denied by: '+', '.join(r['Position'] for r in denied))
                elif status=='Paid':
                    st.success('This voucher has been paid.')
                elif status=='Approved' and not waiting:
                    st.success('All required approvals are complete.')
                elif waiting:
                    st.info('Still waiting for: '+', '.join(r['Position'] for r in waiting))

                st.markdown('#### Request Actions')
                a1,a2=st.columns(2)

                if voucher_can_cancel(v):
                    if a1.button('Cancel Voucher',key=f"cancel_my_voucher_{v['id']}",use_container_width=True):
                        table('reimbursements').update({
                            'status':'Cancelled',
                            'cancelled_at':datetime.now(timezone.utc).isoformat()
                        }).eq('id',v['id']).execute()
                        st.success('Voucher cancelled.')
                        st.rerun()
                else:
                    a1.button('Cancel Voucher',key=f"cancel_disabled_{v['id']}",disabled=True,use_container_width=True)

                if voucher_can_delete(v):
                    confirm=st.checkbox('I understand this permanently deletes this voucher.',key=f"delete_confirm_{v['id']}")
                    if a2.button('Delete Voucher',key=f"delete_my_voucher_{v['id']}",disabled=not confirm,use_container_width=True):
                        table('reimbursements').delete().eq('id',v['id']).execute()
                        st.success('Voucher deleted.')
                        st.rerun()
                else:
                    a2.button('Delete Voucher',key=f"delete_disabled_{v['id']}",disabled=True,use_container_width=True)
                    st.caption('Delete is unavailable after an officer has acted on the voucher. The approval history is retained for financial records.')

    with tabs[2]:
        st.subheader('Grievance / Concern Form')
        st.caption('Private routing to the Advisor, President, or both. Use emergency help instead of this form for immediate danger.')
        with st.form('member_grievance_form',clear_on_submit=True):
            category=st.selectbox('Concern Type',GRIEVANCE_CATEGORIES)
            route=st.selectbox('Send To',['Advisor','President','Advisor & President'])
            subject=st.text_input('Short Subject')
            details=st.text_area('What happened?',height=180)
            prior_steps=st.text_area('What have you already tried, if anything?')
            requested=st.text_area('What outcome or help are you requesting?')
            confidential=st.checkbox('Keep this as limited/confidential as reasonably possible',value=True)
            submit=st.form_submit_button('Submit Grievance / Concern',use_container_width=True)
        if submit:
            if not details.strip():st.warning('Please describe the concern.')
            else:
                guide=grievance_guidance(category,details)
                table('grievances').insert({'submitted_by_member_id':member_id,'submitted_by_name':member_name,'category':category,'route_to':route,'subject':subject.strip(),'details':details.strip(),'prior_steps':prior_steps.strip(),'requested_outcome':requested.strip(),'confidential':confidential,'status':'Submitted','generated_guidance':guide['summary'],'suggested_roles':', '.join(guide['roles']),'created_at':datetime.now(timezone.utc).isoformat()}).execute()
                st.success('Your concern was submitted privately.');st.rerun()
        mine=[g for g in grievance_rows() if g.get('submitted_by_member_id')==member_id]
        st.markdown('### My Submitted Concerns')
        if mine:st.dataframe(pd.DataFrame([{'Subject':g.get('subject') or g.get('category'),'Sent To':g.get('route_to'),'Status':g.get('status'),'Submitted':fmt_dt(g.get('created_at'))} for g in mine]),hide_index=True,use_container_width=True)
        else:st.caption('You have not submitted a grievance/concern form.')

    with tabs[3]:
        events=upcoming_events(100)
        if is_admin and events:
            em={e['id']:f"{fmt_dt(e.get('start_at'))} — {e['title']}" for e in events};eid=st.selectbox('Event',list(em),format_func=lambda x:em[x])
            link=st.text_input('Google Sheets volunteer sign-up link',value=volunteer_link_for_event(eid))
            if st.button('Save Sign-Up Link'):save_volunteer_link_for_event(eid,link);st.success('Saved.')
        for e in events:
            link=volunteer_link_for_event(e['id'])
            if link:st.link_button(f"Volunteer — {e['title']}",link,use_container_width=True)
    with tabs[4]:
        st.subheader('Community Service Hours')
        st.caption('Only events marked Community Service appear here. You can also clock in/out directly from the Event Board.')

        service_events=[
            e for e in upcoming_events(100)
            if event_type_label(e)=='Community Service'
        ]

        if not service_events:
            st.info('No upcoming events are currently marked Community Service.')
        else:
            em={
                e['id']:f"{fmt_dt(e['start_at'])} — {e['title']}"
                for e in service_events
            }
            eid=st.selectbox(
                'Community Service Event',
                list(em),
                format_func=lambda x:em[x],
                key='community_service_event_v34'
            )
            selected_service=next(e for e in service_events if e['id']==eid)
            st.write(f"**{selected_service.get('title')}**")
            if selected_service.get('location'):
                st.caption(f"📍 {selected_service.get('location')}")

            current=open_service_session(member_id,eid)
            c1,c2=st.columns(2)
            if c1.button(
                '🟢 Clock In',
                disabled=current is not None,
                use_container_width=True,
                key='community_tab_clock_in_v34'
            ):
                try:
                    clock_in_service_event(eid)
                    st.rerun()
                except Exception as ex:
                    st.error(str(ex))

            if c2.button(
                '🔴 Clock Out',
                disabled=current is None,
                use_container_width=True,
                key='community_tab_clock_out_v34'
            ):
                try:
                    hours=clock_out_service_event(eid)
                    st.success(f'{hours:.2f} service hour(s) recorded.')
                    st.rerun()
                except Exception as ex:
                    st.error(str(ex))

        st.markdown('### My Service History')
        history=service_sessions_for_member(member_id)
        if history:
            event_map={e['id']:e for e in safe_rows('events')}
            st.dataframe(
                pd.DataFrame([{
                    'Event':event_map.get(x.get('event_id'),{}).get('title') or f"Event #{x.get('event_id')}",
                    'Date':fmt_dt(x.get('check_in')),
                    'Clock In':fmt_dt(x.get('check_in')),
                    'Clock Out':fmt_dt(x.get('check_out')),
                    'Hours':float(x.get('hours') or 0)
                } for x in history]),
                hide_index=True,
                use_container_width=True
            )
            st.metric('Total Community Service Hours',f"{member_service_hours_total(member_id):.2f}")
        else:
            st.caption('No completed community-service sessions recorded yet.')

elif page=='📚 Governance':
    st.title('Governance Center')
    tabs=st.tabs(['Document Library','Recommendations','Parliamentary Helper','Motion Ranking','Version History'])

    with tabs[0]:
        st.subheader('Governing Document Library')
        st.caption('Read the governing documents directly in the Hub. Download/open options remain available when needed.')

        level=st.radio(
            'Document Level',
            ['Local NBS','Northeastern Region','National Philo'],
            horizontal=True,
            key='governance_level_v351'
        )

        if level=='Local NBS':
            st.markdown('### Local NBS Governing Documents')
            local_doc=st.radio(
                'Choose Document',
                ['Local Bylaws','Local SOP / Procedures'],
                horizontal=True,
                key='local_governance_doc_v351'
            )

            bylaws_url=setting('bylaws_url','')
            sop_url=setting('protocol_url','')

            if local_doc=='Local Bylaws':
                st.markdown('#### Current Local Bylaws')
                if bylaws_url:
                    show_governance_url_in_app(bylaws_url,900)
                    st.link_button(
                        'Open Local Bylaws in New Window',
                        bylaws_url,
                        use_container_width=True
                    )
                else:
                    st.warning(
                        'Local Bylaws have not been added yet. '
                        'Advisor/Admin can add the document URL under Google & App Settings.'
                    )

            else:
                st.markdown('#### Current Local SOP / Procedures')
                if sop_url:
                    show_governance_url_in_app(sop_url,900)
                    st.link_button(
                        'Open Local SOP / Procedures in New Window',
                        sop_url,
                        use_container_width=True
                    )
                else:
                    st.warning(
                        'Local SOP / Procedures have not been added yet. '
                        'Advisor/Admin can add the document URL under Google & App Settings.'
                    )

            st.caption(
                'Adopted local amendments and updated versions created through the Hub remain available under Version History.'
            )

        elif level=='Northeastern Region':
            regional_doc=st.radio(
                'Choose Regional Document',
                ['NER Philo Bylaws — April 2022','NER Policies & Procedures — April 2022'],
                horizontal=True,
                key='regional_governance_doc_v351'
            )

            if regional_doc=='NER Philo Bylaws — April 2022':
                st.markdown('### NER Philo Bylaws — April 2022')
                data=local_asset_bytes(Path('assets/governance/NER_Philo_Bylaws_2022.pdf'))
                if data:
                    show_pdf_asset(Path('assets/governance/NER_Philo_Bylaws_2022.pdf'),900)
                    st.download_button(
                        'Download / Print NER Bylaws',
                        data,
                        file_name='NER_Philo_Bylaws_2022.pdf',
                        mime='application/pdf',
                        key='dl_ner_bylaws_v351',
                        use_container_width=True
                    )
                else:
                    st.warning('NER Philo Bylaws file is unavailable.')

            else:
                st.markdown('### NER Policies & Procedures — April 2022')
                data=local_asset_bytes(Path('assets/governance/NER_Philo_Policies_Procedures_2022.pdf'))
                if data:
                    show_pdf_asset(Path('assets/governance/NER_Philo_Policies_Procedures_2022.pdf'),900)
                    st.download_button(
                        'Download / Print NER Policies & Procedures',
                        data,
                        file_name='NER_Philo_Policies_Procedures_2022.pdf',
                        mime='application/pdf',
                        key='dl_ner_pp_v351',
                        use_container_width=True
                    )
                else:
                    st.warning('NER Policies & Procedures file is unavailable.')

        else:
            st.markdown('### National Philo Handbook — 2024 Revision')
            st.info(
                'The National Handbook contains member bylaws, officer duties, committees, '
                'financial obligations, financial procedures, induction, and other national guidance.'
            )
            data=local_asset_bytes(Path('assets/governance/National_Philo_Handbook_2024.pdf'))
            if data:
                show_pdf_asset(Path('assets/governance/National_Philo_Handbook_2024.pdf'),950)
                st.download_button(
                    'Download / Print National Handbook',
                    data,
                    file_name='National_Philo_Handbook_2024.pdf',
                    mime='application/pdf',
                    key='dl_national_handbook_v351',
                    use_container_width=True
                )
            else:
                st.warning('National Philo Handbook file is unavailable.')

    with tabs[1]:
        st.subheader('Recommendation & Alignment Review')
        with st.form('recommendation_v26',clear_on_submit=True):
            title=st.text_input('Recommendation Title')
            doc=st.selectbox('Local Document / Area',['Local Philo Bylaws','Local Philo SOP / Procedures','Operations / Other'])
            section=st.text_input('Article / Section that should be reviewed (if known)')
            rec=st.text_area('Recommendation',height=120)
            rat=st.text_area('Rationale',height=100)
            alignment=st.multiselect('Review Against',ALIGNMENT_REVIEW_OPTIONS)
            submit=st.form_submit_button('Submit Recommendation')
        if submit and title.strip() and rec.strip():
            notes='Alignment review requested: '+', '.join(alignment) if alignment else ''
            table('recommendations').insert({
                'member_id':member_id,'member_name':member_name,'title':title,'category':doc,
                'recommendation':rec,'rationale':rat,'proposed_action':section,
                'meeting_date':date.today().isoformat(),'status':'Submitted',
                'reviewer_notes':notes,'submitted_at':datetime.now(timezone.utc).isoformat()
            }).execute()
            st.success('Recommendation submitted.')

        recs=table('recommendations').select('*').order('submitted_at',desc=True).execute().data or []
        for r in recs:
            with st.expander(f"{r['title']} • {r['status']}"):
                st.write(r['recommendation'])
                st.caption(f"Related local section: {r.get('proposed_action') or 'Not specified'}")
                if r.get('rationale'):st.write(f"**Rationale:** {r.get('rationale')}")
                if r.get('reviewer_notes'):st.info(r['reviewer_notes'])
                if is_admin or president or any(x.get('position')=='Parliamentarian' for x in offices):
                    statuses=['Submitted','Under Review','Approved for Presentation','Adopted by Chapter','Applied to Governing Document','Needs Revision','Closed']
                    status=st.selectbox('Governance status',statuses,index=statuses.index(r['status']) if r['status'] in statuses else 0,key=f"gst26_{r['id']}")
                    notes=st.text_area('Alignment / review notes',value=r.get('reviewer_notes') or '',key=f"gn26_{r['id']}")
                    if st.button('Save Governance Review',key=f"gs26_{r['id']}"):
                        table('recommendations').update({'status':status,'reviewer_notes':notes}).eq('id',r['id']).execute();st.rerun()
                    if r.get('status')=='Adopted by Chapter':
                        st.warning('Apply only after the required adoption/approval has been recorded.')
                        adopted=st.text_input('Adopted at (meeting/date)',key=f"adopt26_{r['id']}")
                        revised=st.text_area('Final adopted replacement text',key=f"rev26_{r['id']}",height=160)
                        if st.button('Adopt & Apply Amendment',key=f"apply26_{r['id']}"):
                            if not adopted.strip() or not revised.strip():
                                st.error('Enter the adoption meeting/date and final adopted text.')
                            else:
                                table('governance_versions').insert({
                                    'document_type':r.get('category'),'recommendation_id':r['id'],
                                    'section_reference':r.get('proposed_action'),'adopted_text':revised,
                                    'effective_date':date.today().isoformat(),'adopted_at':adopted,
                                    'created_by_member_id':member_id
                                }).execute()
                                table('recommendations').update({'status':'Applied to Governing Document'}).eq('id',r['id']).execute()
                                st.success('Official version created and archived.');st.rerun()

    with tabs[2]:
        st.subheader("Parliamentary & Robert's Rules Helper")
        st.info("Order of authority in the Hub: applicable Local Philo governing documents → Northeastern Region governing documents → Sigma Gamma Rho National governing documents → the adopted parliamentary authority when the higher-level documents do not control.")
        action=st.selectbox('I want to…',[
            'Make a motion','Amend a motion','End debate','Reconsider something',
            'Ask a parliamentary question','Challenge a ruling','Bring back previous business',
            'Add something to the agenda','Understand a required vote','Other'
        ])
        q=st.text_area('Describe what is happening in plain language')
        if st.button('Help Me With Procedure'):
            generic={
                'Make a motion':'When recognized, state: “I move that …” Check the governing documents first for any special local or higher-level rule.',
                'Amend a motion':'State the exact words you want to insert, strike, or substitute. The amendment is decided before returning to the main motion.',
                'End debate':'Use the motion for the Previous Question. It is not simply a command to stop discussion; it requires the assembly to act.',
                'Ask a parliamentary question':'Address the chair and request a parliamentary inquiry. The Parliamentarian advises; the presiding officer rules.',
                'Challenge a ruling':'An appeal may be available from a ruling of the chair, subject to the applicable rules.',
                'Understand a required vote':'Check Local, Regional, and National governing documents first. If none establishes a special threshold, consult the adopted parliamentary authority.'
            }
            st.info(generic.get(action,"Document the question, check the governing documents first, and refer uncertain procedure to the Parliamentarian."))
            if q.strip():st.caption('Use the Parliamentarian dashboard or meeting records to document the final ruling if needed.')

    with tabs[3]:
        st.subheader('Ranking / Order of Motions')
        st.caption("Higher-ranked motions take precedence over lower-ranked motions while a lower-ranked question is pending. Incidental motions may arise based on the situation and do not fit this ranking table.")
        st.dataframe(pd.DataFrame(MOTION_RANKING).rename(columns={
            'rank':'Rank','motion':'Motion','second':'Second?','debatable':'Debatable?','amendable':'Amendable?','vote':'Vote / Decision'
        }),hide_index=True,use_container_width=True)
        st.markdown('### Common Incidental Motions / Requests')
        st.write('Point of Order • Appeal • Suspend the Rules • Parliamentary Inquiry • Request for Information • Division of the Assembly • Division of a Question')
        st.caption("This quick guide is a working reference. When a local, Regional, or National rule controls, that governing rule takes priority.")

    with tabs[4]:
        vv=safe_rows('governance_versions')
        if not vv:st.info('No adopted digital local versions have been created yet.')
        for v in vv:
            with st.container(border=True):
                st.markdown(f"**{v.get('document_type')} — {v.get('section_reference') or 'General'}**")
                st.caption(f"Effective {v.get('effective_date')} • Adopted at {v.get('adopted_at')}")
                pdf=governance_pdf(
                    f"{v.get('document_type')} — {v.get('section_reference') or 'Amendment'}",
                    v.get('adopted_text',''),v.get('effective_date',''),v.get('adopted_at','')
                )
                st.download_button('Print Member Version',pdf,file_name=f"NBS_{v.get('document_type','Governance').replace('/','-')}_{v.get('id')}.pdf",mime='application/pdf',key=f"gvpdf26_{v['id']}")
elif page=='👥 Members':
    st.title('Member Directory & My Profile')
    if is_admin:
        with st.expander('➕ Quick Add Member',expanded=False):
            with st.form('directory_quick_add_member',clear_on_submit=True):
                c1,c2=st.columns(2)
                first_q=c1.text_input('First Name')
                last_q=c2.text_input('Last Name')
                email_q=c1.text_input('Email')
                phone_q=c2.text_input('Phone')
                is_philo_q=st.checkbox('Count as Philo',value=True)
                type_q=st.selectbox('Membership Type',['Returning','Inductee','Reactivating'])
                save_q=st.form_submit_button('Add Member',use_container_width=True)
            if save_q:
                try:
                    create_member_account(first_q,last_q,email_q,phone_q,is_philo_q,type_q,'Member')
                    st.success('Member added.')
                    st.rerun()
                except Exception as ex:
                    st.error(str(ex))

    portal_view=st.session_state.pop('member_portal_view',None)

    if portal_view=='profile':
        st.markdown('## My Profile')
        render_member_financial_status(member_id,compact=False,key_prefix='profile_shortcut_v35')
        st.info('Your full profile editor is directly below under the **My Profile** tab.')

    elif portal_view=='reports':
        st.markdown('## Monthly Reports')
        shortcut_reports=visible_non_test_rows(visible_published_reports(member_id))
        if not shortcut_reports:
            st.info('No completed monthly reports are currently published for you.')
        else:
            for rr in shortcut_reports[:20]:
                with st.container(border=True):
                    st.markdown(f"**{rr.get('title') or 'Monthly Report'}**")
                    st.caption(f"{rr.get('report_period') or ''} • {rr.get('published_by_name') or ''}")
                    if rr.get('report_text'):
                        st.write(rr.get('report_text'))
                    if rr.get('file_path'):
                        url=signed_url(rr.get('file_path'),3600)
                        if url:
                            st.link_button('Open Report',url,use_container_width=True)

    elif portal_view=='prayer':
        st.markdown('## Prayer Request')
        st.caption('This request is sent privately to the Chaplain.')
        with st.form('home_shortcut_prayer_v35',clear_on_submit=True):
            prayer_text=st.text_area(
                'Prayer Request',
                height=150,
                placeholder='Share the prayer request you would like the Chaplain to receive.'
            )
            private_prayer=st.checkbox('Keep my name private from general members',value=True)
            send_prayer=st.form_submit_button('Send to Chaplain',use_container_width=True)
        if send_prayer:
            if not prayer_text.strip():
                st.warning('Enter your prayer request.')
            else:
                table('prayer_requests').insert({
                    'member_id':member_id,
                    'member_name':member_name,
                    'request_text':prayer_text.strip(),
                    'private_request':private_prayer,
                    'status':'New',
                    'created_at':datetime.now(timezone.utc).isoformat()
                }).execute()
                st.success('Prayer request sent to the Chaplain.')

    tabs=st.tabs(['Member Directory','My Profile','Monthly Reports','Prayer Request'])

    with tabs[0]:
        selected_member_id=st.session_state.get('directory_selected_member')
        if selected_member_id:
            selected_member=next((m for m in philo_dropdown_people() if m['id']==selected_member_id),None)
            if selected_member:
                profs=safe_rows('member_profiles',member_id=selected_member_id)
                p=profs[0] if profs else {}
                if st.button('← Back to Member Directory'):
                    st.session_state.pop('directory_selected_member',None)
                    st.rerun()
                u=profile_photo_url(selected_member_id)
                c1,c2=st.columns([1,4])
                with c1:
                    if u:st.image(u,width=150)
                    else:st.markdown('## 🌼')
                with c2:
                    st.header(display_member_name(selected_member,p))
                    offices_m=member_offices(selected_member_id)
                    if offices_m:st.caption(", ".join([x['position'] for x in offices_m]))
                    if p.get('show_pronouns') and p.get('pronouns'):
                        st.write(f"**Pronouns:** {p.get('pronouns')}")
                    if p.get('show_phone') and selected_member.get('phone'):
                        st.write(f"**Phone:** {selected_member.get('phone')}")
                    if p.get('show_email') and selected_member.get('email'):
                        st.write(f"**Email:** {selected_member.get('email')}")
                    if p.get('show_address') and p.get('address'):
                        address=p.get('address','')
                        if p.get('city_state_zip'): address += (", " if address else "")+p.get('city_state_zip','')
                        st.write(f"**Address:** {address}")
                    if p.get('show_birthday') and p.get('birthday'):
                        try:
                            bd=pd.to_datetime(p.get('birthday')).strftime('%B %d')
                        except: bd=str(p.get('birthday'))
                        st.write(f"**Birthday:** {bd}")
                    if p.get('show_bio',True) and p.get('bio'):
                        st.markdown('### About Me')
                        st.write(p.get('bio'))
            else:
                st.session_state.pop('directory_selected_member',None)
                st.rerun()
        else:
            st.caption('Click a member’s name to view the information she has chosen to share.')
            for m in philo_dropdown_people():
                profs=safe_rows('member_profiles',member_id=m['id'])
                p=profs[0] if profs else {}
                with st.container(border=True):
                    c1,c2,c3=st.columns([1,4,1])
                    u=profile_photo_url(m['id'])
                    with c1:
                        if u:st.image(u,width=75)
                        else:st.markdown('### 🌼')
                    with c2:
                        st.markdown(f"### {display_member_name(m,p)}")
                        offices_m=member_offices(m['id'])
                        if offices_m:st.caption(", ".join([x['position'] for x in offices_m]))
                    with c3:
                        if st.button('View Profile',key=f"viewprof_{m['id']}",use_container_width=True):
                            st.session_state['directory_selected_member']=m['id']
                            st.rerun()

    with tabs[1]:
        profs=safe_rows('member_profiles',member_id=member_id)
        p=profs[0] if profs else {}

        render_member_financial_status(
            member_id,
            compact=False,
            key_prefix='member_profile_financial_v35'
        )
        st.divider()

        st.subheader('My Profile Information')
        pic=st.file_uploader('Upload / change profile picture',type=['png','jpg','jpeg'])

        c1,c2,c3=st.columns([1,2,2])
        with c1:
            prefix=st.selectbox('Prefix',['','Ms.','Mrs.','Miss','Dr.','Other'],index=(['','Ms.','Mrs.','Miss','Dr.','Other'].index(p.get('prefix','')) if p.get('prefix','') in ['','Ms.','Mrs.','Miss','Dr.','Other'] else 0))
        with c2:
            existing_first=p.get('first_name') or ((member.get('full_name') or '').split(' ')[0] if member.get('full_name') else '')
            first=st.text_input('First Name',value=existing_first)
        with c3:
            existing_last=p.get('last_name') or (' '.join((member.get('full_name') or '').split(' ')[1:]) if member.get('full_name') else '')
            last=st.text_input('Last Name',value=existing_last)

        pronouns=st.text_input('Pronouns',value=p.get('pronouns') or '',placeholder='Optional')
        phone=st.text_input('Phone Number',value=member.get('phone') or '')
        email=st.text_input('Email',value=member.get('email') or '')
        address=st.text_input('Street Address',value=p.get('address') or '')
        cityzip=st.text_input('City, State, Zip',value=p.get('city_state_zip') or '')
        has_birthday=st.checkbox('Add my birthday',value=bool(p.get('birthday')),key='profile_has_birthday')
        birthday_value=date(2000,1,1)
        if p.get('birthday'):
            try:birthday_value=pd.to_datetime(p.get('birthday')).date()
            except Exception:birthday_value=date(2000,1,1)
        birthday=st.date_input('Birthday',value=birthday_value,min_value=date(1900,1,1),max_value=date.today(),disabled=not has_birthday)

        st.markdown('### About Me')
        if st.button('✨ Build My About Me with Dropdowns',use_container_width=True):
            st.session_state['show_bio_questionnaire']=True

        if st.session_state.get('show_bio_questionnaire'):
            with st.container(border=True):
                st.markdown('#### Build My About Me')
                st.caption('Choose from the dropdowns and add optional details in your own words. You can edit the finished draft before saving.')

                hobbies=st.multiselect(
                    'Hobbies & Interests',
                    PROFILE_HOBBY_OPTIONS,
                    key='profile_bio_hobbies'
                )
                hobbies_detail=st.text_input(
                    'Optional hobby detail',
                    placeholder='Example: I especially enjoy mystery novels and weekend trips.',
                    key='profile_bio_hobbies_detail'
                )

                service_reasons=st.multiselect(
                    'Why Community Service Matters to Me',
                    PROFILE_SERVICE_REASON_OPTIONS,
                    key='profile_bio_service'
                )
                service_detail=st.text_input(
                    'Optional service detail',
                    placeholder='Add a short personal reason or experience.',
                    key='profile_bio_service_detail'
                )

                philo_values=st.multiselect(
                    'What I Value About Being a Philo',
                    PROFILE_PHILO_VALUE_OPTIONS,
                    key='profile_bio_philo_values'
                )
                philo_detail=st.text_input(
                    'Optional Philo detail',
                    placeholder='Add what the Philo Affiliate experience means to you.',
                    key='profile_bio_philo_detail'
                )

                strengths=st.multiselect(
                    'Strengths, Talents & Skills',
                    PROFILE_STRENGTH_OPTIONS,
                    key='profile_bio_strengths'
                )
                strengths_detail=st.text_input(
                    'Optional strengths detail',
                    placeholder='Add a skill or strength that is not listed.',
                    key='profile_bio_strengths_detail'
                )

                fun_fact_category=st.selectbox(
                    'Fun Fact Category',
                    ['']+PROFILE_FUN_FACT_OPTIONS,
                    key='profile_bio_fun_category'
                )
                fun_fact_detail=st.text_input(
                    'My Fun Fact',
                    placeholder='Share the actual fun fact here.',
                    key='profile_bio_fun_detail'
                )

                future_goals=st.multiselect(
                    'What I Hope to Accomplish or Contribute',
                    PROFILE_FUTURE_GOAL_OPTIONS,
                    key='profile_bio_goals'
                )
                future_detail=st.text_input(
                    'Optional future-goal detail',
                    placeholder='Add a goal or contribution that is not listed.',
                    key='profile_bio_goal_detail'
                )

                selections_count=sum([
                    len(hobbies),len(service_reasons),len(philo_values),
                    len(strengths),len(future_goals),
                    1 if fun_fact_category else 0
                ])
                if selections_count>12:
                    st.warning(
                        'You selected many items. Consider choosing your strongest 8–12 total selections so your About Me stays easy to read.'
                    )

                q1,q2=st.columns(2)
                if q1.button('Create Bio Draft',use_container_width=True):
                    draft=generate_profile_bio_from_dropdowns(
                        hobbies=hobbies,hobbies_detail=hobbies_detail,
                        service_reasons=service_reasons,service_detail=service_detail,
                        philo_values=philo_values,philo_detail=philo_detail,
                        strengths=strengths,strengths_detail=strengths_detail,
                        fun_fact_category=fun_fact_category,fun_fact_detail=fun_fact_detail,
                        future_goals=future_goals,future_detail=future_detail
                    )
                    if not draft.strip():
                        st.warning('Choose at least one dropdown option or add an optional detail before creating a draft.')
                    else:
                        st.session_state['generated_bio']=draft
                        st.session_state['show_bio_questionnaire']=False
                        st.rerun()
                if q2.button('Close Builder',use_container_width=True):
                    st.session_state['show_bio_questionnaire']=False
                    st.rerun()

        if 'profile_bio_text' not in st.session_state:
            st.session_state['profile_bio_text']=p.get('bio') or ''
        if st.session_state.get('generated_bio') is not None:
            st.session_state['profile_bio_text']=st.session_state.pop('generated_bio')
        bio=st.text_area('About Me',key='profile_bio_text',height=180)

        st.markdown('### What Other Members May See')
        st.caption('Choose what appears when another member opens your profile from the Member Directory.')
        v1,v2=st.columns(2)
        with v1:
            show_pronouns=st.checkbox('Show my pronouns',value=bool(p.get('show_pronouns')))
            show_phone=st.checkbox('Show my phone number',value=bool(p.get('show_phone')))
            show_email=st.checkbox('Show my email',value=bool(p.get('show_email')))
        with v2:
            show_address=st.checkbox('Show my address',value=bool(p.get('show_address')))
            show_birthday=st.checkbox('Show my birthday (month and day only)',value=bool(p.get('show_birthday')))
            show_bio=st.checkbox('Show my About Me',value=p.get('show_bio',True) is not False)

        if st.button('Save My Profile',use_container_width=True):
            try:
                clean_first=str(first or '').strip(); clean_last=str(last or '').strip()
                clean_email=str(email or '').strip().lower(); clean_phone=str(phone or '').strip()
                if not clean_first or not clean_last: raise ValueError('First and last name are required.')
                if not clean_email or '@' not in clean_email: raise ValueError('Enter a valid email address.')
                duplicate=table('members').select('id').eq('email',clean_email).execute().data or []
                if any(str(x.get('id'))!=str(member_id) for x in duplicate):
                    raise ValueError('That email address is already assigned to another account.')

                photo_path=p.get('photo_path','')
                if pic:
                    image_bytes=validate_image_upload(pic,max_mb=8)
                    safe_name=re.sub(r'[^A-Za-z0-9._-]+','_',Path(pic.name or 'profile.jpg').name)
                    new_path=f"profiles/{member_id}/{datetime.now().strftime('%Y%m%d%H%M%S')}_{safe_name}"
                    uploaded_path=upload_private(image_bytes,new_path,pic.type or 'application/octet-stream')
                    if not uploaded_path: raise RuntimeError('Profile photo upload failed. Your existing photo was kept.')
                    photo_path=uploaded_path

                full_name=f"{clean_first} {clean_last}".strip()
                birthday_to_save=birthday.isoformat() if has_birthday else None
                table('members').update({'full_name':full_name,'email':clean_email,'phone':clean_phone}).eq('id',member_id).execute()
                table('member_profiles').upsert({
                    'member_id':member_id,'photo_path':photo_path,'prefix':prefix,'first_name':clean_first,'last_name':clean_last,
                    'pronouns':pronouns,'address':address,'city_state_zip':cityzip,'birthday':birthday_to_save,
                    'bio':bio,'show_pronouns':show_pronouns,'show_email':show_email,'show_phone':show_phone,
                    'show_address':show_address,'show_birthday':show_birthday if has_birthday else False,'show_bio':show_bio,
                    'updated_at':datetime.now(timezone.utc).isoformat()
                }).execute()
                st.success('Profile saved.')
                st.session_state['profile_bio_text']=bio
                st.session_state.pop('generated_bio',None)
                st.rerun()
            except Exception as ex:
                st.error(f'Profile was not saved: {ex}')

        st.markdown("<div class='private-note'><b>Private information:</b> Your Philo ID is visible only to you and the Advisor/Admin.</div>",unsafe_allow_html=True)
        private=safe_rows('member_private',member_id=member_id)
        if private:st.write(f"**My Philo ID:** {private[0].get('philo_id') or 'Not entered yet'}")
    with tabs[2]:
        st.subheader('Monthly Report Library')
        st.caption('Browse finished officer and committee reports that have been published for you.')

        visible_reports=visible_published_reports(member_id)
        if not visible_reports:
            st.info('No finished reports have been published for you yet.')
        else:
            periods=[]
            for rr in visible_reports:
                p=(rr.get('report_period') or '').strip()
                if p and p not in periods:
                    periods.append(p)

            selected_period=st.selectbox(
                'Reporting Month / Period',
                ['All Available']+periods,
                key='member_monthly_report_period'
            )

            filtered=[
                rr for rr in visible_reports
                if selected_period=='All Available' or (rr.get('report_period') or '').strip()==selected_period
            ]

            for rr in filtered:
                with st.container(border=True):
                    st.markdown(f"### {rr.get('title')}")
                    st.caption(
                        f"{rr.get('report_period') or ''} • "
                        f"Published by {rr.get('published_by_name') or ''}"
                    )
                    if rr.get('committee_name'):
                        st.write(f"Committee: {rr.get('committee_name')}")

                    if rr.get('file_path'):
                        furl=signed_url(rr.get('file_path'),3600)
                        st.write(f"**File:** {rr.get('file_name') or 'Completed Report'}")
                        if furl:
                            st.link_button('Open / Download Report',furl,use_container_width=True)
                    elif rr.get('report_text'):
                        with st.expander('Read Report'):
                            st.write(rr.get('report_text'))
                            pdf=report_pdf(
                                rr.get('title') or 'Published Report',
                                rr.get('report_period') or '',
                                [('Report',rr.get('report_text') or '')],
                                rr.get('published_by_name') or ''
                            )
                            st.download_button(
                                'Print / Download Report',
                                pdf,
                                file_name=f"NBS_Report_{rr['id']}.pdf",
                                mime='application/pdf',
                                key=f"monthly_library_pdf_{rr['id']}",
                                use_container_width=True
                            )

    with tabs[3]:
        st.subheader('Prayer Request')
        st.caption('Send a prayer request directly to the Chaplain.')
        with st.form('prayer_request_form',clear_on_submit=True):
            request_text=st.text_area('Prayer request',height=150)
            private_request=st.checkbox('Keep this request private between me, the Chaplain, and Advisor/Admin',value=True)
            submit_prayer=st.form_submit_button('Send Prayer Request',use_container_width=True)
        if submit_prayer:
            if not request_text.strip():
                st.warning('Please enter a prayer request.')
            else:
                table('prayer_requests').insert({
                    'member_id':member_id,
                    'member_name':member_name,
                    'request_text':request_text.strip(),
                    'private_request':private_request,
                    'status':'New',
                    'created_at':datetime.now(timezone.utc).isoformat()
                }).execute()
                st.success('Prayer request sent to the Chaplain.')

elif page=='📸 Historian & Communications':
    st.title('Media, History & Communications');st.caption('Submit photos, communications requests, shoutouts, nominations, and build your Philo service résumé.');st.markdown('<div class="pearl-divider"></div>',unsafe_allow_html=True)
    tabs=st.tabs(['Upload Photos','Request Flyer / Post','Submit Shoutout','Philo of the Month','Build My Philo Resume','My Requests','Events & Celebrations'])
    with tabs[0]:
        with st.form('hist_member_photo',clear_on_submit=True):
            photo=st.file_uploader('Photo',type=['png','jpg','jpeg']);event=st.text_input('Event / Activity');etype=st.selectbox('Event Type',HISTORIAN_EVENT_TYPES);edate=st.date_input('Event Date');caption=st.text_area('Caption / What is happening?');notes=st.text_area('Anything the Historian should know?');save=st.form_submit_button('Submit Photo to Historian',use_container_width=True)
        if save:
            if not photo:st.warning('Choose a photo.')
            elif not event.strip():st.warning('Enter the event/activity name.')
            else:
                safe=re.sub(r'[^A-Za-z0-9._-]+','_',photo.name);path=upload_private(bytes(photo.getbuffer()),f"historian/photos/{member_id}/{datetime.now().strftime('%Y%m%d%H%M%S')}_{safe}",photo.type or 'image/jpeg')
                if path:table('historian_photos').insert({'submitted_by_member_id':member_id,'submitted_by_name':member_name,'event_name':event.strip(),'event_type':etype,'event_date':edate.isoformat(),'photo_path':path,'original_file_name':photo.name,'caption':caption.strip(),'submitter_notes':notes.strip(),'status':'Submitted'}).execute();st.success('Photo submitted.');st.rerun()
    with tabs[1]:
        with st.form('hist_flyer',clear_on_submit=True):
            title=st.text_input('Request Title');event=st.text_input('Event / Program');needed=st.date_input('Needed By',value=date.today()+timedelta(days=7));aud=st.selectbox('Audience',['Philos','Nu Beta Sigma Chapter / Sorors','Philos & Sorors','Community','Social Media','Other']);details=st.text_area('What should it say?',height=140);attach=st.file_uploader('Optional photo/document',type=['png','jpg','jpeg','pdf','docx']);save=st.form_submit_button('Send Request to Historian',use_container_width=True)
        if save:
            ap=''
            if attach:
                safe=re.sub(r'[^A-Za-z0-9._-]+','_',attach.name);ap=upload_private(bytes(attach.getbuffer()),f"historian/communications/{member_id}/{datetime.now().strftime('%Y%m%d%H%M%S')}_{safe}",attach.type or 'application/octet-stream') or ''
            table('communications_requests').insert({'requested_by_member_id':member_id,'requested_by_name':member_name,'request_type':'Flyer/Post','title':title.strip() or event.strip() or 'Flyer/Post Request','event_name':event.strip(),'needed_by':needed.isoformat(),'audience':aud,'details':details.strip(),'attachment_path':ap,'status':'Submitted'}).execute();st.success('Request sent.');st.rerun()
    with tabs[2]:
        mm=historian_member_choice_map();opts=[0]+list(mm)
        with st.form('hist_shoutout',clear_on_submit=True):
            who=st.selectbox('Who is the shoutout for?',opts,format_func=lambda x:'Not Listed — Enter Name' if x==0 else mm[x]);manual=st.text_input('Name if not listed');reason=st.text_area('What should we celebrate?');needed=st.date_input('Preferred Date');save=st.form_submit_button('Submit Shoutout',use_container_width=True)
        if save:
            person=mm.get(who) if who else normalize_historian_name(manual)
            if not person:st.warning('Choose or enter a name.')
            else:table('communications_requests').insert({'requested_by_member_id':member_id,'requested_by_name':member_name,'request_type':'Shoutout','title':f'Shoutout — {person}','needed_by':needed.isoformat(),'audience':'Philos','details':reason.strip(),'subject_member_id':who or None,'subject_name':person,'status':'Submitted'}).execute();st.success('Shoutout submitted.');st.rerun()
    with tabs[3]:
        pm={m['id']:m['full_name'] for m in philo_members()}
        if pm:
            with st.form('philo_month',clear_on_submit=True):
                nominee=st.selectbox('Nominee',list(pm),format_func=lambda x:pm[x]);month=st.date_input('Month to Consider',value=date.today().replace(day=1));reason=st.text_area('Why are you nominating her?');save=st.form_submit_button('Submit Nomination',use_container_width=True)
            if save:table('philo_month_nominations').insert({'nominee_member_id':nominee,'nominee_name':pm[nominee],'submitted_by_member_id':member_id,'submitted_by_name':member_name,'nomination_month':month.replace(day=1).isoformat(),'reason':reason.strip(),'status':'Submitted'}).execute();st.success('Nomination submitted.');st.rerun()
    with tabs[4]:
        if st.button('Generate My Philo Service Resume',use_container_width=True):st.session_state['philo_resume']=historian_resume_text(member_id,member_name)
        resume=st.text_area('Editable Resume Draft',value=st.session_state.get('philo_resume',''),height=620)
        if resume.strip():
            pdf=report_pdf('Philo Service Resume',member_name,[('Service Resume',resume)],member_name);st.download_button('Download My Philo Resume',pdf,file_name=f"{member_name.replace(' ','_')}_Philo_Service_Resume.pdf",mime='application/pdf',use_container_width=True)
    with tabs[5]:
        comm=table('communications_requests').select('*').eq('requested_by_member_id',member_id).order('created_at',desc=True).execute().data or [];photos=table('historian_photos').select('*').eq('submitted_by_member_id',member_id).order('created_at',desc=True).execute().data or [];noms=table('philo_month_nominations').select('*').eq('submitted_by_member_id',member_id).order('created_at',desc=True).execute().data or []
        if comm:st.markdown('#### Communications');st.dataframe(pd.DataFrame([{'Type':r.get('request_type'),'Title':r.get('title'),'Status':r.get('status'),'Historian Note':r.get('historian_notes') or ''} for r in comm]),hide_index=True,use_container_width=True)
        if photos:st.markdown('#### Photos');st.dataframe(pd.DataFrame([{'Event':r.get('event_name'),'Date':r.get('event_date'),'Status':r.get('status'),'Historian Note':r.get('historian_notes') or ''} for r in photos]),hide_index=True,use_container_width=True)
        if noms:st.markdown('#### Nominations');st.dataframe(pd.DataFrame([{'Nominee':r.get('nominee_name'),'Month':r.get('nomination_month'),'Status':r.get('status')} for r in noms]),hide_index=True,use_container_width=True)
        if not(comm or photos or noms):st.info('You have no Historian requests yet.')
    with tabs[6]:
        st.markdown('### Upcoming Events');events_=two_month_events_for_members()
        for e in events_[:12]:
            d_,tm_=calendar_event_date_time(e);st.markdown(f"**{e.get('title','Event')}** — {d_} • {tm_}")
        st.markdown('### Upcoming Birthdays');birthdays=next_birthdays(60)
        if birthdays:st.dataframe(pd.DataFrame(birthdays),hide_index=True,use_container_width=True)

elif page=='📋 Committee Dashboard':
    if not (is_admin or president or chaired_committees(member_id)):
        st.info('Committee Dashboard is available to assigned Committee Chairs, President, and Advisor/Admin.')
        if st.button('Return Home',use_container_width=True):
            set_page('🏠 Dashboard')
        st.stop()

    st.title('Committee Dashboard')

    if is_admin or president:
        st.subheader('Form / Manage Committees')
        philos=philo_members()
        member_map={m['id']:m['full_name'] for m in philos}

        with st.expander('➕ Form a New Committee',expanded=False):
            with st.form('committee_create28',clear_on_submit=True):
                cname=st.text_input('Committee Name')
                description=st.text_area('Committee Purpose / Notes')
                chair=st.selectbox('Committee Chair',list(member_map),format_func=lambda x:member_map[x])
                members_selected=st.multiselect(
                    'Committee Members',
                    list(member_map),
                    default=[chair],
                    format_func=lambda x:member_map[x]
                )
                create=st.form_submit_button('Create Committee',use_container_width=True)

            if create and cname.strip():
                if chair not in members_selected:
                    members_selected.append(chair)
                result=table('committees').insert({
                    'committee_name':cname.strip(),
                    'description':description.strip(),
                    'chair_member_id':chair,
                    'chair_member_name':member_map[chair],
                    'active':True,
                    'created_by_member_id':member_id,
                    'created_at':datetime.now(timezone.utc).isoformat()
                }).execute()
                cid=(result.data or [{}])[0].get('id')
                for mid in members_selected:
                    table('committee_members').insert({
                        'committee_id':cid,
                        'member_id':mid,
                        'member_name':member_map[mid],
                        'active':True,
                        'added_at':datetime.now(timezone.utc).isoformat()
                    }).execute()
                st.success('Committee formed.')
                st.rerun()

    available_committees=active_committees() if (is_admin or president) else chaired_committees(member_id)

    if not available_committees:
        st.info('No committees are currently assigned.')
    else:
        selected_id=st.selectbox(
            'Committee',
            [c['id'] for c in available_committees],
            format_func=lambda cid:next((c.get('committee_name') for c in available_committees if c['id']==cid),str(cid))
        )
        committee=next(c for c in available_committees if c['id']==selected_id)

        st.markdown(f"## {committee.get('committee_name')}")
        st.caption(f"Chair: {committee.get('chair_member_name')}")
        if committee.get('description'):
            st.write(committee.get('description'))

        members_=committee_members(selected_id)
        st.markdown('### Committee Members')
        if members_:
            st.write(", ".join([m.get('member_name','') for m in members_]))

        c1,c2=st.columns(2)
        if c1.button('Open Meeting Center',use_container_width=True):
            set_page('🗓️ Meeting Center')
        if c2.button('Open Report Center',use_container_width=True):
            set_page('📝 Reports')

        if is_admin or president:
            st.divider()
            st.markdown('### Edit Committee')
            philos=philo_members()
            mmap={m['id']:m['full_name'] for m in philos}
            current_ids=[m.get('member_id') for m in members_]
            new_chair=st.selectbox(
                'Chair',
                list(mmap),
                index=(list(mmap).index(committee.get('chair_member_id')) if committee.get('chair_member_id') in mmap else 0),
                format_func=lambda x:mmap[x],
                key=f"committee_edit_chair28_{selected_id}"
            )
            new_members=st.multiselect(
                'Members',
                list(mmap),
                default=[x for x in current_ids if x in mmap],
                format_func=lambda x:mmap[x],
                key=f"committee_edit_members28_{selected_id}"
            )
            if new_chair not in new_members:
                new_members.append(new_chair)

            c1,c2=st.columns(2)
            if c1.button('Save Committee Changes',use_container_width=True):
                table('committees').update({
                    'chair_member_id':new_chair,
                    'chair_member_name':mmap[new_chair]
                }).eq('id',selected_id).execute()

                table('committee_members').update({'active':False}).eq('committee_id',selected_id).execute()
                for mid in new_members:
                    existing=table('committee_members').select('*').eq('committee_id',selected_id).eq('member_id',mid).execute().data or []
                    if existing:
                        table('committee_members').update({
                            'active':True,'member_name':mmap[mid]
                        }).eq('id',existing[0]['id']).execute()
                    else:
                        table('committee_members').insert({
                            'committee_id':selected_id,
                            'member_id':mid,
                            'member_name':mmap[mid],
                            'active':True,
                            'added_at':datetime.now(timezone.utc).isoformat()
                        }).execute()
                st.success('Committee updated.')
                st.rerun()

            confirm=st.checkbox('Confirm deactivate committee',key=f"committee_deactivate_confirm28_{selected_id}")
            if c2.button('Deactivate Committee',disabled=not confirm,use_container_width=True):
                table('committees').update({'active':False}).eq('id',selected_id).execute()
                table('committee_members').update({'active':False}).eq('committee_id',selected_id).execute()
                st.rerun()

elif page=='💬 Messages':
    if not chat_enabled():
        st.info('Philo Chat is currently unavailable.')
        if st.button('Return Home',use_container_width=True): set_page('🏠 Dashboard')
        st.stop()
    st.title('Philo Chat')

    members=philo_dropdown_people()
    others=[m for m in members if m['id']!=member_id]

    if not others:
        st.info('No other members are available to message yet.')
    else:
        member_map={m['id']:m['full_name'] for m in others}
        other_id=st.selectbox(
            'Choose a member',
            list(member_map),
            format_func=lambda x:member_map[x],
            key='chat_member'
        )

        st.markdown('### Conversation')

        try:
            all_msgs=table('messages').select('*').order('sent_at').execute().data or []
        except Exception as ex:
            all_msgs=[]
            st.error(f'Unable to load messages: {ex}')

        convo=[
            m for m in all_msgs
            if (
                m.get('sender_member_id')==member_id and m.get('recipient_member_id')==other_id
            ) or (
                m.get('sender_member_id')==other_id and m.get('recipient_member_id')==member_id
            )
        ]

        if not convo:
            st.caption('No messages yet. Start the conversation below.')
        else:
            for msg in convo[-50:]:
                mine=msg.get('sender_member_id')==member_id
                sender='You' if mine else member_map.get(other_id,'Member')
                sent=msg.get('sent_at') or ''
                try:
                    sent_label=pd.to_datetime(sent).strftime('%b %d, %Y • %I:%M %p')
                except:
                    sent_label=sent
                with st.container(border=True):
                    c1,c2=st.columns([5,1])
                    with c1:
                        st.markdown(f"**{sender}**")
                        st.write(msg.get('message',''))
                    with c2:
                        st.caption(sent_label)

        st.markdown('### New Message')
        with st.form('chat_send_form',clear_on_submit=True):
            new_msg=st.text_area(
                'Write a message',
                height=110,
                placeholder='Type your message here...'
            )
            send=st.form_submit_button('Send Message',use_container_width=True)

        if send:
            if not new_msg.strip():
                st.warning('Type a message before sending.')
            else:
                try:
                    table('messages').insert({
                        'sender_member_id':member_id,
                        'recipient_member_id':other_id,
                        'message':new_msg.strip(),
                        'sent_at':datetime.now(timezone.utc).isoformat()
                    }).execute()
                    st.rerun()
                except Exception as ex:
                    st.error(f'Message could not be sent: {ex}')
elif page=='🌱 Interest & Intake':
    render_interest_intake_center(member_id,is_admin)

elif page=='🏅 Officer Dashboard':
    st.title('Officer Dashboard')
    access=[x.get('position') for x in offices]
    if president:
        access += [x for x in vacant_offices() if x not in access]
    if is_admin and (not access or st.session_state.get('demo_admin_mode')):
        access=OFFICER_POSITIONS

    if not access:
        st.info('You do not currently hold an officer position.')
    else:
        pos=st.selectbox('Office View',access)
        acting=pos not in [x.get('position') for x in offices]
        if acting:
            st.warning(f'Acting access: {pos} is vacant. President access remains until the position is assigned.')

        st.markdown(f"<div class='officer-banner'><b>{pos}</b><br>{position_email(pos)}</div>",unsafe_allow_html=True)

        officer_tabs=st.tabs([
            '🏛️ Role Center',
            '📄 Document Intake',
            '🆘 Advisor Help',
            '📝 Reports & Meetings',
            '🗂️ Records',
            '🎨 Flyer Workflow',
            '📸 Historian'
        ])

        with officer_tabs[0]:
            render_actual_toolkit(pos)
            if pos=='President':
                st.caption('Key workflow: lead chapter operations, review approvals, oversee officers/committees, monitor deadlines, preside over meetings, and maintain chapter-wide accountability.')
            elif pos=='Vice President':
                st.caption('Key workflow: support the President, coordinate committees/programs, follow delegated work, monitor membership follow-up, and step into President responsibilities when required.')
            elif pos=='Financial Secretary':
                st.caption('Key workflow: receive and record incoming monies, issue receipts, forward funds/records to the Treasurer, and assist with financial reporting.')
            elif pos=='Treasurer':
                st.caption('Key workflow: receive funds from the Financial Secretary, maintain electronic financial records, make authorized payments, reconcile accounts, maintain budget-to-actual reporting, and prepare financial reports.')
            elif pos=='Parliamentarian':
                st.caption("Key workflow: interpret governing rules, advise the President/members, support bylaws/policy review, and use Robert's Rules when governing documents do not control.")
            elif pos=='Sergeant-at-Arms':
                st.caption('Key workflow: support meeting order, control entrance/meeting logistics, and handle assigned attendance/order responsibilities.')

            # -------------------------------------------------
            # PRESIDENT LEADERSHIP CENTER
            # -------------------------------------------------
            if pos=='President':
                st.markdown('## President Leadership Center')
                st.caption('Chapter-wide command center for approvals, officers, committees, meetings, tasks, reports, governance, events, and financial oversight.')
                ptabs=st.tabs(['Executive Overview','Approvals','Officers & Committees','Meetings & Events','Tasks & Follow-Up','Reports & Governance'])

                with ptabs[0]:
                    open_tasks=[x for x in safe_rows('event_tasks') if x.get('status')!='Complete']
                    pending_pres=pending_paperwork_for_position('President')
                    upcoming=two_month_events_for_members()
                    snap=finance_snapshot_v3()
                    p1,p2,p3,p4=st.columns(4)
                    p1.metric('Pending Approvals',len(pending_pres))
                    p2.metric('Open Chapter Tasks',len(open_tasks))
                    p3.metric('Upcoming Events',len(upcoming))
                    p4.metric('Current Balance',f"${snap['current_balance']:,.2f}")
                    p1,p2,p3=st.columns(3)
                    p1.metric('Operating Funds',f"${snap['operating']:,.2f}")
                    p2.metric('Committed / Not Paid',f"${snap['committed']:,.2f}")
                    p3.metric('Vacant Offices',len(vacant_offices()))
                    if vacant_offices():
                        st.warning('Vacant offices: '+', '.join(vacant_offices()))
                    st.markdown('### Upcoming Chapter Calendar')
                    for e in upcoming[:10]:
                        d_,tm_=calendar_event_date_time(e)
                        st.markdown(f"**{e.get('title','Event')}** — {d_} • {tm_}")
                    st.markdown('### Items Needing Attention')
                    if pending_pres:
                        st.info(f"{len(pending_pres)} document(s) are waiting for President action.")
                    if open_tasks:
                        overdue=[x for x in open_tasks if x.get('due_date') and str(x.get('due_date'))<date.today().isoformat()]
                        if overdue: st.warning(f"{len(overdue)} open task(s) appear overdue.")

                with ptabs[1]:
                    st.markdown('### President Approval Queue')
                    pending_pres=pending_paperwork_for_position('President')
                    if not pending_pres: st.success('No paperwork is waiting for President action.')
                    for r in pending_pres:
                        with st.expander(f"Voucher #{r['id']} • {r.get('submitted_by_name')} • ${float(r.get('amount') or 0):,.2f}"):
                            st.write(f"**Purpose:** {r.get('description') or ''}")
                            trail=voucher_signature_status(r)
                            if trail: st.dataframe(pd.DataFrame(trail),hide_index=True,use_container_width=True)
                            decision=st.radio('President Decision',['Approved','Needs Information','Denied'],horizontal=True,key=f"presdec{r['id']}")
                            note=st.text_area('President Reason / Note',key=f"presnote{r['id']}")
                            sig=st.text_input('President Typed Signature',key=f"pressig{r['id']}")
                            if st.button('Save President Decision',key=f"pressave{r['id']}",use_container_width=True):
                                if not sig.strip(): st.warning('Enter your typed signature.')
                                else:
                                    save_voucher_decision(r,'President',decision,sig,note)
                                    if decision=='Denied':
                                        table('reimbursements').update({'status':'Denied','status_reason_note':note.strip()}).eq('id',r['id']).execute()
                                    st.success('President action recorded.');st.rerun()

                with ptabs[2]:
                    st.markdown('### Officer Status')
                    officer_rows=safe_rows('officer_assignments')
                    active_officers=[x for x in officer_rows if x.get('active',True)]
                    if active_officers:
                        st.dataframe(pd.DataFrame([{
                            'Office':x.get('position'),
                            'Member':x.get('member_name') or x.get('full_name') or x.get('name') or x.get('member_id'),
                            'Status':'Active'
                        } for x in active_officers]),hide_index=True,use_container_width=True)
                    if vacant_offices():
                        st.markdown('#### Vacancies / Acting Access')
                        for v in vacant_offices(): st.write(f"• {v}")
                    st.markdown('### Committees')
                    committees=safe_rows('committees')
                    if committees:
                        st.dataframe(pd.DataFrame(committees),hide_index=True,use_container_width=True)
                    else:
                        st.caption('No committees are currently listed.')
                    st.caption('Use Meeting Center and Events & Tasks to coordinate committee work and assignments.')

                with ptabs[3]:
                    st.markdown('### Meeting & Event Command Center')
                    m1,m2=st.columns(2)
                    if m1.button('Open Meeting Center',key='pres_open_meeting',use_container_width=True): set_page('🗓️ Meeting Center')
                    if m2.button('Open Events & Tasks',key='pres_open_events',use_container_width=True): set_page('📅 Events & Tasks')
                    st.markdown('### Next Two Months')
                    for e in two_month_events_for_members()[:20]:
                        d_,tm_=calendar_event_date_time(e)
                        st.markdown(f"**{e.get('title','Event')}** — {d_} • {tm_}")

                with ptabs[4]:
                    st.markdown('### Chapter Task Follow-Up')
                    tasks_=safe_rows('event_tasks')
                    if tasks_:
                        st.dataframe(pd.DataFrame([{
                            'Task':x.get('task_name') or x.get('title') or x.get('task'),
                            'Assigned To':x.get('assigned_to_name') or x.get('assigned_member_name') or x.get('assigned_member_id'),
                            'Due':x.get('due_date'),
                            'Status':x.get('status'),
                            'Notes':x.get('progress_notes') or x.get('notes') or ''
                        } for x in tasks_]),hide_index=True,use_container_width=True)
                    else: st.info('No chapter tasks have been created.')
                    if st.button('Manage Tasks in Events & Tasks',key='pres_manage_tasks',use_container_width=True): set_page('📅 Events & Tasks')

                with ptabs[5]:
                    st.markdown('### Chapter Reports & Governance')
                    routed_grievances=[g for g in grievance_rows() if g.get('route_to') in ['President','Advisor & President']]
                    if routed_grievances:
                        st.markdown('#### Grievances / Concerns Routed to President')
                        for g in routed_grievances[:20]:
                            with st.expander(f"{g.get('category')} • {g.get('submitted_by_name')} • {g.get('status')}"):
                                st.write(g.get('details') or '')
                                guide=grievance_guidance(g.get('category'),g.get('details'))
                                st.write('**Suggested handling:** '+guide['summary'])
                                pnote=st.text_area('President Notes',value=g.get('president_notes') or '',key=f"pres_griev_note_{g['id']}")
                                pstatus=st.selectbox('Status',GRIEVANCE_STATUSES,index=GRIEVANCE_STATUSES.index(g.get('status')) if g.get('status') in GRIEVANCE_STATUSES else 0,key=f"pres_griev_status_{g['id']}")
                                if st.button('Save President Review',key=f"pres_griev_save_{g['id']}",use_container_width=True):
                                    table('grievances').update({'president_notes':pnote.strip(),'status':pstatus,'president_reviewed_at':datetime.now(timezone.utc).isoformat()}).eq('id',g['id']).execute();st.rerun()
                    r1,r2=st.columns(2)
                    if r1.button('Open Report Center',key='pres_reports',use_container_width=True): set_page('📝 Reports')
                    if r2.button('Open Bylaws & SOPs',key='pres_governance',use_container_width=True): set_page('📚 Bylaws & SOPs')
                    st.markdown('#### Recommendations / Governance Items')
                    recs=safe_rows('recommendations')
                    if recs:
                        st.dataframe(pd.DataFrame(recs[-25:]),hide_index=True,use_container_width=True)
                    else: st.caption('No recommendations are currently recorded.')

            # -------------------------------------------------
            # VICE PRESIDENT OPERATIONS CENTER
            # -------------------------------------------------
            elif pos=='Vice President':
                st.markdown('## Vice President Operations Center')
                st.caption('Working dashboard for committees, delegated assignments, events, membership follow-up, officer support, and President backup.')
                vtabs=st.tabs(['Operations Overview','Committees','Events & Programs','Tasks & Delegation','Membership & Intake','Recruitment Outreach','Reports & President Backup'])

                with vtabs[0]:
                    mytasks=[x for x in safe_rows('event_tasks') if x.get('assigned_member_id')==member_id and x.get('status')!='Complete']
                    allopen=[x for x in safe_rows('event_tasks') if x.get('status')!='Complete']
                    upcoming=two_month_events_for_members()
                    v1,v2,v3,v4=st.columns(4)
                    v1.metric('My Open Tasks',len(mytasks))
                    v2.metric('Chapter Open Tasks',len(allopen))
                    v3.metric('Upcoming Events',len(upcoming))
                    v4.metric('Committees',len(safe_rows('committees')))
                    st.markdown('### Upcoming Priorities')
                    for e in upcoming[:10]:
                        d_,tm_=calendar_event_date_time(e)
                        st.markdown(f"**{e.get('title','Event')}** — {d_} • {tm_}")
                    if vacant_offices(): st.info('Current office vacancies: '+', '.join(vacant_offices()))

                with vtabs[1]:
                    st.markdown('### Committee Progress Center')
                    committees=safe_rows('committees')
                    if committees:
                        st.dataframe(pd.DataFrame(committees),hide_index=True,use_container_width=True)
                    else: st.caption('No committees are currently listed.')
                    committee_tasks=[x for x in safe_rows('event_tasks') if x.get('status')!='Complete']
                    if committee_tasks:
                        st.markdown('#### Open Committee / Event Work')
                        st.dataframe(pd.DataFrame([{
                            'Task':x.get('task_name') or x.get('title') or x.get('task'),
                            'Assigned To':x.get('assigned_to_name') or x.get('assigned_member_name') or x.get('assigned_member_id'),
                            'Due':x.get('due_date'),'Status':x.get('status'),
                            'Progress Notes':x.get('progress_notes') or x.get('notes') or ''
                        } for x in committee_tasks]),hide_index=True,use_container_width=True)

                with vtabs[2]:
                    st.markdown('### Events & Programs')
                    if st.button('Open Events & Tasks',key='vp_events',use_container_width=True): set_page('📅 Events & Tasks')
                    for e in two_month_events_for_members()[:20]:
                        d_,tm_=calendar_event_date_time(e)
                        st.markdown(f"**{e.get('title','Event')}** — {d_} • {tm_}")

                with vtabs[3]:
                    st.markdown('### Delegated Assignments')
                    tasks_=safe_rows('event_tasks')
                    if tasks_:
                        st.dataframe(pd.DataFrame([{
                            'Task':x.get('task_name') or x.get('title') or x.get('task'),
                            'Assigned To':x.get('assigned_to_name') or x.get('assigned_member_name') or x.get('assigned_member_id'),
                            'Due':x.get('due_date'),'Status':x.get('status'),
                            'Notes':x.get('progress_notes') or x.get('notes') or ''
                        } for x in tasks_]),hide_index=True,use_container_width=True)
                    if st.button('Create / Manage Assignments',key='vp_tasks',use_container_width=True): set_page('📅 Events & Tasks')

                with vtabs[4]:
                    st.markdown('### Membership & Intake')
                    st.caption('The Vice President can manage membership/intake work directly when there are not enough members to form a full Membership Committee.')
                    members_=philo_members()
                    m1,m2,m3=st.columns(3)
                    m1.metric('Active Philos',len(members_))
                    m2.metric('Interest Pool',len(intake_candidate_pool()))
                    m3.metric('Final Voting Pool',len(finalist_candidates()))
                    vote_window=candidate_vote_window()
                    if vote_window['active']:
                        st.success(f"🗳️ Member voting is OPEN until {fmt_vote_time(vote_window['end'])}.")
                    elif vote_window['enabled'] and vote_window['start'] and vote_window['now']<vote_window['start']:
                        st.info(f"🗳️ Voting is scheduled: {fmt_vote_time(vote_window['start'])} – {fmt_vote_time(vote_window['end'])}.")
                    elif vote_window['end']:
                        st.caption(f"Last voting window closed {fmt_vote_time(vote_window['end'])}.")
                    if st.button('Open Full Interest & Intake Center',key='vp_intake_center',use_container_width=True):
                        set_page('🌱 Interest & Intake')
                    st.markdown('#### Current Interested Ladies')
                    pool_=intake_candidate_pool()
                    if pool_:
                        st.dataframe(pd.DataFrame([{
                            'Candidate':f"{c.get('first_name')} {c.get('last_name')}",
                            'Status':c.get('status'),
                            'Finalist':'Yes' if c.get('finalist') else 'No',
                            'Location':c.get('location') or '',
                            'Email':c.get('email')
                        } for c in pool_]),hide_index=True,use_container_width=True)
                    else:st.caption('No Interest Profiles have been created yet.')

                with vtabs[6]:
                    st.markdown('### Recruitment Outreach Planner')
                    st.caption('Starter leads are public campus-life/student-activities contacts. Confirm each school’s current outside-organization/tabling rules before scheduling.')
                    st.dataframe(pd.DataFrame(RECRUITMENT_STARTER_LEADS),hide_index=True,use_container_width=True)
                    st.markdown('#### Suggested Outreach Sequence')
                    st.markdown("""
    1. Email Student Activities/Campus Life with a short affiliate introduction and community-service purpose.
    2. Ask whether outside community organizations may table at club fairs, resource fairs, heritage-month events, women’s programs, or service/leadership events.
    3. Request the external-vendor/community-partner form, insurance requirements, fees, and available dates.
    4. Offer a simple table: affiliate information, community-service examples, upcoming interest event QR code, and contact form.
    5. Follow up within 5–7 business days and track the next action.
                    """)
                    with st.form('vp_recruitment_lead',clear_on_submit=True):
                        school=st.text_input('College / Community Lead')
                        contact=st.text_input('Contact / Office')
                        email=st.text_input('Email / Phone')
                        status=st.selectbox('Status',['Lead Identified','Contacted','Waiting for Reply','Meeting Scheduled','Tabling Approved','Completed','Not a Fit'])
                        next_action=st.text_input('Next Action')
                        follow=st.date_input('Follow-Up Date',value=date.today()+timedelta(days=7))
                        add=st.form_submit_button('Add Outreach Lead',use_container_width=True)
                    if add and school.strip():
                        table('recruitment_outreach').insert({'school_name':school.strip(),'contact_name':contact.strip(),'contact_info':email.strip(),'status':status,'next_action':next_action.strip(),'follow_up_date':follow.isoformat(),'owner_member_id':member_id,'owner_name':member_name,'created_at':datetime.now(timezone.utc).isoformat()}).execute();st.success('Outreach lead added.');st.rerun()
                    leads=safe_rows('recruitment_outreach')
                    if leads:st.dataframe(pd.DataFrame(leads),hide_index=True,use_container_width=True)

                with vtabs[5]:
                    st.markdown('### Reports & President Backup')
                    st.caption('Use this center to prepare the Vice President update and quickly reach chapter-wide tools when supporting or acting for the President.')
                    b1,b2=st.columns(2)
                    if b1.button('Open Report Center',key='vp_reports',use_container_width=True): set_page('📝 Reports')
                    if b2.button('Open Meeting Center',key='vp_meetings',use_container_width=True): set_page('🗓️ Meeting Center')
                    pending_vp=pending_paperwork_for_position('Vice President')
                    if pending_vp:
                        st.markdown('#### Paperwork Assigned to Vice President')
                        st.dataframe(pd.DataFrame([{
                            'ID':x.get('id'),'Submitted By':x.get('submitted_by_name'),
                            'Amount':float(x.get('amount') or 0),'Status':x.get('status')
                        } for x in pending_vp]),hide_index=True,use_container_width=True)
                    else: st.success('No paperwork is currently waiting for Vice President action.')

            # -------------------------------------------------
            # FINANCIAL SECRETARY DASHBOARD
            # -------------------------------------------------
            elif pos=='Financial Secretary':
                tabs=st.tabs(['Incoming Money & Dues','Assessment / Remittance Forms','Voucher Authorization','Reports'])

                with tabs[0]:
                    st.markdown('### Record Money Received')
                    st.caption('Choose who paid, what it was for, enter the amount, then save the payment.')
                    members_=philo_dropdown_people()
                    mopts={0:'Other / Non-member'}
                    mopts.update({m['id']:m['full_name'] for m in members_})

                    with st.form('fs_quick_income',clear_on_submit=True):
                        payer_id=st.selectbox('Who paid?',list(mopts),format_func=lambda x:mopts[x])
                        category=st.selectbox('What was the money for?',['Local Philo Dues','Regional / National Philo Dues','Fundraiser','Event','Donation','Assessment','Other'])

                        blines=budget_lines()
                        if blines:
                            income_budget_line=st.selectbox(
                                'Credit this income to',
                                [x['id'] for x in blines],
                                format_func=lambda x:next((b['name'] for b in blines if b['id']==x),str(x))
                            )
                            income_line_obj=next((b for b in blines if b['id']==income_budget_line),{})
                            income_position=income_line_obj.get('owner_position') or 'General Chapter'
                        else:
                            st.warning('The Treasurer has not created any budget lines yet.')
                            income_budget_line=None
                            income_position='General Chapter'

                        amt=st.number_input('Amount received',min_value=0.0,step=1.0)
                        paid_date=st.date_input('Date received',value=date.today())
                        method=st.selectbox('Payment method',['E-check','Cash','Check','PayPal','Money Order','Other'])
                        note=st.text_input('Optional note')
                        save=st.form_submit_button('Save Payment',use_container_width=True)

                    if save and amt>0:
                        payer_name=mopts[payer_id]
                        possible=finance_possible_duplicate(paid_date.isoformat(),amt,'Income',payer_name,note)
                        if possible:
                            add_duplicate_review(paid_date.isoformat(),amt,'Income',payer_name,note,'Financial Secretary payment resembles an existing ledger row.')
                            st.error('Possible duplicate detected. Nothing was posted. Review it in Treasurer Continuity Center → Duplicate Review.')
                            st.stop()
                        table('finance_transactions').insert({
                            'transaction_date':paid_date.isoformat(),
                            'fiscal_year':finance_fy(),
                            'direction':'Income',
                            'category':category,
                            'position':income_position,
                            'budget_line_id':income_budget_line,
                            'member_id':payer_id if payer_id else None,
                            'payer_payee':payer_name,
                            'amount':amt,
                            'payment_method':method,
                            'notes':note,
                            'entered_by_member_id':member_id,
                            'entered_by_name':member_name
                        }).execute()

                        if category=='Local Philo Dues' and payer_id:
                            fy=setting('current_fiscal_year','2026-2027')
                            table('dues_payments').insert({
                                'member_id':payer_id,
                                'member_name':payer_name,
                                'fiscal_year':fy,
                                'amount':amt,
                                'payment_date':paid_date.isoformat(),
                                'payment_method':method,
                                'entered_by_member_id':member_id
                            }).execute()

                        st.success('Payment recorded and is now visible to the Treasurer.')
                        st.rerun()

                    fy=setting('current_fiscal_year','2026-2027')
                    seed_default_fee_schedule()
                    st.markdown(f'### Philo Financial Standing — {fy}')
                    standing=[]
                    for pm in philo_members():
                        s=member_financial_summary(pm['id'],fy)
                        standing.append({
                            'Philo':pm['full_name'],'Type':s['member_type'],'Required':s['required'],
                            'Paid':s['paid'],'Balance':s['balance'],'Status':s['status']
                        })
                    if standing:
                        st.dataframe(pd.DataFrame(standing),hide_index=True,use_container_width=True)

                    st.markdown('### Current Individual Fee Schedule')
                    inds=[r for r in fee_schedule_rows() if r.get('scope')=='Individual']
                    if inds:
                        st.dataframe(pd.DataFrame([{
                            'Member Type':x.get('member_type'),'Fee':x.get('fee_name'),'Amount':float(x.get('amount') or 0)
                        } for x in inds]),hide_index=True,use_container_width=True)

                with tabs[1]:
                    st.markdown('### Prepare Assessment / Remittance Form')
                    st.caption('The Hub can populate the chapter/officer/roster information for you. Review it, make any corrections, save, then send it to the Treasurer.')

                    if st.button('✨ Populate From Hub Records',use_container_width=True,key='populate_ner_from_hub_v36'):
                        st.session_state['ner_autofill_v36']=assessment_autofill_from_hub()
                        st.success('Chapter, officer, and roster information loaded from the Hub.')
                        st.rerun()

                    auto=st.session_state.get('ner_autofill_v36') or assessment_autofill_from_hub()
                    roster_preview=auto.get('roster') or assessment_roster_from_hub()

                    with st.form('fs_ner_assessment_v36'):
                        c1,c2=st.columns(2)
                        with c1:
                            prep=st.date_input('Date Prepared',value=pd.to_datetime(auto.get('date_prepared') or date.today()).date())
                            affiliate=st.text_input('Affiliate Name',value=auto.get('affiliate_name') or '')
                            addr=st.text_input('Affiliate Address',value=auto.get('affiliate_address') or '')
                            city=st.text_input('City, State, Zip',value=auto.get('city_state_zip') or '')
                            presname=st.text_input('President Name',value=auto.get('president_name') or '')
                            presphone=st.text_input('President Phone',value=auto.get('president_phone') or '')
                            presemail=st.text_input('President Email',value=auto.get('president_email') or '')
                        with c2:
                            fsname=st.text_input('Financial Secretary Name',value=auto.get('fs_name') or '')
                            fsphone=st.text_input('Financial Secretary Phone',value=auto.get('fs_phone') or '')
                            fsemail=st.text_input('Financial Secretary Email',value=auto.get('fs_email') or '')
                            trname=st.text_input('Treasurer Name',value=auto.get('treasurer_name') or '')
                            tremail=st.text_input('Treasurer Email',value=auto.get('treasurer_email') or '')
                            advname=st.text_input('Philo Advisor Name',value=auto.get('advisor_name') or '')
                            advcontact=st.text_input('Philo Advisor Email & Phone',value=auto.get('advisor_contact') or '')

                        st.markdown('#### Chapter Assessments')
                        a=st.number_input('Assessment Fee — number of affiliates ($125 each)',min_value=0,step=1,value=int(auto.get('assessment_count') or 0))
                        f=st.number_input('Fundraiser Fee — number of Philo members ($5 each)',min_value=0,step=1,value=int(auto.get('fundraiser_count') or 0))
                        g=st.number_input('Pageant/Gala Fee — number of affiliates ($150 each)',min_value=0,step=1,value=int(auto.get('gala_count') or 0))
                        l=st.number_input('Late Fee — quantity ($10 each)',min_value=0,step=1,value=int(auto.get('late_count') or 0))
                        d=st.number_input('Delegate Fine — quantity ($50 each)',min_value=0,step=1,value=int(auto.get('delegate_count') or 0))

                        at,ft,gt,lt,dt,total,paypal=assessment_total(a,f,g,l,d)
                        st.write(f"**Total Fees:** ${total:,.2f}  |  **PayPal Amount:** ${paypal:,.2f}  |  **E-check Amount:** ${total:,.2f}")

                        st.markdown('#### Financial Secretary Use Only')
                        rec_date=st.date_input('Date Received',value=date.today())
                        received=st.number_input('Amount Received',min_value=0.0,step=1.0)
                        balance=max(0,total-received);credit=max(0,received-total)
                        st.write(f"**Balance Due:** ${balance:,.2f}  |  **Credit Due:** ${credit:,.2f}")
                        save_form=st.form_submit_button('Save Draft',use_container_width=True)

                    st.markdown('#### Affiliate Roster Pulled From the Hub')
                    st.caption('The Advisor is not included unless she is also marked as a Philo member.')
                    if roster_preview:
                        st.dataframe(pd.DataFrame(roster_preview),hide_index=True,use_container_width=True)
                    else:
                        st.warning('No active Philos were found for the roster.')

                    if save_form:
                        data={
                            'date_prepared':prep.isoformat(),'affiliate_name':affiliate.strip(),'affiliate_address':addr.strip(),'city_state_zip':city.strip(),
                            'president_name':presname.strip(),'president_phone':presphone.strip(),'president_email':presemail.strip(),
                            'fs_name':fsname.strip(),'fs_phone':fsphone.strip(),'fs_email':fsemail.strip(),
                            'treasurer_name':trname.strip(),'treasurer_email':tremail.strip(),'advisor_name':advname.strip(),'advisor_contact':advcontact.strip(),
                            'assessment_count':a,'fundraiser_count':f,'gala_count':g,'late_count':l,'delegate_count':d,
                            'assessment_total':at,'fundraiser_total':ft,'gala_total':gt,'late_total':lt,'delegate_total':dt,
                            'total_fees':total,'paypal_total':paypal,'echeck_total':total,
                            'date_received':rec_date.isoformat(),'amount_received':received,'balance_due':balance,'credit_due':credit,
                            'roster':assessment_roster_from_hub()
                        }
                        table('ner_assessment_forms').insert({
                            'prepared_by_member_id':member_id,'prepared_by_name':member_name,'form_data':data,
                            'total_fees':total,'amount_received':received,'balance_due':balance,'credit_due':credit,
                            'status':'Draft','created_at':datetime.now(timezone.utc).isoformat(),'updated_at':datetime.now(timezone.utc).isoformat()
                        }).execute()
                        st.session_state.pop('ner_autofill_v36',None)
                        st.success('Draft saved.')
                        st.rerun()

                    st.markdown('### My Assessment Forms')
                    forms=[r for r in assessment_forms_for_status() if r.get('prepared_by_member_id')==member_id]
                    if not forms:st.caption('No assessment forms have been saved yet.')
                    for r in forms:
                        with st.expander(assessment_form_label(r)):
                            fd=r.get('form_data') or {}
                            st.write(f"**Status:** {r.get('status','Draft')}")
                            st.write(f"**Total Fees:** ${float(r.get('total_fees') or 0):,.2f}")
                            st.write(f"**Roster Count:** {len(fd.get('roster') or [])}")
                            if r.get('review_notes'):st.warning(f"**Treasurer Note:** {r.get('review_notes')}")
                            if r.get('google_sheet_url'):
                                st.link_button('Open Completed HQ Remittance Sheet',r.get('google_sheet_url'),use_container_width=True)
                            if r.get('status') in ['Draft','Returned to Financial Secretary']:
                                if st.button('Send to Treasurer',key=f"send_assessment_{r['id']}",use_container_width=True):
                                    table('ner_assessment_forms').update({'status':'Submitted to Treasurer','submitted_at':datetime.now(timezone.utc).isoformat(),'updated_at':datetime.now(timezone.utc).isoformat()}).eq('id',r['id']).execute();st.success('Sent to Treasurer.');st.rerun()
                            pdf=ner_assessment_pdf(fd,fd.get('roster') or [])
                            st.download_button('Download / Print Hub Preview',pdf,file_name=f"NBS_NER_Assessment_{r['id']}.pdf",mime='application/pdf',key=f"fs_pdf_{r['id']}",use_container_width=True)

                with tabs[2]:
                    st.markdown('### Voucher Authorization')
                    st.caption('Review vouchers and sign, return, or deny them before President/Treasurer processing.')
                    fs_v=[v for v in safe_rows('reimbursements') if normalized_voucher_status(v) in FINANCE_OPEN_STATUSES]
                    if not fs_v:st.info('No vouchers are waiting for Financial Secretary action.')
                    for v in fs_v:
                        with st.expander(f"Voucher #{v['id']} • {v.get('submitted_by_name')} • ${float(v.get('amount') or 0):,.2f} • {v.get('status')}"):
                            st.write(f"**Purpose:** {v.get('description') or ''}")
                            st.write(f"**Budget:** {budget_line_name(v.get('budget_line_id')) if v.get('budget_line_id') else (v.get('budget_position') or 'General Chapter')}")
                            decision=st.radio('Decision',['Approved','Needs Information','Denied'],horizontal=True,key=f"fsv3d{v['id']}")
                            note=st.text_area('Reason / Note',key=f"fsv3n{v['id']}")
                            sig=st.text_input('Typed Signature',key=f"fsv3s{v['id']}")
                            if st.button('Save Financial Secretary Decision',key=f"fsv3b{v['id']}",use_container_width=True):
                                if not sig.strip():st.warning('Enter your typed signature.')
                                else:
                                    save_voucher_decision(v,'Financial Secretary',decision,sig,note)
                                    if decision=='Approved':
                                        table('reimbursements').update({'status':'Under Review','status_reason_note':note.strip()}).eq('id',v['id']).execute()
                                    st.success('Financial Secretary action recorded.');st.rerun()

                with tabs[3]:
                    st.markdown('### Financial Secretary Reports')
                    month=st.selectbox(
                        'Financial report month',
                        [f"{datetime.now().year}-{m:02d}" for m in range(1,13)],
                        index=datetime.now().month-1,
                        key='fs_report_month'
                    )
                    if st.button('Generate Editable Financial Secretary Report',use_container_width=True):
                        st.session_state['fs_financial_report_draft']=editable_financial_report_text(month,'Financial Secretary')

                    draft=st.text_area(
                        'Editable report draft',
                        value=st.session_state.get('fs_financial_report_draft',''),
                        height=420,
                        key='fs_financial_report_text'
                    )

                    if draft.strip():
                        pdf=report_pdf(
                            'Financial Secretary Financial Report',
                            f"Reporting Period: {month} • Prepared by {member_name}",
                            [('Financial Report',draft)],
                            member_name
                        )
                        st.download_button(
                            'Finalize / Download Financial Report',
                            pdf,
                            file_name=f"NBS_Financial_Secretary_Report_{month}.pdf",
                            mime='application/pdf',
                            use_container_width=True
                        )

            # -------------------------------------------------
            # TREASURER DASHBOARD
            # -------------------------------------------------
            elif pos=='Treasurer':
                st.markdown('## Treasurer Finance Center')
                st.caption('Submitted = no financial impact • Approved = committed budget funds • Paid = actual withdrawal posted once.')
                tabs=st.tabs(['Overview','Voucher Center','Budgets & Fees','Ledger','Assessment / Remittance','Reports & Handoff'])

                with tabs[0]:
                    s=finance_snapshot_v3()
                    st.markdown(f"### Fiscal Year {s['fiscal_year']}")
                    a,b,c=st.columns(3);a.metric('Current Balance',f"${s['current_balance']:,.2f}");b.metric('Operating Funds',f"${s['operating']:,.2f}");c.metric('Approved / Not Paid',f"${s['committed']:,.2f}")
                    a,b,c=st.columns(3);a.metric('YTD Deposits',f"${s['deposits']:,.2f}");b.metric('YTD Withdrawals',f"${s['withdrawals']:,.2f}");c.metric('Protected Minimum',f"${s['protected']:,.2f}")
                    if s['operating']<s['protected']:st.warning('Operating funds are below the protected minimum.')
                    else:st.success(f"Operating above protected minimum: ${s['above_protected']:,.2f}")
                    cfg_=finance_settings()
                    with st.form('trv3_settings'):
                        opening=st.number_input('Opening Bank Balance',min_value=0.0,value=float(cfg_.get('opening_balance') or 0),step=25.0)
                        protected=st.number_input('Protected Minimum',min_value=0.0,value=float(cfg_.get('protected_minimum') or 500),step=25.0)
                        notes=st.text_area('Settings Notes',value=cfg_.get('notes') or '')
                        save=st.form_submit_button('Save Financial Settings',use_container_width=True)
                    if save:
                        table('treasurer_financial_settings').upsert({'fiscal_year':finance_fy(),'opening_balance':opening,'protected_minimum':protected,'notes':notes,'updated_by_member_id':member_id,'updated_at':datetime.now(timezone.utc).isoformat()}).execute()
                        finance_audit('UPDATE_FINANCIAL_SETTINGS','settings',finance_fy(),{'opening':opening,'protected':protected});st.rerun()
                    st.markdown('### Budget Commitments')
                    br=finance_budget_rows()
                    if br:st.dataframe(pd.DataFrame(br),hide_index=True,use_container_width=True)
                    st.markdown('### Monthly / Yearly Totals')
                    mr=monthly_finance_v3()
                    if mr:st.dataframe(pd.DataFrame(mr),hide_index=True,use_container_width=True)

                with tabs[1]:
                    st.markdown('### Treasurer Voucher Center')
                    vouchers=table('reimbursements').select('*').order('submitted_at',desc=True).execute().data or []
                    openv=[v for v in vouchers if normalized_voucher_status(v) in FINANCE_OPEN_STATUSES]
                    if not openv:st.info('No open vouchers.')
                    for v in openv:
                        status=normalized_voucher_status(v)
                        with st.expander(f"Voucher #{v['id']} • {v.get('submitted_by_name')} • ${float(v.get('amount') or 0):,.2f} • {status}"):
                            st.write(f"**Payee:** {v.get('payee_name') or v.get('submitted_by_name')}")
                            st.write(f"**Purpose:** {v.get('description') or ''}")
                            st.write(f"**Budget:** {budget_line_name(v.get('budget_line_id')) if v.get('budget_line_id') else (v.get('budget_position') or 'General Chapter')}")
                            trail=voucher_signature_status(v)
                            if trail:st.dataframe(pd.DataFrame(trail),hide_index=True,use_container_width=True)
                            prior=prior_voucher_approvals(v);ready=voucher_ready_for_treasurer(v)
                            if prior:
                                if ready:st.success('Financial Secretary and President authorization complete.')
                                else:st.warning('Waiting for: '+', '.join(x['Position'] for x in prior if x['Decision']!='Approved'))
                            approved=st.number_input('Amount Approved',min_value=0.0,value=float(v.get('amount_approved') if v.get('amount_approved') is not None else (v.get('amount') or 0)),step=1.0,key=f"trv3amt{v['id']}")
                            note=st.text_area('Treasurer Reason / Note',value=v.get('treasurer_notes') or '',key=f"trv3note{v['id']}")
                            sig=st.text_input('Treasurer Typed Signature',key=f"trv3sig{v['id']}")
                            c1,c2,c3=st.columns(3)
                            if c1.button('Return / Needs Info',key=f"trv3return{v['id']}",use_container_width=True):
                                if not sig.strip():st.warning('Enter your typed signature.')
                                else:save_voucher_decision(v,'Treasurer','Needs Information',sig,note);st.rerun()
                            if c2.button('Deny',key=f"trv3deny{v['id']}",use_container_width=True):
                                if not sig.strip():st.warning('Enter your typed signature.')
                                else:save_voucher_decision(v,'Treasurer','Denied',sig,note);st.rerun()
                            if c3.button('Approve & Commit',key=f"trv3approve{v['id']}",disabled=not ready,use_container_width=True):
                                if not sig.strip():st.warning('Enter your typed signature.')
                                elif approved<=0:st.warning('Approved amount must be greater than $0.')
                                else:
                                    try:treasurer_approve_voucher(v,approved,note,sig);st.success('Approved. Budget is committed; bank balance is unchanged.');st.rerun()
                                    except Exception as ex:st.error(str(ex))
                            if status=='Approved':
                                st.divider();st.markdown('#### Payment — posts the actual withdrawal')
                                pd_=st.date_input('Payment Date',value=date.today(),key=f"trv3pd{v['id']}")
                                pm=st.selectbox('Payment Method',['E-check','Check','Debit/Card','PayPal','Other'],key=f"trv3pm{v['id']}")
                                ref=st.text_input('Confirmation / Reference',value=v.get('payment_reference') or '',key=f"trv3ref{v['id']}")
                                pnote=st.text_area('Payment Note',key=f"trv3pnote{v['id']}")
                                if st.button('Mark Paid & Post Withdrawal',key=f"trv3paid{v['id']}",use_container_width=True):
                                    try:mark_paid_v3(v,pd_,pm,ref,pnote);st.success('Paid and posted exactly once.');st.rerun()
                                    except Exception as ex:st.error(str(ex))
                    closed=[v for v in vouchers if normalized_voucher_status(v) in ['Paid','Denied','Cancelled']][:25]
                    if closed:
                        st.markdown('### Recently Closed')
                        st.dataframe(pd.DataFrame([{'Voucher':v['id'],'Submitted By':v.get('submitted_by_name'),'Amount':float(v.get('amount_approved') if v.get('amount_approved') is not None else (v.get('amount') or 0)),'Status':v.get('status'),'Note':v.get('status_reason_note') or v.get('treasurer_notes') or ''} for v in closed]),hide_index=True,use_container_width=True)

                with tabs[2]:
                    seed_default_fee_schedule();st.markdown('### Dues & Assessment Fee Schedule')
                    for fr in fee_schedule_rows():
                        with st.expander(f"{fr.get('member_type')} • {fr.get('fee_name')} — ${float(fr.get('amount') or 0):,.2f}"):
                            amt=st.number_input('Amount',min_value=0.0,value=float(fr.get('amount') or 0),step=0.25,key=f"trv3fee{fr['id']}")
                            if st.button('Save Fee',key=f"trv3feesave{fr['id']}"):table('fee_schedule').update({'amount':amt,'updated_at':datetime.now(timezone.utc).isoformat()}).eq('id',fr['id']).execute();st.rerun()
                    st.divider();st.markdown('### Budget Manager')
                    with st.form('trv3budgetnew',clear_on_submit=True):
                        name=st.text_input('Budget Name');owner=st.selectbox('Assigned Office / Budget Owner',['General Chapter']+OFFICER_POSITIONS);amount=st.number_input('Budget Amount',min_value=0.0,step=25.0);bn=st.text_area('Description / Notes');create=st.form_submit_button('Create Budget',use_container_width=True)
                    if create and name.strip():
                        table('budget_lines').insert({'name':name.strip(),'owner_position':owner,'budget_amount':amount,'fiscal_year':finance_fy(),'notes':bn,'active':True,'created_by_member_id':member_id}).execute();finance_audit('CREATE_BUDGET','budget','',{'name':name,'amount':amount});st.rerun()
                    br=finance_budget_rows()
                    if br:st.dataframe(pd.DataFrame(br),hide_index=True,use_container_width=True)

                with tabs[3]:
                    st.markdown('### Financial Ledger')
                    st.caption('Record actual deposits and non-voucher withdrawals. Approved vouchers stay commitments until Mark Paid.')
                    with st.form('trv3ledger',clear_on_submit=True):
                        d=st.date_input('Transaction Date',value=date.today());typ=st.selectbox('Transaction Type',['Deposit','Withdrawal']);cat=st.text_input('Category')
                        bl=budget_lines();bid=st.selectbox('Budget Line',[x['id'] for x in bl],format_func=lambda x:next((b['name'] for b in bl if b['id']==x),str(x))) if bl else None
                        obj=next((b for b in bl if b['id']==bid),{}) if bid else {};position=obj.get('owner_position') or 'General Chapter'
                        who=st.text_input('Payer / Payee');amt=st.number_input('Amount',min_value=0.0,step=1.0);method=st.selectbox('Payment Method',['E-check','Check','Cash','Debit/Card','PayPal','Money Order','Other']);ref=st.text_input('Reference / Confirmation');memo=st.text_area('Memo / Notes');save=st.form_submit_button('Post Ledger Transaction',use_container_width=True)
                    if save and amt>0:
                        direction='Income' if typ=='Deposit' else 'Expense'
                        possible=finance_possible_duplicate(d.isoformat(),amt,direction,who or memo,ref)
                        if possible:
                            add_duplicate_review(d.isoformat(),amt,direction,who or memo,ref,'Manual ledger entry resembles an existing ledger row.')
                            st.error('Possible duplicate detected. Nothing was posted. Review it in Treasurer Continuity Center → Duplicate Review.')
                        else:
                            table('finance_transactions').insert({'transaction_date':d.isoformat(),'fiscal_year':finance_fy(),'direction':direction,'category':cat,'position':position,'budget_line_id':bid,'payer_payee':who,'amount':amt,'payment_method':method,'reference_number':ref,'notes':memo,'source_type':'Manual Ledger','entered_by_member_id':member_id,'entered_by_name':member_name}).execute()
                            finance_audit('POST_LEDGER_TRANSACTION','ledger','',{'type':typ,'amount':amt})
                            st.success('Ledger transaction posted.')
                            st.rerun()
                    rows_=finance_year_rows()
                    if rows_:
                        st.dataframe(pd.DataFrame([{'Date':r.get('transaction_date'),'Type':'Deposit' if r.get('direction')=='Income' else 'Withdrawal','Category':r.get('category'),'Budget':budget_line_name(r.get('budget_line_id')) if r.get('budget_line_id') else '','Payer / Payee':r.get('payer_payee'),'Amount':float(r.get('amount') or 0),'Method':r.get('payment_method'),'Reference':r.get('reference_number'),'Source':r.get('source_type')} for r in rows_]),hide_index=True,use_container_width=True)
                        st.download_button('⬇️ QuickBooks-Friendly CSV',quickbooks_csv_bytes(rows_),file_name=f"NBS_QuickBooks_{finance_fy()}.csv",mime='text/csv',use_container_width=True)

                with tabs[4]:
                    st.markdown('### Assessment / Remittance Review Queue')
                    st.caption('Review the Financial Secretary submission. Approving it creates a completed duplicate of the HQ Google Sheet template.')
                    queue=assessment_forms_for_status(['Submitted to Treasurer','Accepted for Payment','Payment Completed','Returned to Financial Secretary'])
                    if not queue:st.info('No assessment forms are waiting for Treasurer review.')
                    for r in queue:
                        with st.expander(assessment_form_label(r)):
                            fd=r.get('form_data') or {}
                            st.write(f"**Prepared by:** {r.get('prepared_by_name')}")
                            st.write(f"**Affiliate:** {fd.get('affiliate_name') or ''}")
                            st.write(f"**Total Fees:** ${float(r.get('total_fees') or 0):,.2f}")
                            st.write(f"**Roster Count:** {len(fd.get('roster') or [])}")
                            if fd.get('roster'):
                                st.dataframe(pd.DataFrame(fd.get('roster')),hide_index=True,use_container_width=True)

                            method=st.radio('Payment Channel',['E-check','PayPal','Other'],key=f"trv36asspm{r['id']}",horizontal=True)
                            correct=float(fd.get('paypal_total') or 0) if method=='PayPal' else float(fd.get('echeck_total') or r.get('total_fees') or 0)
                            st.info(f"Amount to send using {method}: ${correct:,.2f}")
                            note=st.text_area('Treasurer Review Note',value=r.get('review_notes') or '',key=f"trv36assnote{r['id']}")

                            if r.get('google_sheet_url'):
                                st.success('Completed HQ remittance sheet has been created.')
                                st.link_button('Open Completed Google Sheet',r.get('google_sheet_url'),use_container_width=True)

                            c1,c2=st.columns(2)
                            if c1.button('Return to Financial Secretary',key=f"trv36assret{r['id']}",use_container_width=True):
                                table('ner_assessment_forms').update({'status':'Returned to Financial Secretary','review_notes':note or 'Please review and correct this form.','reviewed_at':datetime.now(timezone.utc).isoformat(),'treasurer_member_id':member_id,'updated_at':datetime.now(timezone.utc).isoformat()}).eq('id',r['id']).execute();st.rerun()

                            approve_disabled=bool(r.get('google_sheet_url')) or r.get('status')=='Payment Completed'
                            if c2.button('Approve & Create HQ Google Sheet',key=f"trv36assapprove{r['id']}",disabled=approve_disabled,use_container_width=True):
                                try:
                                    created=create_completed_remittance_sheet(r)
                                    table('ner_assessment_forms').update({
                                        'status':'Accepted for Payment','review_notes':note,'payment_method':method,
                                        'reviewed_at':datetime.now(timezone.utc).isoformat(),'treasurer_member_id':member_id,
                                        'google_sheet_id':created['id'],'google_sheet_url':created['url'],'google_sheet_name':created['name'],
                                        'google_sheet_created_at':datetime.now(timezone.utc).isoformat(),'updated_at':datetime.now(timezone.utc).isoformat()
                                    }).eq('id',r['id']).execute()
                                    st.success('Approved. A completed duplicate of the HQ remittance sheet was created and populated.')
                                    st.rerun()
                                except Exception as ex:
                                    st.error(str(ex))

                            if r.get('status')=='Accepted for Payment':
                                ref=st.text_input('Payment Confirmation / E-check Reference',value=r.get('payment_reference') or '',key=f"trv36assref{r['id']}")
                                if st.button('Mark Payment Completed',key=f"trv36asspaid{r['id']}",use_container_width=True):
                                    table('ner_assessment_forms').update({'status':'Payment Completed','payment_reference':ref.strip(),'paid_at':datetime.now(timezone.utc).isoformat(),'updated_at':datetime.now(timezone.utc).isoformat()}).eq('id',r['id']).execute();st.rerun()

                with tabs[5]:
                    st.markdown('### Financial Reports')
                    months=[f"{datetime.now().year}-{m:02d}" for m in range(1,13)];month=st.selectbox('Reporting Month',months,index=datetime.now().month-1,key='trv3reportmonth')
                    if st.button('Generate Full Treasurer Report',use_container_width=True):st.session_state['trv3report']=treasurer_report_v3(month)
                    draft=st.text_area('Editable Treasurer Report',value=st.session_state.get('trv3report',''),height=600,key='trv3reporttext')
                    if draft.strip():
                        pdf=report_pdf('Treasurer Financial Report',f"Reporting Period: {month} • Prepared by {member_name}",[('Financial Report',draft)],member_name)
                        st.download_button('Finalize / Download Treasurer Report',pdf,file_name=f"NBS_Treasurer_Report_{month}.pdf",mime='application/pdf',use_container_width=True)
                    st.divider();st.markdown('### Treasurer Handoff Summary')
                    if st.button('Generate Handoff Summary',use_container_width=True):st.session_state['trv3handoff']=treasurer_handoff_v3()
                    hand=st.text_area('Editable Handoff Summary',value=st.session_state.get('trv3handoff',''),height=600,key='trv3handofftext')
                    if hand.strip():
                        pdf=report_pdf('Treasurer Handoff Summary',finance_fy(),[('Handoff',hand)],member_name)
                        st.download_button('Download Handoff Summary',pdf,file_name=f"NBS_Treasurer_Handoff_{finance_fy()}.pdf",mime='application/pdf',use_container_width=True)

            # -------------------------------------------------
            # OTHER OFFICERS
            # -------------------------------------------------
            else:
                tasks=[x for x in safe_rows('event_tasks') if x.get('assigned_member_id')==member_id and x.get('status')!='Complete']
                c1,c2,c3=st.columns(3)
                c1.metric('Open Tasks',len(tasks))

                pending=pending_paperwork_for_position(pos)
                c2.metric('Pending Paperwork',len(pending))

                totals,assigned_lines=position_budget_totals(pos)
                c3.metric('Available Budget',f"${totals['available']:,.2f}" if assigned_lines else 'Not set')

                if assigned_lines:
                    st.markdown('### My Budget Lines')
                    st.dataframe(pd.DataFrame([{
                        'Budget Line':x['name'],
                        'Budget':budget_line_summary(x)['budget'],
                        'Income':budget_line_summary(x)['income'],
                        'Spent':budget_line_summary(x)['spent'],
                        'Committed':budget_line_summary(x)['committed'],
                        'Available':budget_line_summary(x)['available']
                    } for x in assigned_lines]),hide_index=True,use_container_width=True)

                st.subheader('Documents Waiting for Action')
                for r in pending[:20]:
                    with st.container(border=True):
                        st.write(f"Voucher #{r['id']} — {r.get('submitted_by_name')} — ${float(r.get('amount') or 0):,.2f}")
                        decision=st.radio('Decision',['Pending','Approved','Denied'],horizontal=True,key=f"od{r['id']}")
                        reason_note=st.text_area('Reason / Note',key=f"orn{r['id']}",placeholder='Add the reason for approval, denial, return, or any instructions for the submitter.')
                        sig=st.text_input('Typed signature',key=f"os{r['id']}")
                        if st.button('Sign / Save Decision',key=f"ob{r['id']}"):
                            try:
                                table('document_approvals').insert({
                                    'document_type':'reimbursement',
                                    'document_id':r['id'],
                                    'position':pos,
                                    'member_id':member_id,
                                    'decision':decision,
                                    'reason_note':reason_note.strip(),
                                    'signature_name':sig,
                                    'signed_at':datetime.now(timezone.utc).isoformat()
                                }).execute()

                                fresh=table('reimbursements').select('*').eq('id',r['id']).single().execute().data
                                if fresh:
                                    trail=voucher_signature_status(fresh)
                                    if decision=='Denied':
                                        table('reimbursements').update({'status':'Denied'}).eq('id',r['id']).execute()
                                    elif trail and all(x.get('Decision')=='Approved' for x in trail):
                                        table('reimbursements').update({'status':'Approved'}).eq('id',r['id']).execute()

                                st.success('Action recorded.')
                                st.rerun()
                            except Exception as ex:
                                st.error(str(ex))

                st.subheader('Position Assistant')
                idea=st.text_input('Ask for ideas or help for your position',placeholder='Example: What should I prepare for this month?')
                if st.button('Ask Position Assistant'):
                    qs=POSITION_QUESTIONS.get(pos,[])
                    st.info(f"For {pos}, start with these monthly checkpoints: " + " ".join([f"{i+1}. {q}" for i,q in enumerate(qs)]))
                    if idea:
                        st.caption('Use the Reports page to turn your answers into the standard NBS report.')

        with officer_tabs[1]:
            st.subheader('📄 Universal Officer Document Intake')
            st.caption(
                'Upload the officer document once. The Hub extracts what it can, you review/correct it, '
                'and only then are draft records created. The existing manual forms remain available as backup.'
            )
            render_universal_document_intake(member_id,pos,is_admin)

        with officer_tabs[2]:
            st.subheader('Advisor Updates & Help')
            st.caption('Updates sent to you by the Advisor appear here. You decide whether to add them to your report, create a vote/action, or handle them separately.')

            incoming_updates=leadership_updates_for_member(member_id,pos)
            if not incoming_updates:
                st.info('No Advisor/leadership updates are assigned to this office right now.')
            for u in incoming_updates:
                with st.expander(f"{u.get('priority') or 'Routine'} • {u.get('update_type') or 'Update'} • {u.get('subject') or 'Leadership Update'}"):
                    st.write(u.get('message') or '')
                    if u.get('suggested_action'):
                        st.info('Suggested action: '+str(u.get('suggested_action')))
                    st.caption(f"Sent by {u.get('created_by_name') or 'Advisor'} • {fmt_dt(u.get('created_at'))}")
                    c1,c2,c3=st.columns(3)
                    if c1.button('➕ Add to My Report',key=f"lead_add_report_{u['id']}_{pos}",use_container_width=True):
                        queue_leadership_update_for_report(u['id'],member_id,member_name)
                        st.success('Added to your report queue.')
                        st.rerun()
                    if c2.button('✅ I’ll Handle This',key=f"lead_handle_{u['id']}_{pos}",use_container_width=True):
                        table('leadership_updates').update({'status':'Handled Separately','updated_at':datetime.now(timezone.utc).isoformat()}).eq('id',u['id']).execute()
                        st.success('Marked as handled separately.')
                        st.rerun()
                    if c3.button('📌 Create Action Item',key=f"lead_action_{u['id']}_{pos}",use_container_width=True):
                        st.session_state[f"lead_action_open_{u['id']}"]=True

                    if u.get('update_type')=='Vote / Decision Needed' or st.session_state.get(f"lead_vote_open_{u['id']}"):
                        pass
                    if st.button('🗳️ Create Member Vote from This Update',key=f"lead_vote_btn_{u['id']}_{pos}",use_container_width=True):
                        st.session_state[f"lead_vote_open_{u['id']}"]=True

                    if st.session_state.get(f"lead_vote_open_{u['id']}"):
                        with st.form(f"lead_vote_form_{u['id']}_{pos}",clear_on_submit=False):
                            vote_title=st.text_input('Vote Title',value=u.get('subject') or 'Philo Affiliate Vote')
                            vote_question=st.text_area('Question Members Will See',value=u.get('message') or '')
                            vc1,vc2=st.columns(2)
                            opens_date=vc1.date_input('Opens Date',value=date.today(),key=f"vote_od_{u['id']}")
                            opens_time=vc1.time_input('Opens Time',value=datetime.now().time().replace(second=0,microsecond=0),key=f"vote_ot_{u['id']}")
                            closes_date=vc2.date_input('Closes Date',value=date.today()+timedelta(days=3),key=f"vote_cd_{u['id']}")
                            closes_time=vc2.time_input('Closes Time',value=datetime.now().time().replace(second=0,microsecond=0),key=f"vote_ct_{u['id']}")
                            make_vote=st.form_submit_button('Activate Vote',use_container_width=True)
                        if make_vote:
                            open_dt=datetime.combine(opens_date,opens_time).astimezone()
                            close_dt=datetime.combine(closes_date,closes_time).astimezone()
                            if close_dt<=open_dt:
                                st.error('Closing time must be after opening time.')
                            else:
                                table('chapter_votes').insert({
                                    'title':vote_title.strip(),'question':vote_question.strip(),
                                    'opens_at':open_dt.isoformat(),'closes_at':close_dt.isoformat(),
                                    'active':True,'source_update_id':u['id'],
                                    'created_by_member_id':member_id,'created_by_name':member_name,
                                    'created_at':datetime.now(timezone.utc).isoformat()
                                }).execute()
                                table('leadership_updates').update({'status':'Vote Created'}).eq('id',u['id']).execute()
                                st.success('Member vote created.')
                                st.rerun()

                    if st.session_state.get(f"lead_action_open_{u['id']}"):
                        with st.form(f"lead_action_form_{u['id']}_{pos}",clear_on_submit=True):
                            action_title=st.text_input('Action Item',value=u.get('subject') or '')
                            action_due=st.date_input('Due Date',value=date.today()+timedelta(days=7))
                            action_note=st.text_area('Notes',value=u.get('suggested_action') or u.get('message') or '')
                            add_action=st.form_submit_button('Create My Action Item',use_container_width=True)
                        if add_action:
                            table('event_tasks').insert({
                                'task_name':action_title.strip(),
                                'assigned_member_id':member_id,
                                'assigned_member_name':member_name,
                                'due_date':action_due.isoformat(),
                                'status':'Not Started',
                                'progress_notes':action_note.strip(),
                                'created_by_member_id':member_id,
                                'created_by_name':member_name
                            }).execute()
                            st.success('Action item created.')
                            st.rerun()

            st.divider()
            st.markdown('### 🆘 Ask the Advisor')
            with st.form(f'advisor_help_form_{pos}',clear_on_submit=True):
                topic=st.text_input('Topic')
                question=st.text_area('Question / Situation')
                priority=st.selectbox('Priority',ADVISOR_HELP_PRIORITIES)
                preferred=st.selectbox('Preferred Response',['Reply in Hub','Talk at Next Meeting','Phone / Call','Either'])
                send=st.form_submit_button('Send to Advisor',use_container_width=True)
            if send:
                if not question.strip():
                    st.warning('Enter your question or situation.')
                else:
                    table('advisor_help_requests').insert({
                        'requesting_member_id':member_id,'requesting_member_name':member_name,
                        'position':pos,'topic':topic.strip(),'question':question.strip(),
                        'priority':priority,'preferred_response':preferred,'status':'Open',
                        'created_at':datetime.now(timezone.utc).isoformat()
                    }).execute()
                    st.success('Your request was sent to the Advisor.')
                    st.rerun()

        with officer_tabs[3]:
            st.subheader('Reports & Meetings')
            c1,c2=st.columns(2)
            if c1.button('📝 Open Report Center',key=f"officer_report_center_623_{pos}",use_container_width=True):
                set_page('📝 Reports')
            if c2.button('🗓️ Open Meeting Center',key=f"officer_meeting_center_623_{pos}",use_container_width=True):
                set_page('🗓️ Meeting Center')

            queued=leadership_report_queue(member_id)
            if queued:
                st.markdown('### Advisor Updates Waiting for Your Report')
                for q in queued:
                    txt=leadership_update_text(q.get('update_id'))
                    with st.container(border=True):
                        st.write(txt)
                        if st.button('Remove from Report Queue',key=f"remove_report_queue_{q['id']}",use_container_width=True):
                            remove_leadership_update_from_report(q.get('update_id'),member_id)
                            st.rerun()

            votes=[v for v in safe_rows('chapter_votes') if v.get('created_by_member_id')==member_id]
            if votes:
                st.markdown('### My Chapter Votes')
                for v in votes[-10:]:
                    results=chapter_vote_results(v['id'])
                    yes=sum(1 for x in results if x.get('vote')=='Yes')
                    no=sum(1 for x in results if x.get('vote')=='No')
                    abstain=sum(1 for x in results if x.get('vote')=='Abstain')
                    st.write(f"**{v.get('title')}** — Yes {yes} • No {no} • Abstain {abstain} • closes {fmt_dt(v.get('closes_at'))}")

        with officer_tabs[4]:
            st.subheader('Officer Records')
            render_previous_sorority_year_records(pos,member_id,is_admin)

        with officer_tabs[5]:
            st.subheader('Flyer Workflow')
            if is_admin or pos in OFFICER_POSITIONS:
                render_philo_flyer_approval_workflow(member_id,is_admin)
            else:
                st.info('Flyer workflow is not assigned to this office.')

        with officer_tabs[6]:
            st.subheader('Historian / Continuity')
            if pos=='Historian' or is_admin:
                render_historian_continuity_center(member_id,is_admin)
            else:
                st.info('Historian continuity tools are available to the Historian and Advisor/Admin.')

elif page=='⚙️ Admin Center':
    if not is_admin:st.error('Administrator access required.');st.stop()
    st.title('Advisor / Admin Center')
    st.subheader('All Account Balances')
    ads=finance_snapshot_v3()
    a1,a2,a3=st.columns(3);a1.metric('Current Balance',f"${ads['current_balance']:,.2f}");a2.metric('Operating Funds',f"${ads['operating']:,.2f}");a3.metric('Approved / Not Paid',f"${ads['committed']:,.2f}")
    balances=finance_budget_rows()
    if balances:
        st.dataframe(pd.DataFrame(balances),hide_index=True,use_container_width=True)
    else:
        st.caption('No Treasurer-created budget lines yet.')
    tabs=st.tabs(['Members','Officers','Private Member Data','Chapter Reports','Activity Log','Advisor Help','Grievances','Paperwork','Test / Preview','Google & App Settings'])

    with tabs[0]:
        st.subheader('Add / Manage Members')

        with st.expander('➕ Add New Member',expanded=True):
            with st.form('admin_add_member',clear_on_submit=True):
                c1,c2=st.columns(2)
                with c1:
                    first=st.text_input('First Name')
                    email=st.text_input('Email')
                    is_philo_new=st.checkbox('Count this person as a Philo member',value=True)
                with c2:
                    last=st.text_input('Last Name')
                    phone=st.text_input('Phone Number')
                    member_type_new=st.selectbox('Membership Type',['Returning','Inductee','Reactivating'])

                role_new=st.selectbox('Account Type',['Member','Admin'],index=0)
                add_member=st.form_submit_button('Add Member',use_container_width=True)

            if add_member:
                try:
                    new_id=create_member_account(
                        first,last,email,phone,
                        is_philo=is_philo_new,
                        member_type=member_type_new,
                        role=role_new
                    )
                    st.success('Member added. She is now available in committee and officer dropdowns.')
                    st.rerun()
                except Exception as ex:
                    st.error(str(ex))

        st.markdown('### Bulk Add Members')
        st.caption('Upload a Google Sheet after downloading it as Excel (.xlsx) or CSV. Preview the rows before importing.')

        template_df=pd.DataFrame([{
            'First Name':'Eleanor',
            'Last Name':'Green',
            'Email':'example@gmail.com',
            'Phone':'',
            'Prefix':'Ms.',
            'Pronouns':'',
            'Address':'',
            'City State Zip':'',
            'Birthday':'',
            'Membership Type':'Returning',
            'Counts as Philo':'Yes',
            'Philo ID':'',
            'Account Type':'Member'
        }])
        st.download_button(
            'Download Bulk Member Template',
            template_df.to_csv(index=False).encode('utf-8'),
            file_name='NBS_Bulk_Member_Template.csv',
            mime='text/csv',
            use_container_width=True
        )

        bulk_file=st.file_uploader(
            'Upload Member Sheet',
            type=['xlsx','csv'],
            key='bulk_member_upload'
        )

        if bulk_file:
            try:
                bulk_df=read_bulk_member_file(bulk_file)
                st.markdown('#### Import Preview')
                st.dataframe(bulk_df,hide_index=True,use_container_width=True)
                st.caption(f"{len(bulk_df)} row(s) found.")

                confirm_bulk=st.checkbox(
                    'I reviewed the preview and want to add these members.',
                    key='confirm_bulk_member_import'
                )
                if st.button(
                    'Import Members',
                    disabled=not confirm_bulk,
                    use_container_width=True
                ):
                    added,skipped,errors=import_bulk_members(bulk_df)
                    st.session_state['bulk_member_results']={
                        'added':added,'skipped':skipped,'errors':errors
                    }
                    st.success(f"Import complete: {len(added)} added, {len(skipped)} skipped, {len(errors)} error(s).")
                    st.rerun()
            except Exception as ex:
                st.error(f'Could not read member sheet: {ex}')

        bulk_results=st.session_state.get('bulk_member_results')
        if bulk_results:
            with st.expander('Last Bulk Import Results',expanded=True):
                if bulk_results.get('added'):
                    st.markdown('**Added**')
                    st.dataframe(pd.DataFrame(bulk_results['added']),hide_index=True,use_container_width=True)
                if bulk_results.get('skipped'):
                    st.markdown('**Skipped**')
                    st.dataframe(pd.DataFrame(bulk_results['skipped']),hide_index=True,use_container_width=True)
                if bulk_results.get('errors'):
                    st.markdown('**Errors**')
                    st.dataframe(pd.DataFrame(bulk_results['errors']),hide_index=True,use_container_width=True)

        st.markdown('### Current People')
        people=all_members_admin()
        if people:
            st.dataframe(
                pd.DataFrame([{
                    'Name':p.get('full_name'),
                    'Email':p.get('email'),
                    'Phone':p.get('phone'),
                    'Account Type':p.get('role'),
                    'Status':'Active' if p.get('active',True) else 'Inactive',
                    'Counts as Philo':'Yes' if is_philo_member(p['id']) else 'No'
                } for p in people]),
                hide_index=True,
                use_container_width=True
            )

        st.caption('After adding a member, use Private Member Data to enter her Philo ID or change membership details.')

        st.markdown('### Edit / Remove Member')
        admin_people=all_members_admin()
        if admin_people:
            admin_map={m['id']:m['full_name'] for m in admin_people}
            manage_mid=st.selectbox(
                'Choose Member',
                list(admin_map),
                format_func=lambda x:admin_map[x],
                key='admin_manage_member_v32'
            )
            manage_row=next(m for m in admin_people if m['id']==manage_mid)
            private_rows=safe_rows('member_private',member_id=manage_mid)
            manage_priv=private_rows[0] if private_rows else {}

            c1,c2=st.columns(2)
            edit_name=c1.text_input('Full Name',value=manage_row.get('full_name') or '',key='admin_edit_full_name_v32')
            edit_email=c2.text_input('Email',value=manage_row.get('email') or '',key='admin_edit_email_v32')
            edit_phone=c1.text_input('Phone',value=manage_row.get('phone') or '',key='admin_edit_phone_v32')
            edit_role=c2.selectbox(
                'Account Type',
                ['Member','Admin'],
                index=0 if manage_row.get('role','Member')!='Admin' else 1,
                key='admin_edit_role_v32'
            )
            edit_active=st.checkbox(
                'Active Account — can log in',
                value=bool(manage_row.get('active',True)),
                key='admin_edit_active_v32'
            )
            edit_is_philo=st.checkbox(
                'Count as Philo',
                value=manage_priv.get('is_philo',True) is not False,
                key='admin_edit_is_philo_v32'
            )
            member_types=['Returning','Inductee','Reactivating']
            current_type=manage_priv.get('member_type','Returning')
            edit_member_type=st.selectbox(
                'Membership Type',
                member_types,
                index=member_types.index(current_type) if current_type in member_types else 0,
                key='admin_edit_member_type_v32'
            )
            status_reason=st.text_area(
                'Reason / Administrative Note',
                placeholder='Required when removing/deactivating an account.',
                key='admin_member_status_reason_v32'
            )

            c1,c2=st.columns(2)
            if c1.button('Save Member Changes',use_container_width=True,key='save_member_changes_v32'):
                old_active=bool(manage_row.get('active',True));old_role=manage_row.get('role') or 'Member'
                if old_active and not edit_active and not status_reason.strip():
                    st.warning('Enter a reason before deactivating/removing a member.')
                else:
                    table('members').update({
                        'full_name':edit_name.strip(),
                        'email':edit_email.strip().lower(),
                        'phone':edit_phone.strip(),
                        'role':edit_role,
                        'active':edit_active
                    }).eq('id',manage_mid).execute()
                    table('member_private').upsert({
                        'member_id':manage_mid,
                        'is_philo':edit_is_philo,
                        'member_type':edit_member_type,
                        'philo_id':manage_priv.get('philo_id') or '',
                        'advisor_notes':manage_priv.get('advisor_notes') or ''
                    }).execute()
                    if old_active!=edit_active or old_role!=edit_role:
                        record_member_status_change(
                            manage_mid,old_active,edit_active,old_role,edit_role,status_reason
                        )
                    st.success('Member account updated.')
                    st.rerun()

            remove_confirm=st.checkbox(
                'Confirm remove/deactivate this member from login',
                key='confirm_remove_member_v32'
            )
            if c2.button(
                'Remove / Deactivate Member',
                disabled=not remove_confirm,
                use_container_width=True,
                key='deactivate_member_v32'
            ):
                if not status_reason.strip():
                    st.warning('Enter a reason before removing/deactivating a member.')
                else:
                    old_role=manage_row.get('role') or 'Member'
                    table('members').update({'active':False}).eq('id',manage_mid).execute()
                    record_member_status_change(
                        manage_mid,bool(manage_row.get('active',True)),False,
                        old_role,old_role,status_reason
                    )
                    st.success('Member deactivated. Historical records were preserved.')
                    st.rerun()

        st.markdown('### Member PIN Reset')
        reset_people=active_members()
        if reset_people:
            reset_map={m['id']:m['full_name'] for m in reset_people}
            reset_mid=st.selectbox(
                'Member',
                list(reset_map),
                format_func=lambda x:reset_map[x],
                key='admin_pin_reset_member'
            )
            current_pin_status='Personal PIN set' if member_has_personal_pin(reset_mid) else 'Needs first-time PIN setup'
            st.caption(current_pin_status)
            confirm_reset=st.checkbox(
                'Confirm reset to first-time setup',
                key='confirm_member_pin_reset'
            )
            if st.button(
                'Reset Member PIN',
                disabled=not confirm_reset,
                use_container_width=True
            ):
                try:
                    reset_member_pin(reset_mid)
                    st.success('PIN reset. That member may use the affiliate first-time access code once, then create a new 4-digit PIN.')
                    st.rerun()
                except Exception as ex:
                    st.error(str(ex))

    with tabs[1]:
        st.subheader('Assign Executive Board Positions')
        members=philo_dropdown_people();mm={m['id']:m['full_name'] for m in members}
        current=safe_rows('officer_assignments',active=True)
        for pos in OFFICER_POSITIONS:
            cur=next((x for x in current if x.get('position')==pos),None)
            options=[0]+list(mm)
            selected=cur.get('member_id') if cur else 0
            mid=st.selectbox(pos,options,index=options.index(selected) if selected in options else 0,format_func=lambda x:'VACANT — President has acting access' if x==0 else mm[x],key=f"oa{pos}")
            email=st.text_input(f'{pos} position email',value=position_email(pos),key=f"oe{pos}")
            if st.button(f'Save {pos}',key=f"osave{pos}"):
                save_setting("position_email_"+pos.lower().replace(" ","_").replace("-","_"),email)
                if cur:table('officer_assignments').update({'active':False,'ended_at':date.today().isoformat()}).eq('id',cur['id']).execute()
                if mid:table('officer_assignments').insert({'position':pos,'member_id':mid,'member_name':mm[mid],'position_email':email,'active':True,'started_at':date.today().isoformat()}).execute()
                st.rerun()
    with tabs[2]:
        msel=st.selectbox('Member',list(mm),format_func=lambda x:mm[x],key='privm')
        priv=safe_rows('member_private',member_id=msel);priv=priv[0] if priv else {}
        is_philo=st.checkbox('Count this person as a Philo member',value=priv.get('is_philo',True) is not False)
        mtype=st.selectbox('Membership Type',['Returning','Inductee','Reactivating'],index=(['Returning','Inductee','Reactivating'].index(priv.get('member_type','Returning')) if priv.get('member_type','Returning') in ['Returning','Inductee','Reactivating'] else 0))
        pid=st.text_input('Philo ID Number',value=priv.get('philo_id') or '')
        note=st.text_area('Advisor-only notes',value=priv.get('advisor_notes') or '')
        if st.button('Save Private Member Information'):
            table('member_private').upsert({'member_id':msel,'philo_id':pid,'advisor_notes':note,'is_philo':is_philo,'member_type':mtype}).execute();st.success('Private information saved.')
    with tabs[3]:
        st.subheader('Advisor Chapter Report Generator')
        st.caption('Build a chapter-facing report from officer/committee reports and live Hub dashboard data.')

        report_type=st.selectbox(
            'Report Type',
            ['Quick Summary','Standard Chapter Report','Detailed Affiliate Report'],
            index=1,
            key='advisor_chapter_report_type_v32'
        )
        default_period=datetime.now().strftime('%B %Y')
        report_period=st.text_input(
            'Reporting Period',
            value=default_period,
            key='advisor_chapter_report_period_v32',
            help='Example: August 2026, 2026-2027, or End of Year 2026-2027.'
        )

        published_count=len(safe_rows('published_reports'))
        monthly_count=len(safe_rows('monthly_reports'))
        st.caption(f"Report sources currently on file: {published_count} published report(s) + {monthly_count} monthly report record(s), plus dashboard data.")

        if st.button('Generate Chapter Report Draft',use_container_width=True,key='generate_advisor_chapter_report_v32'):
            st.session_state['advisor_chapter_report_v32']=advisor_chapter_report_v32(report_period,report_type)

        chapter_draft=st.text_area(
            'Editable Chapter Report',
            value=st.session_state.get('advisor_chapter_report_v32',''),
            height=720,
            key='advisor_chapter_report_text_v32'
        )
        if chapter_draft.strip():
            chapter_pdf=report_pdf(
                report_type,
                f"{report_period} • Prepared by {member_name}, Philo Advisor",
                [('Chapter Report',chapter_draft)],
                member_name
            )
            st.download_button(
                'Download / Print Chapter Report',
                chapter_pdf,
                file_name=f"NBS_{report_type.replace(' ','_')}_{re.sub(r'[^A-Za-z0-9_-]+','_',report_period)}.pdf",
                mime='application/pdf',
                use_container_width=True
            )

    with tabs[4]:
        st.subheader('Member Login & Activity Log')
        st.caption('Duration is exact when a member signs out. If a browser is simply closed, duration is based on the last activity recorded by the Hub.')
        activity=login_activity_rows(750)

        if not activity:
            st.info('No login activity has been recorded yet.')
        else:
            activity_df=pd.DataFrame([{
                'Member':r.get('member_name'),
                'Account Type':r.get('account_type'),
                'Login':fmt_dt(r.get('login_at')),
                'Last Activity':fmt_dt(r.get('last_activity_at')),
                'Logout':fmt_dt(r.get('logout_at')),
                'Status':activity_status(r),
                'Active Duration':activity_duration_label(r),
                'Last Page':r.get('last_page') or ''
            } for r in activity])
            st.dataframe(activity_df,hide_index=True,use_container_width=True)
            st.download_button(
                'Download Activity Log CSV',
                activity_df.to_csv(index=False).encode('utf-8'),
                file_name='NBS_Member_Login_Activity.csv',
                mime='text/csv',
                use_container_width=True
            )

    with tabs[5]:
        atabs=st.tabs(['📣 Send Leadership Update','🆘 Officer Questions'])
        with atabs[0]:
            st.subheader('Send Leadership Update')
            st.caption('Send an update to the President, an officer position, a committee chair, a specific member, or all officers. The recipient can add it to her report, create a task/vote, or handle it separately.')

            recipient_type=st.selectbox(
                'Send To',
                ['Officer Position','Committee','Member','All Officers'],
                key='leadership_update_recipient_type'
            )
            recipient_value=''
            if recipient_type=='Officer Position':
                recipient_value=st.selectbox('Officer Position',OFFICER_POSITIONS,key='leadership_update_position')
            elif recipient_type=='Committee':
                cs=active_committees()
                cmap={c.get('id'):c.get('committee_name') for c in cs}
                if cmap:
                    cid=st.selectbox('Committee / Chair',list(cmap),format_func=lambda x:cmap[x],key='leadership_update_committee')
                    recipient_value=cmap[cid]
                else:
                    st.info('No active committees are available.')
            elif recipient_type=='Member':
                ms=active_members()
                mmap={m.get('id'):m.get('full_name') for m in ms}
                mid_=st.selectbox('Member',list(mmap),format_func=lambda x:mmap[x],key='leadership_update_member')
                recipient_value=str(mid_)
            else:
                recipient_value='All Officers'

            with st.form('admin_send_leadership_update',clear_on_submit=True):
                update_type=st.selectbox('Update Type',LEADERSHIP_UPDATE_TYPES)
                priority=st.selectbox('Priority',LEADERSHIP_UPDATE_PRIORITIES)
                subject=st.text_input('Subject')
                message=st.text_area('What do you need to relay?',height=160)
                suggested_action=st.text_area(
                    'Suggested Action / Guidance',
                    placeholder='Example: Chapter changed its meeting date. Please determine whether the Philo meeting should change too. If member approval is needed, create a vote and communicate the final date.'
                )
                report_section=st.selectbox('Suggested Report Section',['Updates','Action Items','Reminders','Upcoming Events','Accomplishments','Officer decides'])
                send_update=st.form_submit_button('Send Leadership Update',use_container_width=True)
            if send_update:
                if not message.strip():
                    st.warning('Enter the update you want to relay.')
                elif not recipient_value:
                    st.warning('Choose a recipient.')
                else:
                    table('leadership_updates').insert({
                        'recipient_type':recipient_type,'recipient_value':str(recipient_value),
                        'update_type':update_type,'priority':priority,'subject':subject.strip(),
                        'message':message.strip(),'suggested_action':suggested_action.strip(),
                        'suggested_report_section':report_section,'status':'Sent',
                        'created_by_member_id':member_id,'created_by_name':member_name,
                        'created_at':datetime.now(timezone.utc).isoformat(),
                        'updated_at':datetime.now(timezone.utc).isoformat()
                    }).execute()
                    st.success('Leadership update sent.')
                    st.rerun()

            st.markdown('### Sent Updates')
            sent=safe_rows('leadership_updates')
            if sent:
                st.dataframe(pd.DataFrame([{
                    'Recipient':leadership_update_recipient_label(r),
                    'Type':r.get('update_type'),
                    'Subject':r.get('subject'),
                    'Status':r.get('status'),
                    'Sent':fmt_dt(r.get('created_at'))
                } for r in sent[-50:][::-1]]),hide_index=True,use_container_width=True)

        with atabs[1]:
            st.subheader('Advisor Help Requests')

            reqs=advisor_help_rows()
            if not reqs:st.info('No officer help requests are currently open.')
            for r in reqs:
                with st.expander(f"{r.get('priority')} • {r.get('position')} • {r.get('requesting_member_name')} • {r.get('topic') or 'Question'}"):
                    st.write(r.get('question') or '')
                    reply=st.text_area('Advisor Response / Guidance',value=r.get('advisor_response') or '',key=f"advhelp_reply_{r['id']}")
                    status=st.selectbox('Status',['Open','Responded','Waiting for Officer','Closed'],index=['Open','Responded','Waiting for Officer','Closed'].index(r.get('status')) if r.get('status') in ['Open','Responded','Waiting for Officer','Closed'] else 0,key=f"advhelp_status_{r['id']}")
                    if st.button('Save Advisor Response',key=f"advhelp_save_{r['id']}",use_container_width=True):
                        table('advisor_help_requests').update({'advisor_response':reply.strip(),'status':status,'responded_by_member_id':member_id,'responded_by_name':member_name,'responded_at':datetime.now(timezone.utc).isoformat()}).eq('id',r['id']).execute();st.rerun()

    with tabs[6]:
        st.subheader('Grievance / Concern Review')
        gs=grievance_rows()
        if not gs:st.info('No grievances/concerns have been submitted.')
        for g in gs:
            with st.expander(f"{g.get('category')} • {g.get('submitted_by_name')} • {g.get('status')}"):
                st.write(f"**Subject:** {g.get('subject') or ''}")
                st.write(f"**Details:** {g.get('details') or ''}")
                guide=grievance_guidance(g.get('category'),g.get('details'))
                st.markdown('#### Generated Handling Guidance')
                for item in guide['guidance']:st.write(f"• {item}")
                st.write('**Suggested roles/authorities:** '+', '.join(guide['roles']))
                st.caption('Administrative guidance only. Governing documents and applicable law/policy control.')
                note=st.text_area('Leadership Review Notes',value=g.get('leadership_notes') or '',key=f"griev_note_{g['id']}")
                status=st.selectbox('Status',GRIEVANCE_STATUSES,index=GRIEVANCE_STATUSES.index(g.get('status')) if g.get('status') in GRIEVANCE_STATUSES else 0,key=f"griev_stat_{g['id']}")
                if st.button('Save Grievance Review',key=f"griev_save_{g['id']}",use_container_width=True):
                    table('grievances').update({'leadership_notes':note.strip(),'status':status,'reviewed_by_member_id':member_id,'reviewed_by_name':member_name,'reviewed_at':datetime.now(timezone.utc).isoformat()}).eq('id',g['id']).execute();st.rerun()

    with tabs[7]:
        for r in table('reimbursements').select('*').order('submitted_at',desc=True).execute().data or []:
            with st.expander(f"#{r['id']} • {r.get('submitted_by_name')} • ${float(r.get('amount') or 0):,.2f} • {r.get('status')}"):
                st.write(r.get('description'));st.write(f"Budget / Position: {r.get('budget_position') or 'General Chapter'}");st.write(f"Signed by: {r.get('signature_name')}")
                voucher_status_options=['Pending','Approved','Denied','Paid','Cancelled','Returned','Needs Correction']
                current_voucher_status=r.get('status') or 'Pending'
                if current_voucher_status not in voucher_status_options:
                    voucher_status_options.append(current_voucher_status)
                status=st.selectbox(
                    'Overall status',
                    voucher_status_options,
                    index=voucher_status_options.index(current_voucher_status),
                    key=f"arst{r['id']}"
                )
                admin_status_note=st.text_area(
                    'Reason / Note for Status',
                    value=r.get('status_reason_note') or '',
                    key=f"arsnote{r['id']}"
                )
                if st.button('Save Overall Status',key=f"ars{r['id']}"):
                    try:
                        if status=='Paid':
                            fresh=table('reimbursements').select('*').eq('id',r['id']).single().execute().data
                            if not fresh: raise RuntimeError('Voucher could not be reloaded before payment posting.')
                            ensure_paid_voucher_transaction(fresh)
                        table('reimbursements').update({'status':status,'status_reason_note':admin_status_note.strip()}).eq('id',r['id']).execute()
                        st.rerun()
                    except Exception as ex:
                        st.error(f'Status was not changed: {ex}')
    with tabs[8]:
        st.subheader('🧪 Test / Preview Mode')
        st.warning(
            'Test Mode is a read-only sandbox. Database inserts/updates/deletes, file uploads, Google/Drive actions, votes, messages, tasks, forms, and settings changes are blocked.'
        )
        st.markdown('### Candidate Preview')
        st.caption('Preview the prospect/candidate experience without creating a real Interest Profile.')
        if st.button('Open Candidate Test Portal',use_container_width=True,key='admin_open_candidate_preview'):
            st.session_state['test_preview_kind']='candidate'
            st.session_state.pop('test_preview_member_id',None)
            st.rerun()

        st.markdown('### Member Preview')
        st.caption('Choose an existing member and see the Hub with her permissions/navigation. Any actions you try are simulated and are not saved.')
        preview_members=active_members()
        pmap={m['id']:m['full_name'] for m in preview_members if m.get('id')!=member_id}
        if pmap:
            pmid=st.selectbox('Preview As Member',list(pmap),format_func=lambda x:pmap[x],key='admin_preview_member_select')
            if st.button('Open Member Test Mode',use_container_width=True,key='admin_open_member_preview'):
                st.session_state['test_preview_kind']='member'
                st.session_state['test_preview_member_id']=pmid
                st.session_state.pop('test_preview_last_action',None)
                st.rerun()
        else:
            st.info('No other active members are available to preview.')

        st.markdown('### What Test Mode Blocks')
        st.markdown(
            '• Supabase inserts, updates, deletes, and upserts  \n'
            '• File/photo/document uploads  \n'
            '• File deletion  \n'
            '• Google/Drive write actions  \n'
            '• Votes and vote counts  \n'
            '• Messages, tasks, vouchers, forms, reports, and status changes  \n'
            '• Settings changes'
        )
        st.caption('Reads use the live Hub so you can see the same real menus/content a member would see, but Test Mode does not alter that information.')

    with tabs[9]:
        fields={
            'google_calendar_link':'Shared Google Calendar Link or ID',
            'google_calendar_id':'Parsed Google Calendar ID',
            'google_drive_folder_id':'General Google Drive Folder ID',
            'reports_drive_folder_id':'Monthly Reports Google Drive Folder ID',
            'officer_reports_drive_folder_id':'Officer Reports Google Drive Folder ID',
            'committee_reports_drive_folder_id':'Committee Reports Google Drive Folder ID',
            'bylaws_url':'Current Bylaws URL',
            'protocol_url':'Current SOP / Procedures URL',
            'member_card_payment_url':'Member Card Payment Page URL',
            'member_paypal_payment_url':'Member PayPal Payment Page URL',
            'affiliate_name':'Affiliate Name for Remittance Forms',
            'affiliate_address':'Affiliate Mailing Address',
            'affiliate_city_state_zip':'Affiliate City, State, Zip',
            'philo_advisor_name':'Philo Advisor Name',
            'philo_advisor_contact':'Philo Advisor Email & Phone',
            'ner_remittance_template_url':'NER Remittance Google Sheet Template URL',
            'ner_remittance_sheet_name':'NER Remittance Sheet Tab Name',
            'ner_remittance_completed_folder_id':'Completed NER Remittance Drive Folder ID',
            'demo_admin_display_name':'Demo Admin Login Name',
            'demo_admin_code':'Demo Admin Access Code',
            'demo_admin_enabled':'Demo Admin Enabled (true/false)'
        }
        setting_defaults={
            'officer_reports_drive_folder_id':'1k0KcbEbDF87YP76OXJG7vwthKQVxOWlR',
            'committee_reports_drive_folder_id':'15IfzQd916c8xbN4cYSXsS3qMYWJ4F8LA',
            'demo_admin_display_name':'Demo Admin',
            'demo_admin_code':'1922',
            'demo_admin_enabled':'true'
        }
        vals={
            k:st.text_input(
                label,
                value=setting(k,setting_defaults.get(k,'')),
                key=f"set2{k}"
            )
            for k,label in fields.items()
        }
        if st.button('Save App Settings'):
            for k,v in vals.items():save_setting(k,v.strip())
            st.success('Settings saved.')

        st.divider()
        st.markdown('### Remove Testing / Demo Records')
        st.caption(
            'Removes obvious setup records containing Test, Testing, Demo, Sample, or Dummy '
            'from member-facing reports, events, communications, and similar setup data.'
        )
        confirm_cleanup=st.checkbox(
            'I want to permanently remove testing/demo records.',
            key='confirm_testing_cleanup_v35'
        )
        if st.button(
            'Delete Testing / Demo Records',
            disabled=not confirm_cleanup,
            use_container_width=True,
            key='delete_testing_records_v35'
        ):
            removed=cleanup_testing_records()
            if removed:
                st.success(f"Removed {len(removed)} testing/demo record(s).")
                with st.expander('Removed Records'):
                    st.write("\n".join(removed))
            else:
                st.info('No obvious testing/demo records were found.')
        if not google_connected():
            u=google_connect_url() if google_oauth_cfg() else ''
            if u:st.link_button('Connect Advisor Google Account',u,use_container_width=True)
        else:
            st.success('Google connection is active for this browser session.')
            st.caption('If Drive/Sheets says the connection has insufficient permissions, disconnect/reconnect the Advisor Google account once so the updated Drive + Sheets permissions are granted.')

# ============================================================
# GLOBAL BOTTOM CONTROLS
# ============================================================
st.markdown('<div class="pearl-divider"></div>',unsafe_allow_html=True)
st.markdown('### App Controls')

bottom_app,bottom_signout=st.columns(2)

with bottom_app:
    if st.button(
        '📱 Add NBS Philo Hub to Phone',
        key='bottom_add_home_screen',
        use_container_width=True
    ):
        st.session_state['show_bottom_install_help']=not st.session_state.get(
            'show_bottom_install_help',False
        )

with bottom_signout:
    if st.button(
        'Sign Out',
        key='bottom_signout',
        use_container_width=True
    ):
        if not st.session_state.get('demo_admin_mode'):
            close_login_activity()
        st.session_state.clear()
        st.rerun()

if st.session_state.get('show_bottom_install_help'):
    with st.container(border=True):
        st.markdown('#### Add NBS Philo Hub to Your Home Screen')
        st.markdown(
            '**iPhone / iPad:** Open the Hub in Safari → tap **Share** → '
            '**Add to Home Screen** → **Add**.\\n\\n'
            '**Android:** Open the Hub in Chrome → tap the **⋮** menu → '
            '**Add to Home screen** or **Install app**.'
        )
        st.caption('The saved home-screen icon uses the NBS Philo chapter logo.')

